#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
文件脱敏工具 - 支持 Excel (.xlsx/.xls) 和 Word (.docx)
脱敏内容：手机号、身份证、邮箱、固定电话、银行卡、公司名称（模糊匹配）
用法：python desensitize.py <文件路径>
输出：原文件名_脱敏.xlsx (若输入为 .xls，输出自动转为 .xlsx)
"""

import re
import sys
import os
from pathlib import Path

# ---------- 依赖检查 ----------
try:
    import openpyxl
except ImportError:
    print("缺少依赖 openpyxl，请安装：pip install openpyxl")
    sys.exit(1)

try:
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("缺少依赖 python-docx，请安装：pip install python-docx")
    sys.exit(1)

try:
    import pandas as pd
    import xlrd  # pandas 读取 .xls 依赖 xlrd
except ImportError:
    print("缺少依赖 pandas 或 xlrd，请安装：pip install pandas xlrd")
    sys.exit(1)

# ---------- 正则表达式 ----------
PHONE_PATTERN = re.compile(r'(?<!\d)1[3-9]\d{9}(?!\d)')
ID_PATTERN = re.compile(r'\b[1-9]\d{5}(18|19|20)\d{2}(0[1-9]|1[0-2])(0[1-9]|[12]\d|3[01])\d{3}[\dXx]\b')
EMAIL_PATTERN = re.compile(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b')
PHONE_FIXED_PATTERN = re.compile(r'\b0\d{2,3}-\d{7,8}\b')
BANKCARD_PATTERN = re.compile(r'\b\d{4}[\s-]?\d{4}[\s-]?\d{4}[\s-]?\d{4}[\s-]?\d{0,3}\b')

# ---------- 脱敏函数 ----------
def mask_phone(match):
    s = match.group()
    return s[:3] + '****' + s[-4:]

def mask_id(match):
    s = match.group()
    return s[:3] + '***********' + s[-4:]

def mask_email(match):
    s = match.group()
    local, domain = s.split('@')
    if len(local) <= 3:
        masked_local = local[0] + '**' if len(local) > 1 else '**'
    else:
        masked_local = local[:3] + '***'
    return masked_local + '@' + domain

def mask_fixed_phone(match):
    s = match.group()
    parts = s.split('-')
    if len(parts) == 2:
        return parts[0] + '-****' + parts[1][-4:]
    return s

def mask_bankcard(match):
    s = re.sub(r'[\s-]', '', match.group())
    if len(s) >= 16:
        return s[:4] + '********' + s[-4:]
    return s

# 组合规则
REPLACE_RULES = [
    (PHONE_PATTERN, mask_phone),
    (ID_PATTERN, mask_id),
    (EMAIL_PATTERN, mask_email),
    (PHONE_FIXED_PATTERN, mask_fixed_phone),
    (BANKCARD_PATTERN, mask_bankcard),
]

# ---------- 公司名称模糊匹配 ----------
# 匹配2~20个中文字符，以常见企业后缀结尾
COMPANY_PATTERN = re.compile(
    r'([\u4e00-\u9fa5]{2,20}'
    r'(?:有限|集团|股份|合伙|事务所|中心|研究院|大学|医院|公司|厂|社|银行|保险|证券|基金|信托|协会|联合会|委员会))'
)

def mask_company_fuzzy(text):
    """模糊化公司名称：长度≤4保留首尾加*，>4保留前2和后2加***"""
    if not isinstance(text, str):
        return text

    def replace(match):
        name = match.group(1)
        length = len(name)
        if length <= 2:
            return name
        elif length <= 4:
            return name[0] + '*' * (length - 2) + name[-1]
        else:
            return name[:2] + '***' + name[-2:]

    return COMPANY_PATTERN.sub(replace, text)

# ---------- 主脱敏函数 ----------
def mask_text(text):
    """对文本应用所有脱敏规则（公司名称→其他）"""
    if not isinstance(text, str):
        return text
    # 先处理公司名称
    text = mask_company_fuzzy(text)
    # 再处理其他敏感信息
    for pattern, repl_func in REPLACE_RULES:
        text = pattern.sub(repl_func, text)
    return text

def mask_cell(val):
    """处理单元格值（支持 NaN/None）"""
    if pd.isna(val):
        return val
    return mask_text(str(val))

# ---------- 处理 Excel ----------
def process_excel(file_path, output_path):
    ext = Path(file_path).suffix.lower()
    
    if ext == '.xls':
        # 使用 pandas 读取 .xls (所有工作表)
        try:
            df_dict = pd.read_excel(file_path, sheet_name=None, engine='xlrd')
        except Exception as e:
            print(f"❌ 读取 .xls 文件失败: {e}")
            print("提示：请确认文件未损坏，或尝试用 Excel 打开后另存为 .xlsx")
            return
        
        # 将输出路径改为 .xlsx
        output_path = Path(output_path).with_suffix('.xlsx')
        
        with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
            for sheet_name, df in df_dict.items():
                # 对所有单元格脱敏
                df_masked = df.applymap(mask_cell)
                df_masked.to_excel(writer, sheet_name=sheet_name, index=False)
        
        print(f"✅ Excel 脱敏完成（已转为 .xlsx）: {output_path}")
    
    elif ext == '.xlsx':
        # 使用 openpyxl 处理 .xlsx（保留格式）
        wb = openpyxl.load_workbook(file_path, data_only=True)
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value and isinstance(cell.value, str):
                        cell.value = mask_text(cell.value)
        wb.save(output_path)
        print(f"✅ Excel 脱敏完成: {output_path}")
    
    else:
        print(f"❌ 不支持的文件扩展名: {ext}，仅支持 .xls 和 .xlsx")

# ---------- 处理 Word ----------
def process_paragraphs(paragraphs):
    for para in paragraphs:
        full_text = para.text
        if full_text.strip():
            masked = mask_text(full_text)
            if masked != full_text:
                para.clear()
                para.add_run(masked)

def process_word(file_path, output_path):
    doc = Document(file_path)
    # 正文段落
    process_paragraphs(doc.paragraphs)
    # 表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                process_paragraphs(cell.paragraphs)
    # 页眉页脚
    for section in doc.sections:
        process_paragraphs(section.header.paragraphs)
        process_paragraphs(section.footer.paragraphs)
    doc.save(output_path)
    print(f"✅ Word 脱敏完成: {output_path}")

# ---------- 主入口 ----------
def desensitize_file(file_path):
    if not os.path.exists(file_path):
        print(f"❌ 错误: 文件不存在 - {file_path}")
        return

    p = Path(file_path)
    ext = p.suffix.lower()
    
    # 输出路径：若输入是 .xls，输出扩展名会在 process_excel 中改为 .xlsx
    output_path = p.parent / f"{p.stem}_脱敏{p.suffix}"

    if ext in ['.xlsx', '.xls']:
        process_excel(file_path, output_path)
    elif ext == '.docx':
        process_word(file_path, output_path)
    else:
        print(f"❌ 不支持的文件类型: {ext}，目前仅支持 .xlsx, .xls, .docx")

if __name__ == '__main__':
    if len(sys.argv) < 2:
        print("用法: python desensitize.py <文件路径>")
        print("示例: python desensitize.py D:/data/客户信息.xlsx")
        print("      python desensitize.py D:/data/客户信息.xls (自动转为 .xlsx)")
        sys.exit(1)
    file_path = sys.argv[1]
    desensitize_file(file_path)