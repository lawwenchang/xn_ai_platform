#!/usr/bin/env python3
"""Markdown 文档引擎: 解析/模板/生成/改写/转换"""
from __future__ import annotations
import json, re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

@dataclass
class MdNode:
    type: str = ""; level: int = 0; text: str = ""; raw: str = ""
    children: List[MdNode] = field(default_factory=list)
    meta: Dict = field(default_factory=dict)

@dataclass
class MdDocument:
    path: str = ""; title: str = ""
    nodes: List[MdNode] = field(default_factory=list)
    front_matter: Dict = field(default_factory=dict)
    toc: List[Tuple[int, str]] = field(default_factory=list)
    def headings(self): return [n for n in self.nodes if n.type == "heading"]
    def tables(self): return [n for n in self.nodes if n.type == "table"]
    def to_text(self): return "\n\n".join(n.raw for n in self.nodes if n.raw.strip())

def parse_md(filepath: str) -> MdDocument:
    return parse_md_text(Path(filepath).read_text(encoding="utf-8"), filepath)

def parse_md_text(text: str, source: str = "") -> MdDocument:
    doc = MdDocument(path=source); lines = text.split("\n")
    if lines and lines[0].strip() == "---":
        end = 1
        while end < len(lines) and lines[end].strip() != "---": end += 1
        for line in "\n".join(lines[1:end]).split("\n"):
            if ":" in line: k, v = line.split(":", 1); doc.front_matter[k.strip()] = v.strip()
        lines = lines[end + 1:]
    i = 0
    while i < len(lines):
        line = lines[i]
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            lv = len(m.group(1)); t = m.group(2).strip()
            doc.toc.append((lv, t)); doc.nodes.append(MdNode(type="heading", level=lv, text=t, raw=line))
            if lv == 1 and not doc.title: doc.title = t
            i += 1; continue
        if line.strip().startswith("```"):
            buf = [line]; i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"): buf.append(lines[i]); i += 1
            if i < len(lines): buf.append(lines[i]); i += 1
            doc.nodes.append(MdNode(type="code", raw="\n".join(buf), meta={"lang": line.strip()[3:] or "text"}))
            continue
        if line.strip().startswith("|"):
            buf = [line]; i += 1
            while i < len(lines) and lines[i].strip().startswith("|"): buf.append(lines[i]); i += 1
            doc.nodes.append(MdNode(type="table", raw="\n".join(buf))); continue
        if re.match(r'^\s*[-*+]\s+', line) or re.match(r'^\s*\d+[.)]\s+', line):
            buf = [line]; i += 1
            while i < len(lines) and (re.match(r'^\s*[-*+]\s+', lines[i]) or re.match(r'^\s*\d+[.)]\s+', lines[i]) or (lines[i].strip() and re.match(r'^\s{2,}', lines[i]))):
                buf.append(lines[i]); i += 1
            doc.nodes.append(MdNode(type="list", raw="\n".join(buf))); continue
        if line.strip().startswith(">"):
            buf = [line]; i += 1
            while i < len(lines) and lines[i].strip().startswith(">"): buf.append(lines[i]); i += 1
            doc.nodes.append(MdNode(type="blockquote", raw="\n".join(buf))); continue
        if re.match(r'^[-*_]{3,}\s*$', line.strip()):
            doc.nodes.append(MdNode(type="hr", raw=line)); i += 1; continue
        if line.strip():
            buf = [line]; i += 1
            while i < len(lines) and lines[i].strip() and not re.match(r'^(#{1,6}\s|```|\||[-*+]\s|\d+[.)]\s|>)', lines[i]):
                buf.append(lines[i]); i += 1
            doc.nodes.append(MdNode(type="paragraph", text=" ".join(p.strip() for p in buf), raw="\n".join(buf)))
        else: i += 1
    return doc

def chunk_by_headings(filepath: str) -> List[dict]:
    """按标题将 Markdown 文档切分为 chunks，供 RAG 索引使用。异常安全。"""
    try:
        doc = parse_md(filepath)
    except Exception as e:
        print(f"[MD] chunk_by_headings 解析失败 {filepath}: {e}")
        return []
    chunks = []
    cur_h = doc.title or Path(filepath).stem
    cur_t = ""
    for node in doc.nodes:
        if node.type == "heading":
            if cur_t.strip():
                chunks.append({
                    "text": cur_t.strip(),
                    "source": str(Path(filepath).name),
                    "category": "md",
                    "heading": cur_h,
                })
            cur_h = node.text.strip() or cur_h
            # 把标题本身也纳入下一个 chunk，避免标题信息丢失
            cur_t = node.text.strip()
        else:
            raw = getattr(node, "raw", "") or getattr(node, "text", "")
            if raw.strip():
                cur_t += "\n\n" + raw.strip() if cur_t else raw.strip()
    if cur_t.strip():
        chunks.append({
            "text": cur_t.strip(),
            "source": str(Path(filepath).name),
            "category": "md",
            "heading": cur_h,
        })
    return chunks


# ═══════════════ 模板填充 ═══════════════

def fill_template(tmpl: str, vars: Dict[str, str], out: str = "") -> str:
    t = Path(tmpl).read_text(encoding="utf-8")
    for k, v in vars.items(): t = t.replace("{{ " + k + " }}", str(v)).replace("{{" + k + "}}", str(v))
    if out: Path(out).write_text(t, encoding="utf-8")
    return t

def fill_template_text(text: str, vars: Dict[str, str]) -> str:
    for k, v in vars.items(): text = text.replace("{{ " + k + " }}", str(v)).replace("{{" + k + "}}", str(v))
    return text

def intelligent_fill(
    template_path: str,
    instruction: str,
    data: Dict[str, Any] = None,
    out: str = "",
    out_format: str = "md"
) -> str:
    """
    智能填充：上传任意 md 文档 + 自然语言指令 + 数据 → LLM 理解上下文后自动填充。

    不需要在模板里预置 {{ }} 占位符。LLM 会理解文档结构和用户需求，
    自行判断哪些位置需要填入什么内容。

    Args:
        template_path: md 模板路径
        instruction: 自然语言指令（"把第三章的审计结论填上，数据用下面的差异明细"）
        data: 结构化数据字典（可选）
        out: 输出路径（不传则只返回文本）
        out_format: 输出格式 "md" / "docx" / "html"

    Returns:
        填充后的文本（或输出文件路径）

    例：
        intelligent_fill(
            "templates/审计报告模板.md",
            "这是一份医保回款审计报告。请根据数据填写第三章的审计结论和第四章的差异明细表。措辞要符合事务所风格。",
            {"差异明细": [...], "审计结论": "经核对，差异在可接受范围内"},
            out="outputs/报告.md"
        )
    """
    template_text = Path(template_path).read_text(encoding="utf-8")
    return intelligent_fill_text(template_text, instruction, data, out, out_format)


def intelligent_fill_text(
    template_text: str,
    instruction: str,
    data: Dict[str, Any] = None,
    out: str = "",
    out_format: str = "md"
) -> str:
    """同上，传文本而非文件路径"""
    data_str = json.dumps(data or {}, ensure_ascii=False, indent=2)

    prompt = f"""你是资深审计文档撰写专家。请根据用户指令，将数据填入下方的 Markdown 模板中。

【用户指令】
{instruction}

【要填入的数据】
{data_str}

【原始模板】
{template_text}

【要求】
1. 保持模板的整体结构和格式不变
2. 根据指令和数据，填写/完善模板中需要补充的内容
3. 如果指令说了要填哪些章节，就填哪些；没说到的部分保留原样
4. 数据中的表格用 Markdown 表格格式呈现
5. 数字保留两位小数
6. 只输出填充后的完整 Markdown，不要任何解释

【填充后的文档】"""

    try:
        import httpx, os
        r = httpx.post(
            os.environ.get("VLLM_TUNNEL_URL", "http://localhost:18000/v1/chat/completions"),
            headers={"Authorization": "Bearer EMPTY"},
            json={
                "model": os.environ.get("VLLM_MODEL", "qwen3-235b"),
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": max(len(template_text) * 2, 4096),
            },
            timeout=120,
        )
        r.raise_for_status()
        result = r.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        print(f"[MD] 智能填充失败，返回原文: {e}")
        result = template_text

    # 格式转换
    if out:
        _save_with_format(result, out, out_format)
    return result


def _save_with_format(text: str, out: str, fmt: str):
    """按指定格式保存"""
    Path(out).parent.mkdir(parents=True, exist_ok=True)

    if fmt == "docx":
        tmp_md = out + ".tmp.md"
        Path(tmp_md).write_text(text, encoding="utf-8")
        md_to_docx(tmp_md, out)
        Path(tmp_md).unlink()
    elif fmt == "html":
        html = md_to_html_text(text)
        Path(out).write_text(html, encoding="utf-8")
    else:  # md
        Path(out).write_text(text, encoding="utf-8")


def md_to_html_text(text: str) -> str:
    """md 文本转 HTML（不读文件）"""
    try:
        import markdown
        return markdown.markdown(text, extensions=["tables", "fenced_code"])
    except Exception:
        return text

# ═══════════════ 生成 md 报告（智能版） ═══════════════

def generate_report(description: str, data: Dict[str, Any], template: str = "", out: str = "") -> str:
    """
    智能生成 md 审计报告。两种模式：

    模式一（有模板）：template 指向 md 模板文件，用 {{ key }} 占位。
        generate_report("医保回款核对", data, template="templates/医保底稿.md")

    模式二（无模板）：LLM 根据描述 + 数据动态生成报告结构。
        generate_report("帮我生成医保回款核对报告，包含匹配结果和差异明细", data)

    Args:
        description: 自然语言描述报告需求
        data: 结构化数据（匹配结果、差异明细等）
        template: md 模板路径（可选）
        out: 输出路径（可选）
    """
    # 模式一：有模板就用模板
    if template and Path(template).exists():
        return fill_template(template, _flatten_data(data), out)

    # 模式二：LLM 动态生成
    now = __import__("datetime").datetime.now().strftime("%Y-%m-%d")
    data_str = json.dumps(data, ensure_ascii=False, indent=2)

    prompt = f"""你是资深审计报告撰写专家。根据以下数据和用户需求，生成一份专业的 Markdown 审计报告。

【用户需求】
{description}

【数据】
{data_str}

【要求】
1. 包含标题、日期（{now}）、章节结构
2. 关键数字用表格展示
3. 差异/异常用醒目标注
4. 语言专业、符合审计报告风格
5. 只输出 Markdown，不要解释

【报告】"""

    try:
        import httpx, os
        r = httpx.post(
            os.environ.get("VLLM_TUNNEL_URL", "http://localhost:18000/v1/chat/completions"),
            headers={"Authorization": "Bearer EMPTY"},
            json={
                "model": os.environ.get("VLLM_MODEL", "qwen3-235b"),
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.3,
                "max_tokens": 4096,
            },
            timeout=120,
        )
        r.raise_for_status()
        result = r.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        print(f"[MD] LLM 生成失败，降级为纯数据导出: {e}")
        result = _fallback_report(data, now, description)

    if out:
        Path(out).write_text(result, encoding="utf-8")
    return result


def _flatten_data(data: Dict, prefix: str = "") -> Dict[str, str]:
    """将嵌套数据扁平化为 {{ key }} 可用的字符串字典"""
    flat = {}
    for k, v in data.items():
        key = f"{prefix}.{k}" if prefix else k
        if isinstance(v, dict):
            flat.update(_flatten_data(v, key))
        elif isinstance(v, list):
            if v and isinstance(v[0], dict):
                # 列表内是字典，转 markdown 表格
                cols = list(v[0].keys())
                rows = ["| " + " | ".join(cols) + " |", "|" + "|".join(["------"] * len(cols)) + "|"]
                for item in v[:50]:
                    rows.append("| " + " | ".join(str(item.get(c, "")) for c in cols) + " |")
                flat[key] = "\n".join(rows)
            else:
                flat[key] = ", ".join(str(x) for x in v[:20])
        else:
            flat[key] = str(v)
    return flat


def _fallback_report(data: Dict, now: str, desc: str = "") -> str:
    """LLM 不可用时的降级纯数据导出"""
    lines = [f"# 审计报告", f"**日期**: {now}", f"**需求**: {desc}", "", "## 数据"]
    flat = _flatten_data(data)
    for k, v in flat.items():
        lines.append(f"- **{k}**: {v}")
    return "\n".join(lines)

# ═══════════════ LLM 改写 ═══════════════

def rewrite(filepath: str, instruction: str = "优化措辞", out: str = "") -> str:
    return rewrite_text(Path(filepath).read_text(encoding="utf-8"), instruction, out)

def rewrite_text(text: str, instruction: str = "优化措辞", out: str = "") -> str:
    prompt = f"你是审计质控专家。按指令改写文档，保持Markdown格式。\n\n【指令】{instruction}\n\n【原文】\n{text}\n\n【改写后】"
    try:
        import httpx, os
        r = httpx.post(os.environ.get("VLLM_TUNNEL_URL","http://localhost:18000/v1/chat/completions"),
            headers={"Authorization":"Bearer EMPTY"},
            json={"model":os.environ.get("VLLM_MODEL","qwen3-235b"),"messages":[{"role":"user","content":prompt}],"temperature":0.3,"max_tokens":4096}, timeout=60)
        r.raise_for_status()
        result = r.json()["choices"][0]["message"]["content"].strip()
    except Exception as e:
        print(f"[MD] LLM改写失败:{e}"); result = text
    if out: Path(out).write_text(result, encoding="utf-8")
    return result

# ═══════════════ 格式转换 ═══════════════

def _set_run_font(run, name: str = "宋体", size_pt: int = None):
    """设置 Run 字体，同时指定西文字体与东亚（中文）字体，避免中文版 Office 失效。"""
    try:
        from docx.oxml.ns import qn
        run.font.name = name
        rpr = run._element.rPr
        if rpr is not None:
            rFonts = rpr.rFonts
            if rFonts is None:
                from docx.oxml import OxmlElement
                rFonts = OxmlElement("w:rFonts")
                rpr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), name)
        if size_pt is not None:
            from docx.shared import Pt
            run.font.size = Pt(size_pt)
    except Exception:
        pass


def _set_style_font(doc, style_name: str, name: str, size_pt: int = None):
    """设置样式默认字体（含中文字体）。"""
    try:
        from docx.oxml.ns import qn
        style = doc.styles[style_name]
        style.font.name = name
        rpr = style._element.rPr
        if rpr is not None:
            rFonts = rpr.rFonts
            if rFonts is None:
                from docx.oxml import OxmlElement
                rFonts = OxmlElement("w:rFonts")
                rpr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), name)
        if size_pt is not None:
            from docx.shared import Pt
            style.font.size = Pt(size_pt)
    except Exception:
        pass


def _apply_table_borders(table, color: str = "000000", size: str = "4"):
    """为表格设置显式边框，避免中文版 Office 中样式失效导致无框线。"""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tbl = table._tbl
        tblPr = tbl.tblPr
        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        tblBorders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            tblBorders.append(el)
        tblPr.append(tblBorders)
    except Exception:
        pass


def md_to_docx(md_path: str, out: str = "", template: str = "") -> str:
    text = Path(md_path).read_text(encoding="utf-8")
    if not out: out = str(Path(md_path).with_suffix(".docx"))
    try:
        from docx import Document; from docx.shared import Pt
        doc = Document(template) if template else Document()
        _set_style_font(doc, "Normal", "宋体", 11)
        for lvl in range(1, 4):
            _set_style_font(doc, f"Heading {lvl}", "黑体")
        for node in parse_md_text(text).nodes:
            if node.type == "heading":
                h = doc.add_heading(node.text, level=min(node.level, 3))
                for r in h.runs:
                    _set_run_font(r, "黑体")
            elif node.type == "paragraph":
                p = doc.add_paragraph(node.text)
                for r in p.runs:
                    _set_run_font(r, "宋体", 11)
            elif node.type == "table":
                _tbl(doc, node.raw)
            elif node.type == "list":
                for li in node.raw.split("\n"):
                    c = re.sub(r'^[\s\-*+\d.)]+\s*', '', li).strip()
                    if c:
                        p = doc.add_paragraph(c, style="List Bullet")
                        for r in p.runs:
                            _set_run_font(r, "宋体", 11)
            elif node.type == "code":
                p = doc.add_paragraph(); r = p.add_run(node.raw)
                r.font.name = "Consolas"; r.font.size = Pt(9)
            elif node.type == "blockquote":
                p = doc.add_paragraph(node.raw.replace("> ", ""), style="Quote")
                for r in p.runs:
                    _set_run_font(r, "宋体", 11)
        doc.save(out)
    except Exception as e: print(f"[MD] docx:{e}")
    return out


def md_to_html(md_path: str) -> str:
    try:
        import markdown
        return markdown.markdown(Path(md_path).read_text(encoding="utf-8"), extensions=["tables","fenced_code"])
    except Exception: return ""


def _tbl(doc, raw):
    rows = [r for r in raw.split("\n") if r.strip().startswith("|")]
    rows = [r for r in rows if not re.match(r'^\|[\s\-:|]+\|$', r)]
    if not rows: return
    cells = [[c.strip() for c in r.split("|")[1:-1]] for r in rows]
    n_cols = max(len(r) for r in cells)
    tbl = doc.add_table(rows=len(cells), cols=n_cols)
    tbl.style = "Table Grid"
    _apply_table_borders(tbl)
    for ri, rc in enumerate(cells):
        for ci in range(n_cols):
            cell = tbl.cell(ri, ci)
            cell.text = rc[ci] if ci < len(rc) else ""
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_run_font(r, "宋体", 11)

