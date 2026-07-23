#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
智能工具箱 (toolbox.py)
========================
所有工具函数统一注册于此。Tool-Use Agent 根据用户指令自行选择调用。

安全边界：
- 所有文件操作限制在 SANDBOX_ROOT 下
- 沙箱执行器做二次安全校验
- 不直接暴露系统路径给 LLM
"""
from __future__ import annotations
import json, os, re
from pathlib import Path
from typing import Any, Dict, List, Optional

SANDBOX_ROOT = Path(__file__).resolve().parent.parent / "data"
SANDBOX_ROOT.mkdir(parents=True, exist_ok=True)

# ═══════════════ 工具注册表 ═══════════════

TOOLS = [
    {
        "type": "function",
        "function": {
            "name": "read_file",
            "description": "读取用户上传的任意格式文件内容（docx/xlsx/md/txt/csv/pdf），返回文本。只能读取沙箱内的文件。",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "相对于沙箱根目录的文件路径，如 'uploads/报告.md'"}
                },
                "required": ["path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "parse_md_structure",
            "description": "解析md文档的结构：章节标题树、表格数量、段落数。",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "md文件路径"}
                },
                "required": ["path"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "extract_section",
            "description": "从md文档中提取指定标题下的内容。",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "md文件路径"},
                    "heading": {"type": "string", "description": "要提取的章节标题，如'审计结论'"}
                },
                "required": ["path", "heading"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "fill_template",
            "description": "用数据填充Word(.docx)或Excel(.xlsx)模板，保留模板的全部格式（字体/页眉页脚/边框/命名区域）。",
            "parameters": {
                "type": "object",
                "properties": {
                    "template_path": {"type": "string", "description": "模板文件路径"},
                    "data_json": {"type": "string", "description": "要填入的数据，JSON字符串"}
                },
                "required": ["template_path", "data_json"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "normalize_format",
            "description": "以一个文件为格式模板，批量统一其他文件的格式（Excel/Word）。只改格式，不改数据。",
            "parameters": {
                "type": "object",
                "properties": {
                    "template_path": {"type": "string", "description": "格式模板文件路径"},
                    "target_paths": {"type": "string", "description": "要调整格式的文件路径列表，JSON数组"}
                },
                "required": ["template_path", "target_paths"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "generate_report",
            "description": "根据结构化数据和用户需求，生成Markdown格式的审计报告。",
            "parameters": {
                "type": "object",
                "properties": {
                    "description": {"type": "string", "description": "用户对报告的需求描述"},
                    "data_json": {"type": "string", "description": "报告所需数据，JSON字符串"}
                },
                "required": ["description", "data_json"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "rewrite_document",
            "description": "改写文档内容。保持Markdown格式，按用户指令修改措辞/格式/风格。",
            "parameters": {
                "type": "object",
                "properties": {
                    "text": {"type": "string", "description": "要改写的文档原文"},
                    "instruction": {"type": "string", "description": "改写指令，如'改得更正式'、'数字统一用千分位'"}
                },
                "required": ["text", "instruction"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "convert_format",
            "description": "在md、docx、html之间转换文档格式。",
            "parameters": {
                "type": "object",
                "properties": {
                    "path": {"type": "string", "description": "源文件路径"},
                    "target_format": {"type": "string", "enum": ["docx", "html", "md"]}
                },
                "required": ["path", "target_format"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "merge_documents",
            "description": "将多份md文档合并为一份。",
            "parameters": {
                "type": "object",
                "properties": {
                    "paths_json": {"type": "string", "description": "要合并的文件路径列表，JSON数组"},
                    "title": {"type": "string", "description": "合并后的文档标题"}
                },
                "required": ["paths_json", "title"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "diff_documents",
            "description": "对比两份文档内容的差异。",
            "parameters": {
                "type": "object",
                "properties": {
                    "path1": {"type": "string", "description": "第一份文件路径"},
                    "path2": {"type": "string", "description": "第二份文件路径"}
                },
                "required": ["path1", "path2"]
            }
        }
    },
    {
        "type": "function",
        "function": {
            "name": "search_knowledge",
            "description": "从审计知识库（中注协准则/企业会计准则/法律法规）中检索相关内容。用户问准则、法规、专业知识时使用。",
            "parameters": {
                "type": "object",
                "properties": {
                    "query": {"type": "string", "description": "检索查询，如'函证准则 应收账款 金额标准'"}
                },
                "required": ["query"]
            }
        }
    },
]



def _safe_path(path: str) -> Path:
    p = SANDBOX_ROOT / path
    if not str(p.resolve()).startswith(str(SANDBOX_ROOT.resolve())):
        raise PermissionError(f"禁止访问沙箱外: {path}")
    return p.resolve()

def execute_tool(name: str, args: dict) -> dict:
    try:
        if name == "read_file":
            p = _safe_path(args["path"]); ext = p.suffix.lower()
            if ext in (".md", ".txt"): text = p.read_text(encoding="utf-8")
            elif ext in (".docx",):
                from docx import Document
                text = "\n".join(pa.text for pa in Document(str(p)).paragraphs)
            elif ext in (".xlsx", ".xls"):
                import pandas as pd; text = pd.read_excel(str(p)).to_string()
            elif ext == ".csv":
                import pandas as pd; text = pd.read_csv(str(p)).to_string()
            else: text = p.read_text(encoding="utf-8")
            return {"ok": True, "data": text[:8000], "format": ext}

        elif name == "parse_md_structure":
            p = _safe_path(args["path"])
            from core.md_engine import parse_md
            d = parse_md(str(p))
            return {"ok": True, "data": {"title": d.title, "toc": [{"l": lv, "t": t} for lv, t in d.toc], "tables": len(d.tables())}}

        elif name == "extract_section":
            p = _safe_path(args["path"]); h = args["heading"]
            from core.md_engine import parse_md
            d = parse_md(str(p))
            found = [n for n in d.nodes if n.type=="heading" and h in n.text]
            if not found: return {"ok": False, "error": f"未找到'{h}'"}
            idx = d.nodes.index(found[0]); content = []
            for n in d.nodes[idx:]:
                if n.type=="heading" and n!=found[0] and n.level<=found[0].level: break
                content.append(n.raw)
            return {"ok": True, "data": "\n\n".join(content)}

        elif name == "fill_template":
            p = _safe_path(args["template_path"])
            from core.template_engine import fill_template
            out = fill_template(template_path=str(p), data=json.loads(args["data_json"]))
            return {"ok": True, "data": f"模板已填充: {out}", "output_path": out}

        elif name == "normalize_format":
            tmpl = _safe_path(args["template_path"]); tgts = json.loads(args["target_paths"])
            if tmpl.suffix.lower() in (".xlsx",".xls"):
                from core.format_engine import batch_normalize_excel
                r = batch_normalize_excel(str(tmpl), tgts)
            else:
                from core.format_engine import batch_normalize_word
                r = batch_normalize_word(str(tmpl), tgts)
            return {"ok": True, "data": f"批量格式化 {sum(1 for x in r if x['status']=='ok')}/{len(r)} 完成"}

        elif name == "generate_report":
            from core.md_engine import generate_report
            r = generate_report(args["description"], json.loads(args["data_json"]))
            return {"ok": True, "data": r}

        elif name == "rewrite_document":
            from core.md_engine import rewrite_text
            r = rewrite_text(args["text"], args["instruction"])
            return {"ok": True, "data": r}

        elif name == "convert_format":
            p = _safe_path(args["path"]); fmt = args["target_format"]
            if fmt == "docx":
                from core.md_engine import md_to_docx; out = md_to_docx(str(p))
            elif fmt == "html":
                from core.md_engine import md_to_html; out = md_to_html(str(p))
            else: return {"ok": False, "error": f"不支持的格式: {fmt}"}
            return {"ok": True, "data": f"已转换: {out}", "output_path": out}

        elif name == "merge_documents":
            paths = json.loads(args["paths_json"]); parts = [f"# {args['title']}\n"]
            for pp in paths: parts.append(_safe_path(pp).read_text(encoding="utf-8"))
            out = SANDBOX_ROOT / "outputs" / f"merged_{args['title']}.md"
            out.parent.mkdir(parents=True, exist_ok=True)
            out.write_text("\n\n---\n\n".join(parts), encoding="utf-8")
            return {"ok": True, "data": f"合并完成: {out}"}

        elif name == "diff_documents":
            a = _safe_path(args["path1"]).read_text(encoding="utf-8").split("\n")
            b = _safe_path(args["path2"]).read_text(encoding="utf-8").split("\n")
            diffs = [f"行{i+1}: '{l1[:80]}' vs '{l2[:80]}'" for i,(l1,l2) in enumerate(zip(a,b)) if l1!=l2]
            if len(a)>len(b): diffs += [f"行{i+1}: 文档1多出: '{a[i][:80]}'" for i in range(len(b),len(a))]
            elif len(b)>len(a): diffs += [f"行{i+1}: 文档2多出: '{b[i][:80]}'" for i in range(len(a),len(b))]
            return {"ok": True, "data": "\n".join(diffs[:40]) if diffs else "两份文档完全一致"}

        elif name == "search_knowledge":
            from core.rag_engine import hybrid_retrieve
            results = hybrid_retrieve(args["query"], top_k=5)
            return {"ok": True, "data": "\n\n---\n\n".join(f"**{r['source']}**\n{r['text'][:800]}" for r in results) if results else "未找到相关知识"}

        else: return {"ok": False, "error": f"未知工具: {name}"}
    except Exception as e: return {"ok": False, "error": str(e)}
