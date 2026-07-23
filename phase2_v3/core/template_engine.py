
#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
模板引擎 — 直接操作 .docx/.xlsx，保留全部格式。
模板文件永不经过RAG索引。
"""
from __future__ import annotations
import json, os, re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, List, Optional

TEMPLATE_DIRS = [
    Path("D:/审计准则与法规文件整理/03_事务所内部文件/C_底稿模板"),
    Path("D:/审计准则与法规文件整理/06_报告与格式规范/A_报告模板"),
    Path(__file__).resolve().parent.parent / "data" / "templates",
]
OUTPUT_DIR = Path(__file__).resolve().parent.parent / "data" / "outputs"
OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

# ═══════════════ docx/xlsx 通用格式辅助 ═══════════════

def _set_run_font(run, name: str = "宋体", size_pt: int = None):
    """同时设置西文字体与东亚（中文）字体，避免中文版 Office 下中文字体失效。"""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Pt
        run.font.name = name
        rpr = run._element.rPr
        if rpr is not None:
            rFonts = rpr.rFonts
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rpr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), name)
        if size_pt is not None:
            run.font.size = Pt(size_pt)
    except Exception:
        pass


def _set_doc_chinese_fonts(doc):
    """为新建 docx 设置中文默认字体：正文宋体小四、标题黑体。"""
    try:
        from docx.shared import Pt
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(12)
        rpr = style._element.rPr
        if rpr is not None:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            rFonts = rpr.rFonts
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rpr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), "宋体")
    except Exception:
        pass
    for lvl in range(1, 4):
        try:
            style = doc.styles[f"Heading {lvl}"]
            style.font.name = "黑体"
            rpr = style._element.rPr
            if rpr is not None:
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                rFonts = rpr.rFonts
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rpr.append(rFonts)
                rFonts.set(qn("w:eastAsia"), "黑体")
        except Exception:
            pass


def _apply_table_borders(table, color: str = "000000", size: str = "4"):
    """为表格设置显式边框，避免中文版 Office 中表格样式失效导致无框线。"""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tbl = table._tbl
        tblPr = tbl.tblPr
        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        tblBorders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            tblBorders.append(el)
        tblPr.append(tblBorders)
    except Exception:
        pass


def _iter_defined_names(wb):
    """兼容 openpyxl 新旧 API，遍历 workbook/worksheet 级命名区域。"""
    objs = []
    dns = wb.defined_names
    if hasattr(dns, "definedName"):          # 旧版 API
        try:
            objs.extend(list(dns.definedName))
        except Exception:
            pass
    elif hasattr(dns, "values"):             # 新版 dict-like API
        try:
            objs.extend(list(dns.values()))
        except Exception:
            pass
    for ws in getattr(wb, "worksheets", []):
        ws_dns = getattr(ws, "defined_names", None)
        if ws_dns is not None and hasattr(ws_dns, "values"):
            try:
                objs.extend(list(ws_dns.values()))
            except Exception:
                pass
    seen, uniq = set(), []
    for o in objs:
        n = getattr(o, "name", None)
        if n and n not in seen:
            seen.add(n)
            uniq.append(o)
    return uniq


def _set_xlsx_cell_value(cell, value):
    """写入单元格并保留原始类型；空字符串不覆盖模板原值。"""
    if isinstance(value, (int, float, bool)):
        cell.value = value
    elif isinstance(value, str) and value.strip() == "":
        return
    else:
        cell.value = str(value) if value is not None else ""

# ═══════════════ 内部LLM调用 ═══════════════

def _call_llm(prompt: str, max_tokens: int = 4096, temperature: float = 0.3) -> str:
    """内部 vLLM 调用，不走外部API"""
    import httpx
    r = httpx.post(
        os.environ.get("VLLM_TUNNEL_URL", "http://localhost:18000/v1/chat/completions"),
        headers={"Authorization": "Bearer EMPTY"},
        json={
            "model": os.environ.get("VLLM_MODEL", "qwen3-235b"),
            "messages": [{"role": "user", "content": prompt}],
            "temperature": temperature,
            "max_tokens": max_tokens,
        },
        timeout=120,
    )
    r.raise_for_status()
    return r.json()["choices"][0]["message"]["content"].strip()

@dataclass
class TemplateInfo:
    path: str; name: str; format: str; category: str; size_mb: float

def list_templates() -> List[TemplateInfo]:
    """扫描所有可用模板（只读文件名，不读内容）"""
    templates = []
    for td in TEMPLATE_DIRS:
        if not td.exists(): continue
        for f in td.rglob("*"):
            if not f.is_file(): continue
            ext = f.suffix.lower()
            if ext not in (".docx", ".xlsx", ".xls", ".doc"): continue
            cat = str(f.parent.relative_to(td)) if td in f.parents else ""
            templates.append(TemplateInfo(path=str(f), name=f.name,
                format=ext, category=cat,
                size_mb=round(f.stat().st_size / 1048576, 2)))
    return templates

def match_template(user_intent: str) -> Optional[str]:
    """LLM根据意图匹配模板（只传文件名和目录结构，不传文件内容）"""
    templates = list_templates()
    if not templates: return None
    candidates = "\n".join(f"- [{t.format}] {t.category}/{t.name}" for t in templates)
    prompt = f"""你是审计模板匹配专家。根据用户需求，从以下模板列表中选择最合适的一个。

【用户需求】{user_intent}

【可用模板（仅文件名，不含内容）】
{candidates}

【要求】只回复最适合的模板的文件名，如"5.1审计报告附表Excel排版示范案例.xlsx"。没有合适的回复"NONE"。"""
    try:
        import httpx, os
        r = httpx.post(os.environ.get("VLLM_TUNNEL_URL","http://localhost:18000/v1/chat/completions"),
            headers={"Authorization":"Bearer EMPTY"},
            json={"model":os.environ.get("VLLM_MODEL","qwen3-235b"),
                  "messages":[{"role":"user","content":prompt}],
                  "temperature":0.3,"max_tokens":100}, timeout=30)
        r.raise_for_status()
        name = r.json()["choices"][0]["message"]["content"].strip()
        if name.upper() == "NONE": return None
        for t in templates:
            if t.name == name: return t.path
    except Exception as e: print(f"[模板] LLM匹配失败:{e}")
    return None

def fill_docx(template_path: str, data: Dict[str, Any], output_path: str = "") -> str:
    """填充.docx模板，保留全部格式（字体/页眉页脚/边框等）"""
    from docx import Document
    if not output_path:
        output_path = str(OUTPUT_DIR / f"{Path(template_path).stem}_filled.docx")
    doc = Document(template_path)
    flat = _flatten(data)
    _replace_paras(doc, flat)
    _replace_tbls(doc, flat)
    for sec in doc.sections:
        _replace_paras(sec.header, flat)
        _replace_paras(sec.footer, flat)
    doc.save(output_path)
    return output_path

def _replace_paras(container, flat):
    """替换段落中的 {{key}} 占位符；支持跨 Run 拆分的情况。"""
    for para in (container.paragraphs if hasattr(container, "paragraphs") else [container]):
        _replace_text_in_paragraph(para, flat)


def _replace_text_in_paragraph(para, flat):
    """在段落级别合并 Run 文本，完成占位符替换后再写回，解决占位符被拆到多个 Run 时无法替换的问题。"""
    if not para.runs:
        return
    full = "".join(run.text for run in para.runs)
    if not full:
        return
    new_full = full
    for key, val in flat.items():
        for tmpl in (f"{{{{ {key} }}}}", f"{{{{{key}}}}}"):
            new_full = new_full.replace(tmpl, str(val))
    if new_full == full:
        return
    # 写回：保留第一个 Run 的格式，其余清空（跨 Run 时无法安全保留后续格式，优先保证内容完整）
    para.runs[0].text = new_full
    for run in para.runs[1:]:
        run.text = ""


def _replace_tbls(doc, flat):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_text_in_paragraph(para, flat)



def fill_xlsx(template_path: str, data: Dict[str, Any], output_path: str = "") -> str:
    """填充.xlsx模板，保留全部格式（字体/边框/列宽/命名区域等）"""
    from openpyxl import load_workbook
    from openpyxl.utils import range_boundaries
    if not output_path:
        output_path = str(OUTPUT_DIR / f"{Path(template_path).stem}_filled.xlsx")
    wb = load_workbook(template_path)
    flat = _flatten(data)
    raw = _flatten_raw(data)  # 保留原始类型（数字写入命名区域不得转字符串，否则破坏公式勾稽）

    # 命名区域填充：兼容 workbook/worksheet 级、新旧 API、范围引用、含空格/特殊字符 sheet 名
    for name_obj in _iter_defined_names(wb):
        name = getattr(name_obj, "name", None)
        if name is None or name not in flat:
            continue
        try:
            destinations = list(name_obj.destinations)
        except Exception:
            # 部分 worksheet 级命名区域的引用格式 destinations 无法解析，跳过避免崩溃
            continue
        if not destinations:
            continue
        sheet_title, cell_ref = destinations[0]
        if sheet_title not in wb.sheetnames:
            continue
        ws = wb[sheet_title]
        ref = cell_ref.replace("$", "")
        value = raw.get(name, flat[name])
        if ":" in ref:
            min_col, min_row, max_col, max_row = range_boundaries(ref)
            if isinstance(value, (list, tuple)):
                for i, item in enumerate(value):
                    if max_row is not None and min_row + i > max_row:
                        break
                    cell = ws.cell(row=min_row + i, column=min_col)
                    _set_xlsx_cell_value(cell, item)
            else:
                cell = ws.cell(row=min_row, column=min_col)
                _set_xlsx_cell_value(cell, value)
        else:
            cell = ws[ref]
            _set_xlsx_cell_value(cell, value)

    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value and isinstance(cell.value, str):
                    for key, val in flat.items():
                        # 占位符替换后单元格可能已被改为数值，后续 key 无需再判断
                        if not isinstance(cell.value, str):
                            break
                        for tmpl in (f"{{{{ {key} }}}}", f"{{{{{key}}}}}"):
                            if tmpl in cell.value:
                                # 整格恰好是占位符且数据为数值 → 保留数值类型
                                if cell.value.strip() == tmpl and isinstance(raw.get(key), (int, float)):
                                    cell.value = raw[key]
                                    break
                                else:
                                    cell.value = cell.value.replace(tmpl, str(val))
    wb.save(output_path)
    return output_path

def fill_template(
    template_path: str = "",
    data: Dict[str, Any] = None,
    output_path: str = "",
    user_intent: str = "",
) -> str:
    """智能填充模板。自动识别格式，保留全部格式。template_path为空时自动匹配。"""
    data = data or {}
    if not template_path and user_intent:
        template_path = match_template(user_intent)
        if not template_path:
            raise ValueError(f"未找到匹配'{user_intent}'的模板")
    if not template_path:
        raise ValueError("未指定模板路径")
    ext = Path(template_path).suffix.lower()
    if ext in (".docx", ".doc"): return fill_docx(template_path, data, output_path)
    elif ext in (".xlsx", ".xls"): return fill_xlsx(template_path, data, output_path)
    else: raise ValueError(f"不支持的格式: {ext}")

# ═══════════════ 规范驱动报告生成 ═══════════════

def _read_docx_text(path: str) -> str:
    """读取 .docx 文件的纯文本（供LLM理解规范用）"""
    try:
        from docx import Document
        return "\n".join(p.text for p in Document(path).paragraphs if p.text.strip())
    except Exception:
        return ""


def generate_formatted_report(
    data: Dict[str, Any],
    user_instruction: str,
    template_path: str = "",
    style_guides: List[str] = None,
    output_path: str = "",
) -> str:
    """
    规范驱动的报告生成。

    1. 如果提供了 template_path：以模板为骨架，LLM 读取模板结构后智能填充
    2. 如果提供了 style_guides：LLM 读取排版规范（字体/行距/数字格式等），严格遵规则生成
    3. 两者结合：模板提供结构，规范约束格式

    典型调用：
        generate_formatted_report(
            data={"差异明细": [...], "审计结论": "..."},
            user_instruction="生成医保回款专项审计报告",
            template_path="data/templates/3.2（通用）专项审计报告模板.docx",
            style_guides=[
                "data/templates/报告字体、行间距、序号与数字表示等要求.docx",
                "data/templates/5.1审计报告附表Excel排版要求 (2).docx"
            ],
            output_path="data/outputs/医保回款审计报告.docx"
        )
    """
    style_guides = style_guides or []

    # 1. 读取规范文档内容
    style_text = ""
    if style_guides:
        guides = []
        for sg in style_guides:
            content = _read_docx_text(sg)
            if content:
                guides.append(f"## {Path(sg).name}\n{content}")
        if guides:
            style_text = "\n\n---\n\n".join(guides)

    # 2. 确定模板
    if not template_path and user_instruction:
        template_path = match_template(user_instruction)

    # 3. 确定输出路径
    if not output_path:
        ext = ".docx" if (template_path and template_path.endswith(".docx")) else ".docx"
        output_path = str(OUTPUT_DIR / f"report_{_ts()}{ext}")

    # 4. 有模板：以模板为格式骨架，LLM 填充
    if template_path and Path(template_path).exists():
        return _fill_template_with_rules(
            template_path, data, user_instruction, style_text, output_path
        )

    # 5. 无模板：LLM 根据规范从头生成
    return _generate_from_scratch(
        data, user_instruction, style_text, output_path
    )


_FILLABLE_PATTERNS = [
    re.compile(r'XXX+'), re.compile(r'此处填写'), re.compile(r'（[^）]*[填写填入]）'),
    re.compile(r'\([^)]*[填写填入]\)'), re.compile(r'待填'),
    re.compile(r'【[^】]*[填写填入]】'), re.compile(r'〔[^〕]*[填写填入]〕'),
    re.compile(r'［[^］]*[填写填入]］'),
]

def _is_fillable(text: str) -> bool:
    if not text or not text.strip():
        return True
    for pat in _FILLABLE_PATTERNS:
        if pat.search(text):
            return True
    return False

def _extract_template_skeleton(doc) -> Dict:
    """提取模板骨架：段落 + 表格 + 填充位标识"""
    paragraphs = []
    for i, para in enumerate(doc.paragraphs):
        text = para.text
        sn = para.style.name if para.style else "Normal"
        is_heading = "heading" in sn.lower() or "Heading" in sn.lower() or "标题" in sn
        fillable = _is_fillable(text)
        paragraphs.append({"index": i, "style": sn, "is_heading": is_heading,
                           "text": text[:300], "fillable": fillable})
    tables = []
    for ti, table in enumerate(doc.tables):
        rows_info = []
        for ri, row in enumerate(table.rows):
            cells_info = []
            for ci, cell in enumerate(row.cells):
                ct = cell.text.strip()
                cells_info.append({"col": ci, "text": ct[:200], "fillable": _is_fillable(ct)})
            rows_info.append({"index": ri, "cells": cells_info})
        tables.append({"index": ti, "rows": rows_info})
    fillable_p = [p["index"] for p in paragraphs if p["fillable"]]
    has_f = len(fillable_p) > 0 or any(
        any(c["fillable"] for c in r["cells"]) for t in tables for r in t["rows"])
    return {"paragraphs": paragraphs, "tables": tables,
            "fillable_paragraphs": fillable_p, "has_fillable": has_f}

def _build_fill_prompt(skeleton: Dict, data: Dict, instruction: str, style_text: str) -> str:
    para_lines = []
    for p in skeleton["paragraphs"]:
        tag = "[标题]" if p["is_heading"] else "[正文]"
        marker = " <填" if p["fillable"] else ""
        para_lines.append(f"P{p['index']}{tag}({p['style']}): {p['text'][:200]}{marker}")
    tbl_lines = []
    for t in skeleton["tables"]:
        tbl_lines.append(f"T{t['index']}:")
        for r in t["rows"]:
            cells = []
            for c in r["cells"]:
                marker = "<填" if c["fillable"] else ""
                cells.append(f"[{c['text'][:60]}]{marker}")
            tbl_lines.append(f"  R{r['index']}: {' | '.join(cells)}")
    data_str = json.dumps(data, ensure_ascii=False, indent=2)
    return f"""你是资深审计报告专家。根据模板结构、数据和格式规范，填充审计报告模板。

【指令】{instruction}

【模板】([标题]段不修改，[正文]段<填需要你生成)
{chr(10).join(para_lines[:100])}

【表格】
{chr(10).join(tbl_lines[:50])}

【数据】
{data_str}

【规范】{style_text if style_text else '宋体小四正文，1.5倍行距，数字千分位，日期中文格式，措辞专业客观'}

【输出】严格JSON，不解释：
{{"paragraphs": {{"P3": "...", "P7": "..."}}, "tables": {{"T0_R1_C0": "...", ...}}}}"""

def _parse_llm_fill_response(response: str) -> Dict:
    m = re.search(r'\{[\s\S]*\}', response)
    if not m:
        raise ValueError(f"LLM响应无JSON: {response[:200]}")
    fm = json.loads(m.group(0))
    if "paragraphs" not in fm and "tables" not in fm:
        if all(k.startswith("P") for k in fm):
            fm = {"paragraphs": fm}
        else:
            raise ValueError(f"格式无效: {list(fm.keys())[:5]}")
    return fm

def _set_paragraph_text(para, new_text: str) -> None:
    """替换段落文本，保留首个run格式"""
    if para.runs:
        for run in para.runs:
            run.text = ""
        para.runs[0].text = new_text
    else:
        para.add_run(new_text)

def _set_cell_text(cell, new_text: str) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.text = ""
    if cell.paragraphs:
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].text = new_text
        else:
            cell.paragraphs[0].add_run(new_text)
    else:
        cell.add_paragraph(new_text)

def _apply_fill_map(doc, fill_map: Dict) -> None:
    for key, val in fill_map.get("paragraphs", {}).items():
        idx = int(key.lstrip("P"))
        if idx < len(doc.paragraphs):
            _set_paragraph_text(doc.paragraphs[idx], str(val))
    for key, val in fill_map.get("tables", {}).items():
        m = re.match(r'T(\d+)_R(\d+)_C(\d+)', key)
        if m:
            ti, ri, ci = int(m.group(1)), int(m.group(2)), int(m.group(3))
            if ti < len(doc.tables) and ri < len(doc.tables[ti].rows):
                if ci < len(doc.tables[ti].rows[ri].cells):
                    _set_cell_text(doc.tables[ti].rows[ri].cells[ci], str(val))



def _fill_template_with_rules(
    template_path: str, data: Dict, instruction: str,
    style_text: str, output_path: str,
) -> str:
    """
    LLM 驱动的智能模板填充。
    流程: {{key}}确定替换 → 提取骨架 → LLM填充占位符 → 保留格式写入
    LLM不可用时降级为纯{{key}}替换。
    """
    from docx import Document

    doc = Document(template_path)
    flat = _flatten(data)

    # 第一遍: {{key}} 确定性替换
    _replace_paras(doc, flat)
    _replace_tbls(doc, flat)
    for sec in doc.sections:
        _replace_paras(sec.header, flat)
        _replace_paras(sec.footer, flat)

    # 第二遍: 提取骨架
    skeleton = _extract_template_skeleton(doc)
    if not skeleton["has_fillable"]:
        doc.save(output_path)
        return output_path

    # 第三遍: LLM 智能填充
    try:
        prompt = _build_fill_prompt(skeleton, data, instruction, style_text)
        n_para = len(skeleton["fillable_paragraphs"])
        print(f"[模板] LLM填充: {n_para}段+表格待填充")
        response = _call_llm(prompt, max_tokens=4096)
        fill_map = _parse_llm_fill_response(response)
        np = len(fill_map.get("paragraphs", {}))
        nt = len(fill_map.get("tables", {}))
        print(f"[模板] LLM返回 {np}段+{nt}格填充映射")
        _apply_fill_map(doc, fill_map)
    except Exception as e:
        print(f"[模板] LLM填充失败，降级为纯{{key}}替换: {e}")
        doc = Document(template_path)
        _replace_paras(doc, flat)
        _replace_tbls(doc, flat)
        for sec in doc.sections:
            _replace_paras(sec.header, flat)
            _replace_paras(sec.footer, flat)

    doc.save(output_path)
    return output_path


def _add_table_from_text(doc, text: str):
    """根据 LLM 返回的表格文本创建带边框的 docx 表格。"""
    lines = [l.strip() for l in text.split("\n") if l.strip()]
    rows = []
    for line in lines:
        if line.startswith("|"):
            cells = [c.strip() for c in line.split("|")[1:-1]]
        else:
            cells = line.split("\t")
        # 跳过 markdown 分隔线
        if all(re.match(r'^[\s\-:|]+$', c) for c in cells):
            continue
        rows.append(cells)
    if not rows:
        return
    n_cols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=n_cols)
    tbl.style = "Table Grid"
    _apply_table_borders(tbl)
    for ri, rc in enumerate(rows):
        for ci in range(n_cols):
            cell = tbl.cell(ri, ci)
            cell.text = rc[ci] if ci < len(rc) else ""
            for p in cell.paragraphs:
                for r in p.runs:
                    _set_run_font(r, "宋体", 12)


def _ts() -> str:
    import time; return str(int(time.time()))


def _flatten(data: Dict, prefix: str = "") -> Dict[str, str]:
    flat = {}
    for k, v in data.items():
        key = f"{prefix}.{k}" if prefix else k
        if isinstance(v, dict): flat.update(_flatten(v, key))
        elif isinstance(v, list):
            if v and isinstance(v[0], dict):
                cols = list(v[0].keys())
                lines = ["\t".join(cols)]
                for item in v: lines.append("\t".join(str(item.get(c,"")) for c in cols))
                flat[key] = "\n".join(lines)
            else: flat[key] = "\n".join(str(x) for x in v)
        else: flat[key] = str(v)
    return flat


def _flatten_raw(data: Dict, prefix: str = "") -> Dict[str, Any]:
    """同 _flatten，但标量保留原始类型（int/float 不转字符串），供 Excel 数值单元格写入"""
    flat = {}
    for k, v in data.items():
        key = f"{prefix}.{k}" if prefix else k
        if isinstance(v, dict): flat.update(_flatten_raw(v, key))
        elif isinstance(v, list): continue  # 列表仍由 _flatten 的字符串形态处理
        else: flat[key] = v
    return flat


def _generate_from_scratch(
    data: Dict, instruction: str, style_text: str, output_path: str,
) -> str:
    """无模板时，LLM根据规范从头生成.docx"""
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    data_str = json.dumps(data, ensure_ascii=False, indent=2)

    prompt = f"""你是资深审计报告撰写专家。根据以下信息生成一份完整的审计报告正文。

【用户需求】{instruction}

【排版规范（必须严格遵守）】
{style_text if style_text else "默认：宋体小四正文，黑体小二标题，1.5倍行距，首行缩进2字符，数字三位分节"}

【数据】
{data_str}

【要求】
1. 严格遵循排版规范中的每一项要求
2. 生成完整的报告正文文本
3. 用 [标题一]...[标题一结束] [标题二]...[标题二结束] [正文]...[正文结束] [表格开始]...[表格结束] 标记结构
4. 数字使用千分位格式（如 1,250,000.00）
5. 只输出报告内容，不要解释

【报告正文】"""

    text = _call_llm(prompt, max_tokens=4096)

    # 解析标记，写入 docx
    doc = Document()
    _set_doc_chinese_fonts(doc)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    style.paragraph_format.first_line_indent = Cm(0.74)

    # 按标记分段写入，并解析表格内容
    sections = re.split(r'\[(标题[一二三]|正文|表格开始|表格结束)\]', text)
    current_type = None
    table_buf = []
    for part in sections:
        part = part.strip()
        if part == "表格开始":
            current_type = "table"
            table_buf = []
            continue
        if part == "表格结束":
            if table_buf:
                _add_table_from_text(doc, "\n".join(table_buf))
            current_type = None
            table_buf = []
            continue
        if part in ("标题一", "标题二", "标题三", "正文"):
            current_type = part
            continue
        if not part:
            continue

        if current_type == "table":
            table_buf.append(part)
            continue

        if current_type == "标题一":
            p = doc.add_heading(part, level=1)
        elif current_type == "标题二":
            p = doc.add_heading(part, level=2)
        elif current_type == "标题三":
            p = doc.add_heading(part, level=3)
        else:
            p = doc.add_paragraph(part)

        # 为标题/正文 Run 设置中文字体
        font_name = "黑体" if current_type in ("标题一", "标题二", "标题三") else "宋体"
        for r in getattr(p, "runs", []):
            _set_run_font(r, font_name, 12)

    doc.save(output_path)
    return output_path
