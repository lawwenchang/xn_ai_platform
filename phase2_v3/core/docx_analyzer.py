#!/usr/bin/env python3
"""Word文档智能分析器 - vLLM驱动"""
import json
from pathlib import Path
from typing import Dict, Any

def analyze_docx(file_path: str) -> Dict[str, Any]:
    try:
        from docx import Document
        doc = Document(file_path)
    except Exception:
        return {"error": f"无法读取 {file_path}", "file_type": "unknown"}
    full_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    snippet = full_text[:3000]
    if not snippet.strip():
        return {"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables), "file_type": "空文档"}
    analysis = _call_vllm_analyze(snippet)
    return {"paragraph_count": len(doc.paragraphs), "table_count": len(doc.tables), "text_length": len(full_text), **analysis}

def _call_vllm_analyze(text: str) -> dict:
    prompt = f"""分析以下文档内容，提取结构化信息。返回严格JSON。

文档内容：
{text}

JSON格式：
{{
  "file_type": "审计报告/合同协议/会议纪要/董事会决议/底稿文件/制度文件/通用文档",
  "summary": "1-2句概括",
  "key_amounts": [{{"label": "性质", "value": "金额", "context": "原文"}}],
  "key_dates": ["日期"],
  "parties": ["机构/公司名"],
  "opinion_type": "无保留/保留/否定/无法表示/空",
  "clauses": [{{"title": "条款", "content": "摘要"}}]
}}"""
    try:
        import httpx, os
        resp = httpx.post(os.environ.get("VLLM_TUNNEL_URL", "http://localhost:18000/v1/chat/completions"),
            headers={"Authorization": "Bearer EMPTY"},
            json={"model": "qwen3-235b", "messages": [{"role": "user", "content": prompt}], "temperature": 0, "max_tokens": 1000},
            timeout=30)
        resp.raise_for_status()
        raw = resp.json()["choices"][0]["message"]["content"].strip()
        if "```json" in raw: raw = raw.split("```json")[1].split("```")[0]
        elif "```" in raw: raw = raw.split("```")[1].split("```")[0]
        return json.loads(raw)
    except Exception as e:
        return {"file_type": "通用文档", "summary": f"[vLLM不可用: {e}]"}
