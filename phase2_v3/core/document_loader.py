#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
统一文档加载器 (document_loader.py)
====================================
打通平台"只能处理 Excel/CSV"的断层，让主审计流水线同时具备
docx / doc / pdf / md / txt 文档的处理能力。

设计目标：
1. 任意支持格式 → 统一的 LoadedDocument（文本 + 表格列表 + 元数据）
2. 文档中的表格 → pandas DataFrame，与 Excel 数据同权参与 Merge/Diff/对账
3. 文档全文 → 纯文本（供 LLM 摘要、关键词检索、跨文档比对）
4. 失败降级清晰：每个解析步骤独立 try/except，错误写入 errors 字段，绝不静默

依赖（均已在 requirements.txt 中）：
    python-docx（.docx）、PyMuPDF/fitz（.pdf 文本）、pdfplumber（.pdf 表格，可选）、
    chardet（.txt/.md 编码探测）、pywin32（.doc，可选，需本机安装 Word）

注意：fitz==1.20.2 无 page.find_tables()（1.23+ 才有），PDF 表格优先走 pdfplumber。
"""

from __future__ import annotations

import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

import pandas as pd

# ═══════════════════════════════════════════════════════════════
# 常量
# ═══════════════════════════════════════════════════════════════

TABLE_EXTS = {".xlsx", ".xls", ".csv"}
DOC_EXTS = {".docx", ".doc", ".pdf", ".md", ".txt"}
SUPPORTED_EXTS = TABLE_EXTS | DOC_EXTS

MAX_TEXT_CHARS = 200_000          # 单文档文本抽取上限（防爆内存）
MAX_TABLES_PER_DOC = 50           # 单文档表格抽取上限
SNIFF_TEXT_CHARS = 500            # 嗅探预览字符数


@dataclass
class LoadedDocument:
    """统一文档加载结果"""
    path: str = ""
    kind: str = "unknown"         # table / document / mixed
    text: str = ""                # 全文文本（表格文件为空）
    tables: List[pd.DataFrame] = field(default_factory=list)
    meta: Dict[str, Any] = field(default_factory=dict)
    errors: List[str] = field(default_factory=list)

    @property
    def ok(self) -> bool:
        return bool(self.tables) or bool(self.text.strip())


# ═══════════════════════════════════════════════════════════════
# 入口
# ═══════════════════════════════════════════════════════════════

def load_document(file_path, extract_tables: bool = True,
                  max_text_chars: int = MAX_TEXT_CHARS) -> LoadedDocument:
    """加载任意支持的文件，返回统一的 LoadedDocument。"""
    fp = Path(file_path)
    doc = LoadedDocument(path=str(fp))
    if not fp.exists():
        doc.errors.append(f"文件不存在: {fp}")
        return doc
    ext = fp.suffix.lower()
    try:
        if ext in TABLE_EXTS:
            doc.tables = [_read_table(fp)]
            doc.kind = "table"
            doc.meta["rows"] = len(doc.tables[0])
            doc.meta["columns"] = [str(c) for c in doc.tables[0].columns]
        elif ext == ".docx":
            _load_docx(fp, doc, extract_tables, max_text_chars)
        elif ext == ".doc":
            _load_doc(fp, doc, extract_tables, max_text_chars)
        elif ext == ".pdf":
            _load_pdf(fp, doc, extract_tables, max_text_chars)
        elif ext in (".md", ".txt"):
            _load_text(fp, doc, extract_tables, max_text_chars)
        else:
            doc.errors.append(f"不支持的文件格式: {ext}（支持: {sorted(SUPPORTED_EXTS)}）")
    except Exception as e:  # 顶层兜底：任何解析失败都不抛出，写入 errors
        doc.errors.append(f"解析失败: {type(e).__name__}: {e}")
    if doc.tables and doc.text.strip():
        doc.kind = "mixed"
    elif doc.tables:
        doc.kind = "table"
    elif doc.text.strip():
        doc.kind = "document"
    return doc


def load_tables(file_path) -> List[pd.DataFrame]:
    """便捷入口：只取表格（Excel 返回单表，docx/pdf/md 返回全部内嵌表格）。"""
    return load_document(file_path, extract_tables=True).tables



# ═══════════════════════════════════════════════════════════════
# 表格文件（Excel / CSV）
# ═══════════════════════════════════════════════════════════════

def _read_table(fp: Path) -> pd.DataFrame:
    ext = fp.suffix.lower()
    if ext == ".csv":
        for enc in ("utf-8-sig", "gbk", "utf-8"):
            try:
                return pd.read_csv(str(fp), encoding=enc)
            except (UnicodeDecodeError, UnicodeError):
                continue
        return pd.read_csv(str(fp), encoding="utf-8", encoding_errors="replace")
    return pd.read_excel(str(fp), header=_detect_header_row(fp))


def _detect_header_row(fp: Path, max_try: int = 6) -> int:
    """自动检测 Excel 表头行（委托 find_header_row 按内容多样性打分）"""
    try:
        from core.table_normalizer import find_header_row
        raw = pd.read_excel(str(fp), header=None)
        return find_header_row(raw, max_scan=max(max_try, 10))
    except Exception:
        return 0


# ═══════════════════════════════════════════════════════════════
# DOCX
# ═══════════════════════════════════════════════════════════════

def _load_docx(fp: Path, doc: LoadedDocument, extract_tables: bool,
               max_text_chars: int) -> None:
    from docx import Document
    d = Document(str(fp))
    paras = [p.text for p in d.paragraphs if p.text and p.text.strip()]
    doc.text = "\n".join(paras)[:max_text_chars]
    doc.meta["paragraphs"] = len(paras)
    if extract_tables:
        for ti, tbl in enumerate(d.tables[:MAX_TABLES_PER_DOC]):
            try:
                df = _docx_table_to_df(tbl)
                if df is not None and not df.empty:
                    df.attrs["table_name"] = f"{fp.stem}__表{ti + 1}"
                    doc.tables.append(df)
            except Exception as e:
                doc.errors.append(f"DOCX 第{ti + 1}个表格解析失败: {e}")
        doc.meta["tables_count"] = len(doc.tables)


def _docx_table_to_df(tbl) -> Optional[pd.DataFrame]:
    rows = []
    for r in tbl.rows:
        rows.append([c.text.strip() for c in r.cells])
    if not rows:
        return None
    header, data = rows[0], rows[1:]
    seen: Dict[str, int] = {}
    cols = []
    for i, h in enumerate(header):
        name = h or f"列{i + 1}"
        if name in seen:
            seen[name] += 1
            name = f"{name}_{seen[name]}"
        else:
            seen[name] = 0
        cols.append(name)
    return pd.DataFrame(data, columns=cols)


# ═══════════════════════════════════════════════════════════════
# DOC（老式二进制格式，依赖本机 Word COM；不可用时给出明确指引）
# ═══════════════════════════════════════════════════════════════

def _load_doc(fp: Path, doc: LoadedDocument, extract_tables: bool,
              max_text_chars: int) -> None:
    try:
        import win32com.client  # noqa
    except ImportError:
        doc.errors.append(
            "解析 .doc 需要 pywin32 且本机安装 Microsoft Word；"
            "建议将文件另存为 .docx 后重新上传")
        return
    word = None
    tmp_docx = fp.with_suffix(".__conv__.docx")
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        d = word.Documents.Open(str(fp.resolve()), ReadOnly=True)
        d.SaveAs2(str(tmp_docx), FileFormat=16)  # 16 = wdFormatXMLDocument(.docx)
        d.Close(False)
        _load_docx(tmp_docx, doc, extract_tables, max_text_chars)
        doc.meta["converted_from"] = ".doc"
    except Exception as e:
        doc.errors.append(f".doc 转换失败（需本机安装 Word）: {e}")
    finally:
        try:
            if word is not None:
                word.Quit()
        except Exception:
            pass
        if tmp_docx.exists():
            try:
                tmp_docx.unlink()
            except Exception:
                pass


# ═══════════════════════════════════════════════════════════════
# PDF（fitz 抽文本；pdfplumber 抽表格，二者独立降级）
# ═══════════════════════════════════════════════════════════════

def _load_pdf(fp: Path, doc: LoadedDocument, extract_tables: bool,
              max_text_chars: int) -> None:
    # 1) 文本：PyMuPDF
    try:
        import fitz
        texts = []
        with fitz.open(str(fp)) as pdf:
            doc.meta["pages"] = len(pdf)
            for page in pdf:
                t = page.get_text("text")
                if t:
                    texts.append(t)
                if sum(len(x) for x in texts) >= max_text_chars:
                    break
        doc.text = "\n".join(texts)[:max_text_chars]
    except Exception as e:
        doc.errors.append(f"PDF 文本抽取失败: {e}")
    # 2) 表格：pdfplumber（可选依赖，缺失不致命）
    if extract_tables:
        try:
            import pdfplumber
            with pdfplumber.open(str(fp)) as pdf:
                for pi, page in enumerate(pdf.pages):
                    if len(doc.tables) >= MAX_TABLES_PER_DOC:
                        break
                    try:
                        for ti, raw in enumerate(page.extract_tables() or []):
                            df = _pdf_table_to_df(raw)
                            if df is not None and not df.empty:
                                df.attrs["table_name"] = f"{fp.stem}__p{pi + 1}表{ti + 1}"
                                doc.tables.append(df)
                    except Exception as e:
                        doc.errors.append(f"PDF 第{pi + 1}页表格解析失败: {e}")
            doc.meta["tables_count"] = len(doc.tables)
        except ImportError:
            doc.errors.append("未安装 pdfplumber，PDF 表格抽取跳过（文本已抽取）")
        except Exception as e:
            doc.errors.append(f"PDF 表格抽取失败: {e}")
    if not doc.text.strip() and not doc.tables:
        doc.errors.append("PDF 未抽取到任何内容（可能是扫描件，需要 OCR）")


def _pdf_table_to_df(raw: List[List[Optional[str]]]) -> Optional[pd.DataFrame]:
    rows = [[("" if c is None else str(c).replace("\n", "").strip()) for c in r]
            for r in raw if r]
    rows = [r for r in rows if any(r)]
    if len(rows) < 1:
        return None
    header, data = rows[0], rows[1:]
    width = max(len(r) for r in rows)
    header = (header + [""] * width)[:width]
    cols, seen = [], {}
    for i, h in enumerate(header):
        name = h or f"列{i + 1}"
        if name in seen:
            seen[name] += 1
            name = f"{name}_{seen[name]}"
        else:
            seen[name] = 0
        cols.append(name)
    data = [(r + [""] * width)[:width] for r in data]
    return pd.DataFrame(data, columns=cols) if data else None


# ═══════════════════════════════════════════════════════════════
# MD / TXT（chardet 编码探测 + Markdown pipe 表格解析）
# ═══════════════════════════════════════════════════════════════

def _load_text(fp: Path, doc: LoadedDocument, extract_tables: bool,
               max_text_chars: int) -> None:
    raw = fp.read_bytes()
    enc = "utf-8"
    try:
        import chardet
        guess = chardet.detect(raw[:65536])
        if guess.get("encoding"):
            enc = guess["encoding"]
    except ImportError:
        pass
    try:
        doc.text = raw.decode(enc, errors="replace")[:max_text_chars]
    except LookupError:
        doc.text = raw.decode("utf-8", errors="replace")[:max_text_chars]
    doc.meta["encoding"] = enc
    doc.meta["chars"] = len(doc.text)
    if extract_tables and fp.suffix.lower() == ".md":
        doc.tables = _parse_md_tables(doc.text, stem=fp.stem)
        doc.meta["tables_count"] = len(doc.tables)


def _parse_md_tables(text: str, stem: str = "") -> List[pd.DataFrame]:
    """解析 Markdown pipe 表格为 DataFrame 列表"""
    tables: List[pd.DataFrame] = []
    block: List[str] = []

    def flush():
        if len(block) >= 2:
            df = _md_block_to_df(block)
            if df is not None and not df.empty:
                df.attrs["table_name"] = f"{stem}__表{len(tables) + 1}"
                tables.append(df)
        block.clear()

    for line in text.splitlines():
        s = line.strip()
        if s.startswith("|") and s.endswith("|") and s.count("|") >= 2:
            block.append(s)
        else:
            flush()
    flush()
    return tables


def _md_block_to_df(block: List[str]) -> Optional[pd.DataFrame]:
    def split_row(r: str) -> List[str]:
        return [c.strip() for c in r.strip().strip("|").split("|")]

    header = split_row(block[0])
    data_lines = block[1:]
    if data_lines and re.fullmatch(r"[|\s:\-]+", data_lines[0]):
        data_lines = data_lines[1:]
    width = len(header)
    cols, seen = [], {}
    for i, h in enumerate(header):
        name = h or f"列{i + 1}"
        if name in seen:
            seen[name] += 1
            name = f"{name}_{seen[name]}"
        else:
            seen[name] = 0
        cols.append(name)
    data = []
    for r in data_lines:
        cells = split_row(r)
        data.append((cells + [""] * width)[:width])
    return pd.DataFrame(data, columns=cols) if data else None


# ═══════════════════════════════════════════════════════════════
# 嗅探（Data Catalog 集成）与物化（沙箱集成）
# ═══════════════════════════════════════════════════════════════

def sniff_document(file_path) -> Dict[str, Any]:
    """为 Data Catalog 生成文件画像（表格 → 列结构；文档 → 预览/页数/表格数）"""
    fp = Path(file_path)
    ext = fp.suffix.lower()
    info: Dict[str, Any] = {"filename": fp.name, "ext": ext}
    if ext in TABLE_EXTS:
        try:
            df = _read_table(fp).head(1000)
            info.update({
                "kind": "table", "rows": len(df),
                "columns": [str(c) for c in df.columns],
                "dtypes": {str(c): str(df[c].dtype) for c in df.columns},
            })
        except Exception as e:
            info.update({"kind": "table", "error": str(e)})
        return info
    doc = load_document(fp)
    info["kind"] = doc.kind
    info.update(doc.meta)
    if doc.text:
        info["text_preview"] = doc.text[:SNIFF_TEXT_CHARS]
    if doc.tables:
        info["tables_columns"] = [[str(c) for c in t.columns] for t in doc.tables[:5]]
        info["tables_rows"] = [len(t) for t in doc.tables[:5]]
    if doc.errors:
        info["errors"] = doc.errors
    return info


def materialize_document(file_path, out_dir,
                         text_suffix: str = "__text") -> List[Path]:
    """把文档物化为流水线/沙箱可直接消费的文件：

    - 每个内嵌表格 → ``<stem>__表{N}.xlsx``
    - 全文文本 → ``<stem>__text.txt``

    返回生成的文件路径列表。表格文件原样返回空列表。
    """
    fp = Path(file_path)
    out = Path(out_dir)
    out.mkdir(parents=True, exist_ok=True)
    generated: List[Path] = []
    if fp.suffix.lower() in TABLE_EXTS:
        return generated
    doc = load_document(fp)
    for i, df in enumerate(doc.tables):
        name = df.attrs.get("table_name") or f"{fp.stem}__表{i + 1}"
        target = out / f"{name}.xlsx"
        try:
            df.to_excel(str(target), index=False)
            generated.append(target)
        except Exception:
            pass
    if doc.text.strip():
        target = out / f"{fp.stem}{text_suffix}.txt"
        target.write_text(doc.text, encoding="utf-8")
        generated.append(target)
    return generated


# ═══════════════════════════════════════════════════════════════
# 命令行自测
# ═══════════════════════════════════════════════════════════════

if __name__ == "__main__":
    import sys
    p = Path(sys.argv[1]) if len(sys.argv) > 1 else None
    if not p or not p.exists():
        print("用法: python -m core.document_loader <文件路径>")
        sys.exit(1)
    d = load_document(p)
    print(f"kind={d.kind} tables={len(d.tables)} text_chars={len(d.text)}")
    for i, t in enumerate(d.tables):
        print(f"  表{i + 1}: rows={len(t)} cols={list(t.columns)}")
    if d.text:
        print("文本预览:", d.text[:200].replace("\n", " | "))
    for e in d.errors:
        print("  [错误]", e)

