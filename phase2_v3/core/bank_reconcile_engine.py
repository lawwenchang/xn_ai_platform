#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
专业银行对账引擎 (bank_reconcile_engine.py)
==============================================
实现"序时账 × 银行流水"逐笔勾对的专业审计能力，与 matching_engine.py
（医保回款汇总比对专用）并存。

专业规则（依据《中国注册会计师审计准则问题解答第12号——货币资金审计》
及银行存款余额调节表编制惯例）：

1. 方向镜像（第一常识）
   - 企业序时账：借方金额 = 银行存款增加，贷方金额 = 减少
     → 归一化 net = 借方 − 贷方
   - 银行流水：贷方（收入）= 存款增加，借方（支取）= 减少
     → 归一化 net = 贷方收入 − 借方支取
   - 归一化后同号金额方可互为镜像勾对。
2. 逐笔核对精确到分（±0.01 元），百分比容差只用于汇总层面分析性复核。
3. 对账前双方各自勾稽：期初余额 + Σ借方 − Σ贷方 = 期末余额。
4. 未匹配项默认"待人工核查"，只有接近期末且有窗口证据的才列入
   未达账项候选（银收企未收 / 银付企未付 / 企收银未收 / 企付银未付），
   禁止把解释不了的差异一律洗白成"未达账项"。
5. 利息、手续费、冲正不删除、单独成类输出。

匹配层级（确定性规则优先，LLM 不参与自动确认）：
   L1 金额精确(±0.01) + 同日 + 同方向           → 自动确认
   L2 金额精确 + 日期窗口(±N 天可配) + 同方向    → 自动确认（跨期）
   L3 同方向同窗口 n:m 金额合计相等             → 拆分/合并入账
   L4 摘要/对方户名模糊(RapidFuzz)              → 仅"待人工复核"，永不自动确认
"""

from __future__ import annotations

import itertools
import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple
from config.dictionary import FEE_WORDS
import pandas as pd

# ═══════════════════════════════════════════════════════════════
# 常量与列语义注册表
# ═══════════════════════════════════════════════════════════════

JOURNAL = "journal"                 # 企业序时账/银行存款日记账
BANK_STATEMENT = "bank_statement"   # 银行流水/对账单
GENERIC_LEDGER = "generic_ledger"   # 通用台账（格式多样：有日期+金额但无强特征）

AMOUNT_TOLERANCE = 0.01             # 逐笔核对硬容差（精确到分）
DEFAULT_DATE_WINDOW = 3             # 默认日期窗口（天）

STD_COLUMNS = [
    "row_id", "date", "voucher_no", "summary", "counterpart", "account",
    "debit", "credit", "net_amount", "balance", "source_file", "src_index",
]

# 列语义 → 候选列名（按优先级排列；匹配时忽略空格与全半角括号差异）
COLUMN_ROLE_SYNONYMS: Dict[str, List[str]] = {
    "date":        ["日期", "交易日期", "记账日期", "业务日期", "入账日期", "date", "交易时间"],
    "voucher_no":  ["凭证号", "凭证号码", "凭证编号", "凭证字号", "凭证", "voucher_no"],
    "summary":     ["摘要", "摘要信息", "用途", "备注", "说明", "附言", "summary"],
    "counterpart": ["对方户名", "对方客户名称", "对方", "交易对手", "对方单位",
                    "对手方", "对方账号户名", "counterpart"],
    "account":     ["银行账号", "账号", "账户", "银行账户", "开户账号", "本方账号", "account"],
    "debit":       ["借方金额", "借方", "借方(支取)", "借方（支取）", "支取",
                    "支出", "支出金额", "付款金额", "借方发生额", "debit"],
    "credit":      ["贷方金额", "贷方", "贷方(收入)", "贷方（收入）", "收入",
                    "收入金额", "收款金额", "贷方发生额", "credit"],
    "balance":     ["余额", "账户余额", "期末余额", "本次余额", "balance"],
    "amount":      ["交易金额", "金额", "发生额", "交易额", "net_amount", "amount"],
    "subject":     ["科目名称", "会计科目", "科目","科目全称", "subject"],
}

# 序时账特征列（出现即加分）
_JOURNAL_HINTS = {"凭证号", "凭证号码", "凭证编号", "科目编码", "科目名称", "科目全称", "月"}
# 银行流水特征列
_BANK_HINTS = {"对方户名", "对方客户名称", "银行账号", "账号", "余额",
               "对方账号", "开户行"}

# 利息/费用/冲正词表（单独成类输出，不参与噪音删除）
INTEREST_WORDS = ("利息", "结息")

REVERSAL_WORDS = ("冲正", "冲销", "红冲", "撤销")
from config.dictionary import INTERBANK_WORDS as _INTERBANK
# ── v3.3 内容标签层 ──
CONTENT_TAGS = {
    "医保":  ("医保", "医疗保险", "医保结算", "医保回款", "医保中心"),
    "社保":  ("社保", "社会保险", "社保局", "社保回款"),
    "工资":  ("工资", "薪酬", "奖金", "绩效", "劳务费"),
    "税款":  ("税款", "税收", "增值税", "所得税", "营业税", "附加税", "印花税"),
    "现金":  ("现金", "库存现金", "提现"),
    "转账":  ("转账", "汇兑", "网银", "电汇", "划款"),
    "费用":  ("手续费", "短信费", "年费", "账户管理费", "工本费", "服务费", "费用外收", "批量扣费", "收费"),
    "利息":  ("利息", "结息", "利息收入", "利息支出", "批量结息"),
    "冲正":  ("冲正", "冲销", "红冲", "撤销", "调账"),
    "货款":  ("货款", "采购款", "材料款", "货款结算"),
    "往来":  ("往来款", "往来", "借款", "还款", "暂付", "暂收"),
}

_VERB_PREFIX = re.compile(r"^(收到|支付|收|付|转|汇入|汇出|支)")
_COMPANY_SUFFIX = re.compile(r"(.+?(?:有限责任公司|有限公司|集团|公司|厂|店|中心|合作社|事务所|经营部))")

def extract_counterpart(summary: Any) -> str:
    """序时账无对方列形态：从摘要提取对手方主体。
    '收阳信康润商贸有限公司 往来款' → '阳信康润商贸有限公司'
    """
    if not isinstance(summary, str) or not summary.strip():
        return ""
    first = summary.strip().split()[0]          # 摘要习惯：动词+主体 空格 用途
    first = _VERB_PREFIX.sub("", first)
    m = _COMPANY_SUFFIX.match(first)
    return m.group(1) if m else first

def tag_content(df: pd.DataFrame) -> pd.DataFrame:
    """给每行打内容标签（确定性关键词分类器）。

    标签服务于：
    ① 场景筛选——提取式核对按标签过滤
    ② L2/L3 消歧——窗口内同额多笔用标签区分
    ③ 未达分类——利息/手续费自动归类

    返回原 DataFrame，新增 content_tag 列。
    """
    tags = []
    for _, r in df.iterrows():
        summary = str(r.get("summary", ""))
        counterpart = str(r.get("counterpart", ""))
        text = summary + " " + counterpart
        row_tags = []
        for tag, keywords in CONTENT_TAGS.items():
            if any(kw in text for kw in keywords):
                row_tags.append(tag)
        tags.append(",".join(row_tags) if row_tags else "")
    df = df.copy()
    df["content_tag"] = tags
    return df



# ═══════════════════════════════════════════════════════════════
# 步骤 1：文件类型识别与列语义映射
# ═══════════════════════════════════════════════════════════════

def _norm_col(c: Any) -> str:
    """列名归一化：去空格/全角空格，统一括号"""
    s = str(c).replace(" ", "").replace("　", "").strip()
    return s.replace("（", "(").replace("）", ")")


def detect_book_type(df: pd.DataFrame, filename: str = "") -> str:
    """识别账簿类型：journal（序时账）/ bank_statement（流水）/ generic_ledger / unknown

    v3.3: 主路径走 column_semantics 240 词注册表（角色组合判定），
    旧 hints 词表降级为文件名线索辅助。
    """
    cols = {_norm_col(c) for c in df.columns}
    fname = str(filename)

    # ── 主路径：列语义角色组合判定 ──
    try:
        from core.column_semantics import detect_column_roles
        roles = detect_column_roles(df)

        # 序时账信号：有日期 + (借贷双列 或 凭证号)
        has_journal = (
            "date" in roles
            and (("debit" in roles and "credit" in roles)
                 or "voucher_no" in roles or "subject" in roles)
        )
        # 银行流水信号：有日期 + (余额 或 (收支列 + 对方/账号))
        has_bank = (
            "date" in roles and "balance" in roles
            or ("date" in roles
                and ("debit" in roles or "credit" in roles or "amount" in roles)
                and ("counterpart" in roles or "account" in roles))
        )

        if has_journal and not has_bank:
            return JOURNAL
        if has_bank and not has_journal:
            return BANK_STATEMENT
        if has_journal and has_bank:
            # 歧义：用文件名打破
            if any(k in fname for k in ("流水", "对账单", "银行")):
                return BANK_STATEMENT
            if any(k in fname for k in ("序时账", "日记账", "明细账", "台账")):
                return JOURNAL
            return JOURNAL  # 默认序时账
    except Exception:
        pass

    # ── 降级路径：旧 hints + 文件名 ──
    j_score = sum(1 for h in _JOURNAL_HINTS if any(h in c for c in cols))
    b_score = sum(1 for h in _BANK_HINTS if any(h in c for c in cols))
    if any(k in fname for k in ("流水", "对账单", "银行")): b_score += 2
    if any(k in fname for k in ("序时账", "日记账", "明细账")): j_score += 2
    if j_score > b_score: return JOURNAL
    if b_score > 0: return BANK_STATEMENT

    # ── 兜底：date + amount → 通用台账 ──
    try:
        from core.column_semantics import detect_column_roles, infer_amount_columns
        roles2 = detect_column_roles(df) if 'roles' not in dir() else roles
        has_date = "date" in roles2
        has_amount = any(r in roles2 for r in ("amount", "debit", "credit")) \
            or bool(infer_amount_columns(df, max_cols=1))
        if has_date and has_amount:
            return GENERIC_LEDGER
    except Exception:
        pass
    return "unknown"


def auto_map_columns(df: pd.DataFrame,
                     book_type: Optional[str] = None) -> Dict[str, str]:
    """列语义自动映射：返回 {role: 实际列名}。

    银行流水方向镜像关键：列名含"收入/贷方"→ credit；含"支出/支取/借方"→ debit。
    序时账：借方金额 → debit（存款增加）；贷方金额 → credit（存款减少）。
    """
    norm2orig = {_norm_col(c): str(c) for c in df.columns}
    mapping: Dict[str, str] = {}
    for role, candidates in COLUMN_ROLE_SYNONYMS.items():
        for cand in candidates:
            nc = _norm_col(cand)
            if nc in norm2orig:
                mapping[role] = norm2orig[nc]
                break
    return mapping



# ═══════════════════════════════════════════════════════════════
# 步骤 2：方向镜像归一化
# ═══════════════════════════════════════════════════════════════

def _to_float(v: Any) -> float:
    """金额解析：去千分位/中文逗号/货币符号，失败返回 0.0"""
    if isinstance(v, (int, float)) and not pd.isna(v):
        return float(v)
    try:
        s = str(v).replace(",", "").replace("，", "").replace(" ", "")
        s = s.replace("¥", "").replace("￥", "").strip()
        if s in ("", "-", "--", "nan", "None"):
            return 0.0
        return float(s)
    except (ValueError, TypeError):
        return 0.0


def normalize_to_std(df: pd.DataFrame, mapping: Dict[str, str],
                     book_type: str, source_file: str = "") -> pd.DataFrame:
    """归一化为标准结构，核心是按账簿类型做方向镜像：

    - journal（企业序时账）：net = 借方金额 − 贷方金额（正 = 存款增加）
    - bank_statement（银行流水）：net = 贷方收入 − 借方支取（正 = 存款增加）

    归一化后双方 net_amount 同号即表示同向资金流动，可直接勾对。
    """
    out = pd.DataFrame()
    n = len(df)

    def col(role: str) -> pd.Series:
        c = mapping.get(role)
        if c and c in df.columns:
            return df[c]
        return pd.Series([None] * n)
    out["row_id"] = [f"{'B' if book_type == BANK_STATEMENT else 'J'}{i}" for i in range(n)]
    date_raw = col("date")
    if pd.api.types.is_numeric_dtype(date_raw):
        s = date_raw.dropna()
        if len(s) > 0 and s.iloc[:min(5, len(s))].apply(
            lambda x: 19000101 <= int(x) <= 21001231 if pd.notna(x) else False
        ).all():
            date_raw = date_raw.fillna(0).astype("int64").astype(str).replace("0", pd.NA)
    out["date"] = pd.to_datetime(date_raw, errors="coerce")
    out["voucher_no"] = col("voucher_no").astype(str).replace("None", "")
    out["summary"] = col("summary").astype(str).replace("None", "")
    out["counterpart"] = col("counterpart").astype(str).replace("None", "")
    out["account"] = col("account").astype(str).replace("None", "")
    out["subject"] = col("subject").astype(str).replace("None", "")
    out["debit"] = pd.to_numeric(
        col("debit").astype(str).str.replace(",", "").str.replace("，", "").str.replace("¥", "").str.replace("￥", "").str.strip(),
        errors="coerce").fillna(0)
    out["credit"] = pd.to_numeric(
        col("credit").astype(str).str.replace(",", "").str.replace("，", "").str.replace("¥", "").str.replace("￥", "").str.strip(),
        errors="coerce").fillna(0)
    amount_col = mapping.get("amount")
    if book_type == JOURNAL:
        out["net_amount"] = out["debit"] - out["credit"]
        # 通用台账：无借贷双列时回退带符号金额单列（正=流入）
        if (out["net_amount"] == 0).all() and amount_col and amount_col in df.columns:
            out["net_amount"] = df[amount_col].map(_to_float)
    elif book_type == BANK_STATEMENT:
        # 方向镜像：流水借方(支取)=减少，贷方(收入)=增加
        out["net_amount"] = out["credit"] - out["debit"]
        # 只有单列交易金额（正=收入）时直接使用
        if (out["net_amount"] == 0).all() and amount_col and amount_col in df.columns:
            out["net_amount"] = df[amount_col].map(_to_float)
    else:
        if amount_col and amount_col in df.columns:
            out["net_amount"] = df[amount_col].map(_to_float)
        else:
            out["net_amount"] = out["debit"] - out["credit"]
    bal = pd.to_numeric(
        col("balance").astype(str).str.replace(",", "").str.replace("，", "").str.replace("¥", "").str.strip(),
        errors="coerce").fillna(0)
    out["balance"] = bal if not (bal == 0).all() else pd.Series([None] * n)
    out["source_file"] = source_file
    out["src_index"] = list(range(n))
    # 金额转分（int），消除浮点误差
    out["net_cents"] = (out["net_amount"] * 100).round().astype("int64")
    out["abs_cents"] = out["net_cents"].abs()
    # 剔除借贷均为 0 的空行
    out = out[out["net_cents"] != 0].reset_index(drop=True)
    out["src_index"] = list(range(len(out)))
    return out



def recognize_subject_column(df):
    if "subject" not in df.columns: return None, []
    vals = df["subject"].astype(str)
    if vals.str.strip().ne("").mean() >= 0.5 and vals.str.contains(chr(38134)+chr(34892)+chr(23384)+chr(27454), na=False).any():
        return "subject", []
    return None, []

def split_by_bank_subject(book_std, subject_col="subject"):
    import re as _re
    is_bank = book_std[subject_col].astype(str).apply(lambda x: bool(re.match(r'^银行存款(?:$|[^\u4e00-\u9fff])', x.strip())))
    if not is_bank.any(): return {}, book_std
    bank_rows = book_std[is_bank]
    non_bank = book_std[~is_bank]
    subjects = {}
    for subj, grp in bank_rows.groupby(subject_col):
        m = _re.search(r"(\d{4,})", str(subj))
        tag = m.group(1) if m else str(subj)[:20]
        subjects[f"{tag}_{len(grp)}"] = grp.reset_index(drop=True)
    print(f"[split] {len(subjects)} bank subjects: " + ", ".join(f"{k}({len(v)}r)" for k,v in subjects.items()))
    if len(non_bank): print(f"[split] non-bank: {len(non_bank)} rows")
    return subjects, non_bank

def explain_across_accounts(bank_unmatched, other_subjects, book_norm, out_dir=None):
    hints = []
    if not other_subjects or not len(bank_unmatched):
        print("[cross] no other subjects or no unmatched")
        return hints
    import json
    from collections import defaultdict
    kdf = pd.DataFrame(bank_unmatched) if not isinstance(bank_unmatched, pd.DataFrame) else bank_unmatched
    print(f"[cross] ===== 跨账户资金流向探测 =====")
    print(f"[cross] 银方未匹配 {len(kdf)} 笔 vs 其他 {len(other_subjects)} 个银行科目")
    for label, sdf in other_subjects.items():
        for _, row in kdf.iterrows():
            amt = abs(row.get("net_amount", 0))
            if amt == 0: continue
            cp = str(row.get("counterpart", ""))
            hits = sdf[(sdf["net_amount"].abs() - amt).abs() < 0.02]
            if len(hits):
                b_date = str(row.get("date", ""))[:10]
                b_summary = str(row.get("summary", ""))[:30]
                hints.append({
                    "acct": label, "amt": round(amt, 2), "cp": cp[:30],
                    "cand": len(hits),
                    "bank_date": b_date, "bank_summary": b_summary,
                    "book_dates": [str(d)[:10] for d in hits["date"].head(3)],
                    "book_cps": [str(c)[:20] for c in hits.get("counterpart", pd.Series([""]*len(hits))).head(3)],
                })
    if hints:
        hints.sort(key=lambda h: h["amt"], reverse=True)
        by_acct = defaultdict(lambda: {"count": 0, "total_amt": 0.0, "samples": []})
        for h in hints:
            a = by_acct[h["acct"]]
            a["count"] += 1
            a["total_amt"] += h["amt"]
            if len(a["samples"]) < 3:
                a["samples"].append(h)
        print(f"[cross] 命中 {len(hints)} 条，涉及 {len(by_acct)} 个其他账户：")
        for acct, info in sorted(by_acct.items(), key=lambda x: x[1]["total_amt"], reverse=True):
            print(f"[cross]   {acct[:40]}: {info['count']}笔, 合计 {info['total_amt']:,.2f}元")
            for h in info["samples"]:
                bd0 = h['book_dates'][0] if h['book_dates'] else '?'
                bc0 = h['book_cps'][0] if h['book_cps'] else ''
                print(f"[cross]     {h['amt']:>12,.2f}元 | 银:{h['bank_date']} {h['bank_summary']} | 账:{bd0} {bc0}")
        total_cross = sum(h["amt"] for h in hints)
        total_unmatched_bank = kdf["net_amount"].abs().sum()
        if total_unmatched_bank:
            pct = total_cross/total_unmatched_bank*100
            print(f"[cross] 跨账户可解释金额合计: {total_cross:,.2f}元 / 银方未匹配总额: {total_unmatched_bank:,.2f}元 = {pct:.1f}%")
        print(f"[cross] Top 10 明细（金额降序）:")
        for h in hints[:10]:
            book_dates_str = ",".join(h["book_dates"][:3])
            print(f"[cross]   {h['amt']:>12,.2f} | 银:{h['bank_date']} {h['cp'][:25]} | 账方日期:{book_dates_str} | 账户:{h['acct'][:25]}")
        if out_dir:
            with open(Path(out_dir) / "cross_hints.json", "w", encoding="utf-8") as fp:
                json.dump(hints, fp, ensure_ascii=False, indent=2)
    else:
        print("[cross] no hints found")
    return hints
def filter_bank_account(bank_std: pd.DataFrame,
                        account: Optional[str] = None) -> Tuple[pd.DataFrame, str]:
    """按银行账号过滤流水（流水含多个账户、序时账为单账户时必须先过滤）。

    account 为 None 时：若只含一个账号原样返回；含多个账号时返回
    出现次数最多的账号并给出提示（确定性规则，不猜测）。
    """
    if "account" not in bank_std.columns:
        return bank_std, "流水无账号列，未过滤"
    accounts = bank_std["account"].replace("", pd.NA).dropna().unique().tolist()
    if not accounts:
        return bank_std, "流水账号列为空，未过滤"
    if account:
        key = re.sub(r"\s", "", str(account))
        mask = bank_std["account"].astype(str).str.replace(r"\s", "", regex=True)
        hit = bank_std[mask.str.contains(key, na=False)]
        if hit.empty:
            return bank_std, f"⚠ 指定账号 {account} 在流水中未出现，未过滤（请人工确认）"
        hit = hit.reset_index(drop=True)
        if "src_index" in hit.columns:
            hit["src_index"] = list(range(len(hit)))
        return hit, f"已按账号 {account} 过滤：{len(hit)} 笔"
    if len(accounts) == 1:
        return bank_std, f"流水为单一账户 {accounts[0]}，无需过滤"
    top = bank_std["account"].value_counts().idxmax()
    hit = bank_std[bank_std["account"] == top].reset_index(drop=True)
    if "src_index" in hit.columns:
        hit["src_index"] = list(range(len(hit)))
    return hit, (f"⚠ 流水含 {len(accounts)} 个账户，自动取笔数最多的 {top}"
                 f"（{len(hit)} 笔）；如不符请指定账号")




def auto_correct_bank_direction(bank_df, bank_std, bank_map, tie_bank, book_type):
    """银行流水方向自动修正：勾稽不平时尝试交换借贷方向。
    返回 (修正后的 bank_std, bank_map, tie_bank)。"""
    if not tie_bank.get("checked") or tie_bank.get("balanced", False):
        return bank_std, bank_map, tie_bank
    diff = abs(tie_bank.get("difference", 0))
    if diff < 0.02:
        return bank_std, bank_map, tie_bank
    print(f"[方向检测] 银行流水勾稽不平(差{diff:,.0f})，尝试交换收入/支出方向...")
    swapped = dict(bank_map)
    swapped["debit"], swapped["credit"] = bank_map.get("credit"), bank_map.get("debit")
    if not swapped.get("debit") or not swapped.get("credit"):
        print("[方向检测] 缺少借贷列，无法交换")
        return bank_std, bank_map, tie_bank
    std2 = normalize_to_std(bank_df, swapped, book_type)
    std2 = tag_content(std2)
    std2, _ = filter_bank_account(std2, None)
    tie2 = tie_out_balance(std2)
    if tie2.get("balanced", False):
        print(f"[方向检测] 交换后勾稽平衡，已自动修正")
        return std2, swapped, tie2
    else:
        print(f"[方向检测] 交换后仍不平衡(差{tie2.get('difference',0):,.0f})，保留原方向")
        return bank_std, bank_map, tie_bank

# ═══════════════════════════════════════════════════════════════
# 步骤 3：双方勾稽（对账前置校验）
# ═══════════════════════════════════════════════════════════════

def tie_out_balance(std_df: pd.DataFrame, opening: Optional[float] = None,
                    closing: Optional[float] = None) -> Dict[str, Any]:
    """勾稽校验：期初余额 + Σ(net) = 期末余额。

    优先级：显式传入 > 余额列推算（首行余额 − 首行净额 = 期初）。
    不平不阻断对账，但必须在结果中显式报告。
    """
    total = round(float(std_df["net_amount"].sum()), 2)
    result: Dict[str, Any] = {"total_net": total, "checked": False}
    bal = std_df["balance"].dropna()
    if opening is None and len(bal) > 0:
        opening = round(float(bal.iloc[0]) - float(std_df["net_amount"].iloc[0]), 2)
        result["opening_source"] = "余额列推算"
    if closing is None and len(bal) > 0:
        closing = round(float(bal.iloc[-1]), 2)
        result["closing_source"] = "余额列推算"
    result["opening"] = opening
    result["closing"] = closing
    if opening is not None and closing is not None:
        computed = round(opening + total, 2)
        diff = round(computed - closing, 2)
        result.update({
            "checked": True, "computed_closing": computed,
            "difference": diff, "balanced": abs(diff) < AMOUNT_TOLERANCE,
        })
    return result

def _detect_date_granularity(book: pd.DataFrame) -> str:
    """检测账面日期粒度：>90% 集中在月末(日>=28) → 'month'（月末批量记账），否则 'day'。"""
    days = pd.to_datetime(book["date"], errors="coerce").dt.day.dropna()
    if len(days) == 0:
        return "day"
    ratio = float((days >= 28).mean())
    print(f"[日期粒度] 账面日期日>=28占比: {ratio:.1%}")
    return "month" if ratio >= 0.8 else "day"

# ═══════════════════════════════════════════════════════════════
# 步骤 4：逐笔匹配（L1 精确 → L2 窗口 → L3 拆分合并 → L4 模糊待复核）
# ═══════════════════════════════════════════════════════════════

def _match_l1(book: pd.DataFrame, bank: pd.DataFrame,
              book_used: set, bank_used: set,
              tol_cents: int) -> List[Dict[str, Any]]:
    """L1 v3.4：金额精确 + 同日 → pandas merge 向量化 O(n log n)。

    按 (cents_bkt, date_key, dup_n) 三键配对，同额同日多对多各自编号。
    tol_cents 容差：分桶键用 round() 而非 int()，避免浮点 0.1+0.2 假差异。
    """
    # 未匹配行
    left = book[~book.index.isin(book_used)][["date", "net_cents"]].copy()
    right = bank[~bank.index.isin(bank_used)][["date", "net_cents"]].copy()
    if left.empty or right.empty:
        return []

    left["date_key"] = left["date"].apply(lambda d: d.date().isoformat() if pd.notna(d) else "NaT")
    left["cents_bkt"] = (left["net_cents"] / tol_cents).round().astype("int64") * tol_cents
    left["_idx"] = left.index                                          # ← 移到这里（reset 之前）
    left = left.sort_values(["cents_bkt", "date_key"]).reset_index(drop=True)
    left["dup_n"] = left.groupby(["cents_bkt", "date_key"]).cumcount()

    right["date_key"] = right["date"].apply(lambda d: d.date().isoformat() if pd.notna(d) else "NaT")
    right["cents_bkt"] = (right["net_cents"] / tol_cents).round().astype("int64") * tol_cents
    right["_idx"] = right.index                                        # ← 同上
    right = right.sort_values(["cents_bkt", "date_key"]).reset_index(drop=True)
    right["dup_n"] = right.groupby(["cents_bkt", "date_key"]).cumcount()

    m = pd.merge(left, right, on=["cents_bkt", "date_key", "dup_n"], suffixes=("_L", "_R"))
    matches = []
    for _, row in m.iterrows():
        ji = row["_idx_L"]
        bi = row["_idx_R"]
        matches.append({"book_idx": ji, "bank_idx": bi, "level": "L1",
                        "note": "金额精确+同日"})
        book_used.add(ji)
        bank_used.add(bi)
    return matches


def _match_l2(book: pd.DataFrame, bank: pd.DataFrame,
              book_used: set, bank_used: set,
              date_window: int, tol_cents: int) -> List[Dict[str, Any]]:
    """L2 v3.4：金额精确 + 日期窗口 → merge 候选 + 日期差贪心指派。
    按日期差升序贪心一对一配对（drop_duplicates 双侧去重）。
    注意：不做月份限制，月末月初日期窗口内可能跨月匹配。
    """
    left = book[~book.index.isin(book_used)][["date", "net_cents"]].copy()
    right = bank[~bank.index.isin(bank_used)][["date", "net_cents"]].copy()
    if left.empty or right.empty:
        return []

    left["cents_bkt"] = (left["net_cents"] / tol_cents).round().astype("int64") * tol_cents
    left["_idx"] = left.index
    right["cents_bkt"] = (right["net_cents"] / tol_cents).round().astype("int64") * tol_cents
    right["_idx"] = right.index

    cand = pd.merge(left, right, on="cents_bkt", suffixes=("_L", "_R"))
    if cand.empty:
        return []

    cand["dd"] = (cand["date_L"] - cand["date_R"]).dt.days.abs()
    cand = cand[cand["dd"] <= date_window]
    if cand.empty:
        return []

    # 贪心指派：日期差最小优先，双侧去重（一对一）
    cand = cand.sort_values("dd")
    cand = cand.drop_duplicates("_idx_L", keep="first")
    cand = cand.drop_duplicates("_idx_R", keep="first")

    matches = []
    for _, row in cand.iterrows():
        ji = row["_idx_L"]
        bi = row["_idx_R"]
        matches.append({"book_idx": ji, "bank_idx": bi, "level": "L2",
                        "note": f"金额精确+日期差{int(row['dd'])}天"})
        book_used.add(ji)
        bank_used.add(bi)
    return matches

def _match_counterpart_n_to_1(book_m, bank_m, book_used, bank_used, tol_cents):
    matches = []
    for tag, src, dst, src_used, dst_used in [("bank", bank_m, book_m, bank_used, book_used),
                                                ("book", book_m, bank_m, book_used, bank_used)]:
        src_rem = src[~src.index.isin(src_used)]
        dst_rem = dst[~dst.index.isin(dst_used)]
        if len(src_rem) < 2 or len(dst_rem) < 1: continue
        src_cp = src_rem["counterpart"].astype(str).map(normalize_counterpart_name)
        src_m  = src_rem["date"].dt.to_period("M")
        src_abs = src_rem["net_cents"].abs()
        dst_cp = dst_rem["counterpart"].astype(str).map(normalize_counterpart_name)
        dst_m  = dst_rem["date"].dt.to_period("M")
        dst_abs = dst_rem["net_cents"].abs()
        s = pd.DataFrame({"_cp": src_cp, "_m": src_m, "_abs": src_abs, "_si": src_rem.index})
        d = pd.DataFrame({"_cp": dst_cp, "_m": dst_m, "_abs": dst_abs, "_di": dst_rem.index})
        sk = s.groupby(["_cp", "_abs"]).size(); dk = d.groupby(["_cp", "_abs"]).size()
        hot = sk[sk > 5000].index.union(dk[dk > 5000].index)
        if len(hot):
            s = s[~s.set_index(["_cp", "_abs"]).index.isin(hot)]
            d = d[~d.set_index(["_cp", "_abs"]).index.isin(hot)]
        mrg = pd.merge(s, d, on=["_cp", "_abs"], suffixes=("_s", "_d"))
        if mrg.empty: continue
        mrg["_mdiff"] = (mrg["_m_s"] - mrg["_m_d"]).apply(lambda x: abs(x.n) if hasattr(x, "n") else 99)
        mrg = mrg[mrg["_mdiff"] <= 1]
        if mrg.empty: continue
        for (cp, mth, ab), g in mrg.groupby(["_cp", "_m_s", "_abs"]):
            si = [i for i in g["_si"].unique() if i not in src_used]
            di = [i for i in g["_di"].unique() if i not in dst_used]
            if len(si) < 2 or len(di) != 1: continue
            di0 = di[0]
            note = f"{cp[:15]} {mth}{len(si)}p->1" if tag == "bank" else f"{cp[:15]} {mth}1->{len(si)}p"
            matches.append({"book_idxs": [di0] if tag == "bank" else si,
                            "bank_idxs": si if tag == "bank" else [di0],
                            "level": "L3_cp_n1", "note": note})
            dst_used.add(di0)
            src_used.update(si)
    return matches
def _is_fee_row(df, fkw):
    """检查行是否为费用：摘要或对方户名含费用关键词"""
    s = df["summary"].astype(str).str.strip()
    has_fee = s.str.contains(fkw, na=False)
    if "counterpart" in df.columns:
        cp = df["counterpart"].astype(str).str.strip()
        has_fee = has_fee | cp.str.contains(fkw, na=False)
    return has_fee


def _match_l3_month_remainder(book_m, bank_m, book_used, bank_used, tol_cents):
    """月模式专用：某月某方向上，未匹配流水合计 == 单笔账面 → n:1 成组。
    覆盖'整月流水汇总记一笔账'形态。O(n) 分组求和，无组合爆炸。"""
    matches = []
    for side_src, side_dst in (("bank", "book"), ("book", "bank")):
        src = bank_m if side_src == "bank" else book_m
        dst = book_m if side_src == "bank" else bank_m
        src_used = bank_used if side_src == "bank" else book_used
        dst_used = book_used if side_src == "bank" else bank_used
        rem = src[~src.index.isin(src_used)]
        if rem.empty:
            continue
        rem_valid = rem[rem["date"].notna()]
        if rem_valid.empty:
            continue
        grp = rem_valid.groupby([rem_valid["date"].dt.to_period("M"), rem_valid["net_cents"].gt(0)])["net_cents"].sum()
        for di, r in dst[~dst.index.isin(dst_used)].iterrows():
            if pd.isna(r["date"]):
                continue
            key = (r["date"].to_period("M"), r["net_cents"] > 0)
            #if str(key[0]) == "2025-07":
            #    print(f"[L3_month] 2025-07 检查: key={key} 在grp中={key in grp.index} grp值={grp.loc[key] if key in grp.index else 'N/A'} 目标={r['net_cents']}")
            if key in grp.index and abs(grp.loc[key] - r["net_cents"]) <= tol_cents:
                idxs = rem_valid[(rem_valid["date"].dt.to_period("M") == key[0]) &
                           ((rem_valid["net_cents"] > 0) == key[1]) & (~rem_valid.index.isin(src_used))].index.tolist()
                matches.append({"book_idxs": [di] if side_src == "bank" else idxs,
                                "bank_idxs": idxs if side_src == "bank" else [di],
                                "level": "L3_month",
                                "note": f"月末汇总记账：{key[0]}月同方向{len(idxs)}笔合计"})
                src_used.update(idxs); dst_used.add(di)
    # 第三遍：双向月末余额对齐——双侧都是多笔但合计相等
    rem_b = bank_m[~bank_m.index.isin(bank_used)]
    rem_j = book_m[~book_m.index.isin(book_used)]
    if not rem_b.empty and not rem_j.empty:
        rem_bv = rem_b[rem_b["date"].notna()]
        rem_jv = rem_j[rem_j["date"].notna()]
        if not rem_bv.empty and not rem_jv.empty:
            b_grp = rem_bv.groupby([rem_bv["date"].dt.to_period("M"),
                                    rem_bv["net_cents"].gt(0)])["net_cents"].sum()
            j_grp = rem_jv.groupby([rem_jv["date"].dt.to_period("M"),
                                    rem_jv["net_cents"].gt(0)])["net_cents"].sum()
            for key in set(b_grp.index) & set(j_grp.index):
                b_sum = int(round(float(b_grp.loc[key])))
                j_sum = int(round(float(j_grp.loc[key])))

                if abs(b_sum - j_sum) <= tol_cents and b_sum != 0:
                    bidxs = rem_bv[(rem_bv["date"].dt.to_period("M") == key[0]) &
                                   ((rem_bv["net_cents"] > 0) == key[1]) &
                                   (~rem_bv.index.isin(bank_used))].index.tolist()
                    jidxs = rem_jv[(rem_jv["date"].dt.to_period("M") == key[0]) &
                                   ((rem_jv["net_cents"] > 0) == key[1]) &
                                   (~rem_jv.index.isin(book_used))].index.tolist()
                    if len(bidxs) <= 1 and len(jidxs) <= 1:
                        continue
                    # ── 三色风险分层 ──
                    abs_amt = abs(j_sum) / 100  # 元
                    TE = 50000    # 实际执行的重要性水平（默认5万）
                    SAD = TE / 10 # 明显微小错报临界值

                    # 对手方交叉校验
                    cp_b = set(str(rem_bv.loc[i, "counterpart"]) for i in bidxs
                               if pd.notna(rem_bv.loc[i, "counterpart"]) and str(rem_bv.loc[i, "counterpart"]).strip())
                    cp_j = set(str(rem_jv.loc[i, "counterpart"]) for i in jidxs
                               if pd.notna(rem_jv.loc[i, "counterpart"]) and str(rem_jv.loc[i, "counterpart"]).strip())
                    cp_ok = not cp_b or not cp_j or bool(cp_b & cp_j)

                    # 摘要标签覆盖率校验
                    tags_b, tags_j, tagged_b, tagged_j = set(), set(), 0, 0
                    for i in bidxs:
                        t = str(rem_bv.loc[i, "content_tag"]) if "content_tag" in rem_bv.columns else ""
                        if t: tags_b.update(t.split(",")); tagged_b += 1
                    for i in jidxs:
                        t = str(rem_jv.loc[i, "content_tag"]) if "content_tag" in rem_jv.columns else ""
                        if t: tags_j.update(t.split(",")); tagged_j += 1
                    tags_b.discard(""); tags_j.discard("")
                    common_tags = tags_b & tags_j
                    tag_ok = bool(common_tags) and (tagged_b == 0 or tagged_j == 0 or
                               min(tagged_b / max(1, len(bidxs)), tagged_j / max(1, len(jidxs))) >= 0.15)

                    # 高风险关键词
                    RED_FLAG_KW = ("现金", "提现", "借款", "还款", "暂付", "暂收", "投资", "理财",
                                   "划款", "调账", "冲正")
                    has_red_kw = any(
                        kw in str(rem_bv.loc[i, "summary"]) or kw in str(rem_jv.loc[j, "summary"])
                        for i in bidxs for j in jidxs
                        for kw in RED_FLAG_KW
                    )
                    # 化整为零特征（n>=3 拆分/合并形态）
                    split_pattern = (len(bidxs) >= 3 and len(jidxs) == 1) or (len(jidxs) >= 3 and len(bidxs) == 1)

                    # ── 三色判定 ──
                    needs_review = False
                    if (cp_ok or tag_ok) and abs_amt <= SAD and not has_red_kw and not split_pattern:
                        tier, tier_label = 1, "自动通过"
                    elif has_red_kw or split_pattern or (abs_amt > TE and max(1, max(abs(rem_bv.loc[bidxs,"net_amount"]).max() if bidxs else 0, abs(rem_jv.loc[jidxs,"net_amount"]).max() if jidxs else 0)) > 1000):
                        tier, tier_label = 3, "必须人工核查"
                        needs_review = True
                    else:
                        tier, tier_label = 2, "抽样复核"
                        needs_review = True

                    # note
                    j_sample = str(rem_jv.loc[jidxs[0], "summary"])[:20] if jidxs else ""
                    b_sample = str(rem_bv.loc[bidxs[0], "summary"])[:20] if bidxs else ""
                    detail = []
                    if not cp_ok: detail.append("对手方无交集")
                    if common_tags:
                        detail.append("标签一致(%s)" % ",".join(common_tags))
                    elif not tag_ok:
                        detail.append("标签不匹配")
                    if not detail:
                        detail.append("全部校验通过")
                    detail_str = "; ".join(detail)
                    note = (f"[T{tier}] 双向月末对齐：{key[0]}月{'收入' if key[1] else '支出'} "
                            f"账{len(jidxs)}笔 银{len(bidxs)}笔 合计{abs_amt:,.0f}元 "
                            f"| 账:{j_sample} | 银:{b_sample} | {detail_str} | {tier_label}")

                    # 拆分大小额：单笔>1000保持原tier，<=1000降为T2
                    SMALL_CUT = 1000
                    large_b = [i for i in bidxs if abs(rem_bv.loc[i, "net_amount"]) > SMALL_CUT] if bidxs else []
                    small_b = [i for i in bidxs if abs(rem_bv.loc[i, "net_amount"]) <= SMALL_CUT] if bidxs else []
                    large_j = [i for i in jidxs if abs(rem_jv.loc[i, "net_amount"]) > SMALL_CUT] if jidxs else []
                    small_j = [i for i in jidxs if abs(rem_jv.loc[i, "net_amount"]) <= SMALL_CUT] if jidxs else []
                    for b_part, j_part, sub_tier, sub_label in [
                        (large_b, large_j, tier, tier_label),
                        (small_b, small_j, min(tier, 2), "抽样复核" if tier >= 2 else tier_label),
                    ]:
                        if not b_part and not j_part:
                            continue
                        sub_note = note if (not small_b and not small_j) else (
                            note.replace(f"[T{tier}]", f"[T{sub_tier}]").replace(tier_label, sub_label)
                            + f" | 拆分: 单笔{'>' if sub_tier == tier else '<='}1000元部分")
                        book_used.update(j_part)
                        bank_used.update(b_part)
                        matches.append({
                            "book_idxs": j_part, "bank_idxs": b_part,
                            "level": "L3_month",
                            "needs_review": (sub_tier >= 2),
                            "risk_tier": sub_tier,
                            "note": sub_note
                        })
                        if len(b_part) >= 5 and len(j_part) >= 5:
                            print(f"[红旗] {sub_note}")

   
    return matches

def _match_l3_fee_monthly(book_m, bank_m, book_used, bank_used, tol_cents):
    """手续费滚动窗口匹配：逐月累加账款/银端余额，余额对齐时整段匹配。
    允许多月累计对齐(容差=50元)，同时支持逐月精确匹配作为兜底。
    """
    matches = []
    _FKW = "手续费|费用外收|扣费|收费|短信费|年费|工本费|服务费"
    _k_fee_mask = (
        (~bank_m.index.isin(bank_used)) & (
            _is_fee_row(bank_m, _FKW) |
            ((bank_m["summary"].astype(str).str.strip().isin(["", "nan", "None", "nat"]) |
              bank_m["summary"].isna()) & (bank_m["net_cents"].abs() <= 5000))
        )
    )
    fee_b = book_m[(~book_m.index.isin(book_used)) &
                   (_is_fee_row(book_m, _FKW))]
    fee_k = bank_m[_k_fee_mask]
    print(f"[L3_fee_month] 候选池: 账方{len(fee_b)}笔 银方{len(fee_k)}笔 (已用: 账{len(book_used)} 银{len(bank_used)})")
    if fee_b.empty or fee_k.empty:
        return matches
    b_month = fee_b.groupby(fee_b["date"].dt.to_period("M"))["net_cents"].sum()
    k_month = fee_k.groupby(fee_k["date"].dt.to_period("M"))["net_cents"].sum()
    all_months = sorted(set(b_month.index) | set(k_month.index))
    FEE_TOL = tol_cents
    b_acc, k_acc = 0, 0
    b_buf, k_buf = [], []
    seg_count = 0
    for m in all_months:
        b_val = int(b_month.get(m, 0))
        k_val = int(k_month.get(m, 0))
        b_acc += b_val
        k_acc += k_val
        b_now = fee_b[fee_b["date"].dt.to_period("M") == m].index.tolist()
        k_now = fee_k[fee_k["date"].dt.to_period("M") == m].index.tolist()
        b_buf.extend(b_now)
        k_buf.extend(k_now)
        if abs(b_acc - k_acc) <= FEE_TOL and b_buf and k_buf:
            _b_fresh = [i for i in b_buf if i not in book_used]
            _k_fresh = [i for i in k_buf if i not in bank_used]
            if _b_fresh and _k_fresh:
                seg_count += 1
                matches.append({
                    "book_idxs": _b_fresh, "bank_idxs": _k_fresh, "level": "L3_fee_month",
                    "note": f"累计匹配#{seg_count}: {all_months[0]}-{m} 账{len(_b_fresh)}笔 银{len(_k_fresh)}笔 差{abs(b_acc-k_acc)/100:.1f}元"
                })
                book_used.update(_b_fresh)
                bank_used.update(_k_fresh)
            b_acc, k_acc = 0, 0
            b_buf, k_buf = [], []
    b_rem = fee_b[~fee_b.index.isin(book_used)]
    k_rem = fee_k[~fee_k.index.isin(bank_used)]
    if not b_rem.empty and not k_rem.empty:
        br_m = b_rem.groupby(b_rem["date"].dt.to_period("M"))["net_cents"].sum()
        kr_m = k_rem.groupby(k_rem["date"].dt.to_period("M"))["net_cents"].sum()
        m_count = 0
        for m in sorted(set(br_m.index) & set(kr_m.index)):
            if abs(br_m[m] - kr_m[m]) <= FEE_TOL:
                bidx = [i for i in b_rem[b_rem["date"].dt.to_period("M") == m].index if i not in book_used]
                kidx = [i for i in k_rem[k_rem["date"].dt.to_period("M") == m].index if i not in bank_used]
                if bidx and kidx:
                    m_count += 1
                    matches.append({
                        "book_idxs": bidx, "bank_idxs": kidx, "level": "L3_fee_month",
                        "note": f"月聚合: {m} 账{len(bidx)}笔 银{len(kidx)}笔 差{abs(br_m[m]-kr_m[m])/100:.1f}元"
                    })
                    book_used.update(bidx)
                    bank_used.update(kidx)
        if m_count:
            print(f"[L3_fee_month] 月聚合兜底: +{m_count}月")
    b_left = fee_b[~fee_b.index.isin(book_used)]
    k_left = fee_k[~fee_k.index.isin(bank_used)]
    print(f"[L3_fee_month] 累计+月聚合共{len(matches)}段, 剩余: 账{len(b_left)}笔({abs(b_left['net_cents'].sum())/100:.0f}元) 银{len(k_left)}笔({abs(k_left['net_cents'].sum())/100:.0f}元)")
    return matches
def _match_l3_fee_remainder(book_m, bank_m, book_used, bank_used, tol_cents, date_window,
                             max_group=10):
    """手续费剩余匹配：滑动窗口(O(N)) + 小N组合(max 3选1~3)兜底。
    对未匹配账面费用，在同月银方费用中找连续窗口或小组合匹配。
    """
    from itertools import combinations
    matches = []
    _FKW = "手续费|费用外收|扣费|收费|短信费|年费|工本费|服务费|电话费"
    _k_fee_mask = (
        (~bank_m.index.isin(bank_used)) & (
            _is_fee_row(bank_m, _FKW) |
            ((bank_m["summary"].astype(str).str.strip().isin(["", "nan", "None", "nat"]) |
              bank_m["summary"].isna()) & (bank_m["net_cents"].abs() <= 5000))
        )
    )
    b_rem = book_m[(~book_m.index.isin(book_used)) &
                   (_is_fee_row(book_m, _FKW))]
    k_rem = bank_m[_k_fee_mask]
    if b_rem.empty or k_rem.empty:
        return matches
    b_rem = b_rem.copy(); b_rem["_m"] = pd.to_datetime(b_rem["date"], errors="coerce").dt.to_period("M")
    k_rem = k_rem.copy(); k_rem["_m"] = pd.to_datetime(k_rem["date"], errors="coerce").dt.to_period("M")
    k_by_month = {m: g for m, g in k_rem.groupby("_m", sort=False)}
    # 跨月扩展：账面在M月时，也从M-1和M+1月找银方费用
    def _get_cross_pool(bm):
        parts = [k_by_month.get(bm, pd.DataFrame())]
        try: parts.append(k_by_month.get(bm - 1, pd.DataFrame()))
        except: pass
        try: parts.append(k_by_month.get(bm + 1, pd.DataFrame()))
        except: pass
        return pd.concat([p for p in parts if not p.empty]) if parts else pd.DataFrame()
    _b_used = set(int(i) for i in book_used)
    _k_used = set(int(i) for i in bank_used)
    n_matched = 0
    for _, br in b_rem.iterrows():
        bi = int(br.name)
        if bi in _b_used:
            continue
        target = int(br["net_cents"])
        bm = br["_m"]
        pool = _get_cross_pool(bm)
        if pool.empty:
            continue
        pool = pool[(~pool.index.isin(_k_used)) & ((pool["net_cents"] > 0) == (target > 0))]
        if len(pool) < 1:
            continue
        pool = pool.sort_values("date") if "date" in pool.columns else pool
        vals = [(int(pi), int(pc)) for pi, pc in zip(pool.index, pool["net_cents"])]
        # 找最优子集和（不设容差上限，取最小差值）
        best_found, best_diff = None, abs(target)
        # 策略1: 滑动窗口
        for i in range(len(vals)):
            total = 0
            for j in range(i, min(i + max_group, len(vals))):
                total += vals[j][1]
                diff = abs(total - target)
                if diff < best_diff:
                    best_diff = diff
                    best_found = [vals[k][0] for k in range(i, j + 1)]
        # 策略2: 小组合
        closest = sorted(vals, key=lambda x: abs(x[1] - target))[:20]
        for n in range(1, 4):
            for combo in combinations(range(len(closest)), n):
                total = sum(closest[i][1] for i in combo)
                diff = abs(total - target)
                if diff < best_diff:
                    best_diff = diff
                    best_found = [closest[i][0] for i in combo]
        if best_found:
            best_found = [x for x in best_found if x not in _k_used]
            if best_found:
                n_matched += 1
                matches.append({
                    "book_idxs": [bi], "bank_idxs": best_found, "level": "L3_fee_remainder",
                    "note": f"剩余匹配: {bm} 账{target/100:.2f}↔银{len(best_found)}笔 差{best_diff/100:.2f}元"
                })
                book_used.add(bi)
                bank_used.update(best_found)
                _b_used.add(bi)
                _k_used.update(best_found)
    if n_matched:
        print(f"[L3_fee_rem] 剩余匹配: +{n_matched}笔")
    return matches

def _match_l3_fee_embedded(book, bank, book_used, bank_used, date_window, tol_cents,
                           max_fee_cents=5000):
    """手续费伴随匹配：小额定费用（≤50元）允许与同向非费用行 n:1 合并匹配。
    
    场景：银行流水同时产生转账-100,000和费用外收-15两笔，
          企业账只记一笔-100,015"转账"，L3_fee_difference因摘要无"手续费"而漏配。
          本规则将费用行与非费用行合并后匹配对方单条记录。
    """
    import numpy as np

    _FEE_PAT = "|".join(FEE_WORDS)
    matches = []
    # ── 方向一：n笔银行(含1笔费用) ↔ 1笔账面 ──
    _b_used = set(int(i) for i in book_used)
    _k_used = set(int(i) for i in bank_used)
    # 银行侧：分出费用行和非费用行
    k_fee_mask = bank["summary"].astype(str).str.strip().str.contains(_FEE_PAT, na=False)
    k_fee = bank[(~bank.index.isin(_k_used)) & k_fee_mask & (bank["net_cents"].abs() <= max_fee_cents)]
    k_non = bank[(~bank.index.isin(_k_used)) & (~k_fee_mask)]
    if k_fee.empty or k_non.empty or book[~book.index.isin(_b_used)].empty:
        return matches
    # 按月份建非费用行索引
    k_non_m = k_non.copy()
    k_non_m["_m"] = pd.to_datetime(k_non_m["date"], errors="coerce").dt.to_period("M")
    k_non_m["_pos"] = k_non_m["net_cents"] > 0
    k_by_month = {m: g for m, g in k_non_m.groupby(["_m", "_pos"], sort=False)}
    b_rem = book[~book.index.isin(_b_used)].copy()
    b_rem["_m"] = pd.to_datetime(b_rem["date"], errors="coerce").dt.to_period("M")
    b_rem["_pos"] = b_rem["net_cents"] > 0
    b_by_month_cents = {}
    for (mth, pos), g in b_rem.groupby(["_m", "_pos"], sort=False):
        b_by_month_cents[(mth, pos)] = {int(c): list(g.index) for c, idxs in g.groupby("net_cents").groups.items()}
    for _, fr in k_fee.iterrows():
        fc = int(fr["net_cents"])
        fm = pd.to_datetime(fr["date"], errors="coerce").to_period("M") if pd.notna(fr["date"]) else None
        fpos = fc > 0
        if fm is None:
            continue
        pool = k_by_month.get((fm, fpos))
        if pool is None:
            continue
        for _, nr in pool.iterrows():
            nc = int(nr["net_cents"])
            total = fc + nc
            bmap = b_by_month_cents.get((fm, fpos), {})
            b_hits = bmap.get(total, []) + bmap.get(-total, [])
            b_hits = [x for x in b_hits if x not in _b_used]
            if b_hits:
                bi = b_hits[0]
                matches.append({
                    "book_idxs": [bi], "bank_idxs": [int(fr.name), int(nr.name)],
                    "level": "L3_fee_embedded",
                    "note": f"费用伴随: 银行{fc/100:.2f}+{nc/100:.2f}→账面{total/100:.2f}"
                })
                book_used.add(bi)
                bank_used.add(int(fr.name))
                bank_used.add(int(nr.name))
                _b_used.add(bi)
                _k_used.add(int(fr.name))
                _k_used.add(int(nr.name))
                break
    # ── 方向二：n笔账面(含1笔费用) ↔ 1笔银行 ──
    _b_used2 = set(int(i) for i in book_used)
    _k_used2 = set(int(i) for i in bank_used)
    b_fee_mask = book["summary"].astype(str).str.strip().str.contains(_FEE_PAT, na=False)
    b_fee = book[(~book.index.isin(_b_used2)) & b_fee_mask & (book["net_cents"].abs() <= max_fee_cents)]
    b_non = book[(~book.index.isin(_b_used2)) & (~b_fee_mask)]
    if b_fee.empty or b_non.empty or bank[~bank.index.isin(_k_used2)].empty:
        return matches
    b_non_m = b_non.copy()
    b_non_m["_m"] = pd.to_datetime(b_non_m["date"], errors="coerce").dt.to_period("M")
    b_non_m["_pos"] = b_non_m["net_cents"] > 0
    bb_by_month = {m: g for m, g in b_non_m.groupby(["_m", "_pos"], sort=False)}
    k_rem = bank[~bank.index.isin(_k_used2)].copy()
    k_rem["_m"] = pd.to_datetime(k_rem["date"], errors="coerce").dt.to_period("M")
    k_rem["_pos"] = k_rem["net_cents"] > 0
    kb_by_month_cents = {}
    for (mth, pos), g in k_rem.groupby(["_m", "_pos"], sort=False):
        kb_by_month_cents[(mth, pos)] = {int(c): list(g.index) for c, idxs in g.groupby("net_cents").groups.items()}
    for _, fr in b_fee.iterrows():
        fc = int(fr["net_cents"])
        fm = pd.to_datetime(fr["date"], errors="coerce").to_period("M") if pd.notna(fr["date"]) else None
        fpos = fc > 0
        if fm is None:
            continue
        pool = bb_by_month.get((fm, fpos))
        if pool is None:
            continue
        for _, nr in pool.iterrows():
            nc = int(nr["net_cents"])
            total = fc + nc
            kmap = kb_by_month_cents.get((fm, fpos), {})
            k_hits = kmap.get(total, []) + kmap.get(-total, [])
            k_hits = [x for x in k_hits if x not in _k_used2]
            if k_hits:
                ki = k_hits[0]
                matches.append({
                    "book_idxs": [int(fr.name), int(nr.name)], "bank_idxs": [ki],
                    "level": "L3_fee_embedded",
                    "note": f"费用伴随: 账面{fc/100:.2f}+{nc/100:.2f}→银行{total/100:.2f}"
                })
                book_used.add(int(fr.name))
                book_used.add(int(nr.name))
                bank_used.add(ki)
                _b_used2.add(int(fr.name))
                _b_used2.add(int(nr.name))
                _k_used2.add(ki)
                break
    return matches


def _match_l3_fee_difference(book: pd.DataFrame, bank: pd.DataFrame,
                              book_used: set, bank_used: set,
                              date_window: int, tol_cents: int = 1,
                              max_fee: int = 1) -> List[Dict[str, Any]]:
    """P2: 手续费差额规则——流水扣款 - 账面记录 ≈ 小额手续费时自动成组。

    v3.6: 向量化候选过滤（numpy 掩码替代 iterrows 双重循环），语义不变。
    差额 ≤ 10 元且较大一侧的摘要含"手续费/费用"关键词时，自动匹配。
    max_fee=1000 即 ±10.00 元（单位：分）。
    """
    import numpy as np
    matches = []
    FEE_KEYWORDS = ("手续费", "费用", "服务费", "管理费", "短信费", "年费", "工本费")

    # 一次性转 numpy 数组
    b_idx = book.index.to_numpy()
    b_c = pd.to_numeric(book["net_cents"], errors="coerce").fillna(0).to_numpy("int64")
    b_d = pd.to_datetime(book["date"], errors="coerce").to_numpy()
    k_idx = bank.index.to_numpy()
    k_c = pd.to_numeric(bank["net_cents"], errors="coerce").fillna(0).to_numpy("int64")
    k_d = pd.to_datetime(bank["date"], errors="coerce").to_numpy()

    # 流水侧已用位置标记（增量维护，避免反复 np.isin）
    k_pos = {int(v): p for p, v in enumerate(k_idx)}
    k_used = np.zeros(len(k_idx), dtype=bool)
    for i in bank_used:
        p = k_pos.get(int(i))
        if p is not None:
            k_used[p] = True

    b_used_lookup = set(int(i) for i in book_used)

    for n in range(len(b_idx)):
        ji = int(b_idx[n])
        if ji in b_used_lookup:
            continue
        jc = int(b_c[n])

        # 向量化候选：同方向 + 差额≤max_fee + 日期窗口 + 未使用
        m = ((k_c > 0) == (jc > 0)) & (~k_used) & (np.abs(k_c - jc) <= max_fee)
        if pd.notna(b_d[n]):
            a64 = np.datetime64(pd.Timestamp(b_d[n]).to_datetime64())
            dd = np.abs((k_d - a64) / np.timedelta64(1, "D"))
            m &= (dd <= date_window) | np.isnan(dd)
        cand = np.nonzero(m)[0]

        for p in cand:
            bi = int(k_idx[p])
            bc = int(k_c[p])
            # 差额 ≤ 10元，且较大一侧的摘要含手续费关键词
            if abs(jc) >= abs(bc):
                summary = str(book.loc[ji, "summary"])
            else:
                summary = str(bank.loc[bi, "summary"])
            if not any(kw in summary for kw in FEE_KEYWORDS):
                continue
            diff = abs(jc - bc)
            matches.append({
                "book_idxs": [ji], "bank_idxs": [bi], "level": "L3_fee",
                "note": f"手续费差额规则：差额{round(diff/100, 2)}元，{summary[:30]}"
            })
            book_used.add(ji)
            bank_used.add(bi)
            b_used_lookup.add(ji)
            k_used[p] = True
            break
    return matches

def _check_balance_continuity(df: pd.DataFrame) -> List[Dict[str, Any]]:
    """P5: 对账单余额连贯性检查——逐行验证 上日余额+发生=本日余额。

    对于有 balance 列的数据，检查每行余额是否等于上一行余额 + 本行净发生额。
    返回所有不连贯的行。
    """
    flags = []
    if "balance" not in df.columns or "net_amount" not in df.columns:
        return flags
    prev_bal = None
    for i, r in df.iterrows():
        cur_bal = r["balance"]
        net_amt = r["net_amount"]
        if pd.isna(cur_bal):
            continue
        if prev_bal is not None:
            expected = prev_bal + net_amt
            diff = round(abs(cur_bal - expected), 2)
            if diff > 0.02:  # 容差 ±0.02 元（允许舍入）
                flags.append({
                    "type": "余额不连贯", "side": "银行流水",
                    "rows": [r.get("row_id", str(i))],
                    "amount": round(float(cur_bal), 2),
                    "detail": (
                        f"上日余额 {round(float(prev_bal),2)} + 发生 {round(float(net_amt),2)}"
                        f" = {round(float(expected),2)}，但实际余额 {round(float(cur_bal),2)}，差额 {diff}"
                    ),
                })
        prev_bal = cur_bal
    return flags

def _match_l3(book: pd.DataFrame, bank: pd.DataFrame,
              book_used: set, bank_used: set,
              date_window: int, tol_cents: int,
              max_group: int = 3, max_candidates: int = 30) -> List[Dict[str, Any]]:
    """L3 v3.6：按月分桶+持久used掩码。O(n²)全表扫描 → O(n×桶)。
    语义与v3.5完全一致：同方向/同窗口(dd<=window或date为NaT)/容差/
    候选超上限显式记录/NaT锚点不做日期过滤。"""
    import numpy as np
    from collections import defaultdict
    matches: List[Dict[str, Any]] = []
    overflow_log = []

    def _prep(src: pd.DataFrame, used: set) -> dict:
        idx = src.index.to_numpy()
        cents = pd.to_numeric(src["net_cents"], errors="coerce").fillna(0).to_numpy(dtype="int64")
        dts = pd.to_datetime(src["date"], errors="coerce")
        months = np.where(dts.isna(), "NaT", dts.dt.to_period("M").astype(str))
        pos = (cents > 0)
        used_mask = np.isin(idx, list(used)) if used else np.zeros(len(idx), dtype=bool)
        buckets = defaultdict(list)
        for p in range(len(idx)):
            buckets[months[p]].append(p)
        return {"idx": idx, "cents": cents, "d64": dts.to_numpy(),
                "months": months, "pos": pos, "used": used_mask,
                "buckets": buckets, "nat": np.array(buckets.get("NaT", []), dtype=int)}

    def _pool_months(anchor_m):
        if anchor_m == "NaT":
            return None                      # NaT锚点：全表（保持原语义）
        p = pd.Period(anchor_m, "M")
        return [str(p - 1), str(p), str(p + 1)]   # 邻三月覆盖±window跨月边界（如30号↔下月2号）

    def window_pool(S, sign_pos, anchor_dt, anchor_m):
        ms = _pool_months(anchor_m)
        if ms is None:
            p = np.arange(len(S["idx"]))
        else:
            cand = []
            for m in ms:
                b = S["buckets"].get(m)
                if b:
                    cand.extend(b)
            if len(S["nat"]):
                cand.extend(S["nat"].tolist())     # 原语义：date为空的行恒保留
            if not cand:
                return []
            p = np.array(cand, dtype=int)
        m_ = (S["pos"][p] == sign_pos) & (~S["used"][p])
        if pd.notna(anchor_dt):
            a = np.datetime64(pd.Timestamp(anchor_dt).to_datetime64())
            dd = np.abs((S["d64"][p] - a) / np.timedelta64(1, "D"))
            m_ &= (dd <= date_window) | np.isnan(dd)
        p = p[m_]
        return [(int(S["idx"][i]), int(S["cents"][i])) for i in p]

    def subset_hit(target_cents, cand):
        if len(cand) < 2 or len(cand) > max_candidates:
            return None
        for size in range(2, min(max_group, len(cand)) + 1):
            for combo in itertools.combinations(cand, size):
                if abs(sum(c[1] for c in combo) - target_cents) <= tol_cents:
                    return [c[0] for c in combo]
        return None

    def consume(S, row_ids):
        if row_ids:
            S["used"] |= np.isin(S["idx"], list(row_ids))

    B = _prep(bank, bank_used)
    J = _prep(book, book_used)

    # 方向一：1 笔账 ←→ n 笔流水
    for k in range(len(J["idx"])):
        if J["used"][k]:
            continue
        cents = int(J["cents"][k])
        pool = window_pool(B, J["pos"][k], J["d64"][k], J["months"][k])
        if len(pool) > max_candidates:
            overflow_log.append(f"L3 overflow: book J{int(J['idx'][k])} cents={cents} pool={len(pool)}>{max_candidates}")
        hit = subset_hit(cents, pool)
        if hit:
            ji = int(J["idx"][k])
            matches.append({"book_idxs": [ji], "bank_idxs": hit, "level": "L3",
                            "note": f"1笔账面↔{len(hit)}笔流水（拆分/合并入账）"})
            book_used.add(ji)
            bank_used.update(hit)
            J["used"][k] = True
            consume(B, hit)

    # 方向二：n 笔账 ←→ 1 笔流水
    for k in range(len(B["idx"])):
        if B["used"][k]:
            continue
        cents = int(B["cents"][k])
        pool = window_pool(J, B["pos"][k], B["d64"][k], B["months"][k])
        if len(pool) > max_candidates:
            overflow_log.append(f"L3 overflow: bank B{int(B['idx'][k])} cents={cents} pool={len(pool)}>{max_candidates}")
        hit = subset_hit(cents, pool)
        if hit:
            bi = int(B["idx"][k])
            matches.append({"book_idxs": hit, "bank_idxs": [bi], "level": "L3",
                            "note": f"{len(hit)}笔账面↔1笔流水（拆分/合并入账）"})
            book_used.update(hit)
            bank_used.add(bi)
            B["used"][k] = True
            consume(J, hit)

    if overflow_log:
        print("[L3] 超出组合能力（静默漏配→显式记录）:")
        for msg in overflow_log[:5]:
            print(f"  {msg}")
        if len(overflow_log) > 5:
            print(f"  ... 共 {len(overflow_log)} 条")
    return matches

import unicodedata
import re as _re_cp

def normalize_counterpart_name(s) -> str:
    """对手方名称归一化：消灭'同一个户的不同写法'。
    NFKC 处理全/半角（括号、字母、数字），再去掉所有空白和常见标点。"""
    if s is None:
        return ""
    s = unicodedata.normalize("NFKC", str(s))          # （）→()、全角字母数字→半角
    s = _re_cp.sub(r"[\s　]+", "", s)                   # 半角/全角空格、换行
    s = _re_cp.sub(r"[,.，.。、·・''\"\"()（）\[\]【】\-—_*/\\]", "", s)  # 标点全部剥掉
    return s

def _align_counterpart_names(book_names, bank_names, min_len=3, fuzzy_th=92):
    """三级对齐：精确 → 包含 → 模糊兜底。
    模糊兜底红线：只有'唯一高分'才接受（第一名领先第二名≥5分），有歧义就放弃。"""
    from rapidfuzz import fuzz, process
    mapping = {}
    bank_list = [k for k in bank_names if len(k) >= min_len]
    for bn in book_names:
        bn = bn.strip()
        if len(bn) < min_len:
            continue
        # 1) 精确/包含
        hit = next((kn for kn in bank_list if bn == kn or bn in kn or kn in bn), None)
        # 2) 模糊兜底（错别字、个别字差异）
        if hit is None and bank_list:
            res = process.extract(bn, bank_list, scorer=fuzz.ratio,
                                  score_cutoff=fuzzy_th, limit=2)
            if res and (len(res) == 1 or res[0][1] - res[1][1] >= 5):
                hit = res[0][0]
        if hit:
            mapping[bn] = hit
    return mapping

def _match_l3_counterpart(book, bank, book_used, bank_used, tol_cents,
                          max_group=4, max_candidates=40):
    """对手方分区匹配 v2（性能版）：名称分桶字典化+池子按(月,方向/金额)预索引。
    原版每户名全表扫描、池内O(g²)；现按户名O(n)建索引，池查询近O(1)。
    语义与原版一致：同户名+同月 1:1削峰 → 月度闭环 → n:m子集和。"""
    import itertools
    matches = []
    b_rem = book[~book.index.isin(book_used)]
    k_rem = bank[~bank.index.isin(bank_used)]
    _bad = {"", "nan", "none", "nat"}

    def _prep_cp(rem):
        cp = rem[~rem["counterpart"].astype(str).str.strip().str.lower().isin(_bad)]
        cp = cp[cp["counterpart"].astype(str).str.strip().str.len() >= 2].copy()
        cp["_cp"] = cp["counterpart"].map(normalize_counterpart_name)
        cp["_mk"] = cp["date"].dt.to_period("M")
        cp["_pos"] = cp["net_cents"] > 0
        groups = {n: g for n, g in cp.groupby("_cp", sort=False)}
        return cp, groups

    b_cp, b_groups = _prep_cp(b_rem)
    k_cp, k_groups = _prep_cp(k_rem)
    if b_cp.empty or k_cp.empty:
        print("[对手方分区] 一侧无对手方，跳过")
        return matches
    name_map = _align_counterpart_names(
        b_cp["_cp"].unique(),
        k_cp["_cp"].unique())
    print(f"[对手方分区] 账方户名 {b_cp['_cp'].nunique()} 个，"
          f"银方 {k_cp['_cp'].nunique()} 个，对齐成功 {len(name_map)} 个")
    # 诊断（生产日志，保留）
    print("账方未对齐户名示例:", sorted(set(b_cp["_cp"]) - set(name_map.keys()))[:10])
    print("银方未对齐户名示例:", sorted(set(k_cp["_cp"]) - set(name_map.values()))[:10])
    _sizes = k_cp[k_cp["_cp"].isin(name_map.values())].groupby("_cp").size().sort_values(ascending=False)
    print(f"对齐户名的银方池子: >40笔的有 {(_sizes > 40).sum()} 个，最大5个: {_sizes.head(5).to_dict()}")

    def _subset_hit(target, idxs, vals):
        for n in range(1, min(max_group, len(idxs)) + 1):
            for combo in itertools.combinations(range(len(idxs)), n):
                if sum(vals[i] for i in combo) == target:
                    return [idxs[i] for i in combo]
        return None

    for bname, kname in name_map.items():
        bg = b_groups.get(bname)
        kg = k_groups.get(kname)
        if bg is None or kg is None:
            continue
        # 预索引：(_mk, net_cents)→行号 与 (_mk, _pos)→(行号,金额)
        kg_mc = {}
        for key, g in kg.groupby(["_mk", "net_cents"], sort=False):
            kg_mc[key] = g.index.tolist()
        kg_mp = {}
        for key, g in kg.groupby(["_mk", "_pos"], sort=False):
            kg_mp[key] = (g.index.tolist(), [int(x) for x in g["net_cents"]])
        bg_mp = {}
        for key, g in bg.groupby(["_mk", "_pos"], sort=False):
            bg_mp[key] = (g.index.tolist(), [int(x) for x in g["net_cents"]])

        # ↓↓↓ 第一层：同户名+同月+同额 1:1 配对（恰好唯一才自动确认） ↓↓↓
        for ji, jr in bg.iterrows():
            if ji in book_used:
                continue
            lst = kg_mc.get((jr["_mk"], jr["net_cents"]), [])
            cand = [i for i in lst if i not in bank_used]
            if len(cand) == 1:
                bi = cand[0]
                matches.append({"book_idxs": [ji], "bank_idxs": [bi],
                                "level": "L3_counterpart",
                                "note": f"对手方[{kname}]：同月同额1:1"})
                book_used.add(ji)
                bank_used.add(bi)
        # ↓↓↓ 第二层：同户同月总额闭环（账n笔↔银m笔，月度净额相等即整组核销） ↓↓↓
        _bg_av = bg[~bg.index.isin(book_used)]
        _kg_av = kg[~kg.index.isin(bank_used)]
        if not _bg_av.empty and not _kg_av.empty:
            _bm = _bg_av.groupby(_bg_av["_mk"])["net_cents"].sum()
            _km = _kg_av.groupby(_kg_av["_mk"])["net_cents"].sum()
            for _m in _bm.index:
                if _m not in _km.index:
                    continue
                if abs(int(_bm[_m]) - int(_km[_m])) <= tol_cents:
                    _bidx = _bg_av[(_bg_av["_mk"] == _m) & (~_bg_av.index.isin(book_used))].index.tolist()
                    _kidx = _kg_av[(_kg_av["_mk"] == _m) & (~_kg_av.index.isin(bank_used))].index.tolist()
                    if _bidx and _kidx:
                        matches.append({
                            "book_idxs": _bidx, "bank_idxs": _kidx,
                            "level": "L3_counterpart",
                            "note": (f"对手方[{kname}] {str(_m)[:7]}月闭环："
                                     f"账{len(_bidx)}笔↔银{len(_kidx)}笔 总额相等")})
                        book_used.update(_bidx)
                        bank_used.update(_kidx)
        # 方向一：1笔账 ↔ n笔银（同月同向子集和）
        for ji, jr in bg.iterrows():
            if ji in book_used:
                continue
            cents = int(jr["net_cents"])
            got = kg_mp.get((jr["_mk"], cents > 0))
            if not got:
                continue
            idxs_all, vals_all = got
            pool = [(i, v) for i, v in zip(idxs_all, vals_all) if i not in bank_used]
            if len(pool) < 1 or len(pool) > max_candidates:
                continue
            idxs, vals = [p[0] for p in pool], [p[1] for p in pool]
            hit = _subset_hit(cents, idxs, vals)
            if hit:
                matches.append({"book_idxs": [ji], "bank_idxs": hit,
                                "level": "L3_counterpart",
                                "note": f"对手方[{kname}]：账1笔↔银{len(hit)}笔"})
                book_used.add(ji)
                bank_used.update(hit)
        # 方向二：n笔账 ↔ 1笔银
        for bi, br in kg.iterrows():
            if bi in bank_used:
                continue
            cents = int(br["net_cents"])
            got = bg_mp.get((br["_mk"], cents > 0))
            if not got:
                continue
            idxs_all, vals_all = got
            pool = [(i, v) for i, v in zip(idxs_all, vals_all) if i not in book_used]
            if len(pool) < 1 or len(pool) > max_candidates:
                continue
            idxs, vals = [p[0] for p in pool], [p[1] for p in pool]
            hit = _subset_hit(cents, idxs, vals)
            if hit:
                matches.append({"book_idxs": hit, "bank_idxs": [bi],
                                "level": "L3_counterpart",
                                "note": f"对手方[{kname}]：账{len(hit)}笔↔银1笔"})
                book_used.update(hit)
                bank_used.add(bi)
    return matches

def _match_l4(book: pd.DataFrame, bank: pd.DataFrame,
              book_used: set, bank_used: set,
              fuzzy_threshold: int = 85,
              amount_pct: float = 0.01) -> List[Dict[str, Any]]:
    """L4 v3.4：金额分桶预过滤 + RapidFuzz 精排（不再 O(n^2) 全量比较）。"""
    matches = []
    try:
        from rapidfuzz import fuzz
    except ImportError:
        return matches

    left = book[~book.index.isin(book_used)][["summary", "counterpart", "net_cents"]].copy()
    right = bank[~bank.index.isin(bank_used)][["summary", "counterpart", "net_cents"]].copy()
    if left.empty or right.empty:
        return matches

    left["cents_bkt"] = (left["net_cents"] / 100).round().astype("int64")
    right["cents_bkt"] = (right["net_cents"] / 100).round().astype("int64")
    left["_idx"] = left.index
    right["_idx"] = right.index

    cand = pd.merge(left, right, on="cents_bkt", suffixes=("_L", "_R"))
    if cand.empty:
        return matches

    cand["pct"] = abs(cand["net_cents_L"] - cand["net_cents_R"]) / cand[["net_cents_L", "net_cents_R"]].abs().max(axis=1).clip(lower=1)
    cand = cand[cand["pct"] <= amount_pct]
    cand = cand[(cand["net_cents_L"] > 0) == (cand["net_cents_R"] > 0)]
    if cand.empty:
        return matches

    ranked = {}
    for _, row in cand.iterrows():
        ji, bi = int(row["_idx_L"]), int(row["_idx_R"])
        if ji in book_used or bi in bank_used:
            continue
        score = fuzz.token_set_ratio(
            f"{row['summary_L']} {row['counterpart_L']}".strip(),
            f"{row['summary_R']} {row['counterpart_R']}".strip())
        if score >= fuzzy_threshold:
            if ji not in ranked or score > ranked[ji][1]:
                ranked[ji] = (bi, score)

    for ji, (bi, score) in sorted(ranked.items(), key=lambda x: x[1][1], reverse=True):
        if ji in book_used or bi in bank_used:
            continue
        matches.append({"book_idx": ji, "bank_idx": bi, "level": "L4",
                        "score": round(score, 1),
                        "note": "模糊匹配，待人工复核"})
        book_used.add(ji)
        bank_used.add(bi)
    return matches

    return matches


CAT_BANK_RECV = "银收企未收"
CAT_BANK_PAY = "银付企未付"
CAT_ENT_RECV = "企收银未收"
CAT_ENT_PAY = "企付银未付"
CAT_REVIEW = "待人工核查"


def _row_payload(r, side):
    return {"side": side, "row_id": r["row_id"], "src_index": int(r["src_index"]),
            "date": r["date"].date().isoformat() if pd.notna(r["date"]) else "",
            "net_amount": round(float(r["net_amount"]), 2),
            "summary": r["summary"], "counterpart": r["counterpart"],
            "voucher_no": r["voucher_no"], "account": r["account"]}


def classify_unmatched(book, bank, book_used, bank_used, date_window):
    def fw(s):
        for w in ("利息","结息"):
            if w in s: return "利息"
        for w in ("手续费","短信费","年费","账户管理费","工本费","服务费"):
            if w in s: return "银行费用"
        for w in ("冲正","冲销","红冲","撤销"):
            if w in s: return "冲正"
        return ""
    max_bd, max_bb = book["date"].max(), bank["date"].max()
    ub, ubn = [], []
    for ji, r in book.iterrows():
        if ji in book_used: continue
        item = _row_payload(r, "book")
        s = fw(str(r["summary"]))
        if s: item["special"] = s
        if any(w in str(r["summary"]) for w in _INTERBANK):
            item["classification"] = "他行互转-待他行流水"
            item["basis"] = "资金在企业其他银行账户间划转，超出本账户流水范围"
        else:
            ne = pd.notna(r["date"]) and pd.notna(max_bb) and (max_bb - r["date"]).days <= date_window
            item["classification"] = (CAT_ENT_RECV if r["net_cents"] > 0 else CAT_ENT_PAY) if ne else CAT_REVIEW
            item["basis"] = f"接近期末(≤{date_window}天),需期后验证" if ne else "无窗口证据,需人工查明"
        ub.append(item)
    for bi, r in bank.iterrows():
        if bi in bank_used: continue
        item = _row_payload(r, "bank")
        s = fw(str(r["summary"]))
        if s: item["special"] = s
        _txt = str(r["summary"]) + str(r.get("counterpart", ""))
        if any(w in _txt for w in _INTERBANK):
            item["classification"] = "他行互转-待他行流水"
            item["basis"] = "收款/付款方为企业其他银行账户，超出本账户范围"
        else:
            ne = pd.notna(r["date"]) and pd.notna(max_bd) and (max_bd - r["date"]).days <= date_window
            item["classification"] = (CAT_BANK_RECV if r["net_cents"] > 0 else CAT_BANK_PAY) if ne else CAT_REVIEW
            item["basis"] = f"接近期末(≤{date_window}天),需期后验证" if ne else "无窗口证据,需人工查明"
        ubn.append(item)
    return ub, ubn


def detect_duplicates(std_df, side, day_gap=1):
    from collections import defaultdict
    flagged, groups = [], defaultdict(list)
    for i, r in std_df.iterrows():
        d = r["date"]
        groups[(int(r["net_cents"]), d.date().isoformat()[:7] if pd.notna(d) else "NaT")].append(i)
    for idxs in groups.values():
        if len(idxs) < 2: continue
        idxs = sorted(idxs, key=lambda i: (pd.isna(std_df.loc[i,"date"]), std_df.loc[i,"date"]))
        cl = [idxs[0]]
        for p, c in zip(idxs, idxs[1:]):
            d1, d2 = std_df.loc[p,"date"], std_df.loc[c,"date"]
            if pd.notna(d1) and pd.notna(d2) and abs((d2-d1).days) <= day_gap: cl.append(c)
            else:
                if len(cl) >= 2:
                    flagged.append({"side":side,"rows":[std_df.loc[i,"row_id"] for i in cl],
                                    "net_amount":round(float(std_df.loc[cl[0],"net_amount"]),2),
                                    "reason":f"同侧同额{len(cl)}笔相隔≤{day_gap}天"})
                cl = [c]
        if len(cl) >= 2:
            flagged.append({"side":side,"rows":[std_df.loc[i,"row_id"] for i in cl],
                            "net_amount":round(float(std_df.loc[cl[0],"net_amount"]),2),
                            "reason":f"同侧同额{len(cl)}笔相隔≤{day_gap}天"})
    return flagged


def build_balance_reconciliation(bc, sc, ub, ubn):
    def sc_sum(items, cat): return round(sum(i["net_amount"] for i in items if i.get("classification")==cat),2)
    br, bp = sc_sum(ubn, CAT_BANK_RECV), sc_sum(ubn, CAT_BANK_PAY)
    er, ep = sc_sum(ub, CAT_ENT_RECV), sc_sum(ub, CAT_ENT_PAY)
    bc = bc or 0.0; sc = sc or 0.0
    return pd.DataFrame([
        {"项目":"企业账面余额","金额":round(bc,2)},
        {"项目":"加:银收企未收","金额":br},
        {"项目":"减:银付企未付","金额":round(-bp,2)},
        {"项目":"调节后企业余额","金额":round(bc+br+bp,2)},
        {"项目":"银行对账单余额","金额":round(sc,2)},
        {"项目":"加:企收银未收","金额":er},
        {"项目":"减:企付银未付","金额":round(-ep,2)},
        {"项目":"调节后银行余额","金额":round(sc+er+ep,2)},
        {"项目":"调节差异","金额":round((bc+br+bp)-(sc+er+ep),2)},
    ])


def detect_fee_small_monthly_diff(book, bank, max_small_cents=5000):
    """账面费用 vs 流水费用外收：同月两侧都有费用、月度差为固定小额(≤50元)
    → 疑似费用扣自其他账户或漏记，列为审计线索（不核销、不改分类）。"""
    fee_kw = ("手续费", "短信费", "年费", "账户管理费", "工本费", "服务费")
    bf = book[book["summary"].astype(str).str.contains("|".join(fee_kw), na=False)]
    kf = bank[bank["summary"].astype(str).str.contains("|".join(fee_kw), na=False)]
    if bf.empty or kf.empty:
        return []
    bm = bf.groupby(bf["date"].dt.to_period("M"))["net_cents"].sum()
    km = kf.groupby(kf["date"].dt.to_period("M"))["net_cents"].sum()
    flags = []
    for m in bm.index:
        if m not in km.index:
            continue
        diff = int(bm[m] - km[m])
        if 0 < abs(diff) <= max_small_cents:
            flags.append({
                "side": "book", "type": "费用小额月度差异",
                "detail": (f"{m}月账面费用合计{abs(bm[m])/100:.2f}元 vs 流水费用外收"
                           f"{abs(km[m])/100:.2f}元，差{abs(diff)/100:.2f}元——"
                           f"费用或扣自其他账户、或属漏记/跨期，建议索取其他账户流水核对"),
                "rows": bf[(bf["date"].dt.to_period("M") == m)]["row_id"].tolist(),
            })
    return flags

def detect_counterpart_month_gap(book, bank, min_cents=100_000_000):
    """同户名月度透视：一侧整月为零、另一侧≥阈值(默认100万) → 账外收支嫌疑红旗。"""
    def _pv(df):
        d = df[df["counterpart"].astype(str).str.strip().str.len() >= 2].copy()
        d["_cp"] = d["counterpart"].map(normalize_counterpart_name)
        d = d[d["_cp"].str.len() >= 2]
        return d.groupby(["_cp", d["date"].dt.to_period("M")])["net_cents"].sum()
    bm, km = _pv(book), _pv(bank)
    flags = []
    for cp in set(bm.index.get_level_values(0)) | set(km.index.get_level_values(0)):
        months = set(bm.loc[cp].index if cp in bm.index.get_level_values(0) else []) | \
                 set(km.loc[cp].index if cp in km.index.get_level_values(0) else [])
        for m in months:
            b = int(bm.get((cp, m), 0)); k = int(km.get((cp, m), 0))
            if (b == 0) != (k == 0) and max(abs(b), abs(k)) >= min_cents:
                side = "银有账无" if b == 0 else "账有银无"
                flags.append({"side": "both", "type": "对手方整月单边记录",
                              "detail": (f"对手方[{cp}] {m}月 {side}，单边金额"
                                         f"{max(abs(b), abs(k))/100:,.2f}元——"
                                         f"疑似账外收支或跨户混记，需重点核查")})
    return flags

def aggregate_red_flags(red_flags, top_n=20):
    """红旗聚合输出：同类归并为摘要旗（数量+最大若干示例），
    明细不丢——由 export 写进 Excel/底稿，摘要在报告里可读。"""
    from collections import defaultdict
    by_type = defaultdict(list)
    for f in red_flags:
        by_type[f.get("type", "其他")].append(f)
    out = []
    for t, fs in by_type.items():
        if len(fs) <= top_n:
            out.extend(fs)
            continue
        fs_sorted = sorted(fs, key=lambda f: len(str(f.get("detail", ""))), reverse=True)
        out.append({
            "side": "both", "type": t,
            "detail": (f"共{len(fs)}项，示例：{fs_sorted[0].get('detail', '')[:60]}；"
                       f"……明细见《异常资金交易清单》"),
            "count": len(fs)})
    return out

def detect_red_flags(std_df, side, pair_window=3, large_threshold=100000.0, round_unit=10000, burst_days=7, burst_count=3):
    flags = []; df = std_df
    MAX_FLAGS = 5000; MAX_MERGE = 100000; MAX_BUCKET = 200; import numpy as np; seen = set()
    def _add(f):
        if len(flags) < MAX_FLAGS:
            flags.append(f)
        elif len(flags) == MAX_FLAGS:
            flags.append({'type':'红旗超限','side':side,'rows':[],'amount':0,'detail':f'已达上限{MAX_FLAGS}'})

    # 一收一付同额（v3.5: numpy广播向量化，超大桶聚合不展开）
    MAX_BUCKET = 200
    df_p = df[df["net_cents"].notna()].copy()
    df_p["_abs"] = df_p["net_cents"].abs()
    for abs_c, grp in df_p.groupby("_abs"):
        ins = grp[grp["net_cents"] > 0]
        outs = grp[grp["net_cents"] < 0]
        if ins.empty or outs.empty:
            continue
        if len(ins) * len(outs) > MAX_BUCKET * MAX_BUCKET:
            _add({"type": "一收一付同额", "side": side,
                   "amount": round(abs_c / 100, 2),
                   "rows": grp["row_id"].tolist(),
                   "detail": f"高频同额{abs_c/100:,.2f}元出现{len(ins)}收{len(outs)}付，需整体核查"})
            continue
        d_in  = ins["date"].values.astype("datetime64[D]")
        d_out = outs["date"].values.astype("datetime64[D]")
        near = abs(d_in[:, None] - d_out[None, :]).astype("int64") <= pair_window
        for ii, jj in zip(*np.where(near)):
            if len(flags) >= MAX_FLAGS:
                break
            key = tuple(sorted((ins["row_id"].iloc[ii], outs["row_id"].iloc[jj])))
            if key not in seen:
                seen.add(key)
                _add({"type": "一收一付同额", "side": side, "rows": list(key),
                       "amount": round(abs_c / 100, 2),
                       "detail": f"同额资金{int(abs((d_in[ii]-d_out[jj]).astype(int)))}天一进一出"})

# 整数大额（向量化：原全表iterrows）
    _amt_all = pd.to_numeric(df["net_amount"], errors="coerce").abs()
    _m_big = (_amt_all >= large_threshold) & ((_amt_all.astype("int64") % round_unit) == 0)
    for _rid, _a in zip(df.loc[_m_big, "row_id"], _amt_all[_m_big]):
        _add({"type": "整数大额", "side": side, "rows": [_rid],
              "amount": round(float(_a), 2), "detail": f"单笔>{large_threshold:,.0f}且整{round_unit}倍数"})

    # 期末负余额
    bal = df["balance"].dropna()
    if len(bal) > 0 and float(bal.iloc[-1]) < 0:
        _add({"type": "期末负余额", "side": side, "rows": [],
                      "amount": round(float(bal.iloc[-1]), 2), "detail": "透支/未入账负债"})

    # 分次转入转出（v3.5: 滑动窗口+二分查找，O(n log n)）
    cp = df[df["counterpart"].astype(str).str.len() >= 2]
    cp_groups = list(cp.groupby("counterpart"))
    if len(cp_groups) > 500:
        cp_groups = cp_groups[:500]
    for name, grp in cp_groups:
        grp = grp.sort_values("date")
        dates = grp["date"].values.astype("datetime64[D]")
        rights = np.searchsorted(dates, dates + np.timedelta64(burst_days, "D"), side="right")
        counts = rights - np.arange(len(dates))
        hits = np.where(counts >= burst_count)[0]
        if len(hits) == 0:
            continue
        # 合并重叠窗口
        clusters = [[hits[0]]]
        for h in hits[1:]:
            if h <= rights[clusters[-1][0]]:
                clusters[-1].append(h)
            else:
                if len(flags) >= MAX_FLAGS:
                    break
                clusters.append([h])
        for cl in clusters:
            wi = grp.index[cl[0]:rights[cl[0]]].tolist()
            sub = df.loc[wi]
            if (sub["net_cents"] > 0).any() and (sub["net_cents"] < 0).any():
                _add({"type": "分次转入转出", "side": side,
                       "rows": [df.loc[x, "row_id"] for x in wi],
                       "amount": round(float(sub["net_amount"].abs().sum()), 2),
                       "detail": f"对手方[{name}] {burst_days}天内{len(wi)}笔双向资金往来"})

# 大额现金（向量化：原全表iterrows）
    _m_cash = df["summary"].astype(str).str.contains("现金", na=False) & (_amt_all >= 50000)
    for _rid, _a in zip(df.loc[_m_cash, "row_id"], _amt_all[_m_cash]):
        _add({"type": "大额现金", "side": side, "rows": [_rid],
              "amount": round(float(_a), 2), "detail": "大额现金收支"})
    return flags

def run_bank_reconciliation(book_df, bank_df, config=None, progress_callback=None):
    cfg = dict(config or {})
    date_window = int(cfg.get("date_window_days", 3))
    tol = float(cfg.get("amount_tolerance", 0.01))
    tol_cents = max(1, int(round(tol * 100)))
    fuzzy_th = int(cfg.get("fuzzy_threshold", 85))

    def _p(pct, step):
        if progress_callback:
            try: progress_callback(pct, step)
            except: pass

    _p(5, "类型识别与列映射")
    book_type = cfg.get("book_type") or detect_book_type(book_df, cfg.get("book_file", ""))
    bank_type = cfg.get("bank_type") or detect_book_type(bank_df, cfg.get("bank_file", ""))
    if book_type == "unknown": book_type = JOURNAL
    if bank_type == "unknown": bank_type = BANK_STATEMENT
    book_map = cfg.get("book_mapping") or auto_map_columns(book_df, book_type)
    bank_map = cfg.get("bank_mapping") or auto_map_columns(bank_df, bank_type)
    if cfg.get("use_llm_mapping", False):  # v3.4: 默认不调 LLM，列映射规则足够
        try:
            from core.column_semantics import detect_roles_with_llm
            _llm = cfg.get("llm_callable")
            for side, df_, mp in (("book", book_df, book_map), ("bank", bank_df, bank_map)):
                need = ("date" not in mp) or not any(r in mp for r in ("amount", "debit", "credit"))
                if need:
                    extra = detect_roles_with_llm(df_, f"{side}_table", llm_callable=_llm)
                    for role in ("date", "amount", "debit", "credit", "summary", "counterpart", "account", "voucher_no", "balance"):
                        if role not in mp and role in extra: mp[role] = extra[role]
        except Exception: pass
    def _orient(mp, fb):
        if "借" in str(mp.get("debit", "")): return JOURNAL
        return BANK_STATEMENT
    if book_type == GENERIC_LEDGER: book_type = _orient(book_map, JOURNAL)
    if bank_type == GENERIC_LEDGER: bank_type = _orient(bank_map, BANK_STATEMENT)

    _p(10, "方向镜像归一化")
    _bo = book_type if book_type in (JOURNAL, BANK_STATEMENT) else JOURNAL
    _bk = bank_type if bank_type in (JOURNAL, BANK_STATEMENT) else BANK_STATEMENT
    book_std = normalize_to_std(book_df, book_map, _bo, cfg.get("book_file", ""))
    bank_std = normalize_to_std(bank_df, bank_map, _bk, cfg.get("bank_file", ""))
    book_std = tag_content(book_std); bank_std = tag_content(bank_std)

    # v3.4: L4 默认关闭（RapidFuzz 语义匹配 O(n²) 太重，L1-L3+L3_fee 已覆盖确定匹配）
    _enable_l4 = cfg.get("enable_l4", False)  # 需要时显式开启
    bank_std, account_note = filter_bank_account(bank_std, cfg.get("account"))
    bank_std["src_index"] = list(range(len(bank_std)))
    tie_book = tie_out_balance(book_std, cfg.get("book_opening"), cfg.get("book_closing"))
    tie_bank = tie_out_balance(bank_std, cfg.get("bank_opening"), cfg.get("bank_closing"))
    bank_std, bank_map, tie_bank = auto_correct_bank_direction(bank_df, bank_std, bank_map, tie_bank, _bk)
    # v3.10: 账方对手方缺失时，从摘要提取主体（序时账无对方列形态）
    if book_std["counterpart"].fillna("").astype(str).str.strip().eq("").mean() > 0.5:
        print("[对手方] 账方对手方缺失>50%，从摘要提取主体")
        book_std["counterpart"] = book_std["summary"].map(extract_counterpart)
    # v3.7: 日期偏移自适应——月末批量记账形态（账面记账日系统性滞后交易日）
    _p(12, "日期粒度检测")
    _granularity = cfg.get("date_granularity") or _detect_date_granularity(book_std)
    book_m = book_std.copy(); bank_m = bank_std.copy()
    if _granularity == "month":
        print("[日期粒度] 账面为月末批量记账形态，匹配按 年月+金额 对齐（日信息不参与勾对）")
        book_m["date"] = pd.to_datetime(book_m["date"]).dt.to_period("M").dt.to_timestamp()
        bank_m["date"] = pd.to_datetime(bank_m["date"]).dt.to_period("M").dt.to_timestamp()
    book_used, bank_used = set(), set()
    _p(14, "L1 金额精确+同日匹配（全量）")
    m1 = _match_l1(book_m, bank_m, book_used, bank_used, tol_cents)
    _p(15, "L2 日期窗口匹配（全量）")
    m2 = _match_l2(book_m, bank_m, book_used, bank_used, date_window, tol_cents)
    _FEE_KW = "|".join(FEE_WORDS)
    _p(17, "L3_cp_n1 n:1匹配")
    m3_cp_n1 = _match_counterpart_n_to_1(book_m, bank_m, book_used, bank_used, tol_cents)
    _p(18, "L3_counterpart 对手方分区匹配")
    matches_cp = _match_l3_counterpart(book_m, bank_m, book_used, bank_used, tol_cents)
    _p(40, "L3 n:m 拆分合并匹配")
    m3 = _match_l3(book_m, bank_m, book_used, bank_used, date_window, tol_cents)
    _p(45, "L3_month 月末汇总匹配")
    if _granularity == "month":
        m3_month = _match_l3_month_remainder(book_m, bank_m, book_used, bank_used, tol_cents)
    else:
        m3_month = []
    _p(50, "L3_fee 手续费差额匹配")
    m3_fee = _match_l3_fee_difference(book_m, bank_m, book_used, bank_used, date_window)
    _p(55, "L3_fee_month 手续费月度聚合")
    m3_fee_month = _match_l3_fee_monthly(book_m, bank_m, book_used, bank_used, tol_cents)
    _p(57, "L3_month 月末汇总匹配（第二遍：费用剔除后补刀）")
    if _granularity == "month":
        m3_month2 = _match_l3_month_remainder(book_m, bank_m, book_used, bank_used, tol_cents)
        m3_month.extend(m3_month2)
    if not _enable_l4:
        _p(60, "L4 模糊匹配")
        m4 = []
    else:
        _p(60, "L4 模糊匹配")
        m4 = _match_l4(book_m, bank_m, book_used, bank_used, fuzzy_th)
    # 诊断：追踪特定银行费用被谁匹配
    _trace_bank = bank_std[(bank_std["date"].astype(str).str.startswith("2018-12")) &
                           (bank_std["summary"].astype(str).str.contains("短信费"))]
    if not _trace_bank.empty:
        print("[L3诊断] 2018-12 银方短信费匹配追踪:")
        for bi in _trace_bank.index:
            b_amt = _trace_bank.loc[bi, "net_amount"]
            for m in m1 + m2 + m3 + m3_month + m3_fee + m3_cp_n1 + matches_cp + m3_fee_month:
                bank_list = m.get("bank_idxs") or ([m.get("bank_idx")] if "bank_idx" in m else [])
                if bi in bank_list:
                    book_list = m.get("book_idxs") or ([m.get("book_idx")] if "book_idx" in m else [])
                    print(f"  [{m['level']}] 银{len(bank_list)}笔 ↔ 账{len(book_list)}笔 | {m.get('note','')}")                    
                    for ji in book_list:
                        print(f"  [{m['level']}] 银#{bi}({b_amt}) ↔ 账#{ji} {book_std.loc[ji,'summary'][:30]} {book_std.loc[ji,'net_amount']}")

    _p(70, "未匹配项四分类")
    unmatched_book, unmatched_bank = classify_unmatched(book_std, bank_std, book_used, bank_used, date_window)
    duplicates = detect_duplicates(book_std, "book") + detect_duplicates(bank_std, "bank")

    bc = tie_book.get("closing"); sc = tie_bank.get("closing")
    if bc is None and tie_book.get("opening") is not None: bc = round(tie_book["opening"] + tie_book["total_net"], 2)
    if sc is None and tie_bank.get("opening") is not None: sc = round(tie_bank["opening"] + tie_bank["total_net"], 2)
    recon_table = build_balance_reconciliation(bc, sc, unmatched_book, unmatched_bank)

    _p(80, "红旗检测与余额连贯性")
    red_flags = detect_red_flags(bank_std, "bank") + detect_red_flags(book_std, "book")
    red_flags += _check_balance_continuity(bank_std)
    red_flags += detect_fee_small_monthly_diff(book_std, bank_std)
    red_flags += detect_counterpart_month_gap(book_std, bank_std)
    # 费用标签错配检测：银方费用标签 对上了 账方非费用摘要
    for m in m1 + m2:
        if "bank_idx" in m and m["bank_idx"] in bank_std.index and m["book_idx"] in book_std.index:
            bk_tag = str(bank_std.loc[m["bank_idx"], "content_tag"])
            bj_sum = str(book_std.loc[m["book_idx"], "summary"])
            if "费用" in bk_tag and "手续费" not in bj_sum and "费用" not in bj_sum:
                red_flags.append({"type": "费用标签错配", "side": "双方",
                    "detail": f"账{m['book_idx']}({bj_sum[:20]}) 银{m['bank_idx']}({bk_tag}) {bank_std.loc[m['bank_idx'], 'net_amount']:.0f}元"})
    red_flags = aggregate_red_flags(red_flags)


    # ── 手续费匹配诊断 ──
    _fee_book_mask = book_std["summary"].astype(str).str.strip().str.contains(_FEE_KW, na=False)
    _fee_bank_mask = bank_std["summary"].astype(str).str.strip().str.contains(_FEE_KW, na=False)
    _fee_book_total = int(_fee_book_mask.sum())
    _fee_bank_total = int(_fee_bank_mask.sum())
    _fee_book_matched = int((_fee_book_mask & book_std.index.isin(book_used)).sum())
    _fee_bank_matched = int((_fee_bank_mask & bank_std.index.isin(bank_used)).sum())
    # 统计 fee 专用规则匹配到的行数
    _fee_matched_by_fee_rules = set()
    for m in m3_fee_month + m3_fee:
        for bi in m.get("bank_idxs", []):
            _fee_matched_by_fee_rules.add(bi)
    _fee_bank_by_fee_rules = len([i for i in _fee_matched_by_fee_rules if i in bank_std.index and _fee_bank_mask.get(i, False)])
    print(f"[费用诊断] ===== 手续费匹配诊断 =====")
    print(f"[费用诊断] 账方: 费用行 {_fee_book_total} 笔 → 已匹配 {_fee_book_matched} 笔 ({_fee_book_matched/_fee_book_total*100:.1f}%) → 未匹配 {_fee_book_total - _fee_book_matched} 笔" if _fee_book_total else "[费用诊断] 账方: 无费用行")
    print(f"[费用诊断] 银方: 费用行 {_fee_bank_total} 笔 → 已匹配 {_fee_bank_matched} 笔 ({_fee_bank_matched/_fee_bank_total*100:.1f}%) → 未匹配 {_fee_bank_total - _fee_bank_matched} 笔" if _fee_bank_total else "[费用诊断] 银方: 无费用行")
    print(f"[费用诊断] 匹配规则贡献: L3_fee_month={len(m3_fee_month)}组 L3_fee={len(m3_fee)}笔 → 银方由fee规则消耗约{_fee_bank_by_fee_rules}行")
    # 未匹配费用样本
    _fee_ub = [i for i in unmatched_book if any(kw in str(i.get("summary","")) for kw in FEE_WORDS)]
    _fee_uk = [i for i in unmatched_bank if any(kw in str(i.get("summary","")) for kw in FEE_WORDS)]
    if _fee_ub:
        print(f"[费用诊断] 账方未匹配费用 {len(_fee_ub)} 笔 (占未匹配{len(unmatched_book)}笔的 {len(_fee_ub)/max(1,len(unmatched_book))*100:.1f}%):")
        for u in _fee_ub[:5]:
            print(f"[费用诊断]   {u.get('date','')} {str(u.get('summary',''))[:30]} | {u.get('net_amount',0):,.2f} | {u.get('classification','')}")
    if _fee_uk:
        print(f"[费用诊断] 银方未匹配费用 {len(_fee_uk)} 笔 (占未匹配{len(unmatched_bank)}笔的 {len(_fee_uk)/max(1,len(unmatched_bank))*100:.1f}%):")
        for u in _fee_uk[:5]:
            print(f"[费用诊断]   {u.get('date','')} {str(u.get('summary',''))[:30]} | {u.get('net_amount',0):,.2f} | {u.get('classification','')}")
    _fee_book_gap = _fee_book_total - _fee_book_matched
    _fee_bank_gap = _fee_bank_total - _fee_bank_matched
    _gap_msg = []
    if _fee_book_gap: _gap_msg.append(f"账方{_fee_book_gap}笔")
    if _fee_bank_gap: _gap_msg.append(f"银方{_fee_bank_gap}笔")
    print(f"[费用诊断] 剩余缺口: {'，'.join(_gap_msg) if _gap_msg else '全部已匹配 ✅'}")

    n_book, n_bank = len(book_std), len(bank_std)
    stats = {
        "book_rows": n_book, "bank_rows": n_bank,
        "book_type_detected": book_type, "bank_type_detected": bank_type,
        "book_mapping": book_map, "bank_mapping": bank_map,
        "account_filter": account_note,
        "matched_L1": len(m1), "matched_L2": len(m2),
        "matched_L3_groups": len(m3) + len(m3_month) + len(m3_fee) + len(m3_fee_month) + len(matches_cp),
        "matched_L3_fee": len(m3_fee), "matched_L3_subset": len(m3), "review_L4": len(m4),
        "matched_L3_month": len(m3_month), "matched_L3_fee_month": len(m3_fee_month),
        "matched_L3_counterpart": len(matches_cp),
        "matched_L3_cp_n1": len(m3_cp_n1),
        "book_matched": len(book_used), "bank_matched": len(bank_used),
        "book_match_rate": round(len(book_used)/n_book*100,2) if n_book else 0.0,
        "bank_match_rate": round(len(bank_used)/n_bank*100,2) if n_bank else 0.0,
        "unmatched_book": len(unmatched_book), "unmatched_bank": len(unmatched_bank),
        "date_granularity": _granularity,
        "interbank_transfers": sum(1 for i in unmatched_book + unmatched_bank if i.get("classification") == "他行互转-待他行流水"),
        "timing_categories": {
            CAT_BANK_RECV: sum(1 for i in unmatched_bank if i["classification"]==CAT_BANK_RECV),
            CAT_BANK_PAY: sum(1 for i in unmatched_bank if i["classification"]==CAT_BANK_PAY),
            CAT_ENT_RECV: sum(1 for i in unmatched_book if i["classification"]==CAT_ENT_RECV),
            CAT_ENT_PAY: sum(1 for i in unmatched_book if i["classification"]==CAT_ENT_PAY),
            CAT_REVIEW: sum(1 for i in unmatched_book+unmatched_bank if i["classification"]==CAT_REVIEW),
        },
        "red_flag_count": len(red_flags),
        "tolerance": f"{tol}元", "date_window_days": date_window,
    }
    _p(100, "完成")
    return {"stats": stats, "tie_out": {"book": tie_book, "bank": tie_bank},
            "matches_L1": m1, "matches_L2": m2, "groups_L3": m3 + m3_month + m3_fee + m3_fee_month + matches_cp + m3_cp_n1, "review_L4": m4,
            "unmatched_book": unmatched_book, "unmatched_bank": unmatched_bank,
            "all_matches":m1 + m2 + m3 + m3_month + m3_fee + m3_fee_month + matches_cp + m4,
            "duplicates": duplicates, "balance_reconciliation": recon_table.to_dict("records"),
            "red_flags": red_flags, "book_std": book_std, "bank_std": bank_std, "config": cfg}


def _build_detail_workpaper(result):
    book_std = result["book_std"]; status = {}
    for m in result["matches_L1"] + result["matches_L2"]:
        status[m["book_idx"]] = {"状态":"已核对","层级":m["level"],"对方行":result["bank_std"].loc[m["bank_idx"],"row_id"],"备注":m["note"]}
    for m in result["groups_L3"]:
        tier = m.get("risk_tier", 0)
        if tier == 1:       st_label = "已核对(自动通过)"
        elif tier == 3:     st_label = "已核对(必须核查)"
        elif m.get("needs_review"): st_label = "已核对(抽样复核)"
        else:               st_label = "已核对(L3)"
        bank_list = m.get("bank_idxs") or ([m["bank_idx"]] if "bank_idx" in m else [])
        book_list = m.get("book_idxs") or ([m["book_idx"]] if "book_idx" in m else [])
        partners = ",".join(result["bank_std"].loc[b,"row_id"] for b in bank_list)
        for j in book_list: 
            status[j] = {"状态": st_label, "层级": m["level"], "对方行": partners, "备注": m["note"]}
    for m in result["review_L4"]:
        bank_list = m.get("bank_idxs") or ([m["bank_idx"]] if "bank_idx" in m else [])
        book_list = m.get("book_idxs") or ([m["book_idx"]] if "book_idx" in m else [])
        partners = ",".join(result["bank_std"].loc[b,"row_id"] for b in bank_list)
        for j in book_list:
            status[j] = {"状态":"待人工复核","层级":m["level"],"对方行":partners,"备注":m["note"]}
    unmatched_cls = {i["src_index"]:i for i in result["unmatched_book"]}
    rows = []
    for ji, r in book_std.iterrows():
        base = {"行号":r["row_id"],"日期":r["date"].date().isoformat() if pd.notna(r["date"]) else "","凭证号":r["voucher_no"],"摘要":r["summary"],"对方":r["counterpart"],"借方金额":round(float(r["debit"]),2),"贷方金额":round(float(r["credit"]),2),"净额":round(float(r["net_amount"]),2)}
        if ji in status: base.update(status[ji])
        else:
            cls = unmatched_cls.get(ji, {})
            base.update({"状态":cls.get("classification","待人工核查"),"层级":"-","对方行":"-","备注":cls.get("basis","")})
        rows.append(base)
    return pd.DataFrame(rows)




def _build_marked_sheets(result):
    """逐笔标记底稿：银行流水侧+序时账侧，每笔标注完整对账状态"""
    book_std, bank_std = result["book_std"], result["bank_std"]

    # 银行流水侧
    bank_status = {}
    for m in result["matches_L1"] + result["matches_L2"]:
        bank_status[m["bank_idx"]] = "对银成功"
    for m in result["groups_L3"]:
        tier = m.get("risk_tier", 0)
        if tier == 1:       label = "对银成功(自动通过)"
        elif tier == 3:     label = "对银成功(必须核查)"
        elif m.get("needs_review"): label = "对银成功(抽样复核)"
        else:               label = "对银成功(L3)"
        for bi in m.get("bank_idxs") or ([m["bank_idx"]] if "bank_idx" in m else []):
            bank_status[bi] = label
    for m in result["review_L4"]:
        for bi in m.get("bank_idxs") or ([m["bank_idx"]] if "bank_idx" in m else []):
            bank_status[bi] = "对银成功(L4)"
    bank_rows = []
    unmatched_bank_map = {i["src_index"]: i for i in result["unmatched_bank"]}
    for bi, r in bank_std.iterrows():
        row = {"行号": r["row_id"], "日期": str(r["date"])[:10], "摘要": r["summary"],
               "对方": r.get("counterpart", ""), "净额": round(float(r["net_amount"]), 2)}
        if bi in bank_status:
            row["对账状态"] = bank_status[bi]
        else:
            cls = unmatched_bank_map.get(bi, {})
            row["对账状态"] = cls.get("classification", "待人工核查")
            row["依据"] = cls.get("basis", "")
        bank_rows.append(row)
    bank_df = pd.DataFrame(bank_rows)

    # 序时账侧
    book_status = {}
    for m in result["matches_L1"] + result["matches_L2"]:
        book_status[m["book_idx"]] = "对账成功"
    for m in result["groups_L3"]:
        tier = m.get("risk_tier", 0)
        if tier == 1:       label = "对账成功(自动通过)"
        elif tier == 3:     label = "对账成功(必须核查)"
        elif m.get("needs_review"): label = "对账成功(抽样复核)"
        else:               label = "对账成功(L3)"
        for ji in m.get("book_idxs") or ([m["book_idx"]] if "book_idx" in m else []):
            book_status[ji] = label
    for m in result["review_L4"]:
        for ji in m.get("book_idxs") or ([m["book_idx"]] if "book_idx" in m else []):
            book_status[ji] = "对账成功(L4)"

    book_rows = []
    unmatched_book_map = {i["src_index"]: i for i in result["unmatched_book"]}
    for ji, r in book_std.iterrows():
        row = {"行号": r["row_id"], "日期": str(r["date"])[:10], "凭证号": r["voucher_no"],
               "摘要": r["summary"], "对方": r.get("counterpart", ""),
               "借方": round(float(r["debit"]), 2), "贷方": round(float(r["credit"]), 2),
               "净额": round(float(r["net_amount"]), 2)}
        if ji in book_status:
            row["对账状态"] = book_status[ji]
        else:
            cls = unmatched_book_map.get(ji, {})
            row["对账状态"] = cls.get("classification", "待人工核查")
            row["依据"] = cls.get("basis", "")
        book_rows.append(row)
    book_df = pd.DataFrame(book_rows)
    return bank_df, book_df




def _build_counterparty_summary(result):
    """对手方收付汇总表：按对方单位聚合银方/账方流量，一次覆盖往来函证对账、账龄回款核销、关联方识别、截止性测试四大痛点。"""
    book_std = result["book_std"]
    bank_std = result["bank_std"]

    # 构建已匹配索引集合（从 all_matches 提取）
    book_matched_idx = set()
    bank_matched_idx = set()
    for m in result["all_matches"]:
        for ji in m.get("book_idxs") or ([m["book_idx"]] if "book_idx" in m else []):
            book_matched_idx.add(ji)
        for bi in m.get("bank_idxs") or ([m["bank_idx"]] if "bank_idx" in m else []):
            bank_matched_idx.add(bi)

    # 银方按对手方聚合
    bank_cp = bank_std.copy()
    bank_cp["matched"] = bank_cp.index.isin(bank_matched_idx)
    bank_cp["收入"] = bank_cp["net_amount"].clip(lower=0)
    bank_cp["支出"] = (-bank_cp["net_amount"]).clip(lower=0)
    bank_grp = bank_cp.groupby("counterpart").agg(
        银方笔数=("net_amount", "count"),
        银方收入=("收入", "sum"),
        银方支出=("支出", "sum"),
        银方净额=("net_amount", "sum"),
        银方已匹配笔数=("matched", "sum"),
        银方已匹配金额=("net_amount", lambda x: x[bank_cp.loc[x.index, "matched"]].abs().sum()),
        银方未匹配笔数=("matched", lambda x: (~x).sum()),
        银方未匹配金额=("net_amount", lambda x: x[~bank_cp.loc[x.index, "matched"]].abs().sum()),
    ).round(2)

    # 账方按对手方聚合
    book_cp = book_std.copy()
    book_cp["matched"] = book_cp.index.isin(book_matched_idx)
    book_cp["借方"] = book_cp["debit"].clip(lower=0)
    book_cp["贷方"] = book_cp["credit"].clip(lower=0)
    book_grp = book_cp.groupby("counterpart").agg(
        账方笔数=("net_amount", "count"),
        账方借方=("借方", "sum"),
        账方贷方=("贷方", "sum"),
        账方净额=("net_amount", "sum"),
        账方已匹配笔数=("matched", "sum"),
        账方已匹配金额=("net_amount", lambda x: x[book_cp.loc[x.index, "matched"]].abs().sum()),
        账方未匹配笔数=("matched", lambda x: (~x).sum()),
        账方未匹配金额=("net_amount", lambda x: x[~book_cp.loc[x.index, "matched"]].abs().sum()),
    ).round(2)

    # 合并
    df = bank_grp.join(book_grp, how="outer").fillna(0)
    df["差异(银-账)"] = (df["银方净额"] - df["账方净额"]).round(2)
    # 总笔数
    df["银方笔数"] = df["银方笔数"].astype(int)
    df["账方笔数"] = df["账方笔数"].astype(int)
    df["银方已匹配笔数"] = df["银方已匹配笔数"].astype(int)
    df["银方未匹配笔数"] = df["银方未匹配笔数"].astype(int)
    df["账方已匹配笔数"] = df["账方已匹配笔数"].astype(int)
    df["账方未匹配笔数"] = df["账方未匹配笔数"].astype(int)

    # 风险标记
    flags = []
    for cp in df.index:
        f = []
        row = df.loc[cp]
        if row["差异(银-账)"] != 0 and abs(row["差异(银-账)"]) > 100:
            f.append("差异>$100")
        if row["银方未匹配笔数"] >= 5 or row["账方未匹配笔数"] >= 5:
            f.append("未匹配多笔")
        if cp and cp.strip():
            for kw in ("方万鹏", "赵萍", "格林威特", "东晨", "同瑞", "手足"):
                if kw in str(cp):
                    f.append("关联方")
                    break
        flags.append(",".join(f) if f else "")
    df["风险标记"] = flags

    # 集中度：占比和Pareto
    total_abs = df["银方净额"].abs().sum() + df["账方净额"].abs().sum()
    if total_abs > 0:
        df["金额占比%"] = ((df["银方净额"].abs() + df["账方净额"].abs()) / total_abs * 100).round(1)
        cum = 0
        pareto = []
        for v in df["金额占比%"]:
            cum += v
            if cum <= 80:
                pareto.append("A类(80%)")
            elif cum <= 95:
                pareto.append("B类(15%)")
            else:
                pareto.append("C类(5%)")
        df["Pareto分层"] = pareto

    # 排序：风险标记非空在前，差异绝对值大的在前
    df["_sort"] = df["风险标记"].apply(lambda x: 0 if x else 1) * 10000 + df["差异(银-账)"].abs()
    df = df.sort_values("_sort").drop(columns=["_sort"])
    df = df.reset_index().rename(columns={"counterpart": "对方名称"})
    return df



def _build_trend_analysis(result):
    """趋势分析：按月汇总银方/账方收入支出"""
    bank_std = result["bank_std"]
    book_std = result["book_std"]
    bank_cp = bank_std.copy()
    bank_cp["月份"] = pd.to_datetime(bank_cp["date"], errors="coerce").dt.to_period("M").astype(str)
    bank_cp["收入"] = bank_cp["net_amount"].clip(lower=0)
    bank_cp["支出"] = (-bank_cp["net_amount"]).clip(lower=0)
    bm = bank_cp.groupby("月份").agg(银方笔数=("net_amount", "count"), 银方收入=("收入", "sum"), 银方支出=("支出", "sum")).round(2)
    book_cp = book_std.copy()
    book_cp["月份"] = pd.to_datetime(book_cp["date"], errors="coerce").dt.to_period("M").astype(str)
    book_cp["借方"] = book_cp["debit"].clip(lower=0)
    book_cp["贷方"] = book_cp["credit"].clip(lower=0)
    jm = book_cp.groupby("月份").agg(账方笔数=("net_amount", "count"), 账方借方=("借方", "sum"), 账方贷方=("贷方", "sum")).round(2)
    df = bm.join(jm, how="outer").fillna(0).sort_index()
    df["银方净流入"] = (df["银方收入"] - df["银方支出"]).round(2)
    df["账方净发生"] = (df["账方借方"] - df["账方贷方"]).round(2)
    df["差异"] = (df["银方净流入"] - df["账方净发生"]).round(2)
    # 异常标记：差异超过月均2倍标准差
    if len(df) >= 3:
        std = df["差异"].std()
        if std > 0:
            df["异常标记"] = df["差异"].apply(lambda x: "异常" if abs(x) > std * 2 else "")
    return df.reset_index()



def _build_cutoff_test(result, days=7):
    """截止性测试：期末前后大额交易"""
    bank_std = result["bank_std"]
    book_std = result["book_std"]
    max_bd = pd.to_datetime(book_std["date"], errors="coerce").max()
    max_bb = pd.to_datetime(bank_std["date"], errors="coerce").max()
    TE = 50000
    rows = []
    for label, df, max_d in [("银方", bank_std, max_bb), ("账方", book_std, max_bd)]:
        if pd.isna(max_d):
            continue
        cutoff = max_d - pd.Timedelta(days=days)
        df2 = df.copy()
        df2["_dt"] = pd.to_datetime(df2["date"], errors="coerce")
        mask = (df2["_dt"] >= cutoff) & (abs(df2["net_amount"]) > TE)
        for _, r in df2[mask].iterrows():
            rows.append({
                "侧别": label, "日期": str(r["date"])[:10],
                "摘要": r["summary"], "对方": r.get("counterpart", ""),
                "金额": round(float(r["net_amount"]), 2),
                "距期末天数": (max_d - r["_dt"]).days
            })
    return pd.DataFrame(rows).sort_values("金额", key=abs, ascending=False) if rows else pd.DataFrame()



def _build_fee_monthly_analysis(result):
    """费用月度分析：银行手续费按月趋势"""
    bank_std = result["bank_std"]
    FEE_KW = ("手续费", "短信费", "年费", "账户管理费", "工本费", "服务费", "扣费")
    mask = bank_std["summary"].astype(str).apply(lambda s: any(kw in s for kw in FEE_KW))
    fee = bank_std[mask].copy()
    if fee.empty:
        return pd.DataFrame()
    fee["月份"] = pd.to_datetime(fee["date"], errors="coerce").dt.to_period("M").astype(str)
    fm = fee.groupby("月份").agg(笔数=("net_amount", "count"), 费用合计=("net_amount", lambda x: x.abs().sum())).round(2)
    fm["笔均"] = (fm["费用合计"] / fm["笔数"]).round(2)
    # 异常月份：费用超过月均2倍
    if len(fm) >= 3:
        avg = fm["费用合计"].mean()
        fm["异常标记"] = fm["费用合计"].apply(lambda x: "偏高" if x > avg * 2 else ("偏低" if x < avg * 0.3 else ""))
    return fm.sort_index().reset_index()



def _build_aging_by_counterparty(result):
    """账龄分析：按对手方+账龄区间分桶未匹配金额"""
    ub = result["unmatched_book"]
    uk = result["unmatched_bank"]
    max_date = pd.Timestamp.now()
    # 尝试从数据中提取最大日期
    for items in [ub, uk]:
        for it in items:
            d = it.get("date", "")
            if d:
                try:
                    dt = pd.Timestamp(d)
                    if dt > max_date or max_date == pd.Timestamp.now():
                        max_date = dt
                except:
                    pass
    rows = []
    for label, items in [("账方", ub), ("银方", uk)]:
        for it in items:
            cp = str(it.get("counterpart", "")).strip()
            if not cp:
                cp = "(无对手方)"
            amt = abs(it.get("net_amount", 0))
            d = it.get("date", "")
            days = 999
            if d:
                try:
                    days = (max_date - pd.Timestamp(d)).days
                except:
                    pass
            if days <= 90:
                bucket = "0-90天"
            elif days <= 180:
                bucket = "91-180天"
            elif days <= 365:
                bucket = "181-365天"
            else:
                bucket = "365天以上"
            rows.append({"侧别": label, "对方名称": cp, "日期": str(d)[:10] if d else "",
                         "金额": round(amt, 2), "账龄天数": days, "账龄区间": bucket})
    df = pd.DataFrame(rows)
    if df.empty:
        return df, pd.DataFrame()
    # 明细
    detail = df.sort_values(["对方名称", "账龄天数"], ascending=[True, False])
    # 汇总：按对手方+账龄区间聚合
    summary = df.groupby(["对方名称", "账龄区间"]).agg(
        笔数=("金额", "count"), 合计金额=("金额", "sum")
    ).round(2).reset_index()
    # pivot
    pivot = summary.pivot_table(index="对方名称", columns="账龄区间", values="合计金额", aggfunc="sum", fill_value=0)
    for col in ["0-90天", "91-180天", "181-365天", "365天以上"]:
        if col not in pivot.columns:
            pivot[col] = 0
    pivot["合计"] = pivot.sum(axis=1)
    pivot = pivot.sort_values("合计", ascending=False)
    return detail, pivot.reset_index()



def _generate_confirmation_letters(result, out_dir):
    """生成询证函Word文档：按模板填充对手方、金额、期间"""
    from docx import Document
    cp_df = _build_counterparty_summary(result)
    top = cp_df.head(15)
    out = Path(out_dir) / "询证函"
    out.mkdir(parents=True, exist_ok=True)
    st = result["stats"]
    date_range = ""
    try:
        bd = pd.to_datetime(result["book_std"]["date"], errors="coerce")
        date_range = bd.min().strftime("%Y/%m/%d") + " 至 " + bd.max().strftime("%Y/%m/%d")
    except:
        pass
    written = 0
    for _, row in top.iterrows():
        cp_name = str(row["对方名称"]).strip()
        if not cp_name or cp_name == "nan":
            continue
        doc = Document()
        doc.add_heading("企业询证函", level=0)
        doc.add_paragraph(f"编号: XC{written+1:03d}")
        doc.add_paragraph(f"致: {cp_name}")
        doc.add_paragraph("")
        doc.add_paragraph(f"本公司聘请的审计机构正在对本公司{date_range}的财务报表进行审计。")
        doc.add_paragraph("按照中国注册会计师审计准则的要求，应当询证本公司与贵单位的往来账项等事项。")
        doc.add_paragraph("下列数据出自本公司账簿记录，如与贵单位记录相符，请在本函下端'信息证明无误'处签章证明；")
        doc.add_paragraph('如有不符，请在"信息不符"处列明不符项目及金额。回函请直接寄至审计机构。')
        doc.add_paragraph("")
        # 往来明细表
        table = doc.add_table(rows=3, cols=5, style="Table Grid")
        headers = ["项目", "期间", "本公司账面金额", "贵单位账面金额", "差异"]
        for i, h in enumerate(headers):
            table.rows[0].cells[i].text = h
        table.rows[1].cells[0].text = "应收/应付余额"
        table.rows[1].cells[1].text = date_range
        table.rows[1].cells[2].text = f"{row['账方净额']:,.2f}"
        table.rows[1].cells[3].text = ""
        table.rows[1].cells[4].text = ""
        table.rows[2].cells[0].text = "其中: 银行流水净额"
        table.rows[2].cells[1].text = date_range
        table.rows[2].cells[2].text = f"{row['银方净额']:,.2f}"
        table.rows[2].cells[3].text = ""
        table.rows[2].cells[4].text = ""
        doc.add_paragraph("")
        doc.add_paragraph("结论: 1.信息证明无误    2.信息不符(请列明)")
        doc.add_paragraph("")
        doc.add_paragraph(f"经办人: ________    日期: ________")
        doc.add_paragraph(f"回函地址: ________")
        safe_name = cp_name.replace("/", "_").replace("\\", "_")[:30]
        doc.save(str(out / f"询证函_{safe_name}.docx"))
        written += 1
    print(f"[询证函] 已生成{written}份Word文档 -> {out}")
    return written

def export_reconciliation_outputs(result, out_dir):
    out = Path(out_dir); out.mkdir(parents=True, exist_ok=True); written = []
    detail = _build_detail_workpaper(result)
    p = out / "逐笔对账明细底稿.xlsx"; detail.to_excel(str(p), index=False); written.append(p.name)
    recon = pd.DataFrame(result["balance_reconciliation"])
    # ↓↓↓ 新增：日期口径披露（有偏移时才出现）↓↓↓
    if result["stats"].get("date_granularity") == "month":
        recon = pd.concat([recon, pd.DataFrame([
            {"项目": "—— 口径说明 ——", "金额": ""},
            {"项目": "日期口径：账面为月末批量记账，本次按 年月+金额 勾对（日信息不参与）。", "金额": ""},
        ])], ignore_index=True)
    # ↑↑↑ 新增结束 ↑↑↑
    p = out / "银行存款余额调节表.xlsx"; recon.to_excel(str(p), index=False); written.append(p.name)
    un_book = [{**i,"方向":"企业账"} for i in result["unmatched_book"]]
    un_bank = [{**i,"方向":"银行流水"} for i in result["unmatched_bank"]]
    unmatched = pd.DataFrame(un_book + un_bank)
    if not unmatched.empty: unmatched = unmatched.rename(columns={"date":"日期","net_amount":"净额","summary":"摘要","counterpart":"对方","classification":"分类","basis":"依据","special":"标记"})
    p = out / "未达账项与待核查清单.xlsx"; unmatched.to_excel(str(p), index=False); written.append(p.name)
    # v3.9: 逐笔标记底稿
    try:
        bank_marked, book_marked = _build_marked_sheets(result)
        bank_marked.to_excel(str(out / "银行流水匹配情况.xlsx"), index=False)
        book_marked.to_excel(str(out / "序时账匹配情况.xlsx"), index=False)
        written += ["银行流水匹配情况.xlsx", "序时账匹配情况.xlsx"]
    except Exception as _e:
        print(f"[export] 逐笔标记底稿（匹配情况）生成失败: {type(_e).__name__}: {_e}")
    flags = pd.DataFrame(result["red_flags"] + [{**d,"type":"疑似重复入账","detail":d["reason"]} for d in result["duplicates"]])
    if not flags.empty: flags = flags.rename(columns={"type":"类型","side":"侧别","rows":"涉及行","amount":"金额","detail":"说明"})
    p = out / "异常资金交易清单.xlsx"; flags.to_excel(str(p), index=False); written.append(p.name)
    # v3.11: 对手方收付汇总表（含集中度Pareto）
    try:
        cp_df = _build_counterparty_summary(result)
        cp_df.to_excel(str(out / "对手方收付汇总表.xlsx"), index=False)
        written.append("对手方收付汇总表.xlsx")
        n_cp = len(cp_df)
        n_flag = (cp_df["风险标记"] != "").sum()
        print(f"[对手方汇总] {n_cp}个对手方，{n_flag}个有风险标记")
    except Exception as _e:
        print(f"[对手方汇总] 生成失败: {_e}")
    # v3.12: 趋势分析
    try:
        trend = _build_trend_analysis(result)
        trend.to_excel(str(out / "月度趋势分析.xlsx"), index=False)
        written.append("月度趋势分析.xlsx")
        n_abn = (trend.get("异常标记", "") != "").sum() if "异常标记" in trend.columns else 0
        print(f"[趋势分析] {len(trend)}个月，{n_abn}个月异常波动")
    except Exception as _e:
        print(f"[趋势分析] 生成失败: {_e}")
    # v3.12: 截止性测试
    try:
        cutoff = _build_cutoff_test(result, days=7)
        if not cutoff.empty:
            cutoff.to_excel(str(out / "截止性测试.xlsx"), index=False)
            written.append("截止性测试.xlsx")
            print(f"[截止性] 期末前后大额交易{len(cutoff)}笔")
    except Exception as _e:
        print(f"[截止性] 生成失败: {_e}")
    # v3.12: 费用月度分析
    try:
        fee_m = _build_fee_monthly_analysis(result)
        if not fee_m.empty:
            fee_m.to_excel(str(out / "银行费用月度分析.xlsx"), index=False)
            written.append("银行费用月度分析.xlsx")
            n_fee_abn = (fee_m.get("异常标记", "") != "").sum() if "异常标记" in fee_m.columns else 0
            print(f"[费用分析] {len(fee_m)}个月费用记录，{n_fee_abn}个月异常")
    except Exception as _e:
        print(f"[费用分析] 生成失败: {_e}")
    # v3.12: 整数大额清单
    try:
        bank_std = result["bank_std"]
        big_mask = (abs(bank_std["net_amount"]) >= 50000) & (abs(bank_std["net_amount"]) % 10000 == 0)
        big = bank_std[big_mask][["date", "summary", "counterpart", "net_amount"]].copy()
        if not big.empty:
            big.columns = ["日期", "摘要", "对方", "金额"]
            big["日期"] = big["日期"].astype(str).str[:10]
            big.to_excel(str(out / "整数大额交易清单.xlsx"), index=False)
            written.append("整数大额交易清单.xlsx")
            print(f"[整数大额] {len(big)}笔")
    except Exception as _e:
        print(f"[整数大额] 生成失败: {_e}")
    # v3.12: 函证数据（Top对手方）
    try:
        cp_df2 = _build_counterparty_summary(result)
        top_cp = cp_df2.head(20)[["对方名称", "银方净额", "账方净额", "差异(银-账)", "风险标记"]]
        top_cp.columns = ["被询证单位", "银行流水净额", "账面净额", "差异", "注意事项"]
        top_cp.to_excel(str(out / "函证数据清单.xlsx"), index=False)
        written.append("函证数据清单.xlsx")
        print(f"[函证] Top20对手方函证数据已生成")
    except Exception as _e:
        print(f"[函证] 生成失败: {_e}")
    # v3.12: 报告数据提取JSON
    try:
        st = result["stats"]
        cp_df3 = _build_counterparty_summary(result)
        flag_cps = cp_df3[cp_df3["风险标记"] != ""]["对方名称"].tolist()
        report_data = {
            "基本信息": {
                "账方匹配率": f"{st['book_match_rate']}%",
                "银方匹配率": f"{st['bank_match_rate']}%",
                "未匹配_账方": st["unmatched_book"],
                "未匹配_银方": st["unmatched_bank"],
                "待复核_L4": st.get("review_L4", 0),
                "红旗数量": st.get("red_flag_count", 0),
            },
            "未达账项": {
                "银收企未收": st["timing_categories"].get(CAT_BANK_RECV, 0),
                "银付企未付": st["timing_categories"].get(CAT_BANK_PAY, 0),
                "企收银未收": st["timing_categories"].get(CAT_ENT_RECV, 0),
                "企付银未付": st["timing_categories"].get(CAT_ENT_PAY, 0),
                "待人工核查": st["timing_categories"].get(CAT_REVIEW, 0),
            },
            "风险对手方": flag_cps[:20],
            "异常月份": trend[trend.get("异常标记", "") != ""]["月份"].tolist() if "异常标记" in trend.columns else [],
        }
        (out / "报告数据提取.json").write_text(
            json.dumps(report_data, ensure_ascii=False, indent=2, default=str),
            encoding="utf-8")
        written.append("报告数据提取.json")
        print(f"[报告数据] JSON已提取")
    except Exception as _e:
        print(f"[报告数据] 提取失败: {_e}")
    # v3.13: 账龄分析
    try:
        aging_detail, aging_pivot = _build_aging_by_counterparty(result)
        if not aging_detail.empty:
            aging_detail.to_excel(str(out / "账龄分析明细.xlsx"), index=False)
            aging_pivot.to_excel(str(out / "账龄分析汇总.xlsx"), index=False)
            written += ["账龄分析明细.xlsx", "账龄分析汇总.xlsx"]
            print(f"[账龄分析] {len(aging_detail)}笔未达，{len(aging_pivot)}个对手方")
    except Exception as _e:
        print(f"[账龄分析] 生成失败: {_e}")
    # v3.13: 三方圆勾稽
    try:
        st2 = result["stats"]
        tie = result["tie_out"]
        tie_rows = []
        for side, t in [("账方", tie.get("book", {})), ("银方", tie.get("bank", {}))]:
            opening = t.get("opening", 0) or 0
            total_net = t.get("total_net", 0) or 0
            closing = t.get("closing", 0) or 0
            expected = round(opening + total_net, 2)
            diff = round(expected - (closing or 0), 2)
            tie_rows.append({"侧别": side, "期初余额": opening, "本期净发生额": total_net,
                             "计算期末": expected, "账面期末": closing or 0, "差异": diff,
                             "状态": "平衡" if abs(diff) < 0.02 else "异常"})
        tie_df = pd.DataFrame(tie_rows)
        tie_df.to_excel(str(out / "三方圆勾稽.xlsx"), index=False)
        written.append("三方圆勾稽.xlsx")
        abn = (tie_df["状态"] == "异常").sum()
        status_text = "平衡" if abn == 0 else f"{abn}侧异常"; print(f"[三方勾稽] {status_text}")
    except Exception as _e:
        print(f"[三方勾稽] 生成失败: {_e}")
    summary = {"stats":result["stats"],"tie_out":result["tie_out"],"reconciliation":result["balance_reconciliation"]}
    p = out / "reconciliation_summary.json"; p.write_text(json.dumps(summary,ensure_ascii=False,indent=2,default=str),encoding="utf-8"); written.append(p.name)
     # ── v3.9: 报告生成器数据契约（journal_entries.json + analysis_result.csv）──
    st = result["stats"]
    book_std = result["book_std"]
    _tot_amt = float(book_std["net_amount"].abs().sum()) if len(book_std) else 0.0
    _ub_amt = sum(abs(i.get("net_amount", 0.0)) for i in result["unmatched_book"])
    _diff_pct = round(_ub_amt / _tot_amt * 100, 2) if _tot_amt else 0.0

    journal_entries = {
        "total_rows": int(st.get("book_rows", 0)),
        "columns": ["行号", "日期", "凭证号", "摘要", "对方", "借方金额", "贷方金额",
                    "净额", "状态", "层级", "对方行", "备注"],
        "numeric_summary": {
            "净额": {
                "sum": round(float(book_std["net_amount"].sum()), 2),
                "mean": round(float(book_std["net_amount"].mean()), 2),
                "max": round(float(book_std["net_amount"].max()), 2),
                "min": round(float(book_std["net_amount"].min()), 2),
            }
        },
        "match_stats": {
            "total_left": int(st.get("book_rows", 0)),
            "total_right": int(st.get("bank_rows", 0)),
            "matched_count": int(st.get("book_matched", 0)),
            "match_rate": float(st.get("book_match_rate", 0.0)),   # 百分数刻度（阈值>90/>70）
            "diff_percentage": _diff_pct,                          # 未匹配金额占比（阈值5/15/30）
            "interbank_transfers": st.get("interbank_transfers", 0),
        },
    }
    p = out / "journal_entries.json"
    p.write_text(json.dumps(journal_entries, ensure_ascii=False, indent=2), encoding="utf-8")
    written.append(p.name)

    # CSV 预览（报告"数据明细前50行"用）
    detail.head(51).to_csv(str(out / "analysis_result.csv"), index=False, encoding="utf-8-sig")
    written.append("analysis_result.csv")
    # 非银行费用匹配结果
    if result.get("nonbank_fee_hints"):
        pd.DataFrame(result["nonbank_fee_hints"]).to_excel(str(out / "非银行科目费用匹配.xlsx"), index=False)
        written.append("非银行科目费用匹配.xlsx")

    return written


def match_fee_to_nonbank(result, non_bank):
    """银方费用 -> 非银行科目 n:1月度等额匹配。消耗已命中的银方条目。
    """
    from config.dictionary import FEE_WORDS
    from itertools import combinations
    _FKW = "|".join(FEE_WORDS)
    unmatched = result.get("unmatched_bank", [])
    if not unmatched or non_bank is None or non_bank.empty: return []
    kdf = pd.DataFrame(unmatched) if not isinstance(unmatched, pd.DataFrame) else unmatched
    kdf = kdf[
        _is_fee_row(kdf, _FKW) |
        ((kdf["summary"].astype(str).str.strip().isin(["", "nan", "None", "nat"]) |
          kdf["summary"].isna()) & (kdf["net_amount"].abs() <= 50))
    ]
    if kdf.empty: return []
    nb = non_bank[non_bank["net_cents"] < 0].copy()
    nb = nb[_is_fee_row(nb, _FKW)]
    if nb.empty: return []
    kdf = kdf.copy(); kdf["_m"] = pd.to_datetime(kdf["date"], errors="coerce").dt.to_period("M")
    kdf = kdf.reset_index(drop=True)
    nb["_m"] = pd.to_datetime(nb["date"], errors="coerce").dt.to_period("M")
    k_by_month = {m: g for m, g in kdf.groupby("_m", sort=False)}
    hints, n, matched_ids = [], 0, set()
    for km, kg in k_by_month.items():
        parts = []
        for offset in [0, 1]:
            try: parts.append(nb[nb["_m"] == km + offset])
            except: pass
        parts = [p for p in parts if not p.empty]
        if not parts: continue
        pool = pd.concat(parts)
        kg_sorted = kg.sort_values("date") if "date" in kg.columns else kg
        vals = [(int(i), int(round(abs(r["net_amount"]) * 100))) for i, r in kg_sorted.iterrows()]
        for _, nr in pool.sort_values("net_cents").iterrows():
            target = int(round(abs(nr["net_amount"]) * 100))
            found = None
            for w in range(len(vals)):
                total = 0
                for j in range(w, min(w + 20, len(vals))):
                    total += vals[j][1]
                    if total == target: found = [vals[k][0] for k in range(w, j + 1)]; break
                    if total > target: break
                if found: break
            if found is None:
                closest = sorted(vals, key=lambda x: abs(x[1] - target))[:20]
                for m in range(1, 5):
                    for combo in combinations(range(len(closest)), m):
                        if sum(closest[i][1] for i in combo) == target:
                            found = [closest[i][0] for i in combo]; break
                    if found: break
            if found:
                n += 1
                matched_ids.update(found)
                hints.append({
                    "nonbank_date": str(nr.get("date", ""))[:10],
                    "nonbank_amt": round(target / 100, 2),
                    "nonbank_subject": str(nr.get("subject", ""))[:40],
                    "bank_count": len(found),
                    "month": str(km),
                })
                vals = [(i, v) for i, v in vals if i not in found]
    if hints:
        print(f"[非银行费用] 命中 {n} 笔 (仅参考，不消耗条目)")
        for h in hints[:10]:
            print(f"[非银行费用]   {h['nonbank_date']} 账{h['nonbank_amt']:>10,.2f} {h['nonbank_subject'][:25]} <- 银{h['bank_count']}笔")
        result["nonbank_fee_hints"] = hints
    else:
        print("[非银行费用] 无匹配")
    return hints

def reconcile_files(book_path, bank_path, config=None, out_dir=None):
    from core.document_loader import load_tables
    cfg = dict(config or {})
    book_tables = load_tables(book_path); bank_tables = load_tables(bank_path)
    if not book_tables: raise ValueError(f"无法读取: {book_path}")
    if not bank_tables: raise ValueError(f"无法读取: {bank_path}")
    cfg.setdefault("book_file", Path(book_path).name)
    cfg.setdefault("bank_file", Path(bank_path).name)
    progress_cb = cfg.pop("progress_callback", None)
    book_raw = book_tables[0]
    bank_raw = bank_tables[0]
    bt = detect_book_type(book_raw, cfg.get("book_file", ""))
    bm = auto_map_columns(book_raw, bt)
    book_norm = normalize_to_std(book_raw, bm, JOURNAL if bt not in (JOURNAL, BANK_STATEMENT) else bt)

    subj_col, warns = recognize_subject_column(book_norm)
    subjects, non_bank = None, None
    if subj_col: subjects, non_bank = split_by_bank_subject(book_norm, subj_col)
    if not subjects or len(subjects) <= 1:
        result = run_bank_reconciliation(book_raw, bank_raw, cfg, progress_callback=progress_cb)
        if out_dir: result["output_files"] = export_reconciliation_outputs(result, out_dir)
        return result
    main_label = max(subjects, key=lambda k: len(subjects[k]))
    main_df = subjects[main_label]
    other_subjects = {k:v for k,v in subjects.items() if k != main_label}
    print(f"[pipeline] main: {main_label} ({len(main_df)}r), others: {list(other_subjects.keys())}")
    result = run_bank_reconciliation(main_df, bank_raw, cfg, progress_callback=progress_cb)
    if non_bank is not None and not non_bank.empty:
        match_fee_to_nonbank(result, non_bank)
    if out_dir: result["output_files"] = export_reconciliation_outputs(result, out_dir)
    if other_subjects:
        hints = explain_across_accounts(result.get("unmatched_bank", []), other_subjects, book_norm, out_dir)
        result["cross_hints"] = hints
    result["pipeline"] = {"main": main_label, "others": list(other_subjects.keys()),
                          "non_bank": len(non_bank) if non_bank is not None else 0}
    return result


if __name__ == "__main__":
    import sys
    if len(sys.argv) < 3:
        print("用法: python -m core.bank_reconcile_engine <序时账> <流水> [输出目录]")
        sys.exit(1)
    odir = sys.argv[3] if len(sys.argv) > 3 else "outputs/reconcile"
    res = reconcile_files(sys.argv[1], sys.argv[2], out_dir=odir)
    s = res["stats"]
    print(f"账面{s['book_rows']}笔(匹配率{s['book_match_rate']}%)|流水{s['bank_rows']}笔(匹配率{s['bank_match_rate']}%)")
    print(f"L1={s['matched_L1']} L2={s['matched_L2']} L3={s['matched_L3_groups']} L4={s['review_L4']}")
    print("未达四分类:", json.dumps(s["timing_categories"], ensure_ascii=False))
    print("交付物:", res.get("output_files"))
