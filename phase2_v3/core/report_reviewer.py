"""
报告复核工具 v2.0 - 一键审计报告质量检查
用法:
    from core.report_reviewer import review_report
    df = review_report("审计报告.docx", {"trial_balance": "科目余额表.xlsx"})
    # LLM 自动从环境变量 VLLM_TUNNEL_URL 读取，默认 localhost:18000
"""

import re, json, os
import pandas as pd
import requests
from pathlib import Path

VLLM_URL = os.environ.get("VLLM_TUNNEL_URL", "http://localhost:18000/v1/chat/completions")
VLLM_KEY = os.environ.get("VLLM_API_KEY", "EMPTY")
VLLM_MODEL_NAME = os.environ.get("VLLM_MODEL", "qwen3-235b")


class ReportReviewer:

    def __init__(self, report_path=None, data_sources=None, prior_report_docx=None):
        self.report_path = Path(report_path) if report_path else None
        self.prior_report_docx = Path(prior_report_docx) if prior_report_docx else None
        self.src = data_sources or {}
        self.llm_url = VLLM_URL
        self.llm_key = VLLM_KEY
        self.llm_model = VLLM_MODEL_NAME
        self.tables = []
        self.findings = []
        self.coverage = {}

    def set_llm(self, url=None, key=None, model=None):
        if url: self.llm_url = url
        if key: self.llm_key = key
        if model: self.llm_model = model

    def disable_llm(self):
        self.llm_url = None

    def _call_llm(self, prompt, max_tokens=2000):
        resp = requests.post(self.llm_url,
            headers={"Authorization": "Bearer " + self.llm_key, "Content-Type": "application/json"},
            json={"model": self.llm_model, "messages": [{"role": "user", "content": prompt}],
                  "temperature": 0.1, "max_tokens": max_tokens}, timeout=120)
        if resp.status_code == 200:
            return resp.json().get("choices", [{}])[0].get("message", {}).get("content", "")
        return ""

    def _parse_json(self, text):
        for left, right in [("[", "]"), ("{", "}")]:
            s = text.find(left)
            e = text.rfind(right) + 1
            if s >= 0 and e > s:
                try:
                    return json.loads(text[s:e])
                except:
                    pass
        return None

    # -- 表格提取 --
    def _extract_docx_tables(self):
        try:
            from docx import Document
        except ImportError:
            return
        doc = Document(str(self.report_path))
        for ti, table in enumerate(doc.tables):
            rows = [[cell.text.strip().replace(chr(10), " ") for cell in row.cells] for row in table.rows]
            if rows:
                df = pd.DataFrame(rows[1:], columns=rows[0]) if len(rows) > 1 else pd.DataFrame(rows)
                self.tables.append({"index": ti, "data": df})

    # -- 1. 合计=分项和 --
    def _check_table_formulas(self):
        for t in self.tables:
            df = t["data"]
            for col in df.columns:
                vals = pd.to_numeric(df[col], errors="coerce")
                if vals.count() < 3:
                    continue
                total_mask = df.iloc[:, 0].astype(str).str.contains("合计|总计|小计", na=False)
                if total_mask.sum() == 0:
                    continue
                tv = pd.to_numeric(df[total_mask].iloc[0][col], errors="coerce")
                if pd.isna(tv):
                    continue
                detail_sum = vals[~total_mask].dropna().sum()
                diff = round(abs(tv - detail_sum), 2)
                if diff > 0.1:
                    self.findings.append({"类别": "公式校验",
                        "位置": "表{}.{}".format(t["index"] + 1, col),
                        "检查项": "合计!=分项和", "报表值": tv,
                        "计算值": round(detail_sum, 2), "差异": diff, "结果": "异常"})

    # -- 2. 资产负债表勾稽 --
    def _check_balance_sheet_equation(self, tb):
        if tb is None:
            self.coverage["资产负债勾稽"] = "跳过(缺科目余额表)"
            return

        subj_col = amt_col = None
        for c in tb.columns:
            s = str(c)
            if any(kw in s for kw in ["科目", "名称"]):
                subj_col = c
            if any(kw in s for kw in ["期末余额", "余额", "金额"]) and "期初" not in s:
                amt_col = c
        if subj_col is None or amt_col is None:
            return
        a = l = e = 0.0
        for _, row in tb.iterrows():
            name = str(row[subj_col])
            amt = float(pd.to_numeric(row[amt_col], errors="coerce") or 0)
            if name[:4].startswith("1"):
                a += amt
            elif name[:4].startswith("2"):
                l += amt
            elif name[:4].startswith(("3", "4")):
                e += amt
        diff = round(a - l - e, 2)
        self.findings.append({"类别": "报表勾稽", "位置": "资产负债表",
            "检查项": "资产=负债+权益", "资产合计": round(a, 2),
            "负债+权益": round(l + e, 2), "差异": diff,
            "结果": "通过" if abs(diff) < 0.5 else "异常"})

    # -- 3. 财务指标 --
    def _check_financial_ratios(self, tb):
        if tb is None:
            return
        subj_col = amt_col = None
        for c in tb.columns:
            s = str(c)
            if any(kw in s for kw in ["科目", "名称"]):
                subj_col = c
            if any(kw in s for kw in ["期末余额", "余额", "金额"]) and "期初" not in s:
                amt_col = c
        if subj_col is None or amt_col is None:
            return
        bal = {}
        for _, row in tb.iterrows():
            name = str(row[subj_col])
            amt = float(pd.to_numeric(row[amt_col], errors="coerce") or 0)
            for key in ["货币资金", "应收账款", "存货", "流动资产", "固定资产",
                        "资产总计", "流动负债", "负债合计", "所有者权益",
                        "营业收入", "净利润"]:
                if key in name:
                    bal[key] = bal.get(key, 0) + amt
        if "流动资产" in bal and "流动负债" in bal and bal["流动负债"] != 0:
            self.findings.append({"类别": "指标校验", "位置": "流动比率",
                "检查项": "流动资产/流动负债",
                "计算值": round(bal["流动资产"] / bal["流动负债"], 2), "结果": ""})
        if "负债合计" in bal and "资产总计" in bal and bal["资产总计"] != 0:
            self.findings.append({"类别": "指标校验", "位置": "资产负债率",
                "检查项": "负债/资产",
                "计算值": round(bal["负债合计"] / bal["资产总计"] * 100, 1), "结果": ""})

    # -- 4. 交叉校验 --
    def _cross_check(self, tb):
        if tb is None or not self.tables:
            if tb is None: self.coverage["报告交叉校验"] = "跳过(缺科目余额表)"
            elif not self.tables: self.coverage["报告交叉校验"] = "跳过(缺Word报告)"
            return

        subj_col = amt_col = None
        for c in tb.columns:
            s = str(c)
            if any(kw in s for kw in ["科目", "名称"]):
                subj_col = c
            if any(kw in s for kw in ["期末余额", "余额", "金额"]) and "期初" not in s:
                amt_col = c
        if subj_col is None or amt_col is None:
            return
        tb_vals = {}
        for _, row in tb.iterrows():
            tb_vals[str(row[subj_col]).strip()] = float(
                pd.to_numeric(row[amt_col], errors="coerce") or 0)
        for t in self.tables:
            df = t["data"]
            if df.empty or len(df.columns) < 2:
                continue
            name_col = df.columns[0]
            for ci in range(1, len(df.columns)):
                for _, row in df.iterrows():
                    name = str(row[name_col]).strip()
                    amt = pd.to_numeric(
                        str(row[df.columns[ci]]).replace(",", "").replace(chr(65292), ""),
                        errors="coerce")
                    if pd.isna(amt):
                        continue
                    for tn, ta in tb_vals.items():
                        if name in tn or tn in name:
                            diff = round(abs(amt - ta), 2)
                            if diff > 0.5 and abs(ta) > 1:
                                self.findings.append({"类别": "交叉校验",
                                    "位置": "表{} {}".format(t["index"] + 1, name),
                                    "检查项": "报告vs科目余额表", "报表值": amt,
                                    "科目余额表值": ta, "差异": diff, "结果": "异常"})
                            break

    # -- 5. 异动分析 --
    def _variance_analysis(self, tb, prior_tb):
        if tb is None or prior_tb is None:
            self.coverage["异动分析"] = "跳过(缺{}期数据)".format("本" if tb is None else "上")
            return

        subj_col = amt_col = None
        for c in tb.columns:
            if any(kw in str(c) for kw in ["科目", "名称"]):
                subj_col = c
            if any(kw in str(c) for kw in ["期末余额", "余额", "金额"]) and "期初" not in str(c):
                amt_col = c
        if subj_col is None or amt_col is None:
            return
        curr = {}
        prior = {}
        for _, row in tb.iterrows():
            curr[str(row[subj_col]).strip()] = float(
                pd.to_numeric(row[amt_col], errors="coerce") or 0)
        for _, row in prior_tb.iterrows():
            prior[str(row[subj_col]).strip()] = float(
                pd.to_numeric(row[amt_col], errors="coerce") or 0)
        changes = []
        for name in curr:
            c = curr[name]
            p = prior.get(name, 0)
            if abs(c) < 100 and abs(p) < 100:
                continue
            chg = c - p
            pct = round(chg / abs(p) * 100, 1) if abs(p) > 1 else 999
            if abs(pct) > 20:
                changes.append({"科目": name, "本期": round(c, 2), "上期": round(p, 2),
                                "变动额": round(chg, 2), "变动率%": pct})
        changes.sort(key=lambda x: abs(x["变动额"]), reverse=True)
        for ch in changes[:30]:
            self.findings.append({"类别": "异动分析", "位置": ch["科目"],
                "检查项": "同比>20%", "本期": ch["本期"], "上期": ch["上期"],
                "变动额": ch["变动额"], "变动率%": ch["变动率%"], "结果": "需关注"})

        # -- 5b. 期初期末勾稽：本期期初 = 上期期末 --
        beg_col = None
        for c in tb.columns:
            if any(kw in str(c) for kw in ["期初余额", "年初余额", "上年末余额"]):
                beg_col = c
                break
        if beg_col is not None:
            mismatch_count = 0
            for name in curr:
                cur_beg_mask = tb[subj_col].astype(str).str.strip() == name
                cur_beg = float(pd.to_numeric(
                    tb.loc[cur_beg_mask, beg_col], errors="coerce").sum() or 0)
                prior_end = prior.get(name, 0)
                if abs(cur_beg) > 1 and abs(prior_end) > 1:
                    diff = round(abs(cur_beg - prior_end), 2)
                    if diff > 0.5:
                        mismatch_count += 1
                        if mismatch_count <= 20:
                            self.findings.append({"类别": "期初期末勾稽", "位置": name,
                                "检查项": "本期期初vs上期期末", "本期期初": round(cur_beg, 2),
                                "上期期末": round(prior_end, 2), "差异": diff,
                                "结果": "异常"})
            if mismatch_count > 0:
                self.coverage["期初期末勾稽"] = "异常({}处不衔接)".format(mismatch_count)
            else:
                self.coverage["期初期末勾稽"] = "通过"
        else:
            self.coverage["期初期末勾稽"] = "跳过(本期TB无期初余额列)"


    # -- 6. LLM 错别字/标点/术语 --
    def _extract_docx_text(self):
        try:
            from docx import Document
        except ImportError:
            return []
        doc = Document(str(self.report_path))
        return [p.text.strip() for p in doc.paragraphs if len(p.text.strip()) > 10]

    def _check_typos(self):
        if not self.llm_url or not self.report_path:
            return
        paras = self._extract_docx_text()
        if not paras:
            return
        chunks = []
        buf = ""
        for p in paras:
            if len(buf) + len(p) > 3000:
                chunks.append(buf)
                buf = p
            else:
                buf += chr(10) + p
        if buf:
            chunks.append(buf)
        prompt_tpl = chr(10).join([
            "你是审计报告校对专家。检查以下片段，找出:",
            "1.错别字 2.标点错误 3.专有名词不一致 4.会计术语不当",
            "用JSON数组返回: [{位置, 错误类型, 原文, 建议修改, 严重程度}]，无错误返回[]",
            "报告片段:", "{text}", "直接返回JSON数组:"])
        for chunk in chunks:
            try:
                resp = self._call_llm(prompt_tpl.format(text=chunk))
                if resp:
                    errs = self._parse_json(resp)
                    if errs and isinstance(errs, list):
                        for e in errs:
                            self.findings.append({
                                "类别": "错别字/标点",
                                "位置": e.get("位置", ""),
                                "检查项": e.get("错误类型", ""),
                                "原文": e.get("原文", ""),
                                "建议修改": e.get("建议修改", ""),
                                "严重程度": e.get("严重程度", ""),
                                "结果": "需修改"})
            except Exception as ex:
                print("[LLM错别字] " + str(ex))

    # -- 7. LLM 报表公式与勾稽验证 --
    def _check_formulas_llm(self):
        if not self.llm_url or not self.tables:
            return
        for t in self.tables[:5]:
            df = t["data"]
            rows_text = ["|".join(str(c) for c in df.columns)]
            for _, row in df.head(20).iterrows():
                rows_text.append("|".join(str(v) for v in row.values))
            table_text = chr(10).join(rows_text)
            prompt = chr(10).join([
                "你是审计报告复核专家。分析以下财务报表表格，验证:",
                "1. 合计行的数值是否等于各分项之和",
                "2. 是否存在应有的勾稽关系(如资产=负债+权益)",
                "3. 财务指标计算是否正确(如毛利率=毛利/收入)",
                "用JSON数组返回发现的问题: [{表格名, 检查项, 报表值, 应有值, 差异, 说明}]",
                "无问题返回[]",
                "表格:", table_text, "直接返回JSON数组:"])
            try:
                resp = self._call_llm(prompt, max_tokens=3000)
                if resp:
                    errs = self._parse_json(resp)
                    if errs and isinstance(errs, list):
                        for e in errs:
                            self.findings.append({
                                "类别": "LLM公式验证",
                                "位置": e.get("表格名", "表{}".format(t["index"] + 1)),
                                "检查项": e.get("检查项", ""),
                                "报表值": e.get("报表值", ""),
                                "应有值": e.get("应有值", ""),
                                "差异": e.get("差异", ""),
                                "结果": e.get("说明", "需复核")})
            except Exception as ex:
                print("[LLM公式] " + str(ex))


    # -- 8. 报表与附注勾稽 --
    def _check_note_to_statement(self, tb, unaudited_fs, notes):
        """报表项目 vs 对应附注明细的加总一致性"""
        self.coverage["报表与附注勾稽"] = "通过"
        if notes is None and unaudited_fs is None:
            self.coverage["报表与附注勾稽"] = "跳过(缺附注数据)"
            return

        subj_col = amt_col = 0
        tb_vals = {}
        if tb is not None:
            for i, c in enumerate(tb.columns):
                s = str(c)
                if any(kw in s for kw in ["科目", "名称"]): subj_col = i
                if any(kw in s for kw in ["期末余额", "余额", "金额"]) and "期初" not in s: amt_col = i
            for _, row in tb.iterrows():
                name = str(row.iloc[subj_col]).strip()
                tb_vals[name] = float(pd.to_numeric(row.iloc[amt_col], errors="coerce") or 0)

        NOTE_TB_MAP = {
            "应收账款": ["应收账款"], "其他应收款": ["其他应收款"],
            "预付账款": ["预付账款", "预付"], "存货": ["存货"],
            "固定资产": ["固定资产"], "在建工程": ["在建工程"],
            "无形资产": ["无形资产"], "应付账款": ["应付账款"],
            "其他应付款": ["其他应付款"], "预收账款": ["预收账款", "预收"],
            "短期借款": ["短期借款"], "长期借款": ["长期借款"],
            "应交税费": ["应交税费"], "应付职工薪酬": ["应付职工薪酬"],
        }

        note_dict = {}
        if notes is not None:
            if isinstance(notes, dict):
                note_dict.update(notes)
            elif isinstance(notes, pd.DataFrame):
                note_dict["附注表"] = notes
        if unaudited_fs is not None and isinstance(unaudited_fs, dict):
            for k, v in unaudited_fs.items():
                if any(kw in str(k) for kw in ["附注", "账龄", "固定资产", "存货",
                        "应收账款", "应付账款", "借款", "税费", "薪酬", "明细"]):
                    note_dict[str(k)] = v

        for note_name, note_df in note_dict.items():
            if not isinstance(note_df, pd.DataFrame):
                continue
            note_total = 0.0
            for possible_col in ["合计", "合计金额", "期末余额", "期末", "账面价值", "账面净值"]:
                if possible_col in note_df.columns:
                    vals = pd.to_numeric(note_df[possible_col], errors="coerce")
                    total_mask = note_df.iloc[:, 0].astype(str).str.contains(
                        "合计|总计|小计", na=False)
                    # 金额列排除合计行避免重复; "合计"列则直接取合计行
                    if possible_col not in ("合计", "合计金额") and total_mask.any():
                        note_total = float(vals[~total_mask].sum())
                    else:
                        note_total = float(vals.sum())
                    break
            if note_total == 0.0:
                last_col = note_df.columns[-1]
                vals = pd.to_numeric(note_df[last_col], errors="coerce").dropna()
                if len(vals) > 0:
                    total_mask = note_df.iloc[:, 0].astype(str).str.contains(
                        "合计|总计|小计", na=False)
                    note_total = float(vals[~total_mask].sum()) if total_mask.any() else float(vals.sum())
            if abs(note_total) < 1:
                continue

            matched = False
            for note_key, tb_keys in NOTE_TB_MAP.items():
                if note_key in str(note_name):
                    for tb_key in tb_keys:
                        for tn, ta in tb_vals.items():
                            if tb_key in tn and "减值" not in tn and "坏账" not in tn:
                                diff = round(abs(note_total - ta), 2)
                                if diff > 0.5:
                                    self.findings.append({"类别": "报表与附注勾稽",
                                        "位置": str(note_name),
                                        "检查项": "附注合计 vs TB{}".format(tb_key),
                                        "附注合计": round(note_total, 2),
                                        "TB科目值": round(ta, 2), "差异": diff,
                                        "结果": "异常"})
                                    self.coverage["报表与附注勾稽"] = "异常"
                                matched = True; break
                        if matched: break
                    if matched: break


    # -- 9. 附注科目内勾稽 --
    def _check_note_internal(self, notes):
        """同一附注内部的逻辑关系（如期初 + 本期增加 - 本期减少 = 期末）"""
        self.coverage["附注科目内勾稽"] = "通过"
        if notes is None:
            self.coverage["附注科目内勾稽"] = "跳过(缺附注数据)"
            return

        note_dict = {}
        if isinstance(notes, dict):
            note_dict.update(notes)
        elif isinstance(notes, pd.DataFrame):
            note_dict["附注表"] = notes
        if not note_dict:
            self.coverage["附注科目内勾稽"] = "跳过(无附注表)"
            return

        for note_name, note_df in note_dict.items():
            if not isinstance(note_df, pd.DataFrame) or note_df.empty:
                continue

            col_beg = col_add = col_sub = col_end = None
            for c in note_df.columns:
                s = str(c)
                if col_beg is None and any(kw in s for kw in ["期初", "年初", "上年末"]):
                    col_beg = c
                if col_add is None and any(kw in s for kw in ["本期增加", "本年增加", "计提", "新增"]):
                    col_add = c
                if col_sub is None and any(kw in s for kw in ["本期减少", "本年减少", "核销", "转回", "转销", "处置"]):
                    col_sub = c
                if col_end is None and any(kw in s for kw in ["期末", "年末"]):
                    col_end = c

            if col_beg is None or col_end is None:
                continue
            if col_add is None and col_sub is None:
                continue

            name_col = note_df.columns[0]
            err_count = 0
            for _, row in note_df.iterrows():
                beg = float(pd.to_numeric(row[col_beg], errors="coerce") or 0)
                add = float(pd.to_numeric(row[col_add], errors="coerce") or 0) if col_add else 0
                sub = float(pd.to_numeric(row[col_sub], errors="coerce") or 0) if col_sub else 0
                end = float(pd.to_numeric(row[col_end], errors="coerce") or 0)

                if abs(beg) < 1 and abs(add) < 1 and abs(sub) < 1 and abs(end) < 1:
                    continue

                expected = round(beg + add - sub, 2)
                diff = round(abs(end - expected), 2)
                if diff > 0.5:
                    err_count += 1
                    if err_count <= 20:
                        name = str(row[name_col]).strip()
                        self.findings.append({"类别": "附注科目内勾稽",
                            "位置": "{}/{}".format(note_name, name),
                            "检查项": "期初+增加-减少=期末",
                            "期初": beg, "本期增加": add, "本期减少": sub,
                            "期末(报表)": end, "期末(计算)": expected, "差异": diff,
                            "结果": "异常"})

            if err_count > 0:
                self.coverage["附注科目内勾稽"] = "异常({}处)".format(err_count)

    # -- 主流程 --
    def run(self):
        print("[报告复核] 开始...")
        # 加载数据源
        tb = _load_source(self.src.get('trial_balance'))
        prior_tb = _load_source(self.src.get('prior_tb'))
        unaudited_fs = _load_source(self.src.get('unaudited_fs'))
        cash_flow = _load_source(self.src.get('cash_flow'))
        fixed_assets = _load_source(self.src.get('fixed_assets'))
        bank_summary = _load_source(self.src.get('bank_summary'))
        aging = _load_source(self.src.get('aging'))
        notes = self.src.get('notes')  # dict of DataFrames

        # 加载未审报表（可能为多sheet Excel）
        if unaudited_fs is None and 'unaudited_fs' in self.src:
            raw = self.src['unaudited_fs']
            if isinstance(raw, str):
                try:
                    unaudited_fs = pd.read_excel(raw, sheet_name=None)
                except Exception:
                    unaudited_fs = pd.read_excel(raw)

        if self.report_path:
            self._extract_docx_tables()

        # === 类型5：表内项目勾稽 ===
        self._check_table_formulas()
        self._check_balance_sheet_equation(tb)

        # === 类型1：报表间勾稽 ===
        if unaudited_fs is not None or cash_flow is not None:
            self._check_cross_statement(unaudited_fs, cash_flow, tb)

        # === 类型7：账表勾稽 ===
        self._cross_check(tb)
        if bank_summary is not None:
            self._check_bank_consistency(bank_summary, tb)
        if aging is not None:
            self._check_aging_consistency(aging, tb)

        # === 类型2：报表与附注勾稽 ===
        if notes is not None or (isinstance(unaudited_fs, dict) and unaudited_fs):
            self._check_note_to_statement(tb, unaudited_fs, notes)

        # === 类型3：附注科目内勾稽 ===
        self._check_note_internal(notes)

        # === 类型4：附注跨科目勾稽（合理性检验）===
        self._check_cross_note_reasonableness(tb, fixed_assets)

        # === 类型6：期初期末勾稽 + 异动分析 ===
        self._variance_analysis(tb, prior_tb)

        # === 补充校验 ===
        self._check_financial_ratios(tb)

        # LLM（仅当有Word报告且LLM可用时）
        if self.llm_url and self.report_path:
            print("[报告复核] LLM: 错别字检查...")
            self._check_typos()
            print("[报告复核] LLM: 公式验证...")
            self._check_formulas_llm()

        # 汇总
        df = pd.DataFrame(self.findings)
        n = len(df)
        na = (df.get("结果") == "异常").sum() if "结果" in df.columns else 0
        print("[报告复核] {}项: {}异常".format(n, na))

        # 标注未提供的数据源
        for key, label in [("trial_balance","科目余额表"), ("prior_tb","上年TB"),
                           ("unaudited_fs","未审报表"), ("adjustments","调整分录"),
                           ("aging","账龄表"), ("fixed_assets","固资卡片"),
                           ("bank_summary","银行汇总"), ("cash_flow","现金流表"),
                           ("notes","附注明细")]:
            if key not in self.src:
                self.coverage[label + "(未提供)"] = "跳过"
        return df

    def _check_bank_consistency(self, bank_summary, tb):
        """银行流水汇总 vs 科目余额表货币资金"""
        self.coverage["银行资金勾稽"] = "通过"
        if "银方净额" in bank_summary.columns:
            bank_total = bank_summary["银方净额"].sum()
        elif "期末余额" in bank_summary.columns:
            bank_total = pd.to_numeric(bank_summary["期末余额"], errors="coerce").sum()
        else:
            self.coverage["银行资金勾稽"] = "跳过(列名不识别)"
            return
        subj_col = amt_col = 0
        if tb is not None:
            for i, c in enumerate(tb.columns):
                s = str(c)
                if any(kw in s for kw in ["科目", "名称"]): subj_col = i
                if any(kw in s for kw in ["期末余额", "余额", "金额"]) and "期初" not in s: amt_col = i
        cash_amt = 0.0
        if tb is not None:
            for _, row in tb.iterrows():
                if "货币资金" in str(row.iloc[subj_col]):
                    cash_amt += float(pd.to_numeric(row.iloc[amt_col], errors="coerce") or 0)
        diff = round(bank_total - cash_amt, 2)
        self.findings.append({"类别": "银行勾稽", "位置": "货币资金",
            "检查项": "银行汇总 vs 科目余额表", "银行汇总": round(bank_total, 2),
            "科目余额表": round(cash_amt, 2), "差异": diff,
            "结果": "通过" if abs(diff) < 0.5 else "异常"})
        if abs(diff) >= 0.5:
            self.coverage["银行资金勾稽"] = "异常"

    def _check_aging_consistency(self, aging, tb):
        """账龄分析 vs 科目余额表往来科目"""
        self.coverage["账龄勾稽"] = "通过"
        ar_aging = 0.0
        # 仅取"应收账款"行，而非全表求和（账龄表含多个往来科目）
        name_col = aging.columns[0]
        ar_mask = aging[name_col].astype(str).str.contains("应收账款", na=False)
        if "合计" in aging.columns:
            ar_aging = pd.to_numeric(aging.loc[ar_mask, "合计"], errors="coerce").sum()
        elif "合计金额" in aging.columns:
            ar_aging = pd.to_numeric(aging.loc[ar_mask, "合计金额"], errors="coerce").sum()
        subj_col = amt_col = 0
        if tb is not None:
            for i, c in enumerate(tb.columns):
                s = str(c)
                if any(kw in s for kw in ["科目", "名称"]): subj_col = i
                if any(kw in s for kw in ["期末余额", "余额", "金额"]) and "期初" not in s: amt_col = i
        ar_tb = 0.0
        if tb is not None:
            for _, row in tb.iterrows():
                if "应收账款" in str(row.iloc[subj_col]):
                    ar_tb += float(pd.to_numeric(row.iloc[amt_col], errors="coerce") or 0)
        diff = round(ar_aging - ar_tb, 2)
        if ar_aging > 0 or ar_tb > 0:
            self.findings.append({"类别": "账龄勾稽", "位置": "应收账款",
                "检查项": "账龄表 vs 科目余额表", "账龄合计": round(ar_aging, 2),
                "科目余额表": round(ar_tb, 2), "差异": diff,
                "结果": "通过" if abs(diff) < 0.5 else "异常"})
            if abs(diff) >= 0.5:
                self.coverage["账龄勾稽"] = "异常"

    # -- 8. 报表间勾稽 --
    def _check_cross_statement(self, unaudited_fs, cash_flow, tb):
        """报表间勾稽：BS净利润→IS、BS货币资金→CF"""
        self.coverage["报表间勾稽"] = "通过"
        if unaudited_fs is None and tb is None:
            self.coverage["报表间勾稽"] = "跳过(缺报表/TB)"
            return

        # --- 8a. 净利润勾稽：BS未分配利润变动 vs IS净利润 ---
        bs_end_retained = bs_start_retained = is_net_profit = None

        # 尝试从 unaudited_fs 取
        if unaudited_fs is not None:
            bs = is_ = None
            if isinstance(unaudited_fs, dict):
                bs = unaudited_fs.get("资产负债表") or unaudited_fs.get("BS")
                is_ = unaudited_fs.get("利润表") or unaudited_fs.get("IS")
            elif hasattr(unaudited_fs, 'items'):  # pd.ExcelFile / dict of sheets
                for k in unaudited_fs.keys():
                    if "资产负债" in str(k) or "BS" in str(k).upper():
                        bs = unaudited_fs[k]
                    if "利润" in str(k) or "IS" in str(k).upper():
                        is_ = unaudited_fs[k]
            if bs is not None:
                bs = pd.DataFrame(bs) if not isinstance(bs, pd.DataFrame) else bs
                for _, row in bs.iterrows():
                    name = str(row.iloc[0])
                    if "未分配利润" in name or "留存收益" in name:
                        vals = pd.to_numeric(row.iloc[1:], errors="coerce").dropna()
                        if len(vals) >= 2:
                            bs_end_retained = float(vals.iloc[-1])
                            bs_start_retained = float(vals.iloc[-2])
            if is_ is not None:
                is_ = pd.DataFrame(is_) if not isinstance(is_, pd.DataFrame) else is_
                for _, row in is_.iterrows():
                    if "净利润" in str(row.iloc[0]):
                        vals = pd.to_numeric(row.iloc[1:], errors="coerce").dropna()
                        if len(vals) >= 1:
                            is_net_profit = float(vals.iloc[-1])

        # 从 TB 取（兜底）
        if tb is not None:
            subj_col = amt_col = 0
            for i, c in enumerate(tb.columns):
                s = str(c)
                if any(kw in s for kw in ["科目","名称"]): subj_col = i
                if any(kw in s for kw in ["期末余额","余额","金额"]): amt_col = i
            for _, row in tb.iterrows():
                name = str(row.iloc[subj_col])
                amt = float(pd.to_numeric(row.iloc[amt_col], errors="coerce") or 0)
                if is_net_profit is None and "净利润" in name and "未分配" not in name:
                    is_net_profit = amt

        if bs_end_retained is not None and bs_start_retained is not None and is_net_profit is not None:
            bs_change = bs_end_retained - bs_start_retained
            diff = round(bs_change - is_net_profit, 2)
            self.findings.append({"类别": "报表间勾稽", "位置": "BS未分配利润",
                "检查项": "BS未分配利润变动 vs IS净利润", "BS变动": bs_change,
                "IS净利润": is_net_profit, "差异": diff,
                "结果": "通过" if abs(diff) < 0.5 else "异常"})
            if abs(diff) >= 0.5:
                self.coverage["报表间勾稽"] = "异常"

        # --- 8b. 货币资金勾稽：BS变动 vs CF期末 ---
        cf_cash = bs_cash_end = bs_cash_start = None
        if cash_flow is not None:
            cf = pd.DataFrame(cash_flow) if not isinstance(cash_flow, pd.DataFrame) else cash_flow
            for _, row in cf.iterrows():
                for kw in ["期末现金", "现金及现金等价物", "期末余额"]:
                    if kw in str(row.iloc[0]):
                        vals = pd.to_numeric(row.iloc[1:], errors="coerce").dropna()
                        if len(vals) >= 1: cf_cash = float(vals.iloc[-1])
                        break
        if bs is not None:
            for _, row in bs.iterrows():
                if "货币资金" in str(row.iloc[0]):
                    vals = pd.to_numeric(row.iloc[1:], errors="coerce").dropna()
                    if len(vals) >= 2:
                        bs_cash_end = float(vals.iloc[-1])
                        bs_cash_start = float(vals.iloc[-2])

        if cf_cash is not None and bs_cash_end is not None:
            diff = round(cf_cash - bs_cash_end, 2)
            self.findings.append({"类别": "报表间勾稽", "位置": "货币资金",
                "检查项": "CF期末现金 vs BS货币资金期末", "CF期末": cf_cash,
                "BS余额": bs_cash_end, "差异": diff,
                "结果": "通过" if abs(diff) < 0.5 else "异常"})
            if abs(diff) >= 0.5:
                self.coverage["报表间勾稽"] = "异常"

    # -- 9. 附注跨科目勾稽（合理性检验）--
    def _check_cross_note_reasonableness(self, tb, fixed_assets):
        """跨科目合理性：折旧率、利息率是否在合理区间"""
        self.coverage["附注跨科目勾稽"] = "通过"
        subj_col = amt_col = 0
        if tb is not None:
            for i, c in enumerate(tb.columns):
                s = str(c)
                if any(kw in s for kw in ["科目","名称"]): subj_col = i
                if any(kw in s for kw in ["期末余额","余额","金额"]): amt_col = i

        # --- 9a. 折旧率合理区间 ---
        dep = fa_orig = 0.0
        if tb is not None:
            for _, row in tb.iterrows():
                name = str(row.iloc[subj_col])
                amt = abs(float(pd.to_numeric(row.iloc[amt_col], errors="coerce") or 0))
                if "累计折旧" in name: dep = amt
                if "固定资产" in name and "累计" not in name and "减值" not in name: fa_orig = amt
        if fixed_assets is not None:
            fa_df = pd.DataFrame(fixed_assets) if not isinstance(fixed_assets, pd.DataFrame) else fixed_assets
            for c in fa_df.columns:
                if "累计折旧" in str(c):
                    dep = max(dep, pd.to_numeric(fa_df[c], errors="coerce").sum())
                if "原值" in str(c) or "账面原值" in str(c):
                    fa_orig = max(fa_orig, pd.to_numeric(fa_df[c], errors="coerce").sum())

        if dep > 0 and fa_orig > 0:
            ratio = round(dep / fa_orig * 100, 1)
            flag = ""
            if ratio < 1 or ratio > 50:
                flag = "异常(折旧率{}%偏离正常区间)".format(ratio)
                self.coverage["附注跨科目勾稽"] = "异常"
            self.findings.append({"类别": "跨科目勾稽", "位置": "固定资产",
                "检查项": "累计折旧/原值", "累计折旧": round(dep, 2),
                "固定资产原值": round(fa_orig, 2), "折旧覆盖率%": ratio, "结果": flag or "通过"})

        # --- 9b. 利息支出与借款余额匹配 ---
        interest = loan = 0.0
        if tb is not None:
            for _, row in tb.iterrows():
                name = str(row.iloc[subj_col])
                amt = abs(float(pd.to_numeric(row.iloc[amt_col], errors="coerce") or 0))
                if "利息支出" in name or "利息费用" in name: interest = amt
                if "短期借款" in name or "长期借款" in name: loan += amt
        if interest > 0 and loan > 0:
            rate = round(interest / loan * 100, 1)
            flag = ""
            if rate < 1 or rate > 15:
                flag = "异常(推算利率{}%偏离合理区间)".format(rate)
                self.coverage["附注跨科目勾稽"] = "异常"
            self.findings.append({"类别": "跨科目勾稽", "位置": "借款",
                "检查项": "利息支出/借款余额≈利率", "利息支出": round(interest, 2),
                "借款余额": round(loan, 2), "推算年利率%": rate, "结果": flag or "通过"})


def _load_source(src):
    """统一数据源加载：支持路径或DataFrame"""
    if src is None:
        return None
    if isinstance(src, pd.DataFrame):
        return src
    return pd.read_excel(str(src))

def review_report(report_docx=None, data_files=None, data_sources=None,
                  prior_report_docx=None, output_dir=None):
    """报告复核 - 一键审计报告质量检查（七类勾稽关系）

    参数:
        report_docx: 审计报告Word文档路径（可选，无Word报告时仅做非LLM校验）
        data_files:   文件路径列表，自动根据文件名识别类型
        data_sources: 手动指定 {键: 路径或DataFrame}，优先级高于 data_files
            支持的键:
            - trial_balance:  科目余额表(.xlsx)
            - prior_tb:       上年科目余额表(.xlsx)
            - unaudited_fs:   未审报表(.xlsx，多sheet含资产负债表/利润表)
            - cash_flow:      现金流量表(.xlsx)
            - fixed_assets:   固定资产卡片/折旧明细(.xlsx)
            - bank_summary:   银行流水汇总(.xlsx)
            - aging:          账龄分析表(.xlsx)
            - adjustments:    审计调整分录(.xlsx)
            - notes:          附注明细(dict: {附注名: DataFrame})
        prior_report_docx: 上年审计报告路径（暂未使用）
        output_dir:       输出目录，生成 报告复核结果.xlsx 和 校验覆盖状态.xlsx

    返回:
        DataFrame: findings 列表，包含所有检查项结果
    """
    # 自动归类
    auto = {}
    if data_files:
        AUTO_RULES = {
            "trial_balance": ["科目余额表", "TB", "trial_balance", "余额表"],
            "prior_tb":      ["上年科目余额表", "上年余额表", "prior"],
            "unaudited_fs":  ["未审报表", "资产负债表", "利润表", "unaudited"],
            "adjustments":   ["调整分录", "审计调整", "adjust", "调整汇总"],
            "aging":         ["账龄", "aging", "往来明细"],
            "fixed_assets":  ["固定资产", "折旧明细", "卡片", "fixed_asset"],
            "bank_summary":  ["银行流水", "银行汇总", "账户汇总", "bank_summary", "对手方收付"],
            "cash_flow":     ["现金流量", "cash_flow", "现金流"],
            "tax":           ["应交税费", "税费明细", "tax"],
            "payroll":       ["应付职工薪酬", "工资薪酬", "payroll", "薪酬"],
            "notes":         ["附注明细", "坏账准备", "存货跌价", "长期资产减值"],
        }
        for fp in data_files:
            name = Path(fp).stem
            for key, keywords in AUTO_RULES.items():
                if any(kw in name for kw in keywords):
                    auto[key] = fp
                    break
            else:
                print("[报告复核] 未识别文件类型: {}".format(Path(fp).name))
        print("[报告复核] 自动识别: {}".format(list(auto.keys())))

    # 合并：手动指定优先
    src = dict(auto)
    if data_sources:
        src.update(data_sources)

    r = ReportReviewer(report_docx, src, prior_report_docx)

    # 无LLM时禁用（避免连接超时）
    if not os.environ.get("VLLM_TUNNEL_URL"):
        r.disable_llm()

    df = r.run()
    if output_dir:
        out = Path(output_dir); out.mkdir(parents=True, exist_ok=True)
        df.to_excel(str(out / "报告复核结果.xlsx"), index=False)
        cov = pd.DataFrame([{"检查项": k, "状态": v} for k, v in r.coverage.items()])
        cov.to_excel(str(out / "校验覆盖状态.xlsx"), index=False)
    return df
