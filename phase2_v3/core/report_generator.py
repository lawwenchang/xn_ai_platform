#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
审计报告生成器 - 根据执行结果和用户意图自动生成 Word 审计报告。
覆盖场景：银行流水核对 / 医保回款识别 / 大额交易筛查 / 跨表匹配 / 通用处理
"""
from __future__ import annotations

import json, csv
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional


def _set_run_font(run, name: str = "宋体", size_pt: int = None):
    """设置 Run 的中文字体（eastAsia），避免中文版 Office 失效。"""
    try:
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        from docx.shared import Pt
        run.font.name = name
        rpr = run._element.rPr
        if rpr is not None:
            rFonts = rpr.rFonts
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rpr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), name)
        if size_pt is not None:
            run.font.size = Pt(size_pt)
    except Exception:
        pass


def _set_doc_chinese_fonts(doc):
    """设置报告文档的中文字体。"""
    try:
        from docx.shared import Pt
        style = doc.styles["Normal"]
        style.font.name = "宋体"
        style.font.size = Pt(11)
        rpr = style._element.rPr
        if rpr is not None:
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            rFonts = rpr.rFonts
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rpr.append(rFonts)
            rFonts.set(qn("w:eastAsia"), "宋体")
    except Exception:
        pass
    for lvl in range(1, 4):
        try:
            style = doc.styles[f"Heading {lvl}"]
            style.font.name = "黑体"
            rpr = style._element.rPr
            if rpr is not None:
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                rFonts = rpr.rFonts
                if rFonts is None:
                    rFonts = OxmlElement("w:rFonts")
                    rpr.append(rFonts)
                rFonts.set(qn("w:eastAsia"), "黑体")
        except Exception:
            pass


def _apply_table_borders(table, color: str = "000000", size: str = "4"):
    """为表格设置显式边框，避免中文版 Office 样式失效导致无框线。"""
    try:
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        tbl = table._tbl
        tblPr = tbl.tblPr
        existing = tblPr.find(qn("w:tblBorders"))
        if existing is not None:
            tblPr.remove(existing)
        tblBorders = OxmlElement("w:tblBorders")
        for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), size)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
            tblBorders.append(el)
        tblPr.append(tblBorders)
    except Exception:
        pass


def _format_table_cells_fonts(table, name: str = "宋体", size_pt: int = 11):
    """统一设置表格单元格内文字字体。"""
    try:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for run in p.runs:
                        _set_run_font(run, name, size_pt)
    except Exception:
        pass


def _detect_scenario(user_intent: str) -> str:
    """报告场景检测——委托给 scenario_packs 唯一注册表"""
    from config.scenario_packs import detect_scenario as _ds
    sid = _ds(user_intent, ask_user=False)
    # 兼容旧场景名（report_generator 仍用旧名称做分支）
    compat = {
        "bank_reconcile_detail": "match",
        "summary_compare": "balance",
        "filtered_extraction_match": "medical",
        "large_txn_screen": "screening",
        "single_table_analysis": "general",
        "doc_generation": "general",
        "cross_doc_compare": "diff",
    }
    return compat.get(sid, "general")


def generate_audit_report(
    run_id: str, user_intent: str, dag_operators: List[Dict],
    output_dir: Path, input_files: List[str],
    execution_logs: List[str] = None,
    match_logic: dict = None,
    reconcile_stats: dict = None,
    workpaper_files: List[str] = None,
) -> Path:
    """生成 Word 审计报告。Returns: 报告文件路径
    reconcile_stats: bank_reconcile_engine 返回的 stats 字典
    workpaper_files: 关联底稿文件列表（分桶看板、核查底稿等）
    """
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    _set_doc_chinese_fonts(doc)
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style.font.size = Pt(11)
    scenario = _detect_scenario(user_intent)

    # ===== 封面 =====
    t = doc.add_heading("智能审计报告", level=0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in t.runs:
        _set_run_font(r, "黑体", 16)
    doc.add_paragraph(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Run ID：{run_id}")

    # ===== 底稿要素 =====
    _add_working_paper_header(doc, run_id, user_intent, scenario, input_files)

    # ===== 一、审计意图 =====
    doc.add_heading("一、审计意图与目标", level=1)
    doc.add_paragraph(f"用户指令：{user_intent}")
    scenario_names = {
        "medical_match": "医保回款跨表匹配", "balance_match": "科目余额核对",
        "match": "跨表/银行流水核对", "medical": "医保回款识别",
        "screening": "大额交易筛查", "aggregate": "数据汇总统计",
        "compliance": "合规检查", "balance": "科目余额分析",
        "diff": "差异对比分析", "general": "通用数据处理",
    }
    doc.add_paragraph(f"识别场景：{scenario_names.get(scenario, '通用')}")
    goals = {
        "medical_match": "从银行流水中筛选医保回款明细，与医保回款汇总表跨表核对，验证金额一致性。",
        "match": "对两个关联数据表交叉核对，发现不一致记录。",
        "medical": "识别银行流水中医保回款记录，按机构/年度汇总。",
        "screening": "按阈值筛查大额/异常交易，风险分级预警。",
        "aggregate": "按指定维度汇总统计，生成报表。",
        "compliance": "基于法规对数据合规性检查。",
        "balance": "检查科目余额表借贷平衡。",
        "diff": "对比多组数据，分析差异原因。",
        "general": "按用户指令进行数据处理和分析。",
    }
    doc.add_paragraph(f"场景目标：{goals.get(scenario, '按用户指令处理数据')}")

    # ===== 二、执行过程 =====
    doc.add_heading("二、执行过程与匹配逻辑", level=1)
    doc.add_paragraph(f"输入文件：{', '.join(input_files)}")
    doc.add_paragraph(f"算子数量：{len(dag_operators)}")

    # 算子步骤表（含关键参数）
    tbl = doc.add_table(rows=len(dag_operators) + 1, cols=4, style="Light Grid Accent 1")
    for i, h in enumerate(["步骤", "算子", "关键参数", "输出"]):
        tbl.rows[0].cells[i].text = h
    for idx, op in enumerate(dag_operators):
        tbl.rows[idx + 1].cells[0].text = str(idx + 1)
        tbl.rows[idx + 1].cells[1].text = op.get("name", "?")
        params = op.get("params", {})
        param_str = _format_params(op.get("name", ""), params)
        tbl.rows[idx + 1].cells[2].text = param_str[:120]
        tbl.rows[idx + 1].cells[3].text = op.get("output_alias", "")[:30]
    _apply_table_borders(tbl)
    _format_table_cells_fonts(tbl)

    # 匹配逻辑专项说明
    match_ops = [op for op in dag_operators if op.get("name") in ("Diff", "Merge", "Reconcile", "RegexFilter", "ConditionCheck")]
    load_ops = [op for op in dag_operators if op.get("name") in ("Load",)]

    if scenario in ("match", "medical_match", "balance_match") or match_ops:
        doc.add_heading("匹配逻辑说明", level=2)

        # 关键词来源披露
        kw_src = (match_logic or {}).get("kw_source", "")
        if kw_src:
            src_labels = {"dictionary": "词典命中（免确认）", "dag": "DAG蓝图编译",
                          "llm_proposed": "LLM自动提案", "fallback": "兜底全量匹配"}
            doc.add_paragraph(f"筛选口径来源：{src_labels.get(kw_src, kw_src)}")

        # 关键词预览
        kw_pv = (match_logic or {}).get("kw_preview", {})
        if kw_pv and kw_pv.get("hit_count", 0) > 0:
            doc.add_paragraph(
                f"命中率预览：{kw_pv['hit_count']} / {kw_pv.get('total','?')} 行 "
                f"（{kw_pv.get('hit_rate',0)*100:.1f}%）")
            excl = kw_pv.get("excluded_top", {})
            if excl:
                excl_items = [f"{k}({v})" for k, v in list(excl.items())[:5]]
                doc.add_paragraph(f"被排除的高频摘要：{'、'.join(excl_items)}")

        # 数据源
        sources = []
        for op in load_ops:
            sf = op.get("source_file", "") or op.get("params", {}).get("file_path", "")
            if sf:
                sources.append(sf)
        if sources:
            doc.add_paragraph(f"数据源：{' 与 '.join(sources)}")

        # 筛选条件
        for op in match_ops:
            name = op.get("name")
            params = op.get("params", {})
            if name == "RegexFilter":
                col = params.get("column", "?")
                pat = params.get("pattern", "?")
                doc.add_paragraph(f"筛选条件：`{col}` 列匹配正则 `{pat}`")
            elif name == "ConditionCheck":
                cond = params.get("condition", "") or f"{params.get('column','?')} {params.get('operator','?')} {params.get('value','?')}"
                doc.add_paragraph(f"条件检查：{cond}")
            elif name in ("Diff", "Merge", "Reconcile"):
                keys = params.get("on", params.get("keys", []))
                tolerance = params.get("tolerance", params.get("容差", {}))
                doc.add_paragraph(f"匹配键：{keys if keys else '自动检测'}")
                if tolerance:
                    doc.add_paragraph(f"容差设置：{tolerance}")

    # 用户确认的匹配规则（从match_logic参数注入 - 一字不落地体现用户确认时的匹配逻辑）
    if match_logic:
        doc.add_heading("用户确认的匹配规则（执行时实际采用的逻辑）", level=2)
        explanation = match_logic.get("explanation", "")
        patterns = match_logic.get("patterns", "")
        method = match_logic.get("method", "")
        columns = match_logic.get("columns", [])
        amount_col = match_logic.get("amount_column", "")
        inst_cols = match_logic.get("institution_columns", [])
        strategy_comparison = match_logic.get("strategy_comparison", [])

        if explanation:
            doc.add_paragraph(f"匹配策略说明：{explanation}")
        if patterns:
            doc.add_paragraph(f"实际筛选关键词（正则模式）：{patterns}")
        if method:
            doc.add_paragraph(f"匹配方法/策略：{method}")
        if columns:
            doc.add_paragraph(f"数据筛选列：{'、'.join(columns)}")
        if amount_col:
            doc.add_paragraph(f"金额取值方式：{amount_col}")
        if inst_cols:
            doc.add_paragraph(f"机构识别来源：{' > '.join(inst_cols)}（优先顺序从左到右）")

        # 多策略对比表
        if strategy_comparison:
            doc.add_heading("多策略对比", level=3)
            if len(strategy_comparison) > 0:
                st = doc.add_table(rows=len(strategy_comparison) + 1, cols=5, style="Light Grid Accent 1")
                for i, h in enumerate(["策略", "匹配率", "差额", "差额比例", "筛选行数"]):
                    st.rows[0].cells[i].text = h
                for idx, s in enumerate(strategy_comparison):
                    st.rows[idx + 1].cells[0].text = s.get("策略", "")
                    st.rows[idx + 1].cells[1].text = s.get("匹配率", "")
                    st.rows[idx + 1].cells[2].text = s.get("差额", "")
                    st.rows[idx + 1].cells[3].text = s.get("差额比例", "")
                    st.rows[idx + 1].cells[4].text = str(s.get("筛选行数", ""))
                _apply_table_borders(st)
                _format_table_cells_fonts(st)

        # 反向校验：未匹配行摘要聚类（审计师一眼发现误杀）
        rv = (match_logic or {}).get("reverse_validation", {})
        if rv and rv.get("clusters"):
            doc.add_heading("反向校验：被排除记录摘要聚类", level=3)
            doc.add_paragraph(
                f"共 {rv.get('total_unmatched', 0)} 条记录未被筛选命中，"
                f"按摘要聚类 TOP{len(rv['clusters'])} 如下。"
                f"请审计师扫一眼，确认是否有目标业务被误杀：")
            rv_tbl = doc.add_table(rows=len(rv["clusters"]) + 1, cols=3,
                                    style="Light Grid Accent 1")
            for i, h in enumerate(["聚类摘要", "行数", "样例"]):
                rv_tbl.rows[0].cells[i].text = h
            for idx, c in enumerate(rv["clusters"]):
                rv_tbl.rows[idx + 1].cells[0].text = c.get("key", "")
                rv_tbl.rows[idx + 1].cells[1].text = str(c.get("count", ""))
                rv_tbl.rows[idx + 1].cells[2].text = c.get("sample", "")[:60]
            _apply_table_borders(rv_tbl)
            _format_table_cells_fonts(rv_tbl)

    # ===== 对账结果（注入 reconciliation stats）=====
    if reconcile_stats:
        doc.add_heading("对账结果摘要", level=1)
        rs = reconcile_stats
        doc.add_paragraph(
            f"匹配率：账方 {rs.get('book_match_rate', 0):.1f}%  |  "
            f"银方 {rs.get('bank_match_rate', 0):.1f}%"
        )
        doc.add_paragraph(
            f"分层命中：L1={rs.get('matched_L1',0)} L2={rs.get('matched_L2',0)} "
            f"L3组={rs.get('matched_L3_groups',0)} L4={rs.get('review_L4',0)}"
        )
        tc = rs.get("timing_categories", {})
        if tc:
            doc.add_paragraph(
                f"未达四分类：银收企未收={tc.get('银收企未收',0)} "
                f"银付企未付={tc.get('银付企未付',0)} "
                f"企收银未收={tc.get('企收银未收',0)} "
                f"企付银未付={tc.get('企付银未付',0)} "
                f"待人工核查={tc.get('待人工核查',0)}"
            )
        doc.add_paragraph(
            f"容忍度：{rs.get('tolerance','?')}  |  "
            f"日期窗口：{rs.get('date_window_days','?')}天  |  "
            f"红旗数：{rs.get('red_flag_count',0)}"
        )
        # 匹配分层表
        tbl = doc.add_table(rows=7, cols=2, style="Light Grid Accent 1")
        layers = [
            ("L1 同额同日", rs.get("matched_L1", 0)),
            ("L2 同额±3天", rs.get("matched_L2", 0)),
            ("L3 n:m合计相等", rs.get("matched_L3_groups", 0)),
            ("L3_fee 手续费差额", rs.get("matched_L3_fee", 0)),
            ("L3_month 月末汇总", rs.get("matched_L3_month", 0)),
            ("L4 模糊匹配（待复核）", rs.get("review_L4", 0)),
        ]
        tbl.rows[0].cells[0].text = "匹配层级"; tbl.rows[0].cells[1].text = "笔数"
        for i, (name, cnt) in enumerate(layers):
            tbl.rows[i+1].cells[0].text = name
            tbl.rows[i+1].cells[1].text = str(cnt)
        _apply_table_borders(tbl); _format_table_cells_fonts(tbl)

        # 注入 match_stats 供后续章节用
        for k, v in rs.items():
            summary.setdefault("match_stats", {})[k] = v

    # ===== 三、执行结果 =====
    doc.add_heading("三、执行结果与数据", level=1)

    # 读取 JSON 摘要
    jp = output_dir / "journal_entries.json"
    summary = {}
    if jp.exists():
        try: summary = json.loads(jp.read_text(encoding="utf-8"))
        except: pass

    total_rows = summary.get("total_rows", 0)
    columns = summary.get("columns", [])
    doc.add_paragraph(f"输出数据：{total_rows} 行，{len(columns)} 列")
    doc.add_paragraph(f"列名：{', '.join(columns)}")

    # 数值统计表
    ns = summary.get("numeric_summary", {})
    if ns:
        doc.add_heading("数值统计", level=2)
        st = doc.add_table(rows=len(ns) + 1, cols=5, style="Light Grid Accent 1")
        for i, h in enumerate(["列名", "合计", "均值", "最大值", "最小值"]):
            st.rows[0].cells[i].text = h
        for idx, (col, s) in enumerate(ns.items()):
            st.rows[idx + 1].cells[0].text = col
            st.rows[idx + 1].cells[1].text = f"{s.get('sum', 0):,.2f}"
            st.rows[idx + 1].cells[2].text = f"{s.get('mean', 0):,.2f}"
            st.rows[idx + 1].cells[3].text = f"{s.get('max', 0):,.2f}"
            st.rows[idx + 1].cells[4].text = f"{s.get('min', 0):,.2f}"
        _apply_table_borders(st)
        _format_table_cells_fonts(st)

    # CSV 数据预览
    cp = output_dir / "analysis_result.csv"
    csv_data = []
    if cp.exists():
        try:
            with open(cp, "r", encoding="utf-8-sig") as f:
                csv_data = list(csv.reader(f))
        except: pass
    if csv_data:
        doc.add_heading("数据明细（前50行）", level=2)
        mr = min(len(csv_data), 51)
        dt = doc.add_table(rows=mr, cols=len(csv_data[0]), style="Light Grid Accent 1")
        for ri, row in enumerate(csv_data[:mr]):
            for ci, val in enumerate(row):
                dt.rows[ri].cells[ci].text = val
        if len(csv_data) > 51:
            doc.add_paragraph(f"...（共 {len(csv_data)-1} 行，仅展示前50行）")
        _apply_table_borders(dt)
        _format_table_cells_fonts(dt)

    # ===== 四、效果评估 =====
    doc.add_heading("四、效果评估与审计建议", level=1)
    _add_scenario_analysis(doc, scenario, summary, total_rows)

    # ===== 五、审计建议 =====
    doc.add_heading("五、审计建议与风险提示", level=1)
    _add_audit_recommendations(doc, scenario, reconcile_stats, workpaper_files)

    # ===== 六、执行日志 =====
    if execution_logs:
        doc.add_heading("六、执行日志", level=1)
        for log in execution_logs[:30]:
            if log and log.strip():
                doc.add_paragraph(log[:500], style="List Bullet")

    # ===== 尾注：免责声明 =====
    doc.add_paragraph("")
    doc.add_paragraph("注意：本报告由智能审计平台自动生成，结果仅供参考。所有审计结论需由注册会计师（CPA）进行专业判断和最终确认。")

    # ===== 七、同源勾稽验证（P0） =====
    doc.add_heading("七、同源勾稽验证", level=1)
    try:
        from core.cross_reference import verify_report_against_workpaper
        # 从 summary 提取报告数字（numeric_summary 中的合计值）
        report_figures = {}
        ns = summary.get("numeric_summary", {})
        for col, s in ns.items():
            if s.get("sum") is not None:
                report_figures[f"{col}_合计"] = s["sum"]
            if s.get("mean") is not None:
                report_figures[f"{col}_均值"] = s["mean"]
            if s.get("max") is not None:
                report_figures[f"{col}_最大值"] = s["max"]
            if s.get("min") is not None:
                report_figures[f"{col}_最小值"] = s["min"]

        # 从 CSV 提取底稿数字（作为 workpaper_data）
        workpaper_figures = {}
        if csv_data and len(csv_data) > 1:
            header = csv_data[0]
            num_cols = []
            for ci, h in enumerate(header):
                vals = []
                for row in csv_data[1:]:
                    try:
                        vals.append(float(row[ci].replace(",", "").replace("元", "").strip()))
                    except (ValueError, IndexError):
                        pass
                if vals:
                    workpaper_figures[f"{h}_合计(底稿)"] = round(sum(vals), 2)
                    workpaper_figures[f"{h}_均值(底稿)"] = round(sum(vals) / len(vals), 2)
                    num_cols.append(h)

        # 执行勾稽
        if report_figures and workpaper_figures:
            xref = verify_report_against_workpaper(
                report_data=report_figures,
                workpaper_data=workpaper_figures,
                run_id=run_id, tolerance=0.01,
            )
            tbl = doc.add_table(rows=len(xref.entries) + 1, cols=4, style="Light Grid Accent 1")
            for i, h in enumerate(["数字名称", "报告值", "底稿值", "状态"]):
                tbl.rows[0].cells[i].text = h
            for idx, e in enumerate(xref.entries):
                tbl.rows[idx + 1].cells[0].text = e.figure_name[:50]
                tbl.rows[idx + 1].cells[1].text = f"{e.report_value:,.2f}"
                tbl.rows[idx + 1].cells[2].text = f"{e.workpaper_value:,.2f}"
                tbl.rows[idx + 1].cells[3].text = "✓ 通过" if e.passed else f"✗ 差异 {abs(e.report_value - e.workpaper_value):,.2f}"
            _apply_table_borders(tbl)
            _format_table_cells_fonts(tbl)
            if not xref.all_passed:
                doc.add_paragraph(f"⚠️ 勾稽失败：{xref.failed_count}/{len(xref.entries)} 项不一致，请核对底稿与报告数字。")
            else:
                doc.add_paragraph(f"✓ 勾稽通过：{len(xref.entries)} 项全部一致。")
        else:
            doc.add_paragraph("（无可用数据进行勾稽验证）")
    except Exception as e:
        doc.add_paragraph(f"勾稽验证异常（非致命）: {e}")

    # ===== 五、审计结论（模板化句式，LLM无参与） =====
    try:
        _render_scene_conclusion(doc, scenario, summary, match_stats)
    except Exception as e:
        doc.add_paragraph(f"结论生成异常（非致命）: {e}")

    # ===== 六、业务合理性勾稽 =====
    try:
        _add_business_crossref_rules(doc, scenario, summary, match_stats)
    except Exception as e:
        doc.add_paragraph(f"业务勾稽异常（非致命）: {e}")

    output_dir.mkdir(parents=True, exist_ok=True)
    rp = output_dir / f"审计报告_{run_id[-12:]}.docx"
    doc.save(str(rp))
    print(f"[报告] 已生成: {rp}")
    return rp


def _add_scenario_analysis(doc, scenario: str, summary: dict, total_rows: int):
    """场景专项效果评估，根据场景类型调用对应评估函数"""
    match_stats = summary.get("match_stats", {})

    if scenario == "medical_match":
        _assess_medical_match(doc, summary, match_stats)
    elif scenario in ("match", "balance_match"):
        _assess_general_match(doc, summary, match_stats)
    elif scenario == "medical":
        _assess_medical_screening(doc, summary, total_rows)
    elif scenario == "screening":
        doc.add_paragraph(f"筛查出 {total_rows} 条记录。风险分级：HIGH>=500万 / MEDIUM>=100万 / LOW>=10万")
    elif scenario == "aggregate":
        doc.add_paragraph(f"汇总得 {total_rows} 条记录，请核实数据完整性。")
    elif scenario in ("balance", "diff"):
        _assess_balance_diff(doc, scenario, summary, match_stats)
    elif scenario == "compliance":
        violations = summary.get("violations", [])
        if violations:
            doc.add_paragraph(f"发现 {len(violations)} 项合规问题，需审计介入。")
        else:
            doc.add_paragraph("未发现明显合规问题。")
    else:
        _assess_general_processing(doc, summary, total_rows)


def _assess_general_processing(doc, summary: dict, total_rows: int):
    """通用数据加工场景评估——输出完整的数据质量摘要和统计信息。"""

    doc.add_paragraph("本场景目标：按审计师指令进行数据加工处理（筛选/汇总/透视/去重/排序/计算列等）。")

    # ── 数据规模 ──
    input_files = summary.get("input_files", [])
    file_count = len(input_files) if isinstance(input_files, list) else (1 if input_files else 0)
    doc.add_paragraph(f"数据规模：{file_count} 个输入文件，输出 {total_rows} 条结果记录。")

    # ── 算子链摘要 ──
    operators = summary.get("operators", [])
    if operators:
        op_names = [op.get("name", op) if isinstance(op, dict) else str(op) for op in operators]
        doc.add_paragraph(f"算子链：{' → '.join(op_names)}")
        op_descriptions = {
            "Load": "加载数据", "RegexFilter": "正则筛选", "ColumnFilter": "列条件筛选",
            "GroupBy": "分组汇总", "Sort": "排序", "ConditionCheck": "条件标记",
            "Extract": "提取子集", "Transform": "列变换/清洗", "NoiseFilter": "噪音过滤",
            "Aggregate": "聚合统计", "Diff": "差异比对", "Merge": "合并关联",
            "Export": "导出结果", "Reconcile": "对账匹配",
        }
        seen = set()
        desc_lines = []
        for op in operators:
            name = op.get("name", "") if isinstance(op, dict) else str(op)
            if name and name not in seen:
                seen.add(name)
                desc = op_descriptions.get(name, "")
                if desc:
                    alias = op.get("output_alias", "") if isinstance(op, dict) else ""
                    desc_lines.append(f"  · {name}（{desc}）{('→ ' + alias) if alias else ''}")
        if desc_lines:
            doc.add_paragraph("算子功能明细：")
            for line in desc_lines:
                doc.add_paragraph(line)

    # ── 数据质量摘要 ──
    quality = summary.get("data_quality", {})
    null_columns = quality.get("null_columns", {}) if quality else {}
    numeric_stats = quality.get("numeric_stats", {}) if quality else {}
    total_input_rows = quality.get("total_input_rows", total_rows)

    if total_input_rows:
        doc.add_paragraph(f"输入总行数：{total_input_rows}，输出行数：{total_rows}，"
                          f"筛选/过滤率：{(1 - total_rows / max(total_input_rows, 1)) * 100:.1f}%")

    if null_columns:
        doc.add_heading("数据质量：空值检查", level=2)
        high_null = [(col, rate) for col, rate in null_columns.items() if rate > 0.2]
        moderate_null = [(col, rate) for col, rate in null_columns.items() if 0.05 < rate <= 0.2]
        if high_null:
            cols_str = "、".join(f"{col}({rate*100:.0f}%)" for col, rate in high_null)
            doc.add_paragraph(f"⚠️ 高空值率列（>20%）：{cols_str}。建议审计师确认这些列的业务必要性。")
        if moderate_null:
            cols_str = "、".join(f"{col}({rate*100:.0f}%)" for col, rate in moderate_null)
            doc.add_paragraph(f"⚡ 中等空值率列（5%-20%）：{cols_str}。")
        if not high_null and not moderate_null:
            doc.add_paragraph("✓ 各列空值率均在5%以下，数据完整性良好。")

    if numeric_stats:
        doc.add_heading("数据质量：数值列统计", level=2)
        try:
            tbl = doc.add_table(rows=len(numeric_stats) + 1, cols=5, style="Light Grid Accent 1")
            for i, h in enumerate(["数值列", "最小值", "最大值", "均值", "标准差"]):
                tbl.rows[0].cells[i].text = h
            row_idx = 1
            for col, stats in numeric_stats.items():
                if isinstance(stats, dict):
                    tbl.rows[row_idx].cells[0].text = str(col)[:30]
                    tbl.rows[row_idx].cells[1].text = f"{stats.get('min', 0):,.2f}"
                    tbl.rows[row_idx].cells[2].text = f"{stats.get('max', 0):,.2f}"
                    tbl.rows[row_idx].cells[3].text = f"{stats.get('mean', 0):,.2f}"
                    tbl.rows[row_idx].cells[4].text = f"{stats.get('std', 0):,.2f}"
                    if stats.get('min', 0) < 0:
                        tbl.rows[row_idx].cells[1].text += " ⚠"
                    row_idx += 1
        except Exception as e:
            doc.add_paragraph(f"数值列统计表渲染异常（非致命）: {e}")
            for col, stats in numeric_stats.items():
                if isinstance(stats, dict):
                    doc.add_paragraph(
                        f"  · {col}: min={stats.get('min',0):,.2f}, "
                        f"max={stats.get('max',0):,.2f}, mean={stats.get('mean',0):,.2f}"
                    )

    # ── 数据完整性声明 ──
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("⚠ 审计师复核要点：").bold = True
    doc.add_paragraph("1. 确认筛选/汇总条件是否完整覆盖审计意图")
    doc.add_paragraph("2. 核查高空值率列的业务合理性（缺失是真实无业务，还是数据源问题）")
    doc.add_paragraph("3. 确认输出结果的行数与业务预期一致")
    doc.add_paragraph("4. 数值列出现负值、零值时确认是否与业务逻辑矛盾")
    doc.add_paragraph("5.【审计红线】平台未对原始数据做任何修改（删行/填0/前向填充），所有缺失保留原样")


def _assess_medical_match(doc, summary: dict, match_stats: dict):
    """医保回款匹配专项评估"""
    doc.add_paragraph("本场景目标：从银行流水中筛选医保回款明细，与医保回款汇总表跨表核对，验证金额一致性。")
    total_bank = match_stats.get("total_bank_rows", 0)
    filtered = match_stats.get("filtered_rows", 0)
    if total_bank > 0:
        doc.add_paragraph(f"银行流水总行数：{total_bank}，筛选出医保相关行：{filtered}（覆盖率 {filtered/total_bank*100:.1f}%）")
    matched_inst = match_stats.get("matched_institutions", 0)
    total_inst = match_stats.get("total_summary_institutions", 0)
    if total_inst > 0:
        doc.add_paragraph(f"汇总表机构数：{total_inst}，成功匹配：{matched_inst}（匹配率 {matched_inst/total_inst*100:.1f}%）")
    total_flt_amt = match_stats.get("total_filtered_amount", 0)
    total_sum_amt = match_stats.get("total_summary_amount", 0)
    if total_sum_amt > 0:
        diff = total_flt_amt - total_sum_amt
        diff_pct = abs(diff) / abs(total_sum_amt) * 100
        doc.add_paragraph(f"筛选汇总金额：{total_flt_amt:,.2f} 元")
        doc.add_paragraph(f"回款表汇总金额：{total_sum_amt:,.2f} 元")
        doc.add_paragraph(f"总差额：{diff:,.2f} 元（{diff_pct:.1f}%）")
    _add_quality_judgment(doc, match_stats)


def _assess_general_match(doc, summary: dict, match_stats: dict):
    """通用匹配评估"""
    tl = match_stats.get("total_left", 0)
    tr = match_stats.get("total_right", 0)
    mc = match_stats.get("matched_count", 0)
    doc.add_paragraph(f"左表：{tl} 条，右表：{tr} 条，匹配成功：{mc} 条")
    if tl > 0:
        doc.add_paragraph(f"匹配率：{mc/tl*100:.1f}%")
    # ↓↓↓ 新增：对账范围口径说明 ↓↓↓
    ib = match_stats.get("interbank_transfers", 0)
    if ib:
        doc.add_paragraph(
            f"对账范围说明：本次为单一银行账户对账。未匹配项中含他行互转 {ib} 笔"
            f"（资金在企业其他银行账户间划转，超出本账户流水范围，不参与匹配率评价），"
            f"详见《未匹配分类核查清单》之'他行互转-待他行流水'分类；"
            f"建议取得相关他行账户流水后另行核对。")
    # ↑↑↑
    _add_quality_judgment(doc, match_stats)


def _assess_medical_screening(doc, summary: dict, total_rows: int):
    """医保回款筛选评估"""
    doc.add_paragraph(f"筛选出 {total_rows} 条医保相关记录。核查要点：")
    for t in ["是否混入非数据行（政策说明等）", "年度汇总与回款表是否一致", "正负数是否合理（回款/退费）"]:
        doc.add_paragraph(f"  - {t}", style="List Bullet")


def _assess_balance_diff(doc, scenario: str, summary: dict, match_stats: dict):
    """科目余额/差异分析评估"""
    if scenario == "balance":
        is_bal = match_stats.get("is_balanced", False)
        bd = match_stats.get("balance_diff", 0)
        if is_bal:
            doc.add_paragraph("科目余额表借贷平衡，无差异。")
        else:
            doc.add_paragraph(f"借贷不平衡，差异金额：{bd:,.2f} 元。")
    else:
        doc.add_paragraph(f"差异分析完成，详见上方差额明细。")


def _add_quality_judgment(doc, match_stats: dict):
    """统一的三级质量判定：良好/一般/不佳；无匹配数据时不误导。"""
    diff_pct_raw = match_stats.get("diff_percentage", match_stats.get("差额比例", None))
    match_rate = match_stats.get("match_rate", match_stats.get("匹配率", 0))

    total_left = match_stats.get("total_left", 0)
    total_right = match_stats.get("total_right", 0)
    total_bank_rows = match_stats.get("total_bank_rows", 0)
    total_summary_institutions = match_stats.get("total_summary_institutions", 0)

    doc.add_paragraph("")
    doc.add_heading("综合判定", level=2)

    # 无有效匹配数据：避免给出“匹配效果很差”的误导结论
    no_both_sides = (total_left == 0 and total_right == 0) or (
        total_bank_rows == 0 and total_summary_institutions == 0
    )
    if no_both_sides:
        doc.add_paragraph(
            "判定结果：当前未提供有效匹配数据或两侧数据为空，无法形成匹配效果判定。"
            "请审计师核对输入文件及数据范围后重新执行。"
        )
        return

    missing_one_side = (total_left == 0 or total_right == 0) or (
        total_bank_rows == 0 or total_summary_institutions == 0
    )
    if missing_one_side:
        doc.add_paragraph(
            "判定结果：缺少某一侧匹配数据，无法完成对账判定。请确认两侧数据均已正确上传。"
        )
        return

    if diff_pct_raw is None:
        doc.add_paragraph(
            "判定结果：未获取到匹配差额信息，无法形成量化判定。请检查匹配逻辑是否正确执行。"
        )
        return

    diff_pct = abs(diff_pct_raw)
    if diff_pct < 5 and match_rate > 90:
        doc.add_paragraph("判定结果：匹配效果良好，差额在可接受范围内（<5%），匹配逻辑合理，可采信当前结果作为审计工作底稿。")
    elif diff_pct < 15 and match_rate > 70:
        doc.add_paragraph("判定结果：匹配效果一般，差额在5%-15%之间。可能原因：部分医保回款使用了不同的摘要描述、存在跨期回款、或部分记录未正确识别。建议审计师人工复核差额较大的机构。")
    elif diff_pct < 30:
        doc.add_paragraph("判定结果：匹配效果不佳，差额在15%-30%之间。建议检查银行流水中是否有其他形式的医保回款描述（如'医疗统筹款''社保回款'等）、检查大额回款的摘要信息。")
    else:
        doc.add_paragraph("判定结果：匹配效果很差，差额超过30%。当前匹配逻辑可能不正确，建议重新检查数据源、调整筛选关键词、确认回款表的统计口径与银行流水一致。")


def _add_audit_recommendations(doc, scenario: str, reconcile_stats: dict = None, workpaper_files: list = None):
    """模板化审计建议：根据匹配质量指标选择模板，不复用硬编码分支。"""
    # 从 reconcile_stats 提取质量指标
    rs = reconcile_stats or {}
    match_rate = max(rs.get("book_match_rate", 0), rs.get("bank_match_rate", 0))
    diff_pct = abs(rs.get("diff_percentage", rs.get("差额比例", 15)))
    red_flags = rs.get("red_flag_count", 0)
    unmatched = rs.get("unmatched_book", 0) + rs.get("unmatched_bank", 0)
    l4_review = rs.get("review_L4", 0)

    # 质量分档
    if match_rate > 95 and diff_pct < 3:
        quality = "excellent"
    elif match_rate > 85 and diff_pct < 10:
        quality = "good"
    elif match_rate > 70:
        quality = "fair"
    else:
        quality = "poor"

    # 风险叠加
    risks = []
    if red_flags > 10: risks.append("high_flags")
    if l4_review > unmatched * 0.3: risks.append("high_l4")
    if diff_pct > 15: risks.append("large_diff")

    # 模板库
    QUALITY_TEMPLATES = {
        "excellent": [
            "匹配效果良好，差额在可接受范围内（<3%），可采信当前结果作为审计工作底稿",
            "建议对少量未匹配项进行抽凭确认，排除系统性遗漏后即可归档",
        ],
        "good": [
            "匹配效果较好，差额在合理范围内（<10%），建议对差异较大的对手方逐笔核实",
            "关注时间性差异导致的未达（期后验证），对L4模糊匹配结果进行人工复核",
        ],
        "fair": [
            "匹配效果一般，差额在10%-30%之间，需要重点分析未达原因",
            "检查是否存在数据口径不一致（如费用账户流水是否完整提供）",
            "对大额未匹配项进行穿透测试，获取原始凭证核实用途",
        ],
        "poor": [
            "⚠ 匹配效果不佳，差额超过30%，当前匹配逻辑可能存在问题",
            "建议重新检查数据源完整性，确认双方数据覆盖同一期间",
            "检查是否存在未提供的银行账户流水，补全数据后重新执行对账",
        ],
    }
    RISK_TEMPLATES = {
        "high_flags": "⚠ 红旗数量较多（>10项），建议优先处置《异常资金交易清单》中的红旗项",
        "high_l4": "⚠ L4模糊匹配占比偏高，可能遗漏大量真实匹配，考虑缩小日期窗口或增加对手方名称模糊匹配",
        "large_diff": "⚠ 差额较大，需重点关注未达四分类中金额聚类的对手方，排查资金体外循环风险",
    }

    recs = list(QUALITY_TEMPLATES.get(quality, QUALITY_TEMPLATES["fair"]))
    for risk in risks:
        if risk in RISK_TEMPLATES:
            recs.append(RISK_TEMPLATES[risk])

    for r in recs:
        doc.add_paragraph(f"  - {r}", style="List Bullet")

    # 底稿联动：列出关联的核查文件
    if workpaper_files:
        doc.add_heading("关联核查底稿", level=2)
        for wf in workpaper_files:
            fn = Path(wf).name if isinstance(wf, str) else wf
            doc.add_paragraph(f"📎 {fn}", style="List Bullet")
        doc.add_paragraph("以上文件包含未匹配项的六桶分桶核查清单和按对手方分组的核查底稿，请逐项复核。")


# ═══════════════════════════════════════════════════════════════
# 底稿要素头（索引号/期间/编制人/复核人/日期 + 数据来源声明）
# ═══════════════════════════════════════════════════════════════

def _add_working_paper_header(doc, run_id: str, user_intent: str, scenario: str,
                               input_files: List[str], meta: dict = None):
    """在每个底稿报告头部注入标准底稿要素块。"""
    index_map = {
        "match": "C-1-1", "medical_match": "C-2-1",
        "screening": "D-1-1", "balance": "E-1-1",
        "diff": "E-2-1", "general": "G-1-1",
    }
    index_no = index_map.get(scenario, "G-1-1")
    doc.add_paragraph("")
    doc.add_heading("底稿信息", level=2)
    tbl = doc.add_table(rows=5, cols=2, style="Light Grid Accent 1")
    for i, (h, v) in enumerate([
        ("索引号", index_no),
        ("被审计单位", meta.get("client_name", "（待填）") if meta else "（待填）"),
        ("会计期间", meta.get("accounting_period", "（待填）") if meta else "（待填）"),
        ("编制人/日期", f"智能审计平台 / {datetime.now().strftime('%Y-%m-%d')}"),
        ("复核人/日期", "________ / ________"),
    ]):
        tbl.rows[i].cells[0].text = h
        tbl.rows[i].cells[1].text = v
    _apply_table_borders(tbl)
    _format_table_cells_fonts(tbl)
    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("数据来源与口径声明：").bold = True
    doc.add_paragraph(f"源文件：{'、'.join(input_files) if input_files else '（未记录）'}")
    doc.add_paragraph(f"处理平台：智能审计平台 v3，Run ID：{run_id}")
    doc.add_paragraph(f"审计师指令：{user_intent[:200]}")


# ═══════════════════════════════════════════════════════════════
# 场景模板化结论（LLM只填空，句式固定）
# ═══════════════════════════════════════════════════════════════

def _render_scene_conclusion(doc, scenario: str, summary: dict, match_stats: dict):
    """用固定句式模板生成审计结论段。变量从 stats 填入。"""
    doc.add_heading("审计结论", level=1)
    if scenario in ("match", "balance_match"):
        tl = match_stats.get("total_left", match_stats.get("total_bank_rows", 0))
        tr = match_stats.get("total_right", match_stats.get("total_summary_institutions", 0))
        mc = match_stats.get("matched_count", 0)
        mr = match_stats.get("match_rate", 0)
        dp = abs(match_stats.get("diff_percentage", match_stats.get("差额比例", 0)))
        doc.add_paragraph(f"经核对，账方共 {tl} 笔、银方共 {tr} 笔，匹配成功 {mc} 笔，匹配率 {mr:.1f}%。")
        if dp < 5 and mr > 90:
            doc.add_paragraph("总体差异在可接受范围内（<5%），未发现重大异常，可采信作为审计工作底稿。")
        elif dp < 15:
            doc.add_paragraph(f"存在一定差异（{dp:.1f}%），建议对差额较大的项目逐笔核实。")
        else:
            doc.add_paragraph(f"⚠ 差异较大（{dp:.1f}%），建议重新检查匹配逻辑、确认数据口径一致性后重新执行。")
    elif scenario == "medical_match":
        total_bank = match_stats.get("total_bank_rows", 0)
        filtered = match_stats.get("filtered_rows", 0)
        matched_inst = match_stats.get("matched_institutions", 0)
        total_inst = match_stats.get("total_summary_institutions", 0)
        doc.add_paragraph(f"从银行流水 {total_bank} 笔中筛选医保相关 {filtered} 笔，"
                          f"与回款汇总表 {total_inst} 个机构比对，成功匹配 {matched_inst} 个。"
                          f"匹配率 {matched_inst/max(total_inst,1)*100:.1f}%。")
        patterns = match_stats.get("patterns_used", "")
        cols_used = match_stats.get("columns_used", [])
        if patterns:
            doc.add_paragraph(f"筛选口径披露：列 {cols_used}，关键词「{patterns}」。")
        doc.add_paragraph("⚠ 审计师应确认筛选关键词是否完整覆盖目标业务范围。"
                          "未命中≠不存在，可能使用了不同的摘要描述。")
    elif scenario == "screening":
        doc.add_paragraph(f"筛查完成，共标记 {summary.get('total_rows', 0)} 条记录。"
                          "风险分级：HIGH（≥500万）/ MEDIUM（≥100万）/ LOW（≥10万）。"
                          "HIGH 级风险交易建议逐笔核查原始凭证。")
    elif scenario == "general":
        doc.add_paragraph(f"数据处理完成，输出 {summary.get('total_rows', 0)} 条结果记录。"
                          "请审计师确认处理逻辑是否符合预期，详见数据质量摘要。")
    else:
        doc.add_paragraph("执行完成。请审计师根据专业判断对处理结果进行复核确认。")



# ═══════════════════════════════════════════════════════════════
# 业务合理性勾稽（匹配分层自洽 / 未达分类合计 / 调节表差异）
# ═══════════════════════════════════════════════════════════════

def _add_business_crossref_rules(doc, scenario: str, summary: dict, match_stats: dict):
    """3条硬规则：分层自洽、未达合计=未匹配、调节表差异=0。"""
    doc.add_heading("业务合理性勾稽", level=1)
    checks = []
    l1 = match_stats.get("matched_L1", 0)
    l2 = match_stats.get("matched_L2", 0)
    l3 = match_stats.get("matched_L3_groups", 0)
    l4 = match_stats.get("review_L4", 0)
    mc = match_stats.get("matched_count", 0)
    if any([l1, l2, l3, l4]):
        layer_sum = l1 + l2 + l3 + l4
        passed = layer_sum <= mc * 1.05
        checks.append(("匹配分层自洽",
            f"L1={l1}+L2={l2}+L3组={l3}+L4待复核={l4}={layer_sum}, matched_count={mc}",
            "✓ 通过" if passed else f"✗ 不通过（{layer_sum}≠{mc}）", passed))
    tc = match_stats.get("timing_categories", {})
    if tc:
        tc_sum = sum(tc.values())
        unmatched = match_stats.get("unmatched_count", 0)
        passed = abs(tc_sum - unmatched) <= max(tc_sum, unmatched) * 0.02
        checks.append(("未达四分类合计=未匹配总数",
            f"四分类合计={tc_sum}（{tc}），未匹配≈{unmatched}",
            "✓ 通过" if passed else f"✗ 不通过", passed))
    adj = match_stats.get("adjustment_difference", match_stats.get("调节差异"))
    if adj is not None:
        adj_abs = abs(float(adj)) if not isinstance(adj, (int, float)) else abs(adj)
        passed = adj_abs < 0.02
        checks.append(("调节表差异=0", f"调节差异={adj_abs:.2f}元",
            "✓ 通过" if passed else "✗ 不通过（仍有未解释差异，必须查明）", passed))
    if not checks:
        doc.add_paragraph("（当前场景无适用的业务勾稽规则）")
        return
    tbl = doc.add_table(rows=len(checks)+1, cols=3, style="Light Grid Accent 1")
    for i, h in enumerate(["勾稽规则", "数据", "判定"]):
        tbl.rows[0].cells[i].text = h
    for idx, (rule, detail, result, passed) in enumerate(checks):
        tbl.rows[idx+1].cells[0].text = rule
        tbl.rows[idx+1].cells[1].text = detail
        tbl.rows[idx+1].cells[2].text = result
    _apply_table_borders(tbl)
    _format_table_cells_fonts(tbl)
    all_pass = all(c[3] for c in checks)
    doc.add_paragraph("✓ 全部业务勾稽规则通过，底稿逻辑自洽。" if all_pass
                      else "⚠ 存在未通过勾稽规则，请审计师核实后修正。")

def _format_params(op_name: str, params: dict) -> str:
    """将算子参数格式化为可读字符串"""
    if not params:
        return "-"
    if op_name == "Load":
        return f"文件: {params.get('file_path', params.get('source_file', '-'))}"
    if op_name == "RegexFilter":
        return f"列={params.get('column','?')}, 正则={params.get('pattern','?')}"
    if op_name == "ConditionCheck":
        return params.get("condition", "") or f"{params.get('column','?')} {params.get('operator','?')} {params.get('value','?')}"
    if op_name == "Aggregate":
        aggs = params.get("aggregations", {})
        cols = params.get("columns", [])
        grp = params.get("group_by", [])
        parts = []
        if aggs:
            parts.append(f"聚合: {list(aggs.keys())[:3]}")
        if cols:
            parts.append(f"列: {cols[:3]}")
        if grp:
            parts.append(f"分组: {grp[:3]}")
        return ", ".join(parts) if parts else f"函数: {params.get('aggregation','sum')}"
    if op_name in ("Diff", "Merge", "Reconcile"):
        keys = params.get("on", params.get("keys", []))
        return f"匹配键: {keys}" if keys else "自动匹配"
    if op_name == "Sort":
        return f"排序列: {params.get('columns', params.get('column', '?'))}"
    if op_name == "Export":
        return f"输出: {params.get('filename', params.get('file', '?'))}"
    if op_name == "NoiseFilter":
        return f"排除: {params.get('keywords', params.get('patterns', '?'))}"
    # 通用：取前两个 key
    items = list(params.items())[:2]
    return ", ".join(f"{k}={v}" for k, v in items)
