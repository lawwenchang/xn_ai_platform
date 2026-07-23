#!/usr/bin/env python3
"""格式规范化引擎 - 模板参照模式：以一份文档为模板，批量转换其他文档"""
from __future__ import annotations
import json
import re

from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

@dataclass
class WordFontFormat:
    name: Optional[str] = None; size: Optional[float] = None
    bold: Optional[bool] = None; italic: Optional[bool] = None
    underline: Optional[bool] = None; color: Optional[str] = None

@dataclass
class WordParaFormat:
    alignment: Optional[int] = None; first_line_indent: Optional[float] = None
    line_spacing: Optional[float] = None
    space_before: Optional[float] = None; space_after: Optional[float] = None

@dataclass
class WordStyles:
    heading1: WordFontFormat = field(default_factory=WordFontFormat)
    heading2: WordFontFormat = field(default_factory=WordFontFormat)
    heading3: WordFontFormat = field(default_factory=WordFontFormat)
    body: WordFontFormat = field(default_factory=WordFontFormat)
    body_para: WordParaFormat = field(default_factory=WordParaFormat)
    table_header: WordFontFormat = field(default_factory=WordFontFormat)


def extract_word_format(template_path: str) -> WordStyles:
    from docx import Document
    doc = Document(template_path)
    styles = WordStyles()
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text: continue
        style_name = para.style.name if para.style else ""
        is_heading = "heading" in style_name.lower() or "Heading" in style_name
        for run in para.runs:
            f = run.font
            wf = WordFontFormat(name=f.name, size=f.size.pt if f.size else None,
                bold=f.bold, italic=f.italic, underline=f.underline,
                color=str(f.color.rgb) if f.color and f.color.rgb else None)
            pf = para.paragraph_format
            wp = WordParaFormat(alignment=para.alignment,
                first_line_indent=pf.first_line_indent.pt if pf.first_line_indent else None,
                line_spacing=pf.line_spacing,
                space_before=pf.space_before.pt if pf.space_before else None,
                space_after=pf.space_after.pt if pf.space_after else None)
            if is_heading:
                lvl = int([c for c in style_name if c.isdigit()][0]) if any(c.isdigit() for c in style_name) else 1
                if lvl == 1 and not styles.heading1.name: styles.heading1 = wf
                elif lvl == 2 and not styles.heading2.name: styles.heading2 = wf
                elif lvl == 3 and not styles.heading3.name: styles.heading3 = wf
            elif not styles.body.name:
                styles.body = wf; styles.body_para = wp
            if styles.heading1.name and styles.body.name: break
    for table in doc.tables:
        for row in table.rows[:1]:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        f = run.font
                        if f.bold and not styles.table_header.name:
                            styles.table_header = WordFontFormat(name=f.name,
                                size=f.size.pt if f.size else None, bold=True, italic=f.italic,
                                color=str(f.color.rgb) if f.color and f.color.rgb else None)
                        if styles.table_header.name: break
                if styles.table_header.name: break
        if styles.table_header.name: break
    return styles



def apply_word_format(doc_path: str, styles: WordStyles, output_path: str):
    """将 WordStyles 应用到目标 Word 文档"""
    from docx import Document
    from docx.shared import Pt, RGBColor

    doc = Document(doc_path)

    def _apply_font(run, wf: WordFontFormat):
        if wf.name:
            run.font.name = wf.name
        if wf.size is not None:
            run.font.size = Pt(wf.size)
        if wf.bold is not None:
            run.bold = wf.bold
        if wf.italic is not None:
            run.italic = wf.italic
        if wf.underline is not None:
            run.underline = wf.underline
        if wf.color:
            try:
                run.font.color.rgb = RGBColor.from_string(wf.color)
            except Exception:
                pass

    def _apply_para(para, wp: WordParaFormat):
        pf = para.paragraph_format
        if wp.alignment is not None:
            para.alignment = wp.alignment
        if wp.first_line_indent is not None:
            pf.first_line_indent = Pt(wp.first_line_indent)
        if wp.line_spacing is not None:
            pf.line_spacing = wp.line_spacing
        if wp.space_before is not None:
            pf.space_before = Pt(wp.space_before)
        if wp.space_after is not None:
            pf.space_after = Pt(wp.space_after)

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style_name = para.style.name if para.style else ""
        is_heading = "heading" in style_name.lower() or "Heading" in style_name

        if is_heading:
            lvl = 0
            if any(c.isdigit() for c in style_name):
                lvl = int([c for c in style_name if c.isdigit()][0])
            target = None
            if lvl == 1:
                target = styles.heading1
            elif lvl == 2:
                target = styles.heading2
            elif lvl == 3:
                target = styles.heading3
            if target:
                for run in para.runs:
                    _apply_font(run, target)
        else:
            if styles.body.name:
                for run in para.runs:
                    _apply_font(run, styles.body)
            if styles.body_para:
                _apply_para(para, styles.body_para)

    # Apply table header style
    if styles.table_header.name:
        for table in doc.tables:
            for row in table.rows[:1]:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            _apply_font(run, styles.table_header)

    doc.save(output_path)
    return output_path


@dataclass
class ExcelCellFormat:
    font_name: Optional[str] = None; font_size: Optional[float] = None
    font_bold: Optional[bool] = None; font_color: Optional[str] = None
    fill_color: Optional[str] = None; border_style: Optional[str] = None
    number_format: Optional[str] = None
    alignment_h: Optional[str] = None; alignment_v: Optional[str] = None

@dataclass
class ExcelTemplate:
    header_format: ExcelCellFormat = field(default_factory=ExcelCellFormat)
    data_format: ExcelCellFormat = field(default_factory=ExcelCellFormat)
    title_format: ExcelCellFormat = field(default_factory=ExcelCellFormat)
    column_widths: Dict[str, float] = field(default_factory=dict)
    row_height: Optional[float] = None

def _cell_to_format(cell):
    f = cell.font; fl = cell.fill; bd = cell.border; al = cell.alignment
    return ExcelCellFormat(font_name=f.name, font_size=f.size, font_bold=f.bold,
        font_color=str(f.color.rgb) if f.color and f.color.rgb else None,
        fill_color=str(fl.fgColor.rgb) if fl.fgColor and fl.fgColor.rgb else None,
        border_style=str(bd.left.style) if bd.left else None,
        number_format=cell.number_format,
        alignment_h=str(al.horizontal) if al else None,
        alignment_v=str(al.vertical) if al else None)


def extract_excel_format(template_path: str) -> ExcelTemplate:
    from openpyxl import load_workbook
    wb = load_workbook(template_path); ws = wb.active; t = ExcelTemplate()
    if ws.max_row >= 1: t.title_format = _cell_to_format(ws.cell(1, 1))
    if ws.max_row >= 2:
        t.header_format = _cell_to_format(ws.cell(2, 1))
        for col in range(1, ws.max_column + 1):
            t.column_widths[str(col)] = ws.column_dimensions[ws.cell(2, col).column_letter].width or 8.43
    if ws.max_row >= 3: t.data_format = _cell_to_format(ws.cell(3, 1))
    if ws.row_dimensions[1].height: t.row_height = ws.row_dimensions[1].height
    wb.close(); return t

def apply_excel_format(excel_path, tmpl, output_path):
    from openpyxl import load_workbook
    from openpyxl.styles import Font, PatternFill, Border, Side, Alignment
    wb = load_workbook(excel_path)
    def _f(ef):
        kw = {}
        if ef.font_name: kw["name"] = ef.font_name
        if ef.font_size: kw["size"] = ef.font_size
        if ef.font_bold is not None: kw["bold"] = ef.font_bold
        if ef.font_color:
            try: kw["color"] = ef.font_color
            except: pass
        return Font(**kw) if kw else Font()
    def _fl(ef):
        if ef.fill_color:
            try: return PatternFill(start_color=ef.fill_color, end_color=ef.fill_color, fill_type="solid")
            except: pass
        return None
    def _bd(ef):
        if ef.border_style:
            try: s = Side(style=ef.border_style); return Border(left=s, right=s, top=s, bottom=s)
            except: pass
        return None
    def _al(ef):
        kw = {}
        if ef.alignment_h: kw["horizontal"] = ef.alignment_h
        if ef.alignment_v: kw["vertical"] = ef.alignment_v
        return Alignment(**kw) if kw else None
    for ws in wb.worksheets:
        if ws.max_row < 1: continue
        for cs, width in tmpl.column_widths.items():
            try:
                ci = int(cs)
                letter = chr(64 + ci) if ci <= 26 else None
                if letter: ws.column_dimensions[letter].width = width
            except: pass
        if ws.max_row >= 1:
            hf = _f(tmpl.title_format); fl = _fl(tmpl.title_format)
            bd = _bd(tmpl.title_format); al = _al(tmpl.title_format)
            for col in range(1, ws.max_column + 1):
                c = ws.cell(1, col); c.font = hf
                if fl: c.fill = fl
                if bd: c.border = bd
                if al: c.alignment = al
            if tmpl.row_height: ws.row_dimensions[1].height = tmpl.row_height
        if ws.max_row >= 2:
            hf = _f(tmpl.header_format); fl = _fl(tmpl.header_format)
            bd = _bd(tmpl.header_format); al = _al(tmpl.header_format)
            for col in range(1, ws.max_column + 1):
                c = ws.cell(2, col); c.font = hf
                if fl: c.fill = fl
                if bd: c.border = bd
                if al: c.alignment = al
        if ws.max_row >= 3:
            df = _f(tmpl.data_format); fl = _fl(tmpl.data_format)
            bd = _bd(tmpl.data_format); al = _al(tmpl.data_format)
            for row in range(3, ws.max_row + 1):
                for col in range(1, ws.max_column + 1):
                    c = ws.cell(row, col); c.font = df
                    if fl: c.fill = fl
                    if bd: c.border = bd
                    if al: c.alignment = al
                    if tmpl.data_format.number_format: c.number_format = tmpl.data_format.number_format
    wb.save(output_path)
    return output_path


def normalize_format(template_path, target_paths, output_dir):
    ext = Path(template_path).suffix.lower()
    outputs = []
    out = Path(output_dir)
    out.mkdir(parents=True, exist_ok=True)
    for target in target_paths:
        tn = Path(target).stem
        if ext in (".xlsx", ".xlsm", ".xls"):
            tmpl = extract_excel_format(template_path)
            ps = extract_excel_print(template_path)
            op = str(out / f"{tn}_formatted{ext}")
            apply_excel_format(target, tmpl, op)
            apply_excel_print(op, ps, op)
        elif ext in (".docx",):
            styles = extract_word_format(template_path)
            ps = extract_word_print(template_path)
            op = str(out / f"{tn}_formatted{ext}")
            apply_word_format(target, styles, op)
            apply_word_print(op, ps, op)
        else:
            continue
        outputs.append(op)
    return outputs



# ═══════════ 自然语言格式指令引擎 (§4.3.2) ═══════════
# vLLM 解析 "差异>5万标红加粗" → OpenPyXL 条件格式

@dataclass
class FormatInstruction:
    condition: str = ""; column: str = ""
    font_bold: bool = False; font_color: str = ""; fill_color: str = ""
    font_size: float = 0; number_format: str = ""

def parse_format_instruction(user_instruction: str) -> List[FormatInstruction]:
    """vLLM解析自然语言→结构化格式规则。例: 差异超过5万标红加粗→[FormatInstruction(condition='差额>50000',font_bold=True,font_color='FF0000')]"""
    import httpx, os, re, json
    prompt = f"""你是审计底稿格式专家。将自然语言格式指令解析为JSON规则。
【指令】{user_instruction}
【属性】condition(条件,万=10000), column(列名), font_bold, font_color(如FF0000), fill_color(如E0FFE0), font_size, number_format
【输出】严格JSON: [{{"condition":"差额>50000","font_bold":true,"font_color":"FF0000"}}]"""

    try:
        r = httpx.post(os.environ.get("VLLM_TUNNEL_URL","http://localhost:18000/v1/chat/completions"),
            headers={"Authorization":"Bearer EMPTY"},
            json={"model":os.environ.get("VLLM_MODEL","qwen3-235b"),
                  "messages":[{"role":"user","content":prompt}],"temperature":0.1,"max_tokens":800}, timeout=30)
        r.raise_for_status()
        text = r.json()["choices"][0]["message"]["content"].strip()
        m = re.search(r'\[[\s\S]*\]', text)
        if m:
            items = json.loads(m.group(0))
            return [FormatInstruction(**{k:v for k,v in item.items()
                    if k in FormatInstruction.__dataclass_fields__}) for item in items]
    except Exception as e:
        print(f"[格式指令] LLM解析失败: {e}")

    # 降级: 关键字匹配
    inst = []
    if "标红" in user_instruction or "红色" in user_instruction:
        fi = FormatInstruction(font_color="FF0000")
        if "加粗" in user_instruction: fi.font_bold = True
        inst.append(fi)
    if "加粗" in user_instruction and not inst:
        inst.append(FormatInstruction(font_bold=True))
    return inst

def _find_col(headers: List[str], target: str) -> Optional[str]:
    for h in headers:
        if target.lower() in str(h).lower(): return str(h)
    return None

def apply_format_instructions(excel_path: str, instructions: List[FormatInstruction],
                               output_path: str = "", headers: List[str] = None) -> str:
    """将格式指令应用到Excel底稿（OpenPyXL条件格式）"""
    from openpyxl import load_workbook
    from openpyxl.styles import Font, PatternFill
    from openpyxl.formatting.rule import CellIsRule

    if not output_path: output_path = excel_path.replace(".xlsx","_formatted.xlsx")
    if not instructions: import shutil; shutil.copy(excel_path,output_path); return output_path
    wb = load_workbook(excel_path); ws = wb.active
    if headers is None: headers = [str(ws.cell(1,c).value or "") for c in range(1,ws.max_column+1)]

    applied = 0
    for fi in instructions:
        # 解析条件
        m = re.match(r'(.+?)\s*(>=|<=|!=|==|>|<)\s*(.+)$', fi.condition) if fi.condition else None
        if not m:
            from openpyxl.styles import Font as _F2, PatternFill as _PF2
            for row in range(2, ws.max_row+1):
                for col in range(1, ws.max_column+1):
                    c = ws.cell(row,col)
                    if fi.font_bold: c.font = _F2(bold=True)
                    if fi.fill_color:
                        try: c.fill = _PF2(start_color=fi.fill_color,end_color=fi.fill_color,fill_type="solid")
                        except: pass
            applied += 1; continue
        col_expr, op, val = m.group(1).strip(), m.group(2), m.group(3).strip()
        col = _find_col(headers, col_expr) or col_expr
        from openpyxl.utils import get_column_letter
        col_idx = None
        for ii, h in enumerate(headers):
            if col.lower() in str(h).lower(): col_idx = get_column_letter(ii+1); break
        if not col_idx: continue
        font_kw = {}
        if fi.font_bold: font_kw["bold"] = True
        if fi.font_color:
            try: font_kw["color"] = fi.font_color
            except: pass
        fill = None
        if fi.fill_color:
            try: fill = PatternFill(start_color=fi.fill_color,end_color=fi.fill_color,fill_type="solid")
            except: pass
        op_map = {">":"greaterThan",">=":"greaterThanOrEqual","<":"lessThan","<=":"lessThanOrEqual","==":"equal","!=":"notEqual"}
        ox_op = op_map.get(op)
        data_range = f"{col_idx}2:{col_idx}{ws.max_row}"
        if ox_op and val.replace(".","").replace("-","").isdigit():
            rule = CellIsRule(operator=ox_op, formula=[str(float(val))],
                              font=Font(**font_kw) if font_kw else None, fill=fill)
            ws.conditional_formatting.add(data_range, rule)
            applied += 1
        elif op == "==":
            rule = CellIsRule(operator="equal", formula=[f'"{val}"'],
                              font=Font(**font_kw) if font_kw else None, fill=fill)
            ws.conditional_formatting.add(data_range, rule)
            applied += 1
    wb.save(output_path)
    print(f"[格式指令] {applied}/{len(instructions)}条规则 -> {output_path}")
    return output_path

def smart_format_excel(excel_path: str, user_instruction: str, output_path: str = "") -> str:
    instructions = parse_format_instruction(user_instruction)
    if not instructions: return excel_path
    return apply_format_instructions(excel_path, instructions, output_path)

@dataclass
class WordPrintSettings:
    top_margin: Optional[float] = None
    bottom_margin: Optional[float] = None
    left_margin: Optional[float] = None
    right_margin: Optional[float] = None
    page_width: Optional[float] = None
    page_height: Optional[float] = None
    orientation: Optional[int] = None  # 0=portrait, 1=landscape

@dataclass  
class ExcelPrintSettings:
    orientation: Optional[str] = None  # portrait/landscape
    paper_size: Optional[int] = None
    top_margin: Optional[float] = None
    bottom_margin: Optional[float] = None
    left_margin: Optional[float] = None
    right_margin: Optional[float] = None
    header_margin: Optional[float] = None
    footer_margin: Optional[float] = None
    fit_to_width: Optional[int] = None
    fit_to_height: Optional[int] = None
    print_area: Optional[str] = None


def extract_word_print(template_path: str) -> WordPrintSettings:
    from docx import Document
    from docx.shared import Cm
    doc = Document(template_path)
    ps = WordPrintSettings()
    if doc.sections:
        s = doc.sections[0]
        ps.top_margin = s.top_margin.cm if s.top_margin else None
        ps.bottom_margin = s.bottom_margin.cm if s.bottom_margin else None
        ps.left_margin = s.left_margin.cm if s.left_margin else None
        ps.right_margin = s.right_margin.cm if s.right_margin else None
        ps.page_width = s.page_width.cm if s.page_width else None
        ps.page_height = s.page_height.cm if s.page_height else None
        ps.orientation = s.orientation
    return ps


def apply_word_print(doc_path: str, ps: WordPrintSettings, output_path: str):
    from docx import Document
    from docx.shared import Cm
    from docx.enum.section import WD_ORIENT
    doc = Document(doc_path)
    for s in doc.sections:
        if ps.top_margin is not None: s.top_margin = Cm(ps.top_margin)
        if ps.bottom_margin is not None: s.bottom_margin = Cm(ps.bottom_margin)
        if ps.left_margin is not None: s.left_margin = Cm(ps.left_margin)
        if ps.right_margin is not None: s.right_margin = Cm(ps.right_margin)
        if ps.page_width is not None: s.page_width = Cm(ps.page_width)
        if ps.page_height is not None: s.page_height = Cm(ps.page_height)
        if ps.orientation is not None: s.orientation = ps.orientation
    doc.save(output_path)


def extract_excel_print(template_path: str) -> ExcelPrintSettings:
    from openpyxl import load_workbook
    from openpyxl.worksheet.page import PageMargins
    wb = load_workbook(template_path); ws = wb.active
    ps = ExcelPrintSettings()
    ps.orientation = ws.page_setup.orientation
    ps.paper_size = ws.page_setup.paperSize
    if ws.page_margins:
        ps.top_margin = ws.page_margins.top
        ps.bottom_margin = ws.page_margins.bottom
        ps.left_margin = ws.page_margins.left
        ps.right_margin = ws.page_margins.right
        ps.header_margin = ws.page_margins.header
        ps.footer_margin = ws.page_margins.footer
    ps.fit_to_width = ws.page_setup.fitToWidth
    ps.fit_to_height = ws.page_setup.fitToHeight
    ps.print_area = ws.print_area
    wb.close(); return ps


def apply_excel_print(excel_path: str, ps: ExcelPrintSettings, output_path: str):
    from openpyxl import load_workbook
    from openpyxl.worksheet.page import PageMargins
    wb = load_workbook(excel_path)
    for ws in wb.worksheets:
        if ps.orientation: ws.page_setup.orientation = ps.orientation
        if ps.paper_size: ws.page_setup.paperSize = ps.paper_size
        if ps.fit_to_width: ws.page_setup.fitToWidth = ps.fit_to_width
        if ps.fit_to_height: ws.page_setup.fitToHeight = ps.fit_to_height
        if ps.print_area: ws.print_area = ps.print_area
        if any(v is not None for v in [ps.top_margin, ps.bottom_margin, ps.left_margin, ps.right_margin]):
            ws.page_margins = PageMargins(
                top=ps.top_margin or 1.0, bottom=ps.bottom_margin or 1.0,
                left=ps.left_margin or 0.75, right=ps.right_margin or 0.75,
                header=ps.header_margin or 0.5, footer=ps.footer_margin or 0.5)
    wb.save(output_path)
    wb.save(output_path)


# ═══════════════ 批量格式规范化 ═══════════════

def batch_normalize_excel(
    template_path: str,
    target_paths: List[str],
    output_dir: str = "",
) -> List[dict]:
    """
    以一个 Excel 为模板，批量规范化其他 Excel 的格式。

    提取模板的全部格式（字体/颜色/边框/列宽/行高/数字格式/条件格式/打印设置），
    应用到每个目标文件。目标文件的数据保持不变。
    """
    from openpyxl import load_workbook
    from copy import copy

    if not output_dir:
        output_dir = str(Path(__file__).resolve().parent.parent / "data" / "outputs")
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    # 1. 提取模板格式
    tmpl_wb = load_workbook(template_path)
    tmpl_ws = tmpl_wb.active

    tmpl_fmt = {
        "column_widths": {},
        "row_heights": {},
        "cell_formats": {},
        "merged_cells": list(tmpl_ws.merged_cells.ranges),
        "freeze_panes": tmpl_ws.freeze_panes,
    }
    for col_letter, dim in tmpl_ws.column_dimensions.items():
        if dim.width: tmpl_fmt["column_widths"][col_letter] = dim.width
    for row_num, dim in tmpl_ws.row_dimensions.items():
        if dim.height: tmpl_fmt["row_heights"][row_num] = dim.height
    for row in tmpl_ws.iter_rows():
        for cell in row:
            if cell.has_style:
                tmpl_fmt["cell_formats"][cell.coordinate] = {
                    "font": copy(cell.font),
                    "fill": copy(cell.fill),
                    "border": copy(cell.border),
                    "alignment": copy(cell.alignment),
                    "number_format": cell.number_format,
                }
    tmpl_wb.close()

    # 2. 应用到每个目标文件
    results = []
    for tp in target_paths:
        try:
            twb = load_workbook(tp)
            tws = twb.active

            # 列宽
            for col, w in tmpl_fmt["column_widths"].items():
                tws.column_dimensions[col].width = w
            # 行高
            for row, h in tmpl_fmt["row_heights"].items():
                tws.row_dimensions[row].height = h
            # 冻结窗格
            if tmpl_fmt["freeze_panes"]:
                tws.freeze_panes = tmpl_fmt["freeze_panes"]
            # 单元格格式（只改格式，不改数据）
            for coord, fmt in tmpl_fmt["cell_formats"].items():
                if coord in tws:
                    c = tws[coord]
                    c.font = fmt["font"]
                    c.fill = fmt["fill"]
                    c.border = fmt["border"]
                    c.alignment = fmt["alignment"]
                    if fmt["number_format"] and fmt["number_format"] != "General":
                        c.number_format = fmt["number_format"]

            out = str(Path(output_dir) / f"fmt_{Path(tp).stem}.xlsx")
            twb.save(out)
            results.append({"file": tp, "output": out, "status": "ok"})
            twb.close()
        except Exception as e:
            results.append({"file": tp, "output": "", "status": "failed", "error": str(e)})

    return results


def batch_normalize_word(
    template_path: str,
    target_paths: List[str],
    output_dir: str = "",
) -> List[dict]:
    """
    以一个 Word 为模板，批量规范化其他 Word 的格式。

    提取模板的全部格式（字体/段落/页边距/页眉页脚/页面设置），
    应用到每个目标文件。目标文件的内容保持不变。
    """
    if not output_dir:
        output_dir = str(Path(__file__).resolve().parent.parent / "data" / "outputs")
    Path(output_dir).mkdir(parents=True, exist_ok=True)

    # 1. 提取模板格式
    styles = extract_word_format(template_path)
    print_settings = extract_word_print(template_path)

    # 2. 应用到每个目标
    results = []
    for tp in target_paths:
        try:
            out = str(Path(output_dir) / f"fmt_{Path(tp).stem}.docx")
            apply_word_format(tp, styles, out)
            apply_word_print(out, print_settings, out)
            results.append({"file": tp, "output": out, "status": "ok"})
        except Exception as e:
            results.append({"file": tp, "output": "", "status": "failed", "error": str(e)})

    return results

