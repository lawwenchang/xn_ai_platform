#!/usr/bin/env python3
"""
跨文档一致性检查引擎 (§4.3.2)
==============================
vLLM 对比两份审计文档，自动发现数字不一致/措辞矛盾/遗漏披露。
"""
import json, os, re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional

@dataclass
class Discrepancy:
    severity: str = "WARNING"  # CRITICAL/WARNING/INFO
    category: str = ""         # 数字不一致/措辞矛盾/遗漏披露/格式差异
    loc_a: str = ""; val_a: str = ""
    loc_b: str = ""; val_b: str = ""
    suggestion: str = ""; evidence: str = ""

@dataclass
class CrossDocReport:
    doc_a: str = ""; doc_b: str = ""
    items: List[Discrepancy] = field(default_factory=list)
    total: int = 0; critical: int = 0

def _read_doc(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".docx":
        from docx import Document
        doc = Document(path)
        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for t in doc.tables:
            for r in t.rows:
                parts.append(" | ".join(c.text for c in r.cells if c.text.strip()))
        return "\n".join(parts)
    if ext in (".xlsx", ".xls"):
        import pandas as pd
        return pd.read_excel(path).to_string(max_rows=200)
    return Path(path).read_text(encoding="utf-8")[:8000]

def compare_docs(doc_a: str, doc_b: str, instruction: str = "") -> CrossDocReport:
    """vLLM 驱动跨文档对比。返回差异清单。"""
    ta = _read_doc(doc_a)[:5000]
    tb = _read_doc(doc_b)[:5000]
    result = CrossDocReport(doc_a=doc_a, doc_b=doc_b)
    if not ta or not tb:
        result.items = [Discrepancy(severity="WARNING", category="格式差异",
                        loc_a="文档A", val_a="无法提取" if not ta else "已提取",
                        loc_b="文档B", val_b="无法提取" if not tb else "已提取")]
        result.total = 1; return result

    prompt = f"""你是审计质控专家。对比两份文档，找出所有不一致。

【指令】{instruction or '检查数字勾稽/措辞矛盾/遗漏披露'}

【文档A】
{ta[:4000]}

【文档B】
{tb[:4000]}

输出JSON: {{"items":[{{"severity":"CRITICAL|WARNING|INFO","category":"数字不一致|措辞矛盾|遗漏披露|格式差异","loc_a":"...","val_a":"...","loc_b":"...","val_b":"...","suggestion":"...","evidence":"..."}}]}}
两文档一致时返回空数组。只输出JSON。"""

    try:
        import httpx
        r = httpx.post(
            os.environ.get("VLLM_TUNNEL_URL", "http://localhost:18000/v1/chat/completions"),
            headers={"Authorization": "Bearer EMPTY"},
            json={"model": os.environ.get("VLLM_MODEL", "qwen3-235b"),
                  "messages": [{"role": "user", "content": prompt}],
                  "temperature": 0.2, "max_tokens": 2048}, timeout=60)
        r.raise_for_status()
        text = r.json()["choices"][0]["message"]["content"].strip()
        m = re.search(r'\{[\s\S]*\}', text)
        if m:
            data = json.loads(m.group(0))
            for item in data.get("items", []):
                result.items.append(Discrepancy(**{k: v for k, v in item.items()
                    if k in Discrepancy.__dataclass_fields__}))
    except Exception as e:
        # 降级：确定性数字对比
        na = set(re.findall(r'\d[\d,.]*\d', ta))
        nb = set(re.findall(r'\d[\d,.]*\d', tb))
        if na - nb:
            result.items.append(Discrepancy(severity="WARNING", category="数字不一致",
                loc_a="文档A独有", val_a=", ".join(sorted(na - nb)[:8]),
                suggestion="检查文档B是否遗漏"))
        if nb - na:
            result.items.append(Discrepancy(severity="WARNING", category="数字不一致",
                loc_b="文档B独有", val_b=", ".join(sorted(nb - na)[:8]),
                suggestion="检查文档A是否遗漏"))

    result.total = len(result.items)
    result.critical = sum(1 for d in result.items if d.severity == "CRITICAL")
    return result

def cross_check_report_vs_workpaper(report: str, workpaper: str) -> CrossDocReport:
    return compare_docs(report, workpaper, "检查报告数字与底稿审定数是否一致，关注金额和日期差异")

def cross_check_year_over_year(this: str, last: str) -> CrossDocReport:
    return compare_docs(this, last, "对比本年与上年关键数字变动和科目增减")
