#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
docxtpl 报告生成桥接器
读取模板 → 注入数据 → 输出 Word 报告
"""

from pathlib import Path
from typing import Dict, Any

TEMPLATE_DIR = Path(__file__).resolve().parent.parent / "data" / "templates"


class DocxTemplateBridge:
    """基于 docxtpl 的报告生成器"""

    def __init__(self, template_name: str = "3.2（通用）专项审计报告模板.docx"):
        self.template_path = TEMPLATE_DIR / "word" / template_name
        if not self.template_path.exists():
            self.template_path = TEMPLATE_DIR / template_name  # fallback

    def render(self, context: Dict[str, Any], output_path: str) -> str:
        """渲染模板并保存"""
        from docx import Document
        from datetime import datetime

        if not self.template_path.exists():
            # 无模板时生成纯文本报告
            return self._render_fallback(context, output_path)

        doc = Document(str(self.template_path))
        self._replace_in_doc(doc, context)
        doc.save(output_path)
        return output_path

    def _replace_in_doc(self, doc, context: Dict[str, Any]):
        """在文档中替换占位符"""
        import re

        # 上下文预处理
        now = context.get("date", "")
        if not now:
            now = self._current_date()

        replacements = {
            "安徽XX会计师事务所": context.get("firm_name", "安徽XX会计师事务所"),
            "皖XX专审字[2026]XXX号": context.get("report_no", f"皖{context.get('report_prefix','XX')}专审字[2026]{context.get('report_suffix','XXX')}号"),
            "安徽XXX有限公司": context.get("company_name", "安徽XXX有限公司"),
            "XXX等XX个项目": context.get("project_desc", "XXX等项目"),
            "安徽宣城XXXXX管理委员会": context.get("commission", "委托人"),
            "二〇二六年XX月XX日": now,
            "×××": "***",  # 脱敏
        }

        # 替换段落
        for p in doc.paragraphs:
            for old, new in replacements.items():
                if old in p.text:
                    for run in p.runs:
                        if old in run.text:
                            run.text = run.text.replace(old, str(new))

        # 替换表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for old, new in replacements.items():
                        if old in cell.text:
                            for p in cell.paragraphs:
                                for run in p.runs:
                                    if old in run.text:
                                        run.text = run.text.replace(old, str(new))

    def _render_fallback(self, context: Dict[str, Any], output_path: str) -> str:
        """无模板时的纯文本报告"""
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)

        t = doc.add_heading("专项审计报告", level=0)
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph(f"报告编号：{context.get('report_no', '')}")
        doc.add_paragraph(f"被审计单位：{context.get('company_name', '')}")
        doc.add_paragraph(f"审计期间：{context.get('period', '')}")
        doc.add_paragraph("")
        doc.add_paragraph(context.get("conclusion", ""))
        doc.add_paragraph(context.get("data_summary", ""))
        doc.save(output_path)
        return output_path

    @staticmethod
    def _current_date() -> str:
        """当前日期转中文大写"""
        from datetime import datetime
        d = datetime.now()
        nums = "〇一二三四五六七八九"
        year = "".join(nums[int(c)] for c in str(d.year))
        month = f"{nums[d.month] if d.month <= 9 else '十' + (nums[d.month-10] if d.month>10 else '') }"
        day_map = {1: "一", 2: "二", 3: "三", 4: "四", 5: "五", 6: "六", 7: "七",
                   8: "八", 9: "九", 10: "十", 11: "十一", 12: "十二", 13: "十三",
                   14: "十四", 15: "十五", 16: "十六", 17: "十七", 18: "十八",
                   19: "十九", 20: "二十", 21: "二十一", 22: "二十二", 23: "二十三",
                   24: "二十四", 25: "二十五", 26: "二十六", 27: "二十七",
                   28: "二十八", 29: "二十九", 30: "三十", 31: "三十一"}
        return f"二〇{year}年{month.strip()}月{day_map[d.day]}日"


def merge_match_data_to_report(match_result: dict, context: dict) -> dict:
    """将匹配引擎结果合并到报告上下文"""
    stats = match_result.get("match_stats", {})
    diff = match_result.get("diff_summary", [])
    strategy = match_result.get("strategy_name", "")

    total_bank_rows = stats.get("total_bank_rows", 0)
    total_summary_institutions = stats.get("total_summary_institutions", 0)
    total_filtered_amount = stats.get("total_filtered_amount", 0)
    total_summary_amount = stats.get("total_summary_amount", 0)
    no_data = (
        (total_bank_rows == 0 and total_summary_institutions == 0)
        or (total_filtered_amount == 0 and total_summary_amount == 0)
    )

    summary_lines = [
        f"匹配策略：{strategy}",
        f"银行流水总行数：{total_bank_rows}",
        f"筛选命中行数：{stats.get('filtered_rows', 0)}",
        f"汇总机构数：{total_summary_institutions}",
        f"匹配成功机构数：{stats.get('matched_institutions', 0)}",
        f"匹配率：{stats.get('match_rate', 0):.1f}%",
        f"筛选总金额：{total_filtered_amount:,.2f}",
        f"回款表总金额：{total_summary_amount:,.2f}",
        f"差额：{stats.get('total_difference', 0):,.2f}",
        f"差额比例：{stats.get('diff_percentage', 0):.1f}%",
    ]

    diff_lines = ["\n各机构匹配明细："]
    for d in diff[:20]:
        diff_lines.append(f"  {d.get('机构','?')}: 筛选{d.get('筛选金额',0):,.2f} | 回款表{d.get('回款表金额',0):,.2f} | 差额{d.get('差额',0):,.2f}")

    context["data_summary"] = "\n".join(summary_lines)
    if no_data:
        context["conclusion"] = (
            "审计结论：当前未获取到有效匹配数据，无法形成审计结论。"
            "请核对输入数据（银行流水、回款汇总表）及筛选范围后重新执行。"
        )
    else:
        context["conclusion"] = f"审计结论：匹配率 {stats.get('match_rate',0):.1f}%，差额 {stats.get('total_difference',0):,.2f} 元"
    context["diff_detail"] = "\n".join(diff_lines)
    return context
