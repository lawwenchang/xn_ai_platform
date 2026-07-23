#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""document_loader 多格式加载测试：xlsx/csv/docx/pdf/md/txt 全覆盖"""
import sys
import tempfile
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parent.parent))
import pandas as pd
from core.document_loader import (load_document, load_tables, sniff_document,
                                  materialize_document, SUPPORTED_EXTS)

ok = lambda msg: print(f"  [OK] {msg}")
TMP = Path(tempfile.mkdtemp(prefix="docloader_test_"))

# 1) xlsx
fp = TMP / "流水.xlsx"
pd.DataFrame({"日期": ["2026-01-05"], "金额": [100.5]}).to_excel(fp, index=False)
d = load_document(fp)
assert d.kind == "table" and len(d.tables) == 1 and d.tables[0].iloc[0]["金额"] == 100.5
ok("xlsx 表格加载")

# 2) csv（GBK 编码容错）
fp = TMP / "台账.csv"
fp.write_text("客户,金额\n甲公司,200\n", encoding="gbk")
d = load_document(fp)
assert len(d.tables) == 1 and d.tables[0].iloc[0]["客户"] == "甲公司"
ok("csv GBK 编码加载")

# 3) docx（段落 + 表格）
from docx import Document
doc = Document()
doc.add_paragraph("审计报告正文第一段。")
t = doc.add_table(rows=3, cols=2)
t.rows[0].cells[0].text, t.rows[0].cells[1].text = "项目", "金额"
t.rows[1].cells[0].text, t.rows[1].cells[1].text = "差异A", "1000"
t.rows[2].cells[0].text, t.rows[2].cells[1].text = "差异B", "2000"
fp = TMP / "报告.docx"
doc.save(fp)
d = load_document(fp)
assert "审计报告正文" in d.text, "docx 文本缺失"
assert len(d.tables) == 1 and list(d.tables[0].columns) == ["项目", "金额"]
assert d.tables[0].iloc[1]["金额"] == "2000"
ok("docx 段落+表格加载")

# 4) pdf（fitz 生成带文本的 PDF）
import fitz
pdf = fitz.open()
page = pdf.new_page()
page.insert_text((72, 72), "银行询证函回函 账号622848 余额相符", fontsize=12)
fp = TMP / "回函.pdf"
pdf.save(str(fp))
pdf.close()
d = load_document(fp)
assert "622848" in d.text, f"pdf 文本缺失: {d.errors}"
ok("pdf 文本加载")

# 5) md（含 pipe 表格）
fp = TMP / "底稿.md"
fp.write_text("# 审计底稿\n\n说明文字。\n\n| 科目 | 余额 |\n|---|---|\n| 银行存款 | 5000 |\n| 应收 | 300 |\n",
              encoding="utf-8")
d = load_document(fp)
assert "审计底稿" in d.text and len(d.tables) == 1
assert d.tables[0].iloc[0]["余额"] == "5000"
ok("md 文本+pipe表格加载")

# 6) txt
fp = TMP / "说明.txt"
fp.write_text("纯文本内容测试", encoding="utf-8")
d = load_document(fp)
assert d.kind == "document" and "纯文本" in d.text
ok("txt 加载")

# 7) sniff + materialize
info = sniff_document(TMP / "报告.docx")
assert info["kind"] in ("mixed", "document") and "text_preview" in info
ok("sniff_document 画像")
out = TMP / "mat"
gen = materialize_document(TMP / "报告.docx", out)
names = [p.name for p in gen]
assert any(n.endswith(".xlsx") for n in names) and any(n.endswith(".txt") for n in names), names
# 物化后的 xlsx 可回读
back = load_tables(out / [n for n in names if n.endswith(".xlsx")][0])
assert back and list(back[0].columns) == ["项目", "金额"]
ok("materialize_document 物化+回读")

# 8) 不存在文件与不支持格式的错误路径
d = load_document(TMP / "不存在.xlsx")
assert d.errors and not d.ok
bad = TMP / "x.xyz"
bad.write_text("whatever", encoding="utf-8")
d2 = load_document(bad)
assert any("不支持" in e for e in d2.errors), d2.errors
ok("错误路径清晰降级")

print(f"\n全部通过：document_loader 支持 {sorted(SUPPORTED_EXTS)}")
