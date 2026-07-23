#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
新增模块集成冒烟测试 (smoke_test_new_modules.py)
================================================
覆盖：agent / toolbox / md_engine / template_engine / code_corrector（新）
特点：全程离线（VLLM 指向不可达地址，验证所有降级路径），不碰 AutoDL。

运行：cd phase2_v3 && python tests/smoke_test_new_modules.py
"""
import os
import sys
import json
import shutil
import py_compile
import traceback
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

# 强制 LLM 离线（微调期间 vLLM 不可用，验证所有本地降级路径）
os.environ["VLLM_TUNNEL_URL"] = "http://127.0.0.1:9/v1/chat/completions"

TMP = ROOT / "data" / "smoke_tmp"          # toolbox 沙箱内的测试目录
REL = "smoke_tmp"                          # 相对沙箱根的路径前缀

RESULTS = []

def check(name, fn):
    try:
        fn()
        RESULTS.append((name, "PASS", ""))
        print(f"  [PASS] {name}")
    except ImportError as e:
        RESULTS.append((name, "SKIP", str(e)))
        print(f"  [SKIP] {name}: {e}")
    except Exception as e:
        RESULTS.append((name, "FAIL", f"{type(e).__name__}: {e}"))
        print(f"  [FAIL] {name}: {type(e).__name__}: {e}")
        traceback.print_exc(limit=2)

SAMPLE_MD = """---
author: 质控部
---
# 医保回款核对底稿

## 审计结论

经核对，差异 32,000.00 元，在容差范围内。

## 差异明细

| 日期 | 金额 | 状态 |
|------|------|------|
| 2026-01-05 | 12,000.00 | 未达账 |
| 2026-01-18 | 20,000.00 | 未达账 |

- 第一条说明
- 第二条说明

```python
print('code block')
```
"""

# ═══════════ 1. md_engine ═══════════

def t_parse_md():
    from core.md_engine import parse_md_text
    d = parse_md_text(SAMPLE_MD, "sample.md")
    assert d.title == "医保回款核对底稿", d.title
    assert d.front_matter.get("author") == "质控部"
    assert len(d.headings()) == 3
    assert len(d.tables()) == 1
    types = {n.type for n in d.nodes}
    assert {"heading", "table", "list", "code", "paragraph"} <= types, types

def t_fill_template_text():
    from core.md_engine import fill_template_text
    r = fill_template_text("金额：{{ amount }} / {{unit}}", {"amount": "5万", "unit": "元"})
    assert r == "金额：5万 / 元", r

def t_md_to_docx():
    from docx import Document  # noqa 依赖检查
    from core.md_engine import md_to_docx
    src = TMP / "t.md"
    src.write_text(SAMPLE_MD, encoding="utf-8")
    out = md_to_docx(str(src))
    assert Path(out).exists() and Path(out).stat().st_size > 0
    doc = Document(out)
    texts = [p.text for p in doc.paragraphs]
    assert any("医保回款核对底稿" in t for t in texts)
    assert len(doc.tables) == 1

def t_generate_report_fallback():
    from core.md_engine import generate_report
    r = generate_report("生成医保核对报告", {"差异总额": 32000, "结论": "在容差内"})
    assert "32000" in r and "审计报告" in r, r[:200]

def t_rewrite_fallback():
    from core.md_engine import rewrite_text
    r = rewrite_text("原文内容", "改得更正式")
    assert r == "原文内容"  # LLM 离线时应原样返回

# ═══════════ 2. template_engine ═══════════

def t_fill_docx():
    from docx import Document
    from core.template_engine import fill_docx
    tp = TMP / "tmpl.docx"
    doc = Document()
    doc.add_paragraph("科目：{{ subject }}")
    tbl = doc.add_table(rows=1, cols=2)
    tbl.cell(0, 0).text = "{{ amount }}"
    tbl.cell(0, 1).text = "{{nested.value}}"
    doc.save(str(tp))
    out = fill_docx(str(tp), {"subject": "应收账款", "amount": "1,250,000.00",
                              "nested": {"value": "OK"}}, str(TMP / "tmpl_out.docx"))
    d2 = Document(out)
    all_text = "\n".join(p.text for p in d2.paragraphs)
    tbl_text = "\n".join(c.text for t in d2.tables for r in t.rows for c in r.cells)
    assert "应收账款" in all_text, all_text
    assert "1,250,000.00" in tbl_text and "OK" in tbl_text, tbl_text

def t_fill_xlsx():
    from openpyxl import Workbook, load_workbook
    from openpyxl.workbook.defined_name import DefinedName
    from core.template_engine import fill_xlsx
    tp = TMP / "tmpl.xlsx"
    wb = Workbook()
    ws = wb.active
    ws.title = "S1"
    ws["A1"] = "余额：{{ balance }}"
    ws["B2"] = "占位"
    wb.defined_names.add(DefinedName("audited_total", attr_text="S1!$B$2"))
    wb.save(str(tp))
    out = fill_xlsx(str(tp), {"balance": "88.88", "audited_total": 999.99},
                    str(TMP / "tmpl_out.xlsx"))
    wb2 = load_workbook(out)
    assert wb2["S1"]["A1"].value == "余额：88.88", wb2["S1"]["A1"].value
    assert wb2["S1"]["B2"].value == 999.99, wb2["S1"]["B2"].value  # 命名区域写入


# ═══════════ 3. toolbox（11 工具中 LLM 无关的 8 个） ═══════════

def t_toolbox_read_and_parse():
    from core.toolbox import execute_tool
    (TMP / "doc.md").write_text(SAMPLE_MD, encoding="utf-8")
    r = execute_tool("read_file", {"path": f"{REL}/doc.md"})
    assert r["ok"] and "医保回款" in r["data"], r
    r = execute_tool("parse_md_structure", {"path": f"{REL}/doc.md"})
    assert r["ok"] and r["data"]["title"] == "医保回款核对底稿", r

def t_toolbox_extract_section():
    from core.toolbox import execute_tool
    r = execute_tool("extract_section", {"path": f"{REL}/doc.md", "heading": "审计结论"})
    assert r["ok"] and "32,000.00" in r["data"], r
    r2 = execute_tool("extract_section", {"path": f"{REL}/doc.md", "heading": "不存在的章节"})
    assert not r2["ok"], r2

def t_toolbox_merge_diff():
    from core.toolbox import execute_tool
    (TMP / "a.md").write_text("# A\n\n第一行\n共同行", encoding="utf-8")
    (TMP / "b.md").write_text("# B\n\n第一行改了\n共同行", encoding="utf-8")
    r = execute_tool("merge_documents", {"paths_json": json.dumps([f"{REL}/a.md", f"{REL}/b.md"]), "title": "冒烟合并"})
    assert r["ok"], r
    r = execute_tool("diff_documents", {"path1": f"{REL}/a.md", "path2": f"{REL}/b.md"})
    assert r["ok"] and "行1" in r["data"], r

def t_toolbox_convert():
    import docx  # noqa
    from core.toolbox import execute_tool
    r = execute_tool("convert_format", {"path": f"{REL}/doc.md", "target_format": "docx"})
    assert r["ok"] and Path(r["output_path"]).exists(), r

def t_toolbox_fill_template():
    from docx import Document
    from core.toolbox import execute_tool
    tp = TMP / "tb_tmpl.docx"
    d = Document(); d.add_paragraph("客户：{{ client }}"); d.save(str(tp))
    r = execute_tool("fill_template", {"template_path": f"{REL}/tb_tmpl.docx",
                                       "data_json": json.dumps({"client": "XX医院"})})
    assert r["ok"], r

def t_toolbox_security():
    from core.toolbox import execute_tool
    r = execute_tool("read_file", {"path": "../../白皮书越权读取.md"})
    assert not r["ok"], r
    r2 = execute_tool("不存在的工具", {})
    assert not r2["ok"] and "未知工具" in r2["error"], r2

# ═══════════ 4. agent（离线只验证 schema 与导入） ═══════════

def t_agent_schema():
    from core.toolbox import TOOLS
    import core.agent as agent
    assert callable(agent.agent_run)
    names = [t["function"]["name"] for t in TOOLS]
    assert len(names) == len(set(names)), "工具名重复"
    assert len(names) == 11, f"应为11个工具，实际{len(names)}"
    for t in TOOLS:
        f = t["function"]
        assert f["description"] and f["parameters"]["type"] == "object"
        for req in f["parameters"].get("required", []):
            assert req in f["parameters"]["properties"], f"{f['name']} 缺参数定义 {req}"


# ═══════════ 5. code_corrector（新自纠错内核） ═══════════

def t_fix_name_error():
    from engine.code_corrector import rule_based_fix
    code = "df = pd.DataFrame({'a': [1]})\nprint(df)"
    err = "Traceback ...\nNameError: name 'pd' is not defined"
    fixed = rule_based_fix(code, err)
    assert fixed and fixed.startswith("import pandas as pd"), fixed
    exec(compile(fixed, "<t>", "exec"), {})  # 修正后可真实执行

def t_fix_outputs_dir():
    from engine.code_corrector import rule_based_fix
    code = "open('outputs/r.csv', 'w').write('x')"
    err = "FileNotFoundError: [Errno 2] No such file or directory: 'outputs/r.csv'"
    fixed = rule_based_fix(code, err)
    assert fixed and "makedirs('outputs'" in fixed, fixed

def t_fix_unicode_csv():
    from engine.code_corrector import rule_based_fix
    code = "import pandas as pd\ndf = pd.read_csv('流水.csv')"
    err = "UnicodeDecodeError: 'utf-8' codec can't decode byte 0xd6"
    fixed = rule_based_fix(code, err)
    assert fixed and "_robust_read_csv('流水.csv')" in fixed, fixed

def t_fix_keyerror_columns():
    from engine.code_corrector import rule_based_fix
    code = "import pandas as pd\ndf = pd.read_excel('台账.xlsx')\nprint(df['金额'])"
    err = "KeyError: '金额'"
    fixed = rule_based_fix(code, err)
    assert fixed and "str(c).strip()" in fixed, fixed

def t_llm_fix_offline():
    from engine.code_corrector import llm_fix, correct_code
    assert llm_fix("x=1/0", "ZeroDivisionError") is None  # 离线应优雅跳过
    assert correct_code("x=1/0", "ZeroDivisionError: division by zero") is None

def t_corrector_e2e_subprocess():
    """端到端：故意坏代码 → 规则修正 → 子进程重跑成功（模拟 routes.py 闭环）"""
    import subprocess
    from engine.code_corrector import correct_code
    code = "df = pd.DataFrame({'a': [1, 2]})\nprint('SUM=' + str(df['a'].sum()))"
    for attempt in range(1, 4):
        sp = TMP / f"e2e_{attempt}.py"
        sp.write_text(code, encoding="utf-8")
        p = subprocess.run([sys.executable, str(sp)], capture_output=True, text=True, timeout=60)
        if p.returncode == 0:
            assert "SUM=3" in p.stdout, p.stdout
            assert attempt == 2, f"应在第2轮成功，实际第{attempt}轮"
            return
        fixed = correct_code(code, p.stderr, attempt=attempt)
        assert fixed, f"第{attempt}轮无修正方案: {p.stderr[-300:]}"
        code = fixed
    raise AssertionError("3 轮内未自纠成功")

# ═══════════ 6. 接线文件语法完整性 ═══════════

def t_compile_wired_files():
    for f in ["api/routes.py", "engine/sandbox_v3.py", "engine/code_corrector.py",
              "core/agent.py", "core/toolbox.py", "core/md_engine.py", "core/template_engine.py"]:
        py_compile.compile(str(ROOT / f), doraise=True)

# ═══════════ 主流程 ═══════════

if __name__ == "__main__":
    if TMP.exists():
        shutil.rmtree(TMP, ignore_errors=True)
    TMP.mkdir(parents=True, exist_ok=True)
    print("=" * 60)
    print("新增模块集成冒烟测试（LLM 离线模式）")
    print("=" * 60)
    for group, tests in [
        ("md_engine", [t_parse_md, t_fill_template_text, t_md_to_docx, t_generate_report_fallback, t_rewrite_fallback]),
        ("template_engine", [t_fill_docx, t_fill_xlsx]),
        ("toolbox", [t_toolbox_read_and_parse, t_toolbox_extract_section, t_toolbox_merge_diff,
                     t_toolbox_convert, t_toolbox_fill_template, t_toolbox_security]),
        ("agent", [t_agent_schema]),
        ("code_corrector", [t_fix_name_error, t_fix_outputs_dir, t_fix_unicode_csv,
                            t_fix_keyerror_columns, t_llm_fix_offline, t_corrector_e2e_subprocess]),
        ("接线完整性", [t_compile_wired_files]),
    ]:
        print(f"\n[{group}]")
        for t in tests:
            check(t.__name__, t)
    n_pass = sum(1 for _, s, _ in RESULTS if s == "PASS")
    n_fail = sum(1 for _, s, _ in RESULTS if s == "FAIL")
    n_skip = sum(1 for _, s, _ in RESULTS if s == "SKIP")
    print("\n" + "=" * 60)
    print(f"结果: {n_pass} 通过 / {n_fail} 失败 / {n_skip} 跳过（共 {len(RESULTS)}）")
    shutil.rmtree(TMP, ignore_errors=True)
    sys.exit(1 if n_fail else 0)
