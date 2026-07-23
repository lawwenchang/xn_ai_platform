#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Backlog 修复回归测试
覆盖：
  - chunk_by_headings 异常安全
  - fill_docx 跨 Run 占位符替换
  - md_to_docx 中文字体 & 表格显式边框
  - _generate_from_scratch 表格解析 & 字体
  - report_generator 无匹配数据质量判定
  - fill_xlsx 命名区域兼容（含 worksheet 级/特殊 sheet 名）
"""
import os
import re
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

# 强制 LLM 离线，测试降级路径
os.environ["VLLM_TUNNEL_URL"] = "http://127.0.0.1:9/v1/chat/completions"

TMP = ROOT / "data" / "backlog_tmp"


def _clear_tmp():
    import shutil
    if TMP.exists():
        shutil.rmtree(TMP, ignore_errors=True)
    TMP.mkdir(parents=True, exist_ok=True)


def test_chunk_by_headings_safe_and_complete():
    """P0: chunk_by_headings 不应崩溃，且 chunk 包含标题。"""
    from core.md_engine import chunk_by_headings

    p = TMP / "sample.md"
    p.write_text("# 总则\n\n正文 A。\n## 分则\n\n正文 B。\n", encoding="utf-8")
    chunks = chunk_by_headings(str(p))
    assert chunks, chunks
    headings = {c["heading"] for c in chunks}
    assert "总则" in headings, headings
    assert "分则" in headings, headings
    assert all("正文" in c["text"] for c in chunks), chunks

    # 异常输入不崩溃
    assert chunk_by_headings(str(TMP / "not_exist.md")) == []


def test_fill_docx_cross_run_placeholder():
    """P1: fill_docx 能替换被拆到多个 Run 的占位符。"""
    from docx import Document
    from core.template_engine import fill_docx

    tp = TMP / "cross_run.docx"
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("科")
    p.add_run("{{ sub")
    p.add_run("ject }}")
    p.add_run(" 金额")
    doc.save(str(tp))

    out = fill_docx(str(tp), {"subject": "应收账款"}, str(TMP / "cross_run_out.docx"))
    d2 = Document(out)
    full = "\n".join(p.text for p in d2.paragraphs)
    assert "{{" not in full, full
    assert "应收账款" in full, full


def test_md_to_docx_chinese_font_and_table_borders():
    """P1: md_to_docx 设置中文字体且表格带显式边框。"""
    from core.md_engine import md_to_docx
    from docx import Document
    from docx.oxml.ns import qn

    src = TMP / "md_to_docx.md"
    src.write_text("# 标题\n\n正文段落。\n\n| a | b |\n|---|---|\n| 1 | 2 |\n", encoding="utf-8")
    out = md_to_docx(str(src), str(TMP / "md_to_docx_out.docx"))
    doc = Document(out)

    # 中文字体
    normal = doc.styles["Normal"]
    assert normal.font.name == "宋体", normal.font.name
    east = normal._element.rPr.rFonts.get(qn("w:eastAsia"))
    assert east == "宋体", east

    # 表格显式边框
    assert len(doc.tables) == 1
    tbl = doc.tables[0]._tbl
    tblPr = tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    assert borders is not None, "表格缺少显式边框"


def test_generate_from_scratch_table():
    """P0: _generate_from_scratch 正确解析 [表格开始]...[表格结束]。"""
    from core import template_engine as te

    # mock LLM 返回带表格标记的文本
    te._call_llm = lambda *a, **k: (
        "[标题一]审计结果[标题一结束]"
        "[正文]经核对，差异如下。[正文结束]"
        "[表格开始]| 项目 | 金额 |\n|---|---|\n| A | 1,000.00 |\n| B | 2,000.00 |[表格结束]"
    )

    out = TMP / "from_scratch.docx"
    te._generate_from_scratch(
        {"amount": 3000}, "生成测试报告", "", str(out)
    )
    from docx import Document
    doc = Document(str(out))
    full = "\n".join(p.text for p in doc.paragraphs)
    assert "审计结果" in full, full
    assert "差异如下" in full, full
    assert len(doc.tables) == 1, doc.tables
    tbl_text = " ".join(" ".join(cell.text for cell in row.cells) for row in doc.tables[0].rows)
    assert "项目" in tbl_text and "A" in tbl_text, tbl_text

    # 表格带边框
    from docx.oxml.ns import qn
    borders = doc.tables[0]._tbl.tblPr.find(qn("w:tblBorders"))
    assert borders is not None


def test_quality_judgment_no_data():
    """P2: 无匹配数据时质量判定不误导。"""
    from docx import Document
    from core.report_generator import _add_quality_judgment

    doc = Document()
    _add_quality_judgment(doc, {})
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "无法形成匹配效果判定" in text or "缺少某一侧" in text, text
    assert "匹配效果很差" not in text, text


def test_fill_xlsx_named_ranges_compatibility():
    """P2: fill_xlsx 兼容含空格 sheet 名的命名区域及多单元格区域引用。"""
    from openpyxl import Workbook, load_workbook
    from openpyxl.workbook.defined_name import DefinedName
    from core.template_engine import fill_xlsx

    tp = TMP / "named.xlsx"
    wb = Workbook()
    ws1 = wb.active
    ws1.title = "Sheet 1"
    ws1["A1"] = "{{balance}}"
    ws1["B2"] = ""
    ws1["C1"] = ""
    ws2 = wb.create_sheet("数据表")
    ws2["A1"] = ""

    # workbook 级命名区域：含空格 sheet 名单元格
    wb.defined_names.add(DefinedName("balance", attr_text="'Sheet 1'!$B$2"))
    # workbook 级命名区域：多单元格区域（应写入左上角）
    wb.defined_names.add(DefinedName("date_range", attr_text="'Sheet 1'!$C$1:$C$5"))
    # workbook 级命名区域：普通 sheet 名单元格
    wb.defined_names.add(DefinedName("audit_date", attr_text="数据表!$A$1"))
    wb.save(str(tp))

    out = fill_xlsx(
        str(tp),
        {"balance": 1234.56, "date_range": "2026-07-21", "audit_date": "2026-07-22"},
        str(TMP / "named_out.xlsx"),
    )
    wb2 = load_workbook(out)
    assert wb2["Sheet 1"]["B2"].value == 1234.56, wb2["Sheet 1"]["B2"].value
    assert wb2["Sheet 1"]["C1"].value == "2026-07-21", wb2["Sheet 1"]["C1"].value
    assert wb2["数据表"]["A1"].value == "2026-07-22", wb2["数据表"]["A1"].value
    # 占位符单元格也被替换（整格为占位符且数据为数值时保留数值类型）
    assert wb2["Sheet 1"]["A1"].value == 1234.56, wb2["Sheet 1"]["A1"].value


if __name__ == "__main__":
    _clear_tmp()
    tests = [
        test_chunk_by_headings_safe_and_complete,
        test_fill_docx_cross_run_placeholder,
        test_md_to_docx_chinese_font_and_table_borders,
        test_generate_from_scratch_table,
        test_quality_judgment_no_data,
        test_fill_xlsx_named_ranges_compatibility,
    ]
    ok = 0
    for t in tests:
        try:
            t()
            print(f"[PASS] {t.__name__}")
            ok += 1
        except Exception as e:
            print(f"[FAIL] {t.__name__}: {type(e).__name__}: {e}")
            import traceback
            traceback.print_exc()
    print(f"\n结果: {ok}/{len(tests)} 通过")
    import shutil
    shutil.rmtree(TMP, ignore_errors=True)
    sys.exit(0 if ok == len(tests) else 1)
