#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""代码生成端到端测试：_dag_to_python → 子进程执行生成代码 → 校验行为

覆盖：Load 按名匹配 / docx 文档加载 / Merge left_on+日期窗口 /
Diff 按键对齐+容差 / Reconcile 方向镜像 / 序号键拒绝合并 / 缺失不填充
"""
import subprocess
import sys
import tempfile
import types
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))
ok = lambda msg: print(f"  [OK] {msg}")


def _stub(name, **attrs):
    m = types.ModuleType(name)
    for k, v in attrs.items():
        setattr(m, k, v)
    sys.modules[name] = m


# 重依赖打桩（只为导入 api.routes 取 _dag_to_python）
_stub("core.rag_engine", inject_compliance_context=lambda *a, **k: "",
      build_index=lambda *a, **k: None)
_stub("core.run_snapshot", AssetCatalog=object, RunRecord=object,
      RunSnapshotManager=object)
_stub("core.chaos_input", ChaosInputProcessor=object)
_stub("core.privacy_firewall",
      get_firewall=lambda: types.SimpleNamespace(sanitize=lambda t: (t, [])))
_stub("core.pipeline_trace", record=lambda *a, **k: None)
_f = lambda *a, **k: None
_stub("core.format_engine", normalize_format=_f, extract_word_format=_f,
      extract_excel_format=_f, apply_word_format=_f, apply_excel_format=_f,
      extract_word_print=_f, extract_excel_print=_f, apply_word_print=_f,
      apply_excel_print=_f)
_stub("core.template_manager", get_available_templates=list,
      get_rules=lambda *a, **k: [], TEMPLATES_DIR=Path("data/templates"))
_stub("config.fallback_prompts", detect_scenario=lambda *a: "通用",
      get_fallback_prompt=lambda *a: "")
_stub("engine.sandbox_v3", EphemeralSandbox=object, LifecycleHooks=object,
      LifecycleResult=object)

from api.routes import _dag_to_python  # noqa: E402
import pandas as pd  # noqa: E402

# ── 准备运行目录与样本数据 ─────────────────────────────────────
RUN = Path(tempfile.mkdtemp(prefix="e2e_codegen_"))
(RUN / "inputs").mkdir(parents=True)
(RUN / "outputs").mkdir()

pd.DataFrame({
    "序号": [1, 2, 3, 4], "月": [3] * 4,
    "日期": ["2026-03-01", "2026-03-02", "2026-03-05", "2026-03-28"],
    "凭证号": ["记-1", "记-2", "记-3", "记-4"],
    "摘要": ["收甲", "付乙", "收丙", "收丁"],
    "借方金额": [10000, 0, 8000, 6000],
    "贷方金额": [0, 5000, 0, 0],
    "银行账号": ["农行5927"] * 4,
}).to_excel(RUN / "inputs" / "序时账.xlsx", index=False)

pd.DataFrame({
    "序号": [1, 2, 3, 4, 5],
    "银行账号": ["农行5927"] * 5,
    "日期": ["2026-03-01", "2026-03-03", "2026-03-05", "2026-03-05", "2026-03-29"],
    " 摘要": ["收甲", "付乙", "收丙", "收丙", "利息"],
    "借方（支取）": [0, 5000, 0, 0, 0],
    "贷方（收入）": [10000, 0, 5000, 3000, 88.8],
}).to_excel(RUN / "inputs" / "银行流水.xlsx", index=False)

from docx import Document
_doc = Document()
_doc.add_paragraph("合同台账")
_t = _doc.add_table(rows=2, cols=2)
_t.rows[0].cells[0].text, _t.rows[0].cells[1].text = "客户", "金额"
_t.rows[1].cells[0].text, _t.rows[1].cells[1].text = "甲公司", "10000"
_doc.save(RUN / "inputs" / "合同台账.docx")
ok("样本数据就绪（xlsx×2 + docx×1）")


# ── DAG 1：全流程（Load×3 + Reconcile + Merge(日期窗口) + Diff + Export） ──
dag = {
    "objective": "E2E对账", "raw_intent": "序时账和银行流水逐笔对账",
    "operators": [
        {"id": "op_1", "name": "Load", "source_file": "序时账.xlsx",
         "output_alias": "df_j", "params": {}},
        {"id": "op_2", "name": "Load", "source_file": "银行流水.xlsx",
         "output_alias": "df_b", "params": {}},
        {"id": "op_3", "name": "Load", "source_file": "合同台账.docx",
         "output_alias": "df_doc", "params": {}},
        {"id": "op_4", "name": "Reconcile", "input_from": ["op_1", "op_2"],
         "output_alias": "df_rec",
         "params": {"tolerance_abs": 0.01, "date_window_days": 3}},
        {"id": "op_5", "name": "Merge", "input_from": ["op_1", "op_2"],
         "output_alias": "df_m",
         "params": {"how": "inner", "left_on": ["借方金额", "日期"],
                    "right_on": ["贷方（收入）", "日期"], "date_window_days": 3}},
        {"id": "op_6", "name": "Diff", "input_from": ["op_1", "op_2"],
         "output_alias": "df_d",
         "params": {"keys": ["日期"], "col_a": "借方金额", "col_b": "贷方（收入）",
                    "tolerance_abs": 0.01}},
        {"id": "op_7", "name": "Export", "input_from": ["op_6"],
         "output_alias": "df_out", "params": {"output_file": "diff.csv"}},
    ],
}
code = _dag_to_python(dag, None)
assert "ffill" not in code, "生成代码仍含 ffill 伪造逻辑"
assert "_load_any_document" in code and "_reconcile_lite" in code
(RUN / "sandbox_exec.py").write_text(code, encoding="utf-8")
import os as _os
_env = {**_os.environ, "PYTHONIOENCODING": "utf-8"}
r = subprocess.run([sys.executable, "sandbox_exec.py"], cwd=str(RUN),
                   capture_output=True, text=True, encoding="utf-8",
                   errors="replace", env=_env, timeout=120)
out = r.stdout + r.stderr
assert r.returncode == 0, f"生成代码执行失败:\n{out[-3000:]}"
ok("DAG 1 生成代码执行成功")

assert "[Load] 按名匹配: 序时账.xlsx" in out, out[-2000:]
ok("Load 按真实文件名精确匹配")
assert "[Load] 文档 合同台账.docx" in out and "表格" in out
ok("docx 文档表格经 _load_any_document 加载")
assert "[Reconcile]" in out and "已核对" in out
ok("Reconcile 算子真实执行（方向镜像归一）")
assert "[Merge] 日期窗口±3天" in out, out[-2000:]
ok("Merge left_on/right_on + 日期窗口真实生效")
assert "[Diff] 键=" in out and "仅右=" in out
ok("Diff 按键对齐（indicator 三态）")
assert "[数据质量]" in out and "ffill" not in out
ok("防御层只标记不伪造")
assert (RUN / "outputs" / "analysis_result.csv").exists()
assert (RUN / "outputs" / "reconcile_仅右表有.csv").exists()
ok("成果物落盘（分析结果 + 未匹配流水清单）")

# ── DAG 2：Merge 回退拒绝按序号合并 ────────────────────────────
pd.DataFrame({"序号": [1, 2], "特有A": [10, 20]}).to_excel(
    RUN / "inputs" / "甲表.xlsx", index=False)
pd.DataFrame({"序号": [1, 2], "特有B": [30, 40]}).to_excel(
    RUN / "inputs" / "乙表.xlsx", index=False)
dag2 = {
    "objective": "E2E序号", "raw_intent": "两表合并",
    "operators": [
        {"id": "op_1", "name": "Load", "source_file": "甲表.xlsx",
         "output_alias": "df_a", "params": {}},
        {"id": "op_2", "name": "Load", "source_file": "乙表.xlsx",
         "output_alias": "df_b2", "params": {}},
        {"id": "op_3", "name": "Merge", "input_from": ["op_1", "op_2"],
         "output_alias": "df_m2", "params": {"how": "outer"}},
    ],
}
code2 = _dag_to_python(dag2, None)
(RUN / "sandbox_exec2.py").write_text(code2, encoding="utf-8")
r2 = subprocess.run([sys.executable, "sandbox_exec2.py"], cwd=str(RUN),
                    capture_output=True, text=True, encoding="utf-8",
                    errors="replace", env=_env, timeout=120)
out2 = r2.stdout + r2.stderr
assert r2.returncode == 0, out2[-2000:]
assert "拒绝按行号合并" in out2, out2[-2000:]
ok("Merge 无有意义公共键时拒绝按序号合并（修复'对账变乱对'）")

print("\n全部通过：代码生成端到端行为符合设计")
