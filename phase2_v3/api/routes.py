#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
异步生灭路由网关 (routes.py) - 高防御修复版
=============================
"全无状态生命周期与语义编译版"白皮书 §3.3 + §5 的核心实现

核心设计理念：
- 任何新动作 → 毫秒级分发唯一 Run_ID
- 生命周期钩子 → 执行完毕 + 成果物落盘 → 触发容器销毁
- 时序完全解耦 → 前后 Run 互不冲突
- 状态查询 → 通过 Run_ID 查询元数据库
"""

from __future__ import annotations

import asyncio
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import time
import uuid
import zipfile
import sqlite3
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import httpx
from fastapi import (
    APIRouter, BackgroundTasks, FastAPI, File, Form, HTTPException,
    UploadFile, Request
)
from pydantic import BaseModel, Field

import re as _re

# ── 文件上传安全常量 ──────────────────────────────────
MAX_UPLOAD_SIZE = 100 * 1024 * 1024           # 100MB
ALLOWED_EXTENSIONS = {".xlsx", ".xls", ".csv", ".docx", ".doc",
                      ".pdf", ".md", ".txt", ".zip", ".rar", ".7z"}
ALLOWED_MIME_TYPES = {
    "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
    "application/vnd.ms-excel",
    "application/pdf",
    "application/msword",
    "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "text/plain", "text/markdown",
    "text/csv",
    "application/zip",
    "application/x-rar-compressed",
    "application/x-7z-compressed",
    "application/octet-stream",
}
MAX_COMPRESSION_RATIO = 100                   # zip bomb 检测阈值
MAX_DAG_JSON_SIZE = 100 * 1024                # DAG JSON 最大 100KB

# ── 代码安全清洗常量 ──────────────────────────────────
_VALID_OPERATORS_FOR_CODE = {
    "Load", "RegexFilter", "ColumnFilter", "GroupBy", "Merge",
    "Sort", "ConditionCheck", "Extract", "Transform", "NoiseFilter",
    "Aggregate", "Diff", "Export", "Reconcile", "AuditAdjustment",
}
_VALID_COMPARISON_OPS = {">", "<", ">=", "<=", "==", "!="}

# ── 内部导入 ──────────────────────────────────────────
sys.path.append(str(Path(__file__).parent.parent))

from core.chaos_input import ChaosInputProcessor
from core.dag_compiler import DAGBlueprint, DAGParser, Operator, DAGToCodeDescription
from core.run_snapshot import AssetCatalog, RunRecord, RunSnapshotManager
from core.privacy_firewall import get_firewall  # 隐私防火墙
from core.constraint_engine import parse_constraints, format_constraint_report  # 约束引擎
from core.rag_engine import inject_compliance_context, build_index  # RAG
try:
    from core.pipeline_trace import record as trace_record  # B2 全链路可观测
except Exception:  # 埋点模块缺失不影响业务
    def trace_record(*a, **k):
        pass

from core.format_engine import normalize_format
from core.template_manager import TEMPLATES_DIR
from config.fallback_prompts import detect_scenario, get_fallback_prompt
from engine.sandbox_v3 import EphemeralSandbox, LifecycleHooks, LifecycleResult

# ═══════════════════════════════════════════════════════════════
# 性能优化：共享 HTTP 客户端（连接复用 + 连接池）
# ═══════════════════════════════════════════════════════════════
_shared_http_client: Optional[httpx.AsyncClient] = None
_client_loop_id: Optional[int] = None

def _get_http_client() -> httpx.AsyncClient:
    """共享 httpx.AsyncClient（连接池复用；事件循环切换时自动重建，防止 'Event loop is closed'）"""
    global _shared_http_client, _client_loop_id
    try:
        loop_id = id(asyncio.get_running_loop())
    except RuntimeError:
        loop_id = None
    if (_shared_http_client is None or _shared_http_client.is_closed
            or loop_id != _client_loop_id):
        _shared_http_client = httpx.AsyncClient(
            timeout=httpx.Timeout(60.0, connect=10.0),
            limits=httpx.Limits(max_keepalive_connections=10, max_connections=50),
            http2=False,  # 环境未装 h2 包，且 vLLM/Dify 均为 HTTP/1.1，开启反而全链路报错
        )
        _client_loop_id = loop_id
    return _shared_http_client

# 沙箱并发控制信号量（防止 Docker 瞬时过载）
_SANDBOX_SEMAPHORE = asyncio.Semaphore(3)  # 最多 3 个并发容器

# ═══════════════════════════════════════════════════════════════
# 配置常量
# ═══════════════════════════════════════════════════════════════
BASE_DIR = Path(__file__).resolve().parent.parent
DIFY_BASE_URL = os.environ.get("DIFY_BASE_URL", "http://localhost:5001")
DIFY_API_KEY = os.environ.get("DIFY_API_KEY", "app-default-key-replace-me")

# SSH 隧道：AutoDL vLLM → 本地 localhost:18000
VLLM_TUNNEL_URL = "http://localhost:18000/v1/chat/completions"
VLLM_API_KEY = "EMPTY"
VLLM_MODEL = os.environ.get("VLLM_MODEL", "qwen3-235b")  # 可切换: qwen3-235b / qwen2.5-32b / deepseek-r1

# 审计领域知识（内置，增强模型对业务的理解）
AUDIT_DOMAIN_KNOWLEDGE = """
## 审计领域知识（系统内置）

### 银行对账第一常识（方向镜像，必须遵守）
1. 企业序时账/日记账：借方金额=银行存款增加，贷方金额=减少 → 净额=借方−贷方
2. 银行流水：贷方（收入）=存款增加，借方（支取）=减少 → 净额=贷方收入−借方支取
3. 双方记账方向互为镜像：序时账"借方金额"对应流水"贷方（收入）"，序时账"贷方金额"对应流水"借方（支取）"
4. 流水含多个银行账户、序时账为单账户时：必须先按银行账号过滤流水再对账

### 列名跨表映射（Merge 时必须正确配对）
序时账与银行流水的列名映射是方向镜像，不是同名映射：
- 序时账.借方金额 ←→ 银行流水.贷方（收入）   【存款增加方向】
- 序时账.贷方金额 ←→ 银行流水.借方（支取）   【存款减少方向】
- 序时账.银行账号 ←→ 银行流水.银行账号       【同名】
- 序时账.日期     ←→ 银行流水.日期           【同名】
- 序时账.摘要     ←→ 银行流水.摘要           【同名，但表述可能不同】
⚠️ 关键：Merge 的 left_on/right_on 必须使用以上映射，严禁按同名列名硬匹配！

### 会计实务记账方式（对账差异根因）
1. 【汇总记账】多笔银行流水合并为1笔会计分录：
   当天10笔支付宝到账→会计只做1笔汇总"借：银行存款/贷：主营业务收入"
   → 金额之和相等但笔数不对，需按日期+金额窗口做 n:1 匹配
2. 【拆分入账】1笔银行流水拆成多笔分录：
   一笔100万付款对应3个部门费用→拆成40+35+25三笔
   → 需按日期窗口做 1:n 合计匹配
3. 【跨期到账】银行已扣款但企业次月才记账（或反之）：
   常见于月末最后几天的交易，日期差≤3天属于正常窗口
4. 【小额费用】银行扣手续费/利息/账户管理费，企业可能月末统一记账或不记账
   → 这类差异金额通常很小（<100元），单独成类输出，不混入未达账项
5. 【同名异户】同一笔款从农行5927出但会计错记到农行8310
   → 需按银行账号过滤后分别对账
6. 【借贷方向颠倒】银行流水"收入"在企业账上误记到借方
   → 需方向感知匹配（自动取绝对值/取负号后再比）

### 未达账项四分类（时间差导致的正常差异）
1. 银行已收、企业未收：钱到账了但会计还没记（如利息收入、POS隔日到账）
2. 银行已付、企业未付：钱扣了但会计还没记（如银行扣手续费、贷款利息）
3. 企业已收、银行未收：会计记了但银行还没到账（如在途支票、跨行转账）
4. 企业已付、银行未付：会计记了但银行还没扣（如开出支票对方未兑付）

### 逐笔核对与容差（不得混淆）
1. 银行存款逐笔核对必须精确到分（±0.01 元），差一分钱都必须查明原因
2. 百分比容差（如 1%）只适用于汇总层面的分析性复核，绝不可用于逐笔核对
3. 匹配策略优先级：
   L1: 金额精确+同日 → 自动确认
   L2: 金额精确+日期窗口(±3天) → 自动确认
   L3: 同日同方向 n:m 合计相等（合并/拆分入账）→ 标记为"合并入账待确认"
   L4: 摘要/对方户名模糊 → 仅进人工复核，不自动确认
4. 未匹配项默认"待人工核查"；只有接近期末且有窗口证据的才能列为未达账项候选
5. 利息、手续费、冲正不删除、单独成类输出

### 跨表数据匹配原则
1. 优先识别两表共有的有业务含义的关键列；严禁使用"序号/编号/行号"作为连接键
2. 列名以 Data Catalog 中的真实列名为准，严禁照抄示例中的文件名/列名
3. 机构名称"去市县区中心管理后缀"的标准化仅用于医保回款场景，其他场景禁止
4. 金额汇总比对容差由用户指定（如"差异控制在5万以内"），逐笔核对一律 ±0.01 元
5. 支持文档格式输入（docx/doc/pdf/md/txt），文档中的表格与 Excel 同权参与比对
6. ⚠️ Merge 的 left_on/right_on 必须使用本知识中的列名跨表映射，不要简单地按同名匹配

### DAG 算子使用指南
- Load：每个文件一个 Load，必须设置 source_file
- NoiseFilter：按"摘要"列过滤噪音（利息、手续费、冲正等），过滤后单独输出噪音清单
- Sort：按日期+金额列排序，为 Merge 的日期窗口匹配做准备
- Merge：两表比对核心算子，left_on和right_on按列名跨表映射填写
- Diff：按金额列比对差异，tolerance_abs=0.01
- Export：导出差异表，标注差异来源文件
- ConditionCheck：差异超过用户设定阈值时标记

### 对账Playbook（审计实务铁律，来自真实审计经验）

#### P1 双向核对铁律
从账到流水（查记录真实性）+ 从流水到账（查完整性），两方向缺一不可。
高危输出"流水有、账上无"→ 疑似收入不入账/出借账户/挪用资金（先入一笔再分次转出是典型形态）。
高危输出"一借一贷金额相等的大额发生额"→ 疑似过桥/资金划转，已销户账户尤其要查。
→ DAG 规划时 Diff 必须同时输出"仅左表有"和"仅右表有"两类差异。

#### P2 n:m 匹配规则（L3 层）
- 拆分收款：一笔账面对应多笔流水（总额相等，日期窗口±3天）
- 多笔合并支付：多笔账面对应一笔流水
- 手续费单独扣款：流水 10,000 ↔ 账面 9,995 + 手续费 5（差额<10元且摘要含"手续费"→自动成组）
- 票据跨期兑付：账面日期与流水日期跨期，金额相等
- 结息/利息：单独成类，不参与匹配、不删除

#### P3 未达账项四分类与舞弊信号
四类：企业已收银行未收 / 企业已付银行未付 / 银行已收企业未收 / 银行已付企业未付。
长期未达（>30天）+ 大额未达 → 核实原因，必要时函证交易性质。
期后验证：资产负债表日后对账单是否记录了该未达项。
舞弊典型：银行已收企业未收 + 等额支出配合 → 调节表看不出；长期挪用。

#### P4 大额标准动态化
大额阈值 = 固定金额（用户指定，如≥10万）或分位数（如该账户发生额 top 5%），交互确认后留痕。
→ Merge/Diff 的 params 中可设 threshold: {type: fixed|quantile, value: 用户指定}

#### P5 对账单真实性检查（红旗规则）
- 对账单编号是否重复/不连续
- 结息日、结息金额是否合理（季度结息：3/6/9/12月21日前后）
- 存款余额是否连贯（上日余额+发生=本日余额，逐行可验）
- 同账户不同阶段户名/账号/开户行是否一致
- 对公账户是否含"积分"等异常信息

#### P6 账户选择优先级（多账户场景）
优先对：基本户、余额较大户、发生额大且频繁户。
高危优先：发生额大但余额小/零余额/当期销户户（舞弊高发）。
→ 流水含多账户时，DAG 规划应先注入账户过滤步骤。
"""

DIFY_DAG_WORKFLOW_ID = "c46f68e0-a757-424c-868a-1144eb5a4260"
DIFY_REFINE_WORKFLOW_ID = "f273495a-28e4-498f-a017-389f3e3b37b8"
DIFY_SINGLE_TABLE_ID = "ca45a52f-1dc0-41aa-88c7-cc3046f66b0c"
DIFY_REPORT_GEN_ID = "49dda78f-1ae2-4c16-84d1-2c306c0a22b3"
DIFY_REPORT_REVIEW_ID = "d0cb295d-4545-4afd-956c-059992fdc63f"
DIFY_KNOWLEDGE_QA_ID = "f97a2880-92a1-45b2-9fbd-0479a6c81255"

CELERY_BROKER_URL = "redis://localhost:6379/0"
CELERY_RESULT_BACKEND = "redis://localhost:6379/0"


# ═══════════════════════════════════════════════════════════════
# Pydantic 请求/响应模型
# ═══════════════════════════════════════════════════════════════

class CreateRunRequest(BaseModel):
    project_code: str = Field(..., description="项目编号，如 A2025001")
    subject: str = Field(..., description="审计科目/主题")
    user_intent: str = Field(..., description="大白话审计意图")
    preset_button: Optional[str] = Field(None, description="预设按钮名称")
    parent_run_id: Optional[str] = Field(None, description="继承的父 Run ID")


class CreateRunResponse(BaseModel):
    run_id: str
    status: str
    message: str
    compile_task_id: Optional[str] = None


class ExecuteRequest(BaseModel):
    confirmed: bool = Field(True, description="审计师已确认 DAG")


class StatusResponse(BaseModel):
    run_id: str
    status: str
    progress: int
    current_step: str
    elapsed_seconds: float
    retry_count: int
    output_files: List[str]


class DownloadResponse(BaseModel):
    task_id: str
    status: str
    message: str


# ═══════════════════════════════════════════════════════════════
# 生命周期钩子实现
# ═══════════════════════════════════════════════════════════════

class AuditLifecycleHooks(LifecycleHooks):
    def __init__(self, snapshot_mgr: RunSnapshotManager):
        self.snapshot = snapshot_mgr

    def on_born(self, run_id: str, container_id: str, run_dir: Path) -> None:
        pass

    def on_complete(self, run_id: str, container_id: str, result: LifecycleResult, run_dir: Path) -> None:
        output_dir = run_dir / "outputs"
        if output_dir.exists():
            files = [str(f.name) for f in output_dir.iterdir() if f.is_file()]
            result.output_files = files

    def on_destroy(self, run_id: str, container_id: str) -> None:
        self.snapshot.cleanup_temp(run_id)


# ═══════════════════════════════════════════════════════════════
# 路由定义
# ═══════════════════════════════════════════════════════════════

router = APIRouter(prefix="/api/v3", tags=["Run 管理"])

# 💡 修复一：使用兼容性更好的 Optional 替代 | 符号，防止低版本 Python 启动报错
_snapshot_mgr: Optional[RunSnapshotManager] = None
_chaos_processor: Optional[ChaosInputProcessor] = None
_sandbox: Optional[EphemeralSandbox] = None


def _get_snapshot_mgr() -> RunSnapshotManager:
    global _snapshot_mgr
    if _snapshot_mgr is None:
        _snapshot_mgr = RunSnapshotManager()
    return _snapshot_mgr


def _get_chaos_processor() -> ChaosInputProcessor:
    global _chaos_processor
    if _chaos_processor is None:
        _chaos_processor = ChaosInputProcessor()
    return _chaos_processor


def _get_sandbox() -> EphemeralSandbox:
    global _sandbox
    if _sandbox is None:
        _sandbox = EphemeralSandbox(hooks=AuditLifecycleHooks(_get_snapshot_mgr()))
        # 确保沙箱镜像存在（自动构建）
        import docker
        try:
            client = docker.from_env()
            try:
                client.images.get("audit-sandbox:alpine-v2")
                print("[Sandbox] 镜像 audit-sandbox:alpine-v2 已存在")
            except docker.errors.ImageNotFound:
                print("[Sandbox] 镜像不存在，正在构建 audit-sandbox:alpine-v2 ...")
                if _sandbox.build_image():
                    print("[Sandbox] 镜像构建成功")
                else:
                    print("[Sandbox] 镜像构建失败！将跳过 Docker 沙箱，使用直接执行模式")
            finally:
                client.close()
        except Exception as e:
            print(f"[Sandbox] 镜像检查/构建失败: {e}，将在首次执行时尝试")
    return _sandbox


# ── 提取核对快车道（词典命中→免LLM组装DAG）────────────

def _detect_extract_pair(input_dir: Path) -> Optional[Tuple]:
    """检测银行流水+汇总表文件对。返回 (bank_path, ledger_path, bank_df, ledger_df) 或 None。"""
    try:
        from core.matching_engine import detect_file_type
        excel_files = [f for f in input_dir.glob("*")
                       if f.suffix.lower() in (".xlsx", ".xls", ".csv")]
        if len(excel_files) < 2:
            return None
        import pandas as _pd
        bank_f, ledger_f = None, None
        bank_df, ledger_df = None, None
        for f in excel_files:
            df = _pd.read_csv(f, encoding="utf-8-sig", nrows=1000) \
                if f.suffix.lower() == ".csv" \
                else _pd.read_excel(f, nrows=1000)
            ft = detect_file_type(df, f.name)
            if ft == "bank_statement" and bank_f is None:
                bank_f, bank_df = f, df
            elif ft in ("summary_table", "journal") and ledger_f is None:
                ledger_f, ledger_df = f, df
        if not bank_f and not ledger_f and len(excel_files) >= 2:
            bank_f, ledger_f = excel_files[0], excel_files[1]
            bank_df = _pd.read_csv(bank_f, encoding="utf-8-sig", nrows=1000) \
                if bank_f.suffix.lower() == ".csv" \
                else _pd.read_excel(bank_f, nrows=1000)
            ledger_df = _pd.read_csv(ledger_f, encoding="utf-8-sig", nrows=1000) \
                if ledger_f.suffix.lower() == ".csv" \
                else _pd.read_excel(ledger_f, nrows=1000)
        # 一方已确定、另一方未知 → 未知方直接作为另一方
        if bank_f and not ledger_f:
            other = [f for f in excel_files if f != bank_f]
            if other:
                ledger_f = other[0]
                ledger_df = _pd.read_csv(ledger_f, encoding="utf-8-sig", nrows=1000) \
                    if ledger_f.suffix.lower() == ".csv" \
                    else _pd.read_excel(ledger_f, nrows=1000)
        elif ledger_f and not bank_f:
            other = [f for f in excel_files if f != ledger_f]
            if other:
                bank_f = other[0]
                bank_df = _pd.read_csv(bank_f, encoding="utf-8-sig", nrows=1000) \
                    if bank_f.suffix.lower() == ".csv" \
                    else _pd.read_excel(bank_f, nrows=1000)
        if not bank_f or not ledger_f:
            return None
        return (bank_f, ledger_f, bank_df, ledger_df)
    except Exception as e:
        print(f"[提取快车道] 文件配对失败: {e}")
        return None


def _assemble_extraction_dag(bank_file: Path, ledger_file: Path,
                              kw: dict, b_cols: dict, l_cols: dict) -> dict:
    """固定拓扑：Load×2 → RegexFilter → GroupBy → Merge → Diff → Export。
    与 LLM 生成的 DAG 同构，下游执行/勾稽/报告链路零改动。"""
    desc_col = b_cols.get("desc_col", "摘要")
    amt_col = b_cols.get("amount_col") or b_cols.get("income_col") or "金额"
    name_col = l_cols.get("name_col", "机构名称")
    total_col = l_cols.get("total_col", "合计")
    patterns = kw["patterns"]
    exclude = kw.get("exclude", "")
    agg_key = b_cols.get("counterparty_col") or b_cols.get("name_col") or name_col

    operators = [
        {"id": "op_1", "name": "Load", "source_file": bank_file.name,
         "params": {}, "output_alias": "df_bank"},
        {"id": "op_2", "name": "Load", "source_file": ledger_file.name,
         "params": {}, "output_alias": "df_ledger"},
        {"id": "op_3", "name": "RegexFilter", "input_from": ["op_1"],
         "output_alias": "df_filtered",
         "params": {"columns": kw["columns"],
                    "pattern": patterns, "case_sensitive": False}},
    ]
    next_input = ["op_3"]
    next_alias = "df_filtered"
    if exclude:
        operators.append(
            {"id": "op_3b", "name": "NoiseFilter", "input_from": ["op_3"],
             "output_alias": "df_excluded",
             "params": {"columns": kw["columns"],
                        "noise_patterns": [p.strip() for p in exclude.split("|") if p.strip()]}},
        )
        next_input = ["op_3b"]
        next_alias = "df_excluded"
    operators += [
        {"id": "op_4", "name": "GroupBy", "input_from": next_input,
         "output_alias": "df_bank_agg",
         "params": {"by": [agg_key], "aggregations": {amt_col: "sum"}}},
        {"id": "op_5", "name": "Merge", "input_from": ["op_4", "op_2"],
         "output_alias": "df_merged",
         "params": {"how": "outer", "left_on": [agg_key],
                    "right_on": [name_col]}},
        {"id": "op_6", "name": "Diff", "input_from": ["op_5"],
         "output_alias": "df_result",
         "params": {"col_a": f"{amt_col}_sum", "col_b": total_col,
                    "tolerance_pct": 1.0, "output_mode": "all"}},
        {"id": "op_7", "name": "Export", "input_from": ["op_6"],
         "params": {"filename": "提取核对结果.xlsx"}},
    ]

    return {
        "objective": f"提取式核对: {kw['dict_key']}",
        "compiled_by": "fastlane_template",
        "kw_source": "dictionary",
        "dict_key": kw["dict_key"],
        "operators": operators,
    }


def _try_extraction_fastlane(input_dir: Path, user_intent: str,
                              run_dir: Path) -> Optional[dict]:
    """提取核对快车道：词典命中 + 文件类型确定 + 列映射齐全 → 免LLM组装DAG。
    任一条件不满足 → 返回 None，回退常规LLM编译链。"""
    # ① 词典必须命中
    from config.extraction_dictionary import resolve_patterns_full
    kw = resolve_patterns_full(user_intent)
    if not kw or kw.get("source") not in ("dictionary", "dictionary_fuzzy"):
        return None  # LLM提案路径要走确认，不能快进
    # ② 一流水一汇总表，类型确定
    pair = _detect_extract_pair(input_dir)
    if not pair:
        return None
    bank_f, ledger_f, bank_df, ledger_df = pair
    # ③ 列映射齐全
    from core.matching_engine import identify_columns
    b_cols = identify_columns(bank_df, "bank_statement")
    l_cols = identify_columns(ledger_df, "summary_table")
    if not (b_cols.get("desc_col") and
            (b_cols.get("amount_col") or b_cols.get("income_col"))
            and l_cols.get("name_col") and l_cols.get("total_col")):
        return None
    # 组装
    return _assemble_extraction_dag(bank_f, ledger_f, kw, b_cols, l_cols)


# ── 后台编译任务 ─────────────────────────────────────

async def _background_compile(
        run_id: str,
        catalog: AssetCatalog,
        user_intent: str,
        preset_button: Optional[str],
        parent_run_id: Optional[str],
) -> None:
    print(f"DEBUG: 开始编译 Run {run_id}")  # 确认任务是否启动
    from time import time as _now
    _t0 = _now()
    try:
        _get_snapshot_mgr().update_status(run_id, "COMPILING")

        # ── 提取核对快车道：词典命中→免LLM组装DAG ──
        dag_blueprint = None
        record = _get_snapshot_mgr().get_run(run_id)
        if record:
            fast_dag = _try_extraction_fastlane(
                record.input_dir, user_intent, record.run_dir)
            if fast_dag:
                print(f"[提取快车道] 词典命中，免LLM组装DAG（{len(fast_dag['operators'])}算子）")
                blueprint_dict = fast_dag
                _get_snapshot_mgr().update_blueprint(run_id, blueprint_dict)
                dag_path = record.run_dir / "dag_blueprint.json"
                with open(dag_path, "w", encoding="utf-8") as f:
                    json.dump(blueprint_dict, f, ensure_ascii=False, indent=2)
                plan_path = record.run_dir / "execution_plan.txt"
                plan_text = f"提取核对·快车道模板: {fast_dag['objective']}\n" \
                            f"筛选关键词: {blueprint_dict.get('kw_source','?')}:{blueprint_dict.get('dict_key','?')}\n" \
                            f"算子链: Load→RegexFilter→GroupBy→Merge→Diff→Export"
                with open(plan_path, "w", encoding="utf-8") as f:
                    f.write(plan_text)
                _get_snapshot_mgr().update_status(run_id, "PENDING_REVIEW")
                trace_record(run_id, "dag_fastlane", "OK", (_now() - _t0) * 1000)
                print(f"[提取快车道] Run {run_id} 编译完成（模板），等待人工确认")
                return  # ⚠ 关键：直接return，别滑进LLM编译段

        dag_blueprint = None
        is_single = (catalog.total_files == 1 and any(k in user_intent for k in ["筛选", "筛查", "分类"]) and not any(k in user_intent for k in ["核对", "对账", "匹配", "比对"]))
        if is_single:
            catalog_text = _format_catalog_for_prompt(catalog)
            single_result = await _call_dify_single_table(catalog_text, user_intent)
            if single_result:
                print(f"[场景路由] 单表筛选模式")
                dag_blueprint = _parse_single_table_result(single_result, catalog)
        if dag_blueprint is None:
            dag_blueprint = await _call_dify_compiler(
                catalog=catalog,
            user_intent=user_intent,
            preset_button=preset_button,
            parent_run_id=parent_run_id,
        )
        print(f"DEBUG: 编译完成，结果: {dag_blueprint is not None}")
        if dag_blueprint is None:
            raise ValueError("编译器返回了空结果")

        # ── 📚 分层知识注入（V2.7 新增）──────────────────
        # 遍历 DAG 中的算子类型，为每种算子注入其专属知识子集
        try:
            from core.rag_engine import inject_layered_context, OPERATOR_KNOWLEDGE_MAP
            operators = (
                dag_blueprint.operators
                if hasattr(dag_blueprint, "operators")
                else (dag_blueprint if isinstance(dag_blueprint, dict) else {}).get("operators", [])
            )
            op_types = set()
            for op in operators:
                name = op.name if hasattr(op, "name") else op.get("name", "")
                if name in OPERATOR_KNOWLEDGE_MAP and OPERATOR_KNOWLEDGE_MAP[name]:
                    op_types.add(name)

            if op_types:
                layered_parts = []
                for ot in sorted(op_types):
                    ctx = inject_layered_context(user_intent, ot)
                    if ctx:
                        layered_parts.append(ctx)

                if layered_parts:
                    layered_text = (
                        "## 算子专项知识（分层注入）\n"
                        + f"涉及算子: {', '.join(sorted(op_types))}\n\n"
                        + "\n---\n".join(layered_parts)
                    )
                    # 嵌入到 DAG 的 context 中
                    if hasattr(dag_blueprint, "context"):
                        existing = dag_blueprint.context or {}
                        existing["_layered_knowledge"] = layered_text
                    else:
                        dag_blueprint["context"] = dag_blueprint.get("context", {})
                        dag_blueprint["context"]["_layered_knowledge"] = layered_text
                    print(f"[分层注入] 为 {len(op_types)} 类算子({', '.join(sorted(op_types))})注入了专项知识")
        except Exception as e:
            print(f"[分层注入] 跳过（非致命）: {e}")

        record = _get_snapshot_mgr().get_run(run_id)
        if record is None:
            raise ValueError(f"Run {run_id} 不存在")

        # 高防御字典化
        blueprint_dict = dag_blueprint.to_dict() if hasattr(dag_blueprint, "to_dict") else dag_blueprint.__dict__
        blueprint_dict = _ensure_essential_operators(blueprint_dict, user_intent)
        _get_snapshot_mgr().update_blueprint(run_id, blueprint_dict)

        # 保存 DAG JSON 文件
        dag_path = record.run_dir / "dag_blueprint.json"
        with open(dag_path, "w", encoding="utf-8") as f:
            if hasattr(dag_blueprint, "to_json"):
                f.write(dag_blueprint.to_json())
            else:
                json.dump(blueprint_dict, f, ensure_ascii=False, indent=2)

        # 尝试生成自然语言计划（防报错）
        plan_path = record.run_dir / "execution_plan.txt"
        try:
            plan_text = DAGToCodeDescription.describe(dag_blueprint)
        except Exception as e:
            plan_text = f"执行计划解析降级: 包含 {len(blueprint_dict.get('operators', []))} 个物理算子"

        with open(plan_path, "w", encoding="utf-8") as f:
            f.write(plan_text)

        # 生成自然语言匹配逻辑说明
        try:
            explanation = await _explain_dag(blueprint_dict, catalog, user_intent)
            if explanation:
                explanation_path = record.run_dir / "match_explanation.txt"
                with open(explanation_path, "w", encoding="utf-8") as f:
                    f.write(explanation)
                blueprint_dict["match_explanation"] = explanation
                _get_snapshot_mgr().update_blueprint(run_id, blueprint_dict)
            else:
                # DAG 为空时（Dify 不可用），从场景注册表生成匹配计划
                _generate_fallback_plan(record.run_dir, user_intent, preset_button, catalog)
        except Exception as e:
            print(f"[后台编译] 自然语言说明生成失败（非致命）: {e}")
            _generate_fallback_plan(record.run_dir, user_intent, preset_button, catalog)

        _get_snapshot_mgr().update_status(run_id, "PENDING_REVIEW")
        trace_record(run_id, "dag_compile", "OK", (_now() - _t0) * 1000)
        print(f"[后台编译] Run {run_id} 编译完成，等待人工确认")

    except Exception as e:
        import traceback
        error_detail = traceback.format_exc()
        print(f"[后台编译] Run {run_id} 编译失败: {error_detail}")

        # 最终兜底：从 catalog 直接生成最小 DAG
        try:
            trace_record(run_id, "dag_fallback", "WARN",
                         (_now() - _t0) * 1000, str(e)[:200])
            print(f"[后台编译] 激活最终兜底：从数据目录生成基础 DAG")
            dag_blueprint = _build_fallback_dag(catalog, user_intent)
            if dag_blueprint:
                record = _get_snapshot_mgr().get_run(run_id)
                if record:
                    blueprint_dict = dag_blueprint.to_dict() if hasattr(dag_blueprint, "to_dict") else dag_blueprint.__dict__
                    _get_snapshot_mgr().update_blueprint(run_id, blueprint_dict)
                    dag_path = record.run_dir / "dag_blueprint.json"
                    with open(dag_path, "w", encoding="utf-8") as f:
                        json.dump(blueprint_dict, f, ensure_ascii=False, indent=2)
                    _get_snapshot_mgr().update_status(run_id, "PENDING_REVIEW")
                    _generate_fallback_plan(record.run_dir, user_intent, preset_button, catalog)
                    # 同步到蓝图，前端从 dag_blueprint.match_explanation 读取
                    try:
                        _plan_file = record.run_dir / "match_explanation.txt"
                        if _plan_file.exists():
                            blueprint_dict["match_explanation"] = _plan_file.read_text(encoding="utf-8")
                            _get_snapshot_mgr().update_blueprint(run_id, blueprint_dict)
                    except Exception:
                        pass
                    print(f"[后台编译] Run {run_id} 兜底编译完成")
                    return
        except Exception as e2:
            print(f"[后台编译] 兜底编译也失败: {e2}")

        trace_record(run_id, "dag_compile", "FAIL",
                     (_now() - _t0) * 1000, str(e)[:200])
        _get_snapshot_mgr().update_status(run_id, "FAILED")

        # 💡 修复二：解决 sqlite 报错问题，动态探查 DB_PATH
        try:
            from core.run_snapshot import DB_PATH
            db_target = DB_PATH
        except ImportError:
            db_target = Path("data/audit_platform.db")

        try:
            with sqlite3.connect(str(db_target)) as conn:
                conn.execute(
                    "UPDATE runs SET execution_logs = ? WHERE run_id = ?",
                    (json.dumps([f"编译失败: {str(e)}", error_detail], ensure_ascii=False), run_id)
                )
                conn.commit()
        except Exception as db_err:
            print(f"保存错误日志失败: {db_err}")


# ── 表格自动归一化（用户无感，最佳努力） ──────────────────

def _normalize_input_tables(input_dir: Path, catalog: Any) -> None:
    """对上传的 Excel/CSV 文件自动做表头定位+清洗+归一化。
    
    不阻断主流程，失败静默降级。银行对账文件会自动跳过（对账引擎自处理）。
    归一化后的文件保存在 {filename}_normalized.csv。
    """
    try:
        from core.table_normalizer import normalize_and_validate, clean_dataframe
    except Exception as e:
        print(f"[归一化] 模块不可用，跳过: {e}")
        return

    excel_exts = {".xlsx", ".xls", ".csv"}
    for item in sorted(input_dir.iterdir()):
        if not item.is_file():
            continue
        if item.suffix.lower() not in excel_exts:
            continue
        # 跳过已归一化的文件（防重复处理）
        if "_normalized" in item.stem:
            continue

        try:
            import pandas as pd
            if item.suffix.lower() == ".csv":
                df = pd.read_csv(item, encoding="utf-8-sig", nrows=None)
            else:
                df = pd.read_excel(item, nrows=None)

            if df.empty or len(df.columns) == 0:
                continue

            # 清洗：自动找表头、丢垃圾行、空格转空值
            cleaned = clean_dataframe(df, auto_header=True, strip_blank_rows=True)
            if cleaned.empty:
                continue

            # 判断是否真的有变化（无变化就跳过，不产生多余文件）
            same_shape = (len(cleaned) == len(df) and
                          list(cleaned.columns) == list(df.columns))
            if same_shape:
                continue  # 标准长表，无需归一化

            # 保存归一化版本
            out_path = item.parent / f"{item.stem}_normalized.csv"
            cleaned.to_csv(out_path, index=False, encoding="utf-8-sig")

            # 更新 catalog：把归一化后的列名同步到 AssetCatalog
            if hasattr(catalog, "files") and catalog.files:
                for f_info in catalog.files:
                    if f_info.get("filename") == item.name:
                        f_info["normalized_file"] = out_path.name
                        f_info["normalized_columns"] = [
                            {"name": str(c), "dtype": str(cleaned[c].dtype)}
                            for c in cleaned.columns
                        ]
                        break

            print(f"[归一化] {item.name}: {len(df)}行→{len(cleaned)}行, "
                  f"列: {list(df.columns)[:5]}→{list(cleaned.columns)[:5]}")

        except Exception as e:
            print(f"[归一化] {item.name} 处理失败（非阻断）: {e}")
            continue


def _reverse_validation_for_report(match_result: dict, input_dir: Path) -> dict:
    """从匹配结果中提取反向校验数据（未匹配行摘要聚类）。"""
    try:
        from core.matching_engine import reverse_validate_unmatched
        # 尝试加载未匹配数据
        unmatched_path = input_dir.parent / "outputs" / "unmatched_detail.csv"
        if not unmatched_path.exists():
            return {}
        import pandas as pd
        df = pd.read_csv(unmatched_path)
        if df.empty:
            return {"total_unmatched": 0, "clusters": []}
        desc_col = None
        for c in ["摘要", "对方客户名称", "附言"]:
            if c in df.columns:
                desc_col = c
                break
        if desc_col:
            return reverse_validate_unmatched(df, desc_col=desc_col)
        return {"total_unmatched": len(df), "clusters": [],
                "note": "未匹配明细无摘要列，无法聚类"}
    except Exception as e:
        return {"error": str(e)}


# ── 1. 创建新 Run ─────────────────────────
@router.post("/runs", response_model=CreateRunResponse, summary="创建新 Run（异步）")
async def create_run(
        project_code: Optional[str] = Form(None),
        subject: Optional[str] = Form("未命名审计"),
        user_intent: Optional[str] = Form(None),
        files: List[UploadFile] = File(..., description="混沌输入（Excel/ZIP/文件夹）"),
        preset_button: Optional[str] = Form(None),
        parent_run_id: Optional[str] = Form(None),
        background_tasks: BackgroundTasks = BackgroundTasks(),
) -> CreateRunResponse:
    # 【核心调试代码】

    if not subject:
        raise HTTPException(status_code=400, detail=f"后端未收到 subject 字段。")
    print(f"DEBUG: 接收到的参数 -> project_code={project_code}, subject={subject}, user_intent={user_intent}, files={len(files)}")

    from core.run_snapshot import TEMP_BASE
    TEMP_BASE.mkdir(parents=True, exist_ok=True)
    temp_dir = TEMP_BASE / f"upload_{uuid.uuid4().hex[:8]}"
    temp_dir.mkdir(parents=True, exist_ok=True)

    # 保存所有上传文件
    upload_paths = []
    for file in files:
        safe_filename = _validate_upload(file)
        upload_path = temp_dir / safe_filename
        with open(upload_path, "wb") as f:
            content = await file.read()
            f.write(content)
        upload_paths.append(upload_path)

    # 逐个处理文件，合并 catalog
    all_files = []
    global_hash = hashlib.md5()
    all_flat_dirs = []
    for up in upload_paths:
        try:
            # 使用唯一临时 run_id 而非硬编码 "PENDING"，避免不同上传的文件串扰
            temp_rid = f"proc_{uuid.uuid4().hex[:8]}"
            flat_dir, catalog = _get_chaos_processor().process(
                str(up), run_id=temp_rid,
            )
            all_files.extend(catalog.files)
            global_hash.update(catalog.global_hash.encode())
            all_flat_dirs.append(flat_dir)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"混沌输入处理失败 ({Path(up).name}): {e}")

    from core.chaos_input import AssetCatalog
    catalog = AssetCatalog(
        files=all_files,
        total_files=len(all_files),
        global_hash=global_hash.hexdigest(),
    )

    try:
        record = _get_snapshot_mgr().create_run(
            project_code=project_code,
            subject=subject,
            user_intent=user_intent,
            input_catalog=catalog,
            parent_run_id=parent_run_id,
            preset_button=preset_button,
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Run 创建失败: {e}")

    run_input_dir = record.input_dir
    if run_input_dir.exists():
        shutil.rmtree(str(run_input_dir), ignore_errors=True)
    run_input_dir.mkdir(parents=True, exist_ok=True)
    # 拷贝所有文件的扁平化目录到 inputs
    for fd in all_flat_dirs:
        for item in fd.iterdir():
            dest = run_input_dir / item.name
            if item.is_dir():
                shutil.copytree(str(item), str(dest), dirs_exist_ok=True)
            else:
                shutil.copy2(str(item), str(dest))
    _get_snapshot_mgr().lock_readonly(run_input_dir)

    # ── 表格归一化（自动清洗表头偏移/空格/合计行，用户无感） ──
    _normalize_input_tables(run_input_dir, catalog)

    # 拷贝完成后清理临时目录
    try:
        shutil.rmtree(str(temp_dir), ignore_errors=True)
        for fd in all_flat_dirs:
            parent = fd.parent
            if parent != temp_dir and parent.name.startswith("upload_"):
                shutil.rmtree(str(parent), ignore_errors=True)
    except Exception:
        pass

    # ── B4 数据质量预检（编译前注入质量报告卡） ──
    quality_report = []
    try:
        from core.data_quality import inspect_catalog
        quality_report = inspect_catalog(str(run_input_dir))
        print(f"[质量门] 扫描 {len(quality_report)} 个文件，"
              f"WARNING/POOR: {sum(1 for q in quality_report if q['overall'] in ('WARNING','POOR'))}")
    except Exception as e:
        print(f"[质量门] 预检降级（非阻断）: {e}")

    # （quality_report 需从 Run 记录的扩展字段查询；前端 /runs/{id} 已自动返回 dag_blueprint；
    #  暂以独立的 /runs/{id}/quality 端点提供，不影响现有链路）
    try:
        qp = record.run_dir / "quality_report.json"
        with open(qp, "w", encoding="utf-8") as f:
            json.dump(quality_report, f, ensure_ascii=False)
    except Exception:
        pass

    background_tasks.add_task(
        _background_compile,
        run_id=record.run_id,
        catalog=catalog,
        user_intent=user_intent,
        preset_button=preset_button,
        parent_run_id=parent_run_id,
    )

    return CreateRunResponse(
        run_id=record.run_id,
        status="COMPILING",
        message="文件已接收，DAG 正在后台编译中，请稍后刷新查看结果...",
        compile_task_id=record.run_id,
    )

@router.get("/runs", summary="获取所有 Run 列表")
async def list_runs():
    """获取最近的 Run 记录列表"""
    try:
        records = _get_snapshot_mgr().get_recent_runs(limit=50)
        return {"runs": [r.to_dict() for r in records]}
    except Exception as e:
        return {"runs": []}


@router.delete("/runs/{run_id}", summary="删除单个 Run")
async def delete_run(run_id: str):
    """删除指定 Run 的数据库记录 + 磁盘目录"""
    import shutil
    from core.run_snapshot import RUNS_BASE, DB_PATH
    try:
        # 删除磁盘目录
        disk_dir = RUNS_BASE / run_id
        if disk_dir.exists():
            shutil.rmtree(str(disk_dir), ignore_errors=True)
        # 删除数据库记录
        with sqlite3.connect(str(DB_PATH)) as conn:
            conn.execute("DELETE FROM runs WHERE run_id = ?", (run_id,))
            conn.commit()
        return {"success": True, "message": f"已删除 {run_id}"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/runs/cleanup", summary="批量清理历史 Run 和临时文件")
async def cleanup_runs(keep_latest: int = 5, clean_temp: bool = True):
    """
    批量清理：
    - keep_latest: 保留最近 N 个 Run，其余删除
    - clean_temp: 是否同时清理 data/temp 残留目录
    """
    import shutil
    from core.run_snapshot import RUNS_BASE, DB_PATH, TEMP_BASE
    result = {"deleted_runs": [], "cleaned_temp": 0}

    try:
        # 清理历史 Run
        records = _get_snapshot_mgr().get_recent_runs(limit=200)
        if len(records) > keep_latest:
            to_delete = records[keep_latest:]
            with sqlite3.connect(str(DB_PATH)) as conn:
                for r in to_delete:
                    rid = r.run_id if hasattr(r, 'run_id') else r.get('run_id', '')
                    disk_dir = RUNS_BASE / rid
                    if disk_dir.exists():
                        shutil.rmtree(str(disk_dir), ignore_errors=True)
                    conn.execute("DELETE FROM runs WHERE run_id = ?", (rid,))
                    result["deleted_runs"].append(rid)
                conn.commit()

        # 清理 temp 目录
        if clean_temp and TEMP_BASE.exists():
            for d in TEMP_BASE.iterdir():
                if d.is_dir():
                    shutil.rmtree(str(d), ignore_errors=True)
                    result["cleaned_temp"] += 1

        return {"success": True, "data": result}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ── 2. 查询 Run（💡 修复三：注入高防御物理探测逻辑）─────────────────────────
@router.get("/runs/{run_id}", summary="查询 Run 详情（含编译状态）")
async def get_run(run_id: str) -> Dict[str, Any]:
    record = _get_snapshot_mgr().get_run(run_id)
    if not record:
        # 防御性回退：DB 中无记录但磁盘目录存在时，从 run_meta.json 重建
        from core.run_snapshot import RUNS_BASE
        disk_dir = RUNS_BASE / run_id
        meta_file = disk_dir / "run_meta.json"
        if meta_file.exists():
            try:
                with open(meta_file, "r", encoding="utf-8") as f:
                    result = json.load(f)
                result["_source"] = "disk_fallback"
                result.setdefault("execution_plan", "未生成执行计划")
                result.setdefault("outputs", [])
                out_dir = disk_dir / "outputs"
                if out_dir.exists():
                    try:
                        result["outputs"] = [f.name for f in out_dir.iterdir() if f.is_file()]
                    except Exception:
                        pass
                if not result.get("dag_blueprint"):
                    dag_file = disk_dir / "dag_blueprint.json"
                    if dag_file.exists():
                        try:
                            with open(dag_file, "r", encoding="utf-8") as f:
                                result["dag_blueprint"] = json.load(f)
                        except Exception:
                            pass
                return result
            except Exception:
                pass
        raise HTTPException(status_code=404, detail="Run 不存在")

    result = record.to_dict()

    # 物理防崩溃探测
    run_dir = Path(record.run_dir) if hasattr(record, "run_dir") else None
    result["execution_plan"] = "未生成执行计划"
    result["outputs"] = []

    if run_dir and run_dir.exists():
        plan_path = run_dir / "execution_plan.txt"
        if plan_path.exists():
            try:
                with open(plan_path, "r", encoding="utf-8") as f:
                    result["execution_plan"] = f.read()
            except Exception:
                result["execution_plan"] = "执行计划文件读取失败"

        out_dir = run_dir / "outputs"
        if out_dir.exists():
            try:
                result["outputs"] = [f.name for f in out_dir.iterdir() if f.is_file()]
            except Exception:
                pass

        # 防御性回退：如果 DB 中 dag_blueprint 为空但磁盘文件存在，从磁盘读取
        if not result.get("dag_blueprint"):
            dag_file = run_dir / "dag_blueprint.json"
            if dag_file.exists():
                try:
                    with open(dag_file, "r", encoding="utf-8") as f:
                        result["dag_blueprint"] = json.load(f)
                except Exception:
                    pass

    if record.status == "COMPILING":
        result["compile_progress"] = "编译中，请稍后..."
        result["poll_interval"] = 2
    elif record.status == "PENDING_REVIEW":
        result["compile_progress"] = "编译完成"
    elif record.status == "FAILED":
        result["compile_progress"] = "编译失败"
        # 提取真实错误信息
        logs = record.execution_logs or []
        result["error_msg"] = logs[-1] if logs else (
            record.sandbox_code and "沙箱执行失败，请检查算子逻辑" or "编译失败，请检查 DAG 蓝图"
        )

    return result


# ── 3. 获取 DAG 蓝图 ──────────────────────────────────
@router.get("/runs/{run_id}/dag", summary="获取 DAG 蓝图")
async def get_dag(run_id: str) -> Dict[str, Any]:
    record = _get_snapshot_mgr().get_run(run_id)
    if not record:
        raise HTTPException(status_code=404, detail="Run 不存在")
    if record.status == "COMPILING":
        raise HTTPException(status_code=425, detail="DAG 正在编译中")
    if not record.dag_blueprint:
        # 防御性回退：从磁盘文件读取
        dag_file = record.run_dir / "dag_blueprint.json"
        if dag_file.exists():
            try:
                with open(dag_file, "r", encoding="utf-8") as f:
                    return json.load(f)
            except Exception:
                pass
        raise HTTPException(status_code=404, detail="DAG 蓝图不存在")
    return record.dag_blueprint


# ── 4. 触发执行 ───────────────────────────────────────
@router.post("/runs/{run_id}/execute", summary="确认 DAG 并触发执行")
async def execute_run(
        run_id: str,
        request: ExecuteRequest,
        background_tasks: BackgroundTasks,
) -> StatusResponse:
    record = _get_snapshot_mgr().get_run(run_id)
    if not record:
        raise HTTPException(status_code=404, detail="Run 不存在")
    if record.status == "COMPILING":
        raise HTTPException(status_code=425, detail="DAG 正在编译中")
    if record.status not in ("PENDING_REVIEW", "QUEUED"):
        raise HTTPException(status_code=400, detail=f"当前状态 {record.status} 不允许执行")

    _get_snapshot_mgr().update_status(run_id, "RUNNING")

    try:
        dag_source = record.dag_blueprint or {}
        code = _dag_to_python(dag_source, record)
        _write_audit_trace(run_id, record.user_intent or "", dag_source)
        _append_approval_hash(run_id, record.user_intent or "", dag_source, request.confirmed)
    except Exception as e:
        _get_snapshot_mgr().update_status(run_id, "FAILED")
        raise HTTPException(status_code=500, detail=f"底层算子编译失败: {str(e)}")

    code_path = record.run_dir / "sandbox_code.py"
    with open(code_path, "w", encoding="utf-8") as f:
        f.write(code)

    # 同步写入数据库，方便前端展示真实错误
    _get_snapshot_mgr().update_sandbox_code(run_id, code)

    background_tasks.add_task(
        _execute_in_sandbox,
        run_id=run_id,
        code=code,
        run_dir=record.run_dir,
    )

    return StatusResponse(
        run_id=run_id, status="RUNNING", progress=0,
        current_step="正在启动 Docker 沙箱...", elapsed_seconds=0.0,
        retry_count=0, output_files=[],
    )


# ── 4b. 确认关键词提案 ────────────────────────────────

class KeywordConfirmBody(BaseModel):
    action: str = "approve"
    patterns: str = ""
    category: str = ""
    approved_by: str = "审计师"


@router.post("/runs/{run_id}/confirm_keywords", summary="确认/修订关键词提案")
async def confirm_keywords(run_id: str, body: KeywordConfirmBody, background_tasks: BackgroundTasks):
    from core.keyword_resolver import (
        get_proposal, clear_proposal, approve_and_intake,
        backtest_patterns as _kw_backtest,
    )
    record = _get_snapshot_mgr().get_run(run_id)
    if not record:
        raise HTTPException(status_code=404, detail="Run 不存在")
    if record.status != "PENDING_KEYWORD_CONFIRM":
        raise HTTPException(status_code=400, detail=f"当前状态 {record.status} 不允许确认关键词")
    proposal = get_proposal(run_id)
    final_patterns = body.patterns or (proposal.get("patterns", "") if proposal else "")
    if not final_patterns:
        raise HTTPException(status_code=400, detail="关键词不能为空")
    preview = {}
    try:
        import pandas as _pd3
        input_dir = record.input_dir
        input_files = list(input_dir.iterdir()) if input_dir.exists() else []
        bank_files = [f for f in input_files if f.is_file() and f.suffix.lower() in (".xlsx", ".xls", ".csv")]
        if bank_files:
            _sample = _pd3.read_excel(bank_files[0], nrows=5000) if bank_files[0].suffix.lower() != ".csv" else _pd3.read_csv(bank_files[0], nrows=5000, encoding="utf-8-sig")
            preview = _kw_backtest(final_patterns, _sample, ["摘要", "对方客户名称", "附言", "用途"])
    except Exception as _be:
        print(f"[confirm_keywords] backtest 失败: {_be}")
    kw_source, kw_version = "", ""
    if body.action == "approve":
        category = body.category or (proposal.get("category", "用户提案") if proposal else "用户提案")
        meta = proposal or {}
        kw_version = approve_and_intake(category, final_patterns, meta, body.approved_by)
        kw_source = "dictionary_" + kw_version
    else:
        category = body.category or "用户修订"
        meta = proposal or {}
        meta["依据摘要"] = "用户手动修订关键词"
        kw_version = approve_and_intake(category, final_patterns, meta, body.approved_by)
        kw_source = "user_approved@" + _now_date()
    clear_proposal(run_id)
    _get_snapshot_mgr().update_status(run_id, "RUNNING")
    background_tasks.add_task(_execute_keyword_confirmed, run_id=run_id, record=record, final_patterns=final_patterns, kw_source=kw_source, kw_version=kw_version, preview=preview)
    return {"run_id": run_id, "status": "RUNNING", "patterns": final_patterns, "kw_source": kw_source, "kw_version": kw_version, "preview": preview}


def _now_date() -> str:
    from datetime import date
    return date.today().isoformat()


async def _execute_keyword_confirmed(run_id, record, final_patterns, kw_source, kw_version, preview):
    from core.matching_engine import run_matching_pipeline
    logs = []
    try:
        input_dir = record.input_dir
        output_dir = record.run_dir / "outputs"
        output_dir.mkdir(parents=True, exist_ok=True)
        print(f"[Standalone] Run {run_id} 确认后执行: pattern={final_patterns[:60]} source={kw_source}")
        match_result = run_matching_pipeline(input_dir, output_dir, patterns=final_patterns, kw_source=kw_source, kw_version=kw_version)
        logs.append("[匹配引擎] 匹配流水线执行完成")
        output_files_final = [f.name for f in output_dir.iterdir() if f.is_file()]
        if output_files_final:
            _get_snapshot_mgr().update_outputs(run_id=run_id, output_files=output_files_final, validation_results=[{"check": "matching_engine", "passed": True}], all_passed=True)
        dag_ops = (record.dag_blueprint or {}).get("operators", [])
        input_names = [f.name for f in input_dir.iterdir() if f.is_file()]
        from core.report_generator import generate_audit_report
        dag_bp = record.dag_blueprint or {}
        explanation = dag_bp.get("match_explanation", "") if isinstance(dag_bp, dict) else ""
        engine_match_logic = match_result.get("match_logic", {}) if match_result else {}
        match_info = {"patterns": final_patterns, "columns": engine_match_logic.get("筛选列", []), "kw_source": kw_source, "kw_preview": preview or {}, "method": match_result.get("strategy_name", "多列联合匹配") if match_result else "多列联合匹配", "explanation": explanation}
        rp = generate_audit_report(run_id=run_id, user_intent=record.user_intent or "", dag_operators=dag_ops, output_dir=output_dir, input_files=input_names, execution_logs=logs, match_logic=match_info)
        output_files_final.append(rp.name)
        _get_snapshot_mgr().update_outputs(run_id=run_id, output_files=output_files_final, validation_results=[{"check": "matching_engine", "passed": True}], all_passed=True)
        _get_snapshot_mgr().update_status(run_id, "COMPLETED")
        print(f"[Standalone] Run {run_id} 完成")
    except Exception as e:
        import traceback
        print(f"[Standalone] Run {run_id} 失败: {traceback.format_exc()}")
        _get_snapshot_mgr().update_status(run_id, "FAILED")



# ── 5. 查询执行状态 ───────────────────────────────────
@router.get("/runs/{run_id}/status", summary="查询执行状态")
async def get_execution_status(run_id: str) -> StatusResponse:
    record = _get_snapshot_mgr().get_run(run_id)
    if not record:
        raise HTTPException(status_code=404, detail="Run 不存在")

    output_files = record.output_files or []
    output_dir = record.run_dir / "outputs"
    if output_dir.exists():
        output_files = [str(f.name) for f in output_dir.iterdir() if f.is_file()]

    return StatusResponse(
        run_id=run_id, status=record.status,
        progress=50 if record.status == "RUNNING" else (
            100 if record.status in ("COMPLETED", "FAILED", "SUCCESS") else 0),
        current_step=record.status, elapsed_seconds=0.0,
        retry_count=record.retry_count, output_files=output_files,
    )


@router.get("/projects/{project_code}/tree", summary="获取版本树")
async def get_version_tree(project_code: str, subject: Optional[str] = None) -> List[Dict]:
    mgr = _get_snapshot_mgr()
    if subject:
        records = mgr.get_version_tree(project_code, subject)
    else:
        records = mgr.get_recent_runs(limit=100)
        records = [r for r in records if r.project_code == project_code]
    return [r.to_dict() for r in records]


@router.get("/runs/{run_id}/download", summary="异步打包下载")
async def download_run(run_id: str) -> DownloadResponse:
    record = _get_snapshot_mgr().get_run(run_id)
    if not record:
        raise HTTPException(status_code=404, detail="Run 不存在")

    task_id = f"dl_{run_id}_{uuid.uuid4().hex[:6]}"
    output_dir = record.run_dir / "outputs"

    if not output_dir.exists():
        raise HTTPException(status_code=404, detail="成果物不存在")

    download_dir = Path("data/downloads")
    download_dir.mkdir(parents=True, exist_ok=True)
    zip_path = download_dir / f"{task_id}.zip"

    with zipfile.ZipFile(str(zip_path), 'w', zipfile.ZIP_DEFLATED) as zf:
        for file_path in output_dir.iterdir():
            if file_path.is_file():
                zf.write(str(file_path), arcname=file_path.name)
        dag_path = record.run_dir / "dag_blueprint.json"
        if dag_path.exists():
            zf.write(str(dag_path), arcname="dag_blueprint.json")
        plan_path = record.run_dir / "execution_plan.txt"
        if plan_path.exists():
            zf.write(str(plan_path), arcname="execution_plan.txt")

    return DownloadResponse(task_id=task_id, status="COMPLETED", message=f"打包完成: {zip_path.name}")


@router.get("/download/status/{task_id}", summary="查询下载状态")
async def get_download_status(task_id: str) -> Dict[str, Any]:
    zip_path = Path("data/downloads") / f"{task_id}.zip"
    if zip_path.exists():
        return {
            "task_id": task_id, "status": "COMPLETED",
            "download_url": f"/api/v3/download/file/{task_id}.zip",
            "file_size": zip_path.stat().st_size,
        }
    return {"task_id": task_id, "status": "PROCESSING", "message": "打包中..."}


@router.get("/download/file/{filename:path}", summary="下载成果物文件")
async def download_file(filename: str):
    from fastapi.responses import FileResponse
    file_path = Path("data/downloads") / filename
    if not file_path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")
    return FileResponse(str(file_path), media_type="application/zip", filename=filename)


# ═══════════════════════════════════════════════════════════════
# 内部函数
# ═══════════════════════════════════════════════════════════════

# ── 上传安全校验 ──────────────────────────────────────

def _validate_upload(file: UploadFile) -> str:
    """
    安全校验上传文件，返回清洗后的安全文件名。
    
    校验规则：
    1. 去路径遍历字符（仅保留纯文件名）
    2. 扩展名白名单
    3. 文件大小限制
    """
    # 1. 文件名去路径遍历
    safe_name = Path(file.filename).name if file.filename else ""
    if not safe_name:
        raise HTTPException(status_code=400, detail="文件名为空")
    if safe_name != file.filename:
        raise HTTPException(status_code=400, detail="文件名包含非法路径字符")

    # 2. 扩展名白名单
    ext = Path(safe_name).suffix.lower()
    if ext not in ALLOWED_EXTENSIONS:
        raise HTTPException(status_code=400,
                           detail=f"不支持的文件类型: {ext}，允许的类型: {', '.join(ALLOWED_EXTENSIONS)}")

    # 3. 内容大小校验
    if hasattr(file, "size") and file.size is not None and file.size > MAX_UPLOAD_SIZE:
        raise HTTPException(status_code=400,
                           detail=f"文件过大（{file.size} 字节），最大允许 {MAX_UPLOAD_SIZE // 1024 // 1024}MB")

    return safe_name


# ── 格式规范化 API ──────────────────────────────────

@router.post("/format/normalize", summary="格式规范化：按模板统一排版")
async def normalize_document_format(
    files: List[UploadFile] = File(..., description="待排版文件 + 可选模板文件"),
    template_index: Optional[int] = Form(None, description="指定第几个文件为模板（0-based），不指定则使用内部默认模板"),
    output_format: Optional[str] = Form("docx", description="输出格式: docx / xlsx / auto"),
) -> Dict[str, Any]:
    """格式规范化：以模板为准，批量转换其他文件的格式（含打印设置）"""
    if not files:
        raise HTTPException(status_code=400, detail="请至少上传1个文件")

    from core.run_snapshot import TEMP_BASE
    TEMP_BASE.mkdir(parents=True, exist_ok=True)
    temp_dir = TEMP_BASE / f"fmt_{uuid.uuid4().hex[:8]}"
    temp_dir.mkdir(parents=True, exist_ok=True)

    saved_files = []
    for f in files:
        safe_name = _validate_upload(f)
        save_path = temp_dir / safe_name
        content = await f.read()
        if len(content) > MAX_UPLOAD_SIZE:
            raise HTTPException(status_code=400, detail=f"文件 {safe_name} 过大")
        with open(save_path, "wb") as fout:
            fout.write(content)
        saved_files.append(str(save_path))

    # 确定模板：用户指定 > 内部模板 > 第一个文件推测
    if template_index is not None and 0 <= template_index < len(saved_files):
        template_path = saved_files[template_index]
        targets = [p for i, p in enumerate(saved_files) if i != template_index]
        template_source = f"用户上传（第{template_index + 1}个文件）"
    else:
        # 用内部默认模板
        first_ext = Path(saved_files[0]).suffix.lower()
        if first_ext in (".xlsx", ".xlsm", ".xls"):
            internal = list((TEMPLATES_DIR / "excel").glob("*.xlsx"))
        else:
            internal = list((TEMPLATES_DIR / "word").glob("*.docx"))
        if internal:
            template_path = str(internal[0])
            targets = saved_files
            template_source = f"内部默认模板（{Path(template_path).name}）"
        else:
            # 回退：用第一个文件当模板
            template_path = saved_files[0]
            targets = saved_files[1:]
            template_source = "自动（第一个文件）"

    if not targets:
        return {"status": "error", "message": "没有需要转换的目标文件（至少需要模板之外的1个文件）",
                "template": template_source}

    output_dir = str(temp_dir / "formatted")
    try:
        results = normalize_format(template_path, targets, output_dir)
        # 将结果复制到 downloads 目录供下载
        import shutil
        dl_dir = Path("data/downloads")
        dl_dir.mkdir(parents=True, exist_ok=True)
        dl_name = f"fmt_{uuid.uuid4().hex[:6]}.zip"
        dl_path = dl_dir / dl_name
        import zipfile
        with zipfile.ZipFile(str(dl_path), "w", zipfile.ZIP_DEFLATED) as zf:
            for r in results:
                zf.write(r, Path(r).name)
        return {
            "status": "success",
            "template": template_source,
            "converted_count": len(results),
            "converted_files": [Path(r).name for r in results],
            "download_url": f"/api/v3/download/file/{dl_name}",
        }
    except Exception as e:
        return {"status": "error", "message": f"格式规范化失败: {str(e)}"}




# ── 代码安全清洗 ──────────────────────────────────────

def _sanitize_code_param(value: str, max_len: int = 200) -> str:
    """
    安全清洗拼入沙箱代码的字符串参数值。
    
    规则：
    - 转义反斜杠和单引号（防止字符串逃逸）
    - 截断超长值（防止缓冲区类攻击）
    - 过滤控制字符（保留常用空白符 \\t \\n \\r）
    """
    if not isinstance(value, str):
        value = str(value)
    # 转义反斜杠必须先于单引号
    value = value.replace("\\", "\\\\")
    value = value.replace("'", "\\'")
    # 截断
    if len(value) > max_len:
        value = value[:max_len]
    # 过滤控制字符（保留常用的空白符 0x09=TAB, 0x0A=LF, 0x0D=CR）
    value = _re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', value)
    return value


# ── DAG JSON 大小校验 ────────────────────────────────

def _validate_dag_json_size(dag_json: str) -> None:
    """校验 DAG JSON 大小，防止超大数据注入"""
    if len(dag_json.encode('utf-8')) > MAX_DAG_JSON_SIZE:
        raise ValueError(f"DAG JSON 过大（>{MAX_DAG_JSON_SIZE}字节），可能存在注入风险")


# ── 后台编译 ─────────────────────────────────────────

async def _call_dify_compiler(catalog: AssetCatalog, user_intent: str, preset_button: Optional[str],
                              parent_run_id: Optional[str]) -> Any:
    catalog_text = _format_catalog_for_prompt(catalog)

    # 🛡️ 隐私防火墙：数据出本地前脱敏
    fw = get_firewall()
    catalog_text, _ = fw.sanitize(catalog_text)
    user_intent, _ = fw.sanitize(user_intent)

    # 📚 RAG 合规注入：检索相关法规片段
    compliance_context = inject_compliance_context(user_intent)
    if compliance_context:
        catalog_text = compliance_context + "\n\n" + catalog_text

    # 🌐 联网搜索触发器：检测是否需要最新法规/政策信息（在已有事件循环中异步执行）
    try:
        from core.search_trigger import search_and_inject
        search_context = await search_and_inject(user_intent)
        if search_context:
            catalog_text = search_context + "\n\n" + catalog_text
            print(f"[编译] 已注入联网搜索结果")
    except Exception as e:
        print(f"[编译] 搜索触发器降级（非致命）: {e}")

    # 🧠 领域知识注入：审计最佳实践
    catalog_text = AUDIT_DOMAIN_KNOWLEDGE + "\n\n" + catalog_text

    # 📖 Few-shot 示例注入（B3 动态版）：按意图关键词选 top-3 相关范例；
    # 未命中关键词时自动回退全量 7 条（行为兼容旧版）
    from config.few_shot_examples import build_dynamic_few_shot
    few_shot_text = build_dynamic_few_shot(user_intent, max_examples=3)
    catalog_text = few_shot_text + "\n\n" + catalog_text

    # 📖 合成 Few-shot 注入（50条对账场景）：从 JSONL 池选匹配案例
    try:
        from config.few_shot_examples import build_synthetic_few_shot
        synth_text = build_synthetic_few_shot(user_intent or "", max_examples=3)
        if synth_text:
            catalog_text = synth_text + "\n\n" + catalog_text
    except Exception as e:
        print(f"[编译] 合成 Few-shot 降级（非致命）: {e}")

    # 🎯 场景知识包：规划检查单注入（RAG 知识 → 规划约束，覆盖逐笔/汇总/提取式等场景）
    try:
        from config.scenario_packs import detect_scenario as _detect_scn, build_scenario_prompt
        _scn = _detect_scn(user_intent or "")
        _scn_prompt = build_scenario_prompt(_scn)
        if _scn_prompt:
            catalog_text = _scn_prompt + "\n\n" + catalog_text
            print(f"[编译] 场景知识包: {_scn}")
    except Exception as e:
        print(f"[编译] 场景知识包降级（非致命）: {e}")

    _MAX_CATALOG_CHARS = 35000
    if len(catalog_text) > _MAX_CATALOG_CHARS:
        _truncated = catalog_text[:_MAX_CATALOG_CHARS]
        _last_nl = _truncated.rfind("\n")
        if _last_nl > _MAX_CATALOG_CHARS * 0.8:
            _truncated = _truncated[:_last_nl]
        catalog_text = _truncated + "\n\n[提示：数据目录过长已自动截断]"
        print(f"[编译] Token 预算保护")

    parent_summary = ""
    if parent_run_id:
        parent_data = _get_snapshot_mgr().extract_summary_for_inheritance(parent_run_id)
        if parent_data:
            parent_summary = (
                f"\n## 前序操作摘要\n"
                f"前序意图: {parent_data.get('previous_intent', '')}\n"
                f"前序目标: {parent_data.get('previous_objective', '')}\n"
                f"前序算子: {parent_data.get('previous_operators', '')}\n"
            )

    # 幻觉校验基准：Data Catalog 中的真实文件名
    _known_files = [f["filename"] for f in catalog.files] if catalog and catalog.files else []

    try:
        client = _get_http_client()
        response = await client.post(
            f"{DIFY_BASE_URL}/v1/workflows/run",
            headers={"Authorization": f"Bearer {DIFY_API_KEY}"},
            json={
                "inputs": {
                    "catalog_text": catalog_text, "user_intent": user_intent,
                    "preset_button": preset_button or "", "parent_summary": parent_summary,
                },
                "response_mode": "blocking", "user": "audit_platform",
            },
        )
        response.raise_for_status()
        result = response.json()
        print(f"DEBUG Dify response keys: {list(result.keys())}")
        if "data" in result:
            print(f"DEBUG data keys: {list(result['data'].keys())}")
            if "outputs" in result["data"]:
                print(f"DEBUG outputs keys: {list(result['data']['outputs'].keys())}")
                print(f"DEBUG dag_json value (first 200): {str(result['data']['outputs'].get('dag_json', ''))[:200]}")
        dag_json_str = result.get("data", {}).get("outputs", {}).get("dag_json", "")

        if dag_json_str:
            # --- 🛡️ 防御性 Patch：解析前强行注入 source_file ---
            try:
                dag_data = json.loads(dag_json_str)
                # 获取第一个文件的名称作为默认值（如果有文件）
                default_file = catalog.files[0]["filename"] if catalog.files else "input.xlsx"

                if "operators" in dag_data:
                    for op in dag_data["operators"]:
                        # 规范化：Dify 有时用 "operator" / "type" 代替 "name"
                        if "operator" in op and "name" not in op:
                            op["name"] = op.pop("operator")
                        if "name" not in op and op.get("type"):
                            op["name"] = op["type"]
                        # 规范化：output → output_alias
                        if "output" in op and "output_alias" not in op:
                            op["output_alias"] = op.pop("output")
                        # 规范化：input_from 字符串→列表
                        if "input_from" in op and isinstance(op["input_from"], str):
                            op["input_from"] = [op["input_from"]]
                        # Load 算子注入 source_file
                        if op.get("name") in ("Load", "load"):
                            if not op.get("source_file"):
                                op["source_file"] = op.get("file")
                            if not op.get("source_file"):
                                op["source_file"] = catalog.files[0]["filename"] if catalog.files else "input.xlsx"
                dag_json_str = json.dumps(dag_data)

            except Exception as e:
                print(f"JSON 预处理 Patch 失败: {e}")
            # ---------------------------------------------

            try:
                dag_for_return = DAGParser.parse(dag_json_str, known_files=_known_files or None)
            except ValueError:
                # 幻觉熔断：确定性规则修复（difflib 就近纠正文件名）后重校验，
                # 仍失败则抛错进入上级熔断链（DAG 精修 → vLLM 直连 → 本地兜底）
                from core.dag_compiler import rule_fix_dag
                _fixed = rule_fix_dag(dag_json_str, known_files=_known_files or None)
                if _fixed:
                    dag_for_return = DAGParser.parse(_fixed, known_files=_known_files or None)
                else:
                    raise
            return dag_for_return

    except Exception as e:
        print(f"Dify 主链路失败: {e}")
        # 尝试 vLLM 降级；vLLM 也不可用时返回 None，由上层 _build_fallback_dag 接管
        try:
            return await _fallback_compiler(catalog_text, user_intent, preset_button)
        except Exception as e2:
            print(f"vLLM 降级也失败: {e2}，交由本地兜底 DAG 接管")
            return None


def _extract_patterns_from_dag(dag_blueprint: dict) -> str:
    """从 DAG 蓝图的 RegexFilter 算子中提取筛选关键词"""
    if not dag_blueprint:
        return ""
    ops = dag_blueprint.get("operators", [])
    for op in ops:
        name = op.get("name", "").lower()
        if name in ("regexfilter", "regex_filter", "regex", "filter"):
            pattern = op.get("params", {}).get("pattern", "")
            if pattern:
                return pattern
    # 兜底：从第一个 Load 的 source_file 和常见关键词推断
    return ""


async def _explain_dag(blueprint: dict, catalog: Any, user_intent: str) -> str:
    """用 vLLM 生成匹配逻辑的自然语言说明"""
    ops = blueprint.get("operators", [])
    if not ops:
        return ""
    ops_text = "\n".join(
        f"{i+1}. {op.get('name')}: params={json.dumps(op.get('params',{}), ensure_ascii=False)[:300]}" 
        for i, op in enumerate(ops)
    )
    # 提取实际的筛选关键词
    patterns = _extract_patterns_from_dag(blueprint) or "（未指定具体关键词）"
    prompt = f"""你是审计助手。请用自然语言向审计师解释以下匹配逻辑：

用户意图：{user_intent}

执行步骤（含参数）：
{ops_text}

实际筛选关键词：{patterns}

请用3-5句详细说明：
1) 平台从哪些文件加载数据
2) 用什么关键词或规则筛选（列出具体关键词）
3) 按什么维度汇总和比对
4) 匹配策略流程
直接回复说明，不要JSON。"""
    try:
        client = _get_http_client()
        response = await client.post(
            VLLM_TUNNEL_URL,
            headers={"Authorization": f"Bearer {VLLM_API_KEY}"},
            json={"model": VLLM_MODEL, "messages": [{"role": "user", "content": prompt}], "temperature": 0.3, "max_tokens": 2048},
        )
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"].strip()
    except Exception:
        # vLLM 不可用时用模板生成（增强版：遍历所有算子生成详细描述）
        names = [op.get("name") for op in ops]
        files = [op.get("source_file", "") for op in ops if op.get("source_file")]
        unique_files = list(dict.fromkeys(files))  # 去重保持顺序

        parts = []
        # 文件加载
        if unique_files:
            file_list = "、".join([f"「{f}」" for f in unique_files])
            parts.append(f"平台将加载 {len(unique_files)} 个文件：{file_list}")
        else:
            parts.append(f"平台将加载 {len(files)} 个文件")

        # 遍历每种算子生成具体描述
        for op in ops:
            name = op.get("name", "")
            params = op.get("params", {})
            if name == "NoiseFilter":
                cols = params.get("columns", params.get("column", []))
                noises = params.get("noise_patterns", [])
                col_str = "、".join(cols) if isinstance(cols, list) else str(cols)
                noise_str = "、".join(noises) if isinstance(noises, list) else str(noises)
                parts.append(f"过滤「{col_str}」列中的噪音：{noise_str}")
            elif name == "ColumnFilter":
                col = params.get("column", "")
                op_val = params.get("operator", "=")
                val = params.get("value", "")
                parts.append(f"筛选「{col}」{op_val}「{val}」的记录")
            elif name == "RegexFilter":
                pattern = params.get("pattern", "")
                parts.append(f"通过正则「{pattern}」筛选相关记录")
            elif name == "Sort":
                by = params.get("by", [])
                by_str = "、".join(by) if isinstance(by, list) else str(by)
                parts.append(f"按「{by_str}」排序")
            elif name == "Merge":
                how = params.get("how", "outer")
                left_on = params.get("left_on", [])
                right_on = params.get("right_on", [])
                l_str = "+".join(left_on) if isinstance(left_on, list) else str(left_on)
                r_str = "+".join(right_on) if isinstance(right_on, list) else str(right_on)
                parts.append(f"以「{l_str}」↔「{r_str}」{how} 连接两表")
            elif name == "Diff":
                col_a = params.get("col_a", "")
                col_b = params.get("col_b", "")
                tol = params.get("tolerance_abs", params.get("tolerance", 0.01))
                parts.append(f"比对「{col_a}」与「{col_b}」，容差 ±{tol}")
            elif name == "Reconcile":
                tol = params.get("tolerance_abs", 0.01)
                window = params.get("date_window_days", "")
                parts.append(f"逐笔对账匹配，容差 ±{tol}" + (f"，日期窗口 ±{window}天" if window else ""))
            elif name == "Aggregate":
                group_by = params.get("group_by", params.get("by", []))
                gb_str = "、".join(group_by) if isinstance(group_by, list) else str(group_by)
                parts.append(f"按「{gb_str}」汇总统计")
            elif name == "GroupBy":
                by = params.get("by", params.get("columns", []))
                by_str = "、".join(by) if isinstance(by, list) else str(by)
                agg = params.get("agg", "count")
                parts.append(f"按「{by_str}」分组，聚合方式：{agg}")

        # 导出
        export_files = [op.get("params", {}).get("filename", "") for op in ops if op.get("name") == "Export"]
        if export_files:
            parts.append("最后导出「" + "、".join(export_files) + "」")
        else:
            parts.append("最后导出匹配结果和审计报告")

        return "。".join(parts) + "。"


def _generate_fallback_plan(run_dir, user_intent, preset_button, catalog):
    """Dify 不可用时从场景注册表生成匹配计划"""
    try:
        from config.scenario_packs import detect_scenario, SCENARIO_PACKS
        from config.presets import normalize_preset_key

        sid = normalize_preset_key(preset_button) if preset_button else None
        if not sid:
            sid = detect_scenario(user_intent or "", ask_user=False)
        # 兼容旧场景名（中文→英文key）
        _name_map = {"银行对账":"bank_reconcile_detail","数据比对":"summary_compare",
                     "提取式核对":"filtered_extraction_match","大额交易筛查":"large_txn_screen",
                     "智能筛选":"single_table_analysis","文档生成":"doc_generation",
                     "跨文件对比":"cross_doc_compare"}
        sid = _name_map.get(sid, sid)
        print(f"[匹配计划] 场景={sid} 存在={sid in SCENARIO_PACKS}")
        pack = SCENARIO_PACKS.get(sid, {})

        files = [f["filename"] for f in catalog.files] if catalog and catalog.files else []
        plan = f"## {pack.get('name', '数据处理')}\n\n"
        plan += f"### 数据源\n"
        for fn in files:
            plan += f"- {fn}\n"
        plan += f"\n### 检查单\n"
        for c in pack.get("checklist", [])[:5]:
            plan += f"- {c}\n"
        plan += f"\n### 容差纪律\n{pack.get('tolerance_rule', '不适用')}\n"
        plan += f"\n### 应交付\n{'、'.join(pack.get('deliverables', []))}\n"
        plan += f"\n> 当前 Dify 编译服务不可用，以上为场景默认匹配逻辑，执行时将按此标准进行。"

        (run_dir / "match_explanation.txt").write_text(plan, encoding="utf-8")
    except Exception as e:
        print(f"[匹配计划] 生成失败: {e}")


async def _fallback_compiler(catalog_text: str, user_intent: str, preset_button: Optional[str]) -> Any:
    # 预设按钮优先（别名归一），未指定则按意图推断场景
    scenario = None
    if preset_button:
        try:
            from config.presets import normalize_preset_key
            scenario = normalize_preset_key(preset_button)
        except Exception:
            scenario = None
    if not scenario:
        scenario = detect_scenario(user_intent)
    system_prompt = get_fallback_prompt(scenario)
    print(f"[降级编译] 场景: {scenario}（preset_button={preset_button or '无'}）")
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"## 数据目录\n{catalog_text}\n\n## 审计意图\n{user_intent}"},
    ]

    client = _get_http_client()
    response = await client.post(
        VLLM_TUNNEL_URL,
        headers={"Authorization": f"Bearer {VLLM_API_KEY}"},
        json={
            "model": VLLM_MODEL, "messages": messages,
            "temperature": 0.3, "max_tokens": 4096,
        },
    )
    if response.status_code != 200:
        print(f"[vLLM] 错误响应 ({response.status_code}): {response.text[:500]}")
    response.raise_for_status()
    raw_content = response.json()["choices"][0]["message"]["content"]

    import re
    json_match = re.search(r'```(?:json)?\s*(\{[\s\S]*?\})\s*```', raw_content)
    if json_match:
        dag_json = json_match.group(1)
    else:
        json_match = re.search(r'(\{[\s\S]*?("operators"|"nodes")[\s\S]*?\})', raw_content)
        if json_match:
            dag_json = json_match.group(1)
        else:
            raise ValueError("无法从 vLLM 响应中提取 JSON。")

    # 🔒 安全校验：限制 DAG JSON 大小
    _validate_dag_json_size(dag_json)

    # 🛡️ 防御性 Patch：为缺少 source_file 的 Load 算子自动注入
    try:
        dag_data = json.loads(dag_json)
        default_file = "input.xlsx"
        if "operators" in dag_data:
            for op in dag_data["operators"]:
                if op.get("name") in ("Load", "load") or op.get("type") in ("Load", "load"):
                    if not op.get("source_file"):
                        op["source_file"] = op.get("file") or default_file
            dag_json = json.dumps(dag_data)
    except Exception:
        pass

    # 幻觉校验基准：从 catalog_text 还原真实文件名（"文件: xxx" 行）
    _kf = re.findall(r"^文件[:：]\s*(.+?)\s*$", catalog_text, flags=re.M)
    try:
        return DAGParser.parse(dag_json, known_files=_kf or None)
    except Exception as e:
        from core.dag_compiler import rule_fix_dag
        _fixed = rule_fix_dag(dag_json, known_files=_kf or None)
        if _fixed:
            try:
                return DAGParser.parse(_fixed, known_files=_kf or None)
            except Exception:
                pass
        raise ValueError(f"DAG 解析失败: {e}")


def _build_fallback_dag(catalog: Any, user_intent: str) -> Any:
    """最终兜底：从 catalog 生成最小 DAG（不依赖任何 LLM）"""
    from core.dag_compiler import DAGBlueprint, Operator

    files = getattr(catalog, "files", [])
    if not files:
        return None

    scenario = detect_scenario(user_intent)
    ops = []

    # 每个文件一个 Load
    for i, f in enumerate(files):
        fname = f.get("filename", f"file_{i}")
        ops.append(Operator(
            id=f"load_{i}", name="Load", description=f"读取 {fname}",
            input_from=[], source_file=fname, params={"file_path": fname},
            output_alias=f"df_load_{i}",
        ))

    # 根据场景加算子
    if "匹配" in user_intent or "核对" in user_intent:
        ops.append(Operator(
            id="merge_0", name="Merge", description="数据合并",
            input_from=[f"load_{i}" for i in range(len(files))],
            params={"how": "outer"},
            output_alias="df_merge_0",
        ))
        ops.append(Operator(
            id="export_0", name="Export", description="导出匹配结果",
            input_from=["merge_0"],
            params={"filename": "match_result.xlsx"},
            output_alias="df_export_0",
        ))
    else:
        ops.append(Operator(
            id="export_0", name="Export", description="导出结果",
            input_from=[f"load_{i}" for i in range(len(files))],
            params={"filename": "analysis_result.csv"},
            output_alias="df_export_0",
        ))

    return DAGBlueprint(
        blueprint_id="fallback_001",
        generated_at=datetime.now().isoformat(),
        operators=ops,
        objective=f"兜底编译: {user_intent[:50]}",
        raw_intent=user_intent,
    )


def _format_catalog_for_prompt(catalog: AssetCatalog) -> str:
    """Data Catalog → LLM 提示文本。

    富化内容（让 LLM 看着真实语义规划，而不是照抄示例幻觉）：
    - 每文件：列名(dtype) + 语义角色（确定性识别）+ 文档文本预览/内嵌表格列
    - 跨表：连接键建议（已排除序号/行号等无意义键）
    - 顶部硬约束：source_file 与列名必须来自本目录
    """
    lines = [
        f"文件总数: {catalog.total_files}",
        "【硬约束】Load 的 source_file 和算子引用的列名必须逐字来自下方文件清单，"
        "严禁使用任何示例中的文件名/列名；严禁把'序号/编号/行号'用作连接键。",
        "=== 文件清单 ===",
    ]
    # 语义角色（列名驱动，确定性；与实际业务含义不符时 LLM 可按需修正并说明）
    try:
        import pandas as _pd
        from core.column_semantics import (detect_column_roles,
                                           is_meaningless_key, suggest_join_keys)
        _sem_ok = True
    except Exception:
        _sem_ok = False
    frames = {}
    for f in catalog.files:
        lines.append(f"\n文件: {f['filename']}")
        if f.get("kind") and f["kind"] != "table":
            lines.append(f"  类型: {f['kind']}")
        if "columns" in f and f["columns"]:
            cols = [f"{col['name']}({col['dtype']})" for col in f["columns"]]
            lines.append(f"  列: {', '.join(cols)}")
            if _sem_ok:
                try:
                    _df = _pd.DataFrame(
                        columns=[str(col["name"]) for col in f["columns"]])
                    frames[f["filename"]] = _df
                    roles = detect_column_roles(_df)
                    if roles:
                        lines.append("  语义角色: " + ", ".join(
                            f"{r}→{c}" for r, c in roles.items()))
                    mk = [c for c in _df.columns if is_meaningless_key(c)]
                    if mk:
                        lines.append(f"  无意义键（禁止连接）: {', '.join(mk)}")
                except Exception:
                    pass
        if f.get("text_preview"):
            lines.append(f"  文本预览: {str(f['text_preview'])[:200]}")
        if f.get("tables_columns") and f["tables_columns"]:
            lines.append(f"  文档内嵌表格列: {', '.join(str(c) for c in f['tables_columns'][0])}")
    # 跨表连接键建议（两两组合，最多 5 组）
    if _sem_ok and len(frames) >= 2:
        try:
            names = list(frames.keys())
            hints = []
            for i in range(len(names)):
                for j in range(i + 1, len(names)):
                    keys = suggest_join_keys(frames[names[i]], frames[names[j]])
                    if keys:
                        hints.append(f"{names[i]} × {names[j]}: "
                                     + ", ".join(f"{a}↔{b}" for a, b in keys))
                if len(hints) >= 5:
                    break
            if hints:
                lines.append("\n=== 跨表连接键建议（供参考） ===")
                lines.extend(hints[:5])
        except Exception:
            pass
    # 跨表列名映射提示（序时账×银行流水时注入方向镜像规则）
    if _sem_ok and len(frames) >= 2:
        try:
            from core.column_semantics import detect_column_roles
            _all_roles = {}
            for _fname, _fdf in frames.items():
                _all_roles[_fname] = detect_column_roles(_fdf)
            # 检测是否有序时账+银行流水的组合
            _has_journal = any(
                ("debit" in r and "credit" in r) for r in _all_roles.values())
            _has_bank = any(
                ("debit" in r or "credit" in r) and "balance" in r
                for r in _all_roles.values())
            if _has_journal and _has_bank:
                lines.append("\n=== ⚠️ 跨表列名方向镜像映射（Merge/Diff 必须遵守） ===")
                for _fn, _roles in _all_roles.items():
                    _maps = []
                    if "debit" in _roles:
                        _maps.append(f"借方金额={_roles['debit']}")
                    if "credit" in _roles:
                        _maps.append(f"贷方金额={_roles['credit']}")
                    if _maps:
                        lines.append(f"  {_fn}: {', '.join(_maps)}")
                lines.append("  规则：序时账.借方金额 ←→ 银行流水.贷方（收入）  【方向镜像】")
                lines.append("       序时账.贷方金额 ←→ 银行流水.借方（支取）  【方向镜像】")
                lines.append("   Merge 的 left_on/right_on 必须用以上映射，禁止按同名列名硬匹配！")
        except Exception:
            pass
    return "\n".join(lines)


# 💡 修复四：彻底解耦严格面向对象，对 dict 和 object 提供极致的防御包容
def _dag_to_python(dag: Any, record: RunRecord) -> str:
    # 1. 安全提取头部元数据
    is_dict = isinstance(dag, dict)
    objective = (dag.get("objective") if is_dict else getattr(dag, "objective", None)) or "默认合并对账"
    raw_intent = (dag.get("raw_intent") if is_dict else getattr(dag, "raw_intent", None)) or "无预设意图"
    blueprint_id = (dag.get("blueprint_id") if is_dict else getattr(dag, "blueprint_id",
                                                                    None)) or f"bp_{uuid.uuid4().hex[:8]}"

    # 2. 安全提取算子列表
    raw_ops = (dag.get("operators", []) if is_dict else getattr(dag, "operators", []))

    # 3. 标准化清洗算子
    normalized_ops = []
    for op in raw_ops:
        op_dict = op if isinstance(op, dict) else op.__dict__ if hasattr(op, "__dict__") else {}

        n_op = {
            "id": op_dict.get("id") or f"op_{uuid.uuid4().hex[:6]}",
            "name": op_dict.get("name") or op_dict.get("type", "UnknownOperator"),
            "description": op_dict.get("description", ""),
            "params": op_dict.get("params") or {},
            "source_file": op_dict.get("source_file", ""),
            "output_alias": op_dict.get("output_alias", "df")
        }
        # 关键修复：op 级 input_from 注入 params.depends_on——否则 Merge/Diff/
        # Reconcile 只能"取最后两个变量"猜输入，接线顺序一乱就张冠李戴
        _inp = op_dict.get("input_from") or op_dict.get("depends_on") or []
        if isinstance(_inp, str):
            _inp = [_inp]
        if _inp and isinstance(n_op["params"], dict) \
                and "depends_on" not in n_op["params"]:
            n_op["params"]["depends_on"] = _inp
        normalized_ops.append(n_op)

    # 4. 提取执行顺序
    order = []
    if not is_dict and hasattr(dag, "get_execution_order") and callable(dag.get_execution_order):
        try:
            order = dag.get_execution_order()
        except Exception:
            pass
    if not order:
        order = [op["id"] for op in normalized_ops]

    op_map = {op["id"]: op for op in normalized_ops}

    # 5. 代码拼装（带数据流追踪）
    code_lines = [
        "import pandas as pd", "import json", "import os", "import sys",
        "try:",
        "    sys.stdout.reconfigure(encoding='utf-8', errors='replace')",
        "    sys.stderr.reconfigure(encoding='utf-8', errors='replace')",
        "except Exception:",
        "    pass", "",
        "os.makedirs('outputs', exist_ok=True)",
        "",
        "# === 文件追踪：按序分配 inputs 目录中的文件 ===",
        "_inputs_dir = os.path.join(os.path.dirname(__file__), 'inputs')",
        "_used_files = []",
        "",
        "def _load_any_document(path):",
        "    # 文档 → DataFrame：docx/pdf 表格 → 纯文本段落表（库缺失时清晰告警）",
        "    import os as _os",
        "    ext = _os.path.splitext(path)[1].lower()",
        "    tables, text = [], ''",
        "    try:",
        "        if ext == '.docx':",
        "            from docx import Document as _D",
        "            _d = _D(path)",
        "            text = '\\n'.join(p.text for p in _d.paragraphs if p.text.strip())",
        "            for _t in _d.tables:",
        "                _rows = [[c.text.strip() for c in r.cells] for r in _t.rows]",
        "                if len(_rows) > 1:",
        "                    tables.append(pd.DataFrame(_rows[1:], columns=[h or '列%d' % (i+1) for i, h in enumerate(_rows[0])]))",
        "        elif ext == '.pdf':",
        "            try:",
        "                import pdfplumber",
        "                with pdfplumber.open(path) as _pdf:",
        "                    for _pg in _pdf.pages:",
        "                        text += (_pg.extract_text() or '') + '\\n'",
        "                        for _raw in (_pg.extract_tables() or []):",
        "                            _rows = [[('' if c is None else str(c).strip()) for c in r] for r in _raw if r]",
        "                            _rows = [r for r in _rows if any(r)]",
        "                            if len(_rows) > 1:",
        "                                tables.append(pd.DataFrame(_rows[1:], columns=[h or '列%d' % (i+1) for i, h in enumerate(_rows[0])]))",
        "            except Exception as _e:",
        "                print('[Load] PDF 解析告警: ' + str(_e))",
        "        elif ext in ('.md', '.txt'):",
        "            text = open(path, 'rb').read().decode('utf-8', errors='replace')",
        "    except Exception as _e:",
        "        print('[Load] 文档解析告警: ' + str(_e))",
        "    if tables:",
        "        print('[Load] 文档 ' + _os.path.basename(path) + ' → ' + str(len(tables)) + ' 个表格，取第 1 个')",
        "        return tables[0]",
        "    if text.strip():",
        "        print('[Load] 文档 ' + _os.path.basename(path) + ' 无表格，转为段落表')",
        "        return pd.DataFrame({'段落': [l for l in text.splitlines() if l.strip()]})",
        "    print('[Load] ⚠ 文档未提取到内容: ' + _os.path.basename(path))",
        "    return pd.DataFrame()",
        "",
        "def _is_meaningless_key(col, df=None):",
        "    # 序号/行号等无业务含义的键，禁止作为连接键（防止'按行号对账'）",
        "    n = str(col).replace(' ', '').lower()",
        "    if n in ('序号', '编号', '行号', 'no', 'no.', 'id', 'index', 'idx', '#', '顺序号', '排名', 'code'):",
        "        return True",
        "    if df is not None and col in df.columns:",
        "        try:",
        "            v = pd.to_numeric(df[col], errors='coerce').dropna()",
        "            if len(v) >= 3 and v.nunique() == len(v):",
        "                d = v.sort_values().diff().dropna().unique()",
        "                if len(d) == 1 and d[0] == 1:",
        "                    return True",
        "        except Exception:",
        "            pass",
        "    return False",
        "",
        "def _reconcile_lite(dfA, dfB, dw=3, tol=0.01):",
        "    # 轻量逐笔对账：方向感知净额（序时账 借-贷 / 流水 收入-支出，互为镜像）",
        "    # L1 金额精确(±tol)+同日 → L2 金额精确+日期窗口±dw；输出对齐状态",
        "    def _pick(df, kws):",
        "        for c in df.columns:",
        "            if any(k in str(c) for k in kws): return c",
        "        return None",
        "    def _kind(df):",
        "        cs = ' '.join(str(c) for c in df.columns)",
        "        if ('凭证' in cs) or ('科目' in cs): return 'journal'",
        "        if ('对方' in cs) or ('账号' in cs) or ('余额' in cs): return 'bank'",
        "        return 'journal'",
        "    def _net(df, kind):",
        "        if kind == 'journal':",
        "            inc = _pick(df, ['借方金额', '借方发生额', '借方'])",
        "            exp = _pick(df, ['贷方金额', '贷方发生额', '贷方'])",
        "        else:",
        "            inc = _pick(df, ['收入', '贷方（收入）', '贷方(收入)', '贷方金额'])",
        "            exp = _pick(df, ['支出', '借方（支取）', '借方(支取)', '借方金额'])",
        "        if inc is not None or exp is not None:",
        "            a = pd.to_numeric(df[inc], errors='coerce').fillna(0) if inc else 0",
        "            b = pd.to_numeric(df[exp], errors='coerce').fillna(0) if exp else 0",
        "            return a - b",
        "        amt = _pick(df, ['交易金额', '金额', '发生额'])",
        "        if amt is not None: return pd.to_numeric(df[amt], errors='coerce').fillna(0)",
        "        nums = [c for c in df.columns if pd.api.types.is_numeric_dtype(df[c]) and not _is_meaningless_key(c, df)]",
        "        return pd.to_numeric(df[nums[0]], errors='coerce').fillna(0) if nums else pd.Series([0] * len(df))",
        "    ka, kb = _kind(dfA), _kind(dfB)",
        "    da = _pick(dfA, ['日期', '交易日期', '记账日期', '业务日期'])",
        "    db = _pick(dfB, ['日期', '交易日期', '记账日期', '业务日期'])",
        "    A = pd.DataFrame({'d': pd.to_datetime(dfA[da], errors='coerce') if da else pd.NaT, 'c': (_net(dfA, ka) * 100).round().astype('int64')})",
        "    B = pd.DataFrame({'d': pd.to_datetime(dfB[db], errors='coerce') if db else pd.NaT, 'c': (_net(dfB, kb) * 100).round().astype('int64')})",
        "    a_match, b_match = {}, {}",
        "    tol_c = max(1, int(round(tol * 100)))",
        "    pool = {}",
        "    for j, r in B.iterrows():",
        "        key = (r['d'].date().isoformat() if pd.notna(r['d']) else 'NaT', int(r['c']))",
        "        pool.setdefault(key, []).append(j)",
        "    for i, r in A.iterrows():",
        "        key = (r['d'].date().isoformat() if pd.notna(r['d']) else 'NaT', int(r['c']))",
        "        hit = None",
        "        for delta in range(-tol_c, tol_c + 1):",
        "            q = pool.get((key[0], key[1] + delta))",
        "            if q: hit = q.pop(0); break",
        "        if hit is not None:",
        "            a_match[i] = (hit, 0, 'L1'); b_match[hit] = i",
        "    bc = {}",
        "    for j, r in B.iterrows():",
        "        if j not in b_match: bc.setdefault(int(r['c']), []).append(j)",
        "    for i, r in A.iterrows():",
        "        if i in a_match: continue",
        "        best, bd = None, dw + 1",
        "        for delta in range(-tol_c, tol_c + 1):",
        "            for j in bc.get(int(r['c']) + delta, []):",
        "                if j in b_match: continue",
        "                dd = abs((r['d'] - B.loc[j, 'd']).days) if pd.notna(r['d']) and pd.notna(B.loc[j, 'd']) else dw",
        "                if dd <= dw and dd < bd: best, bd = j, dd",
        "        if best is not None:",
        "            a_match[i] = (best, bd, 'L2'); b_match[best] = i",
        "    out = dfA.copy()",
        "    out['对账状态'] = ['已核对' if i in a_match else '仅左表有' for i in range(len(dfA))]",
        "    out['对账层级'] = [a_match[i][2] if i in a_match else '-' for i in range(len(dfA))]",
        "    out['日期差(天)'] = [a_match[i][1] if i in a_match else None for i in range(len(dfA))]",
        "    out['对方行号'] = [int(a_match[i][0]) if i in a_match else None for i in range(len(dfA))]",
        "    un_b = dfB.loc[[j for j in dfB.index if j not in b_match]].copy()",
        "    print('[Reconcile] 左表类型=' + ka + ' 右表类型=' + kb + '（方向镜像已归一）')",
        "    print('[Reconcile] 已核对=' + str(len(a_match)) + ' 仅左=' + str(len(dfA) - len(a_match)) + ' 仅右=' + str(len(un_b)) + '（容差 ±' + str(tol) + ' 元，窗口 ±' + str(dw) + ' 天）')",
        "    if len(un_b) > 0:",
        "        try:",
        "            un_b.to_csv(os.path.join('outputs', 'reconcile_仅右表有.csv'), index=True, encoding='utf-8-sig')",
        "        except Exception: pass",
        "    return out",
        "",
        "# === DAG 执行代码 ===",
        f"# 目标: {objective}",
        f"# 原始意图: {raw_intent}",
        f"# 算子数: {len(order)}",
        "", ""
    ]

    # 数据流追踪：output_alias → 对应的 Python 变量名
    alias_vars = {}
    alias_formulas = {}
    alias_formula_file = {}
    last_df_var = None

    for op_id in order:
        op = op_map.get(op_id)
        if not op: continue
        code_lines.append(f"# Step: {op['name']} - {op['description']}")

        op_name = op['name']
        op_alias = op.get('output_alias', f'df_{op_id}')
        params = op.get('params', {}) or {}
        # 兼容 Dify 嵌套格式: {"parameters": {...}} → 解包
        if isinstance(params, dict) and "parameters" in params and isinstance(params["parameters"], dict):
            params = params["parameters"]

        if op_name == "Load":
            source = op.get('source_file') or "input.xlsx"
            var_name = f"df_{op_alias}" if not op_alias.startswith('df_') else op_alias
            code_lines.extend([
                f"# 智能文件匹配（优先按 DAG 指定的真实文件名精确匹配，禁止乱序分配）",
                f"_dag_file = '{source}'",
                f"_all_inputs = os.listdir(_inputs_dir) if os.path.exists(_inputs_dir) else []",
                f"_LOAD_EXTS = ('.xlsx','.xls','.csv','.docx','.doc','.pdf','.md','.txt')",
                f"if _dag_file in _all_inputs and _dag_file not in _used_files:",
                f"    _pick = _dag_file",
                f"    print('[Load] 按名匹配: ' + _pick)",
                f"else:",
                f"    _avail = [f for f in _all_inputs if f not in _used_files and f.lower().endswith(_LOAD_EXTS)]",
                f"    _pick = _avail[0] if _avail else _dag_file",
                f"    if _avail and _pick != _dag_file:",
                f"        print('[Load] ⚠ 指定文件 {{}} 不存在，回退到 {{}}'.format(_dag_file, _pick))",
                f"source_file = os.path.join(_inputs_dir, _pick) if os.path.exists(os.path.join(_inputs_dir, _pick)) else 'data/readonly/' + _dag_file",
                f"_used_files.append(_pick)",
                f"if not os.path.exists(source_file):",
                f"    print('[Load] 跳过: 文件不存在 ' + source_file)",
                f"    {var_name} = pd.DataFrame()",
                f"elif source_file.lower().endswith(('.docx','.doc','.pdf','.md','.txt')):",
                f"    {var_name} = _load_any_document(source_file)",
                f"else:",
                f"    # 自动检测表头行",
                f"    {var_name} = None",
                f"    _best_score = -1",
                f"    _best_hr = 0",
                f"    for _hr in range(6):",
                f"        try:",
                f"            _tmp = pd.read_excel(source_file, header=_hr, nrows=0)",
                f"            _cols = list(_tmp.columns)",
                f"            _score = 0",
                f"            for _c in _cols:",
                f"                _cs = str(_c)",
                f"                if _cs.startswith('Unnamed'): _score -= 1",
                f"                elif len(_cs) > 2 and not any('\\u4e00' <= ch <= '\\u9fff' for ch in _cs): _score -= 1",
                f"                else: _score += 1",
                f"            # 加分：列数多的行更可能是真正的表头",
                f"            _score += len(_cols) * 0.5",
                f"            if _score > _best_score:",
                f"                _best_score = _score",
                f"                _best_hr = _hr",
                f"        except Exception:",
                f"            continue",
                f"    if _best_score > 0:",
                f"        {var_name} = pd.read_excel(source_file, header=_best_hr)",
                f"        print('[Load] ' + os.path.basename(source_file) + ' -> {var_name}, h=' + str(_best_hr) + ', rows=' + str(len({var_name})) + ', cols=' + str(list({var_name}.columns)))",
                f"    else:",
                f"        {var_name} = pd.read_excel(source_file, header=None)",
                f"        {var_name}.columns = [f'Col_{{i}}' for i in range(len({var_name}.columns))]",
                f"        print('[Load] ' + os.path.basename(source_file) + ' -> {var_name} (无表头), rows=' + str(len({var_name})))",
            ])
            alias_vars[op_alias] = var_name
            last_df_var = var_name

        elif op_name == "ColumnFilter":
            columns = params.get("columns", [])
            dep_ids = params.get("depends_on", [])
            src_var = last_df_var
            for dep_id in dep_ids:
                dep_op = op_map.get(dep_id)
                if dep_op:
                    src_var = alias_vars.get(dep_op.get('output_alias'), last_df_var)
                    break
            src_var = src_var or 'df'
            cols_safe = [_sanitize_code_param(str(c), max_len=200) for c in columns]
            code_lines.extend([
                f"if '{src_var}' in dir() and {src_var} is not None:",
                f"    {op_alias} = {src_var}[{cols_safe}].copy()",
                f"    print('[ColumnFilter] cols=' + str({cols_safe}) + ', rows=' + str(len({op_alias})))",
                f"else:",
                f"    {op_alias} = pd.DataFrame()",
            ])
            alias_vars[op_alias] = op_alias
            last_df_var = op_alias

        elif op_name == "RegexFilter":
            col_raw = params.get("column", "")
            pattern_raw = params.get("pattern", "")
            dep_ids = params.get("depends_on", [])
            src_var = last_df_var
            for dep_id in dep_ids:
                dep_op = op_map.get(dep_id)
                if dep_op:
                    src_var = alias_vars.get(dep_op.get('output_alias'), last_df_var)
                    break
            src_var = src_var or 'df'
            col_safe = _sanitize_code_param(col_raw, max_len=200)
            pattern_safe = _sanitize_code_param(pattern_raw, max_len=1000)
            code_lines.extend([
                f"if '{src_var}' in dir() and {src_var} is not None and '{col_safe}' in {src_var}.columns:",
                f"    {op_alias} = {src_var}[{src_var}['{col_safe}'].astype(str).str.contains('{pattern_safe}', na=False)]",
                f"    print(f'[RegexFilter] pattern={pattern_safe}, rows={{len({op_alias})}}')",
                f"else:",
                f"    {op_alias} = {src_var}.copy() if '{src_var}' in dir() and {src_var} is not None else pd.DataFrame()",
            ])
            alias_vars[op_alias] = op_alias
            last_df_var = op_alias

        elif op_name == "NoiseFilter":
            # 噪音过滤器：从指定列中排除匹配噪音关键词的行
            # params.columns: 要过滤的列名列表
            # params.noise_patterns: 噪音关键词列表（如 ['利息', '手续费', '冲正']）
            columns = params.get("columns", [])
            if isinstance(columns, str):
                columns = [columns]
            noise_patterns = params.get("noise_patterns", [])
            if isinstance(noise_patterns, str):
                noise_patterns = [noise_patterns]
            dep_ids = params.get("depends_on", [])
            src_var = last_df_var
            for dep_id in dep_ids:
                dep_op = op_map.get(dep_id)
                if dep_op:
                    src_var = alias_vars.get(dep_op.get('output_alias'), last_df_var)
                    break
            src_var = src_var or 'df'
            if columns and noise_patterns:
                cols_safe = [_sanitize_code_param(str(c), max_len=200) for c in columns]
                patterns_safe = [_sanitize_code_param(str(p), max_len=200) for p in noise_patterns]
                patterns_repr = repr(noise_patterns)
                code_lines.extend([
                    f"if '{src_var}' in dir() and {src_var} is not None and not {src_var}.empty:",
                    f"    _mask = pd.Series(False, index={src_var}.index)",
                    f"    _cols_exist = [c for c in {cols_safe} if c in {src_var}.columns]",
                    f"    if _cols_exist:",
                    f"        for _col in _cols_exist:",
                    f"            for _pat in {patterns_safe}:",
                    f"                _mask = _mask | {src_var}[_col].astype(str).str.contains(_pat, na=False)",
                    f"    {op_alias} = {src_var}[~_mask].copy()",
                    f"    _removed = len({src_var}) - len({op_alias})",
                    f"    print(f'[NoiseFilter] cols={cols_safe}, patterns={patterns_repr}, 过滤 {{_removed}} 行, 保留 {{len({op_alias})}} 行')",
                    f"else:",
                    f"    {op_alias} = {src_var}.copy() if '{src_var}' in dir() and {src_var} is not None else pd.DataFrame()",
                    f"    print('[NoiseFilter] 源数据为空，直通')",
                ])
            else:
                # 无过滤条件 → 直通
                code_lines.extend([
                    f"{op_alias} = {src_var}.copy() if '{src_var}' in dir() and {src_var} is not None else pd.DataFrame()",
                    f"print('[NoiseFilter] 无过滤条件，数据直通, rows=' + str(len({op_alias})))",
                ])
            alias_vars[op_alias] = op_alias
            last_df_var = op_alias

        elif op_name == "GroupBy":
            by_cols = params.get("by", params.get("columns", params.get("group_by_columns", [])))
            aggs = params.get("aggregations", {})
            # 保护：如果 aggs 为空，跳过聚合
            if not aggs:
                code_lines.append(f"# [GroupBy] 跳过：aggregations 为空")
                code_lines.append(f"{op_alias} = {last_df_var or 'df'}")
                alias_vars[op_alias] = op_alias
                last_df_var = op_alias
                code_lines.append("")
                continue
            # 去重：如果某列同时在 by 和 agg 里，从 by 中移除（避免 reset_index 列名冲突）
            agg_keys = set(aggs.keys()) if isinstance(aggs, dict) else set()
            by_cols = [c for c in by_cols if c not in agg_keys]
            dep_ids = params.get("depends_on", [])
            src_var = last_df_var
            for dep_id in dep_ids:
                dep_op = op_map.get(dep_id)
                if dep_op:
                    src_var = alias_vars.get(dep_op.get('output_alias'), last_df_var)
                    break
            src_var = src_var or 'df'
            if by_cols:
                by_safe = [_sanitize_code_param(str(c), max_len=200) for c in by_cols]
                agg_auto = isinstance(aggs, dict) and list(aggs.keys()) == ["auto"]
                aggs_render = "{c: 'sum' for c in _num_cols[:3]}" if agg_auto else repr(aggs)
                code_lines.extend([
                    f"if '{src_var}' in dir() and {src_var} is not None and not {src_var}.empty:",
                    f"    if {by_safe} == ['auto']:",
                    f"        _objs = [c for c in {src_var}.columns if {src_var}[c].dtype == object and not _is_meaningless_key(c, {src_var})]",
                    f"        print('[GroupBy] 自动选择分组列: ' + str(_objs[:1]))",
                    f"    _by_valid = _objs[:1] if {by_safe} == ['auto'] else [c for c in {by_safe} if c in {src_var}.columns]",
                    f"    _num_cols = [c for c in {src_var}.columns if pd.api.types.is_numeric_dtype({src_var}[c]) and not _is_meaningless_key(c, {src_var})]",
                    f"    _by_missing = [c for c in {by_safe} if c not in {src_var}.columns]",
                    f"    if _by_missing:",
                    f"        print('[GroupBy] 警告：以下列不存在，已自动跳过: ' + str(_by_missing))",
                    f"    if _by_valid:",
                    f"        {op_alias} = {src_var}.groupby(_by_valid).agg({aggs_render}).reset_index()",
                    f"        print('[GroupBy] by=' + str(_by_valid) + ', rows=' + str(len({op_alias})))",
                    f"    else:",
                    f"        print('[GroupBy] 错误：没有有效的分组列！可用列: ' + str(list({src_var}.columns)))",
                    f"        {op_alias} = pd.DataFrame()",
                    f"else:",
                    f"    {op_alias} = pd.DataFrame()",
                ])
            else:
                # by_cols 全部在 agg 里，直接聚合
                code_lines.extend([
                    f"if '{src_var}' in dir() and {src_var} is not None:",
                    f"    {op_alias} = pd.DataFrame([{src_var}.agg({aggs})])",
                    f"    print('[GroupBy] agg-only, rows=' + str(len({op_alias})))",
                    f"else:",
                    f"    {op_alias} = pd.DataFrame()",
                ])
            alias_vars[op_alias] = op_alias
            last_df_var = op_alias

        elif op_name == "Merge":
            on_cols = params.get("on", [])
            if isinstance(on_cols, str):
                on_cols = [on_cols]
            how = params.get("how", "outer")
            lo = params.get("left_on", []) or []
            ro = params.get("right_on", []) or []
            lo_list = lo if isinstance(lo, list) else [lo]
            ro_list = ro if isinstance(ro, list) else [ro]
            date_window = params.get("date_window_days")
            dep_ids = params.get("depends_on", [])
            src_vars = []
            for dep_id in dep_ids:
                dep_op = op_map.get(dep_id)
                if dep_op:
                    src_vars.append(alias_vars.get(dep_op.get('output_alias'), f'df_{dep_id}'))
            if len(src_vars) < 2:
                all_vars = list(alias_vars.values())
                src_vars = all_vars[-2:] if len(all_vars) >= 2 else (src_vars + ['df'])
            left_var, right_var = (src_vars[0], src_vars[1]) if len(src_vars) >= 2 else ('df', 'df')
            # v3.1: LLM 可能生成相同上游（如 Merge input_from ['op_4','op_4']），
            # 导致两边变量相同 → 自交。检测并回退到不同数据源。
            if left_var == right_var:
                all_aliases = list(alias_vars.values())
                alt = None
                for v in all_aliases:
                    if v != left_var:
                        alt = v
                        break
                if alt:
                    right_var = alt
                    print(f"[代码生成] Merge 两边同源({left_var})，自动切换到 {alt} 作为右表")
            how_safe = how if how in ("inner", "left", "right", "outer") else "outer"
            dw_repr = repr(int(date_window)) if isinstance(date_window, (int, float)) else "None"
            code_lines.extend([
                f"if '{left_var}' in dir() and '{right_var}' in dir() and not {left_var}.empty and not {right_var}.empty:",
                f"    _lo, _ro, _on, _dw = {lo_list}, {ro_list}, {on_cols}, {dw_repr}",
                f"    _on = [c for c in _on if c != 'auto']",
                f"    _lo_v = [c for c in _lo if c in {left_var}.columns]",
                f"    _ro_v = [c for c in _ro if c in {right_var}.columns]",
                f"    _miss_keys = [c for c in _lo if c not in {left_var}.columns] + [c for c in _ro if c not in {right_var}.columns]",
                f"    if _miss_keys:",
                f"        print('[Merge] ⚠ 指定连接键在数据中不存在: ' + str(_miss_keys))",
                f"        print('[Merge] 左表列: ' + str(list({left_var}.columns)) + ' | 右表列: ' + str(list({right_var}.columns)))",
                f"    if _lo_v and len(_lo_v) == len(_ro_v) and _dw:",
                f"        _dpr = [(a, b) for a, b in zip(_lo_v, _ro_v) if ('日期' in str(a) + str(b) or '时间' in str(a) + str(b))]",
                f"        _nd_l = [a for a, b in zip(_lo_v, _ro_v) if (a, b) not in _dpr]",
                f"        _nd_r = [b for a, b in zip(_lo_v, _ro_v) if (a, b) not in _dpr]",
                f"        if _nd_l:",
                f"            _m = pd.merge({left_var}, {right_var}, left_on=_nd_l, right_on=_nd_r, how='inner', suffixes=('_L', '_R'))",
                f"        else:",
                f"            _m = {left_var}.assign(_xj=1).merge({right_var}.assign(_xj=1), on='_xj', suffixes=('_L', '_R')).drop(columns=['_xj'])",
                f"            if len(_m) > 200000:",
                f"                print('[Merge] ⚠ 日期窗口连接笛卡尔积 ' + str(len(_m)) + ' 行，截断至 200000（请增加非日期键）')",
                f"                _m = _m.head(200000)",
                f"        if _dpr:",
                f"            _dl = _dpr[0][0] if _dpr[0][0] in _m.columns else _dpr[0][0] + '_L'",
                f"            _dr = _dpr[0][1] if _dpr[0][1] in _m.columns else _dpr[0][1] + '_R'",
                f"            {op_alias} = _m[(pd.to_datetime(_m[_dl], errors='coerce') - pd.to_datetime(_m[_dr], errors='coerce')).dt.days.abs() <= _dw].reset_index(drop=True)",
                f"            print('[Merge] 日期窗口±' + str(_dw) + '天(' + str(_dl) + ' vs ' + str(_dr) + '), rows=' + str(len({op_alias})))",
                f"        else:",
                f"            print('[Merge] ⚠ 连接键中无日期列对，date_window_days 被忽略')",
                f"            {op_alias} = _m",
                f"    elif _lo_v and len(_lo_v) == len(_ro_v):",
                f"        {op_alias} = pd.merge({left_var}, {right_var}, left_on=_lo_v, right_on=_ro_v, how='{how_safe}')",
                f"        print('[Merge] left_on/right_on=' + str(list(zip(_lo_v, _ro_v))) + ', rows=' + str(len({op_alias})))",
                f"    else:",
                f"        _common = [c for c in (_on or list({left_var}.columns)) if c in {left_var}.columns and c in {right_var}.columns and not _is_meaningless_key(c, {left_var})]",
                f"        if _on and not _common:",
                f"            print('[Merge] ⚠ 指定的 on 键在双方不存在或无意义: ' + str(_on))",
                f"        if not _common:",
                f"            _common = [c for c in list({left_var}.columns) if c in {right_var}.columns and not _is_meaningless_key(c, {left_var})]",
                f"            print('[Merge] 回退到公共有意义键（已排除序号/行号）: ' + str(_common))",
                f"        if _common:",
                f"            {op_alias} = pd.merge({left_var}, {right_var}, on=_common, how='{how_safe}')",
                f"            print('[Merge] on=' + str(_common) + ', rows=' + str(len({op_alias})))",
                f"        else:",
                f"            print('[Merge] ❌ 无有意义公共键（已排除序号/行号），拒绝按行号合并！左列: ' + str(list({left_var}.columns)) + ' | 右列: ' + str(list({right_var}.columns)))",
                f"            {op_alias} = pd.DataFrame()",
                f"else:",
                f"    print('[Merge] 输入数据为空或不存在')",
                f"    {op_alias} = pd.DataFrame()",
            ])
            alias_vars[op_alias] = op_alias
            last_df_var = op_alias

        elif op_name == "Sort":
            by_cols = params.get("by", params.get("columns", []))
            ascending = params.get("ascending", True)
            dep_ids = params.get("depends_on", [])
            src_var = last_df_var
            for dep_id in dep_ids:
                dep_op = op_map.get(dep_id)
                if dep_op:
                    src_var = alias_vars.get(dep_op.get('output_alias'), last_df_var)
                    break
            src_var = src_var or 'df'
            by_safe = [_sanitize_code_param(str(c), max_len=200) for c in by_cols]
            code_lines.extend([
                f"if '{src_var}' in dir() and {src_var} is not None and not {src_var}.empty:",
                f"    _nums_auto = [c for c in {src_var}.columns if pd.api.types.is_numeric_dtype({src_var}[c]) and not _is_meaningless_key(c, {src_var})]",
                f"    _sort_valid = (_nums_auto[:1] if {by_safe} == ['auto'] else [c for c in {by_safe} if c in {src_var}.columns])",
                f"    if {by_safe} == ['auto'] and _sort_valid:",
                f"        print('[Sort] 自动选择排序列: ' + str(_sort_valid))",
                f"    _sort_missing = [c for c in {by_safe} if c not in {src_var}.columns]",
                f"    if _sort_missing:",
                f"        print('[Sort] 警告：以下列不存在，已自动跳过: ' + str(_sort_missing))",
                f"    if _sort_valid:",
                f"        {op_alias} = {src_var}.sort_values(by=_sort_valid, ascending={ascending})",
                f"        print('[Sort] by=' + str(_sort_valid) + ', asc=' + str({ascending}) + ', rows=' + str(len({op_alias})))",
                f"    else:",
                f"        print('[Sort] 错误：没有有效排序列！可用列: ' + str(list({src_var}.columns)))",
                f"        {op_alias} = {src_var}.copy()",
                f"else:",
                f"    {op_alias} = pd.DataFrame()",
            ])
            alias_vars[op_alias] = op_alias
            last_df_var = op_alias
        elif op_name == "ConditionCheck":
            col_raw = params.get("column", "")
            operator_raw = params.get("operator", ">")
            value_raw = params.get("value", 0)
            col_safe = _sanitize_code_param(col_raw, max_len=200)
            operator_safe = operator_raw if operator_raw in _VALID_COMPARISON_OPS else "=="
            try:
                value_safe = float(value_raw)
            except (ValueError, TypeError):
                value_safe = 0.0
            dep_ids = params.get("depends_on", [])
            src_var = last_df_var
            for dep_id in dep_ids:
                dep_op = op_map.get(dep_id)
                if dep_op:
                    src_var = alias_vars.get(dep_op.get('output_alias'), last_df_var)
                    break
            src_var = src_var or 'df'
            code_lines.extend([
                f"if '{src_var}' in dir() and {src_var} is not None and '{col_safe}' in {src_var}.columns:",
                f"    mask = {src_var}['{col_safe}'] {operator_safe} {value_safe}",
                f"    {op_alias}_passed = {src_var}[mask]",
                f"    {op_alias}_failed = {src_var}[~mask]",
            ])

        elif op_name == "Transform":
            operation = params.get("operation", "")
            dep_ids = params.get("depends_on", [])
            src_var = last_df_var
            for dep_id in dep_ids:
                dep_op = op_map.get(dep_id)
                if dep_op:
                    src_var = alias_vars.get(dep_op.get('output_alias'), last_df_var)
                    break
            src_var = src_var or 'df'

            if operation == "compute_columns":
                formulas = params.get("formulas", {})
                output_as_formula = params.get("output_as_formula", False)

                if output_as_formula:
                    # 公式模式：不计算，只记录公式→Export 阶段用 openpyxl 写
                    code_lines.append(f"if '{src_var}' in dir() and {src_var} is not None:")
                    code_lines.append(f"    {op_alias} = {src_var}.copy()")
                    formula_dict = {}
                    for col, formula in formulas.items():
                        safe_col = _sanitize_code_param(str(col), max_len=100)
                        safe_f = _sanitize_code_param(str(formula), max_len=500)
                        # 公式里引用源列名（如 D 列→第4列，B 列→第2列）
                        formula_dict[safe_col] = safe_f
                    alias_formulas[op_alias] = formula_dict
                    alias_formula_file[op_alias] = params.get("source_file", "")
                    code_lines.append(f"    print('[Transform] formula mode: ' + str(list({op_alias}.columns)) + ' + formulas={len(formula_dict)} cols')")
                else:
                    # 值模式：pandas 直接计算
                    code_lines.append(f"if '{src_var}' in dir() and {src_var} is not None:")
                    code_lines.append(f"    {op_alias} = {src_var}.copy()")
                    for col, formula in formulas.items():
                        safe_formula = str(formula).replace("df", src_var)
                        safe_col = _sanitize_code_param(str(col), max_len=100)
                        code_lines.append(f"    try:")
                        code_lines.append(f"        {op_alias}['{safe_col}'] = {safe_formula}")
                        code_lines.append(f"    except Exception as _e:")
                        code_lines.append(f"        print('[Transform] 计算列 {safe_col} 失败: ' + str(_e))")
                        code_lines.append(f"        {op_alias}['{safe_col}'] = None")
                    code_lines.append(f"    print('[Transform] compute_columns: ' + str(list({op_alias}.columns)))")

            elif operation == "extract_date_part":
                columns = params.get("columns", [])
                date_part = params.get("date_part", "year")
                new_col = params.get("new_column", f"日期_{date_part}")
                cols_safe = [_sanitize_code_param(str(c), max_len=100) for c in (columns if isinstance(columns, list) else [columns])]
                date_col = cols_safe[0] if cols_safe else "'date'"
                code_lines.append(f"if '{src_var}' in dir() and {src_var} is not None and {date_col} in {src_var}.columns:")
                code_lines.append(f"    {op_alias} = {src_var}.copy()")
                code_lines.append(f"    {op_alias}['{new_col}'] = pd.to_datetime({op_alias}[{date_col}], errors='coerce').dt.{date_part}")
                code_lines.append(f"    print('[Transform] extract_date_part: {new_col} from {date_col}')")

            elif operation == "standardize_name":
                columns = params.get("columns", [])
                cols_safe = [_sanitize_code_param(str(c), max_len=100) for c in (columns if isinstance(columns, list) else [columns])]
                code_lines.append(f"if '{src_var}' in dir() and {src_var} is not None:")
                code_lines.append(f"    {op_alias} = {src_var}.copy()")
                for col in cols_safe or []:
                    code_lines.append(f"    if '{col}' in {op_alias}.columns:")
                    code_lines.append(f"        {op_alias}['{col}'] = {op_alias}['{col}'].astype(str).str.replace(r'[（(].*?[）)]', '', regex=True)")
                    code_lines.append(f"        {op_alias}['{col}'] = {op_alias}['{col}'].str.replace(r'[有限公司|有限责任公司|股份有限公司]$', '', regex=True)")
                    code_lines.append(f"        {op_alias}['{col}'] = {op_alias}['{col}'].str.strip()")
                code_lines.append(f"    print('[Transform] standardize_name: ' + str({cols_safe}))")

            else:
                # Generic fallback: copy through
                code_lines.append(f"if '{src_var}' in dir() and {src_var} is not None:")
                code_lines.append(f"    {op_alias} = {src_var}.copy()")
                code_lines.append(f"else:")
                code_lines.append(f"    {op_alias} = pd.DataFrame()")

            alias_vars[op_alias] = op_alias
            last_df_var = op_alias



        elif op_name == "Export":
            output_file = params.get("output_file", params.get("output_file_path", "analysis_result.csv"))
            dep_ids = params.get("depends_on", [])
            src_var = last_df_var
            for dep_id in dep_ids:
                dep_op = op_map.get(dep_id)
                if dep_op:
                    src_var = alias_vars.get(dep_op.get('output_alias'), last_df_var)
                    break
            src_var = src_var or 'df'
            _has_fmts = src_var in str(alias_formulas)
            if _has_fmts:
                code_lines.extend([
                    f"if '{src_var}' in dir() and {src_var} is not None and not {src_var}.empty:",
                    f"    _out = os.path.join('outputs', '{output_file}')",
                    f"    from openpyxl import Workbook",
                    f"    _wb = Workbook(); _ws = _wb.active",
                    f"    for _ci, _cn in enumerate({src_var}.columns, 1):",
                    f"        _ws.cell(row=1, column=_ci, value=str(_cn))",
                    f"    _fmts = _FORMULA_MAP.get('{src_var}', {{}})",
                    f"    for _ri, (_, _row) in enumerate({src_var}.iterrows(), 2):",
                    f"        for _ci, _cn in enumerate({src_var}.columns, 1):",
                    f"            if _cn in _fmts:",
                    f"                _f = _fmts[_cn].replace('{{row}}', str(_ri))",
                    f"                _ws.cell(row=_ri, column=_ci, value=_f)",
                    f"            else:",
                    f"                _v = _row[_cn]",
                    f"                if pd.isna(_v): _v = ''",
                    f"                _ws.cell(row=_ri, column=_ci, value=_v)",
                    f"    _wb.save(_out)",
                    f"    print('[Export] ' + '{output_file}' + ' (openpyxl+formulas)')",
                    f"else:",
                    f"    print('[Export] \xe8\xb7\xb3\xe8\xbf\x87\xef\xbc\x9a\xe6\x95\xb0\xe6\x8d\xae\xe4\xb8\xba\xe7\xa9\xba')",
                ])
                fmts_repr = repr(alias_formulas.get(src_var, {}))
                code_lines.insert(-14, f"    _FORMULA_MAP = {{}}")
                code_lines.insert(-13, f"    _FORMULA_MAP['{src_var}'] = {fmts_repr}")
            else:
                code_lines.extend([
                    f"if '{src_var}' in dir() and {src_var} is not None and not {src_var}.empty:",
                    f"    {src_var}.to_csv(os.path.join('outputs', '{output_file}'), index=False, encoding='utf-8-sig')",
                    f"    print('[Export] ' + '{output_file}' + ', rows=' + str(len({src_var})))",
                    f"else:",
                    f"    print('[Export] \xe8\xb7\xb3\xe8\xbf\x87\xef\xbc\x9a\xe6\x95\xb0\xe6\x8d\xae\xe4\xb8\xba\xe7\xa9\xba\xef\xbc\x8c\xe4\xb8\x8d\xe5\xaf\xbc\xe5\x87\xba\xe7\xa9\xba\xe6\x96\x87\xe4\xbb\xb6')",
                ])

        elif op_name == "Aggregate":
            aggs = params.get("aggregations", {})
            # 格式1: {"columns": ["col"], "aggregation": "sum"}（扁平）
            if not aggs and params.get("columns"):
                flat_cols = params.get("columns", [])
                flat_func = params.get("aggregation", "sum")
                if isinstance(flat_cols, list) and flat_cols:
                    aggs = {f"{c}_{flat_func}": {"column": c, "agg_func": flat_func} for c in flat_cols}
            # 格式2: {"columns": ["col"], "aggregations": ["sum"]}（列表）
            if isinstance(aggs, list) and params.get("columns"):
                cols = params.get("columns", [])
                func = aggs[0] if aggs else "sum"
                if isinstance(cols, list) and cols:
                    aggs = {f"{c}_{func}": {"column": c, "agg_func": func} for c in cols}
            if not aggs:
                code_lines.append(f"# [Aggregate] 跳过：aggregations 为空")
                code_lines.append(f"{op_alias} = {last_df_var or 'df'}")
                alias_vars[op_alias] = op_alias
                last_df_var = op_alias
                code_lines.append("")
                continue
            dep_ids = params.get("depends_on", [])
            src_var = last_df_var
            for dep_id in dep_ids:
                dep_op = op_map.get(dep_id)
                if dep_op:
                    src_var = alias_vars.get(dep_op.get('output_alias'), last_df_var)
                    break
            src_var = src_var or 'df'
            # 转换 LLM 输出的嵌套格式为兼容 pandas>=0.25 的语法
            # 输入: {"col_sum": {"column": "col", "agg_func": "sum"}}
            # 输出: pd.DataFrame({'col_sum': [df['col'].sum()], ...})
            agg_items = []
            for new_name, agg_spec in aggs.items():
                col = agg_spec.get("column", "") if isinstance(agg_spec, dict) else new_name
                func = agg_spec.get("agg_func", "sum") if isinstance(agg_spec, dict) else agg_spec
                safe_col = repr(str(col))
                safe_name = repr(str(new_name))
                agg_items.append((new_name, col, func, safe_name, safe_col))
            if agg_items:
                group_by = params.get("group_by", [])
                if group_by:
                    # 有分组列：用 groupby + named agg（现代 pandas 语法）
                    gp_safe = [_sanitize_code_param(str(c), max_len=200) for c in group_by]
                    code_lines.extend([
                        f"if '{src_var}' in dir() and {src_var} is not None and not {src_var}.empty:",
                        f"    _gp_valid = [c for c in {gp_safe} if c in {src_var}.columns]",
                        f"    if _gp_valid:",
                    ])
                    gp_parts = [f"{repr(n)}: (pd.to_numeric({src_var}[{repr(c)}], errors='coerce'), '{fn}')" for n, c, fn, _, _ in agg_items]
                    code_lines.append(
                        f"        {op_alias} = {src_var}.groupby(_gp_valid).agg(**{{{', '.join(gp_parts)}}}).reset_index()"
                    )
                    code_lines.extend([
                        f"    else:",
                        f"        {op_alias} = pd.DataFrame()",
                        f"else:",
                        f"    {op_alias} = pd.DataFrame()",
                    ])
                else:
                    # 无分组：用 pd.DataFrame 构造（避免嵌套 renamer 不兼容）
                    item_lines = [f"            {sn}: [pd.to_numeric({src_var}[{sc}], errors='coerce').{fn}()]" for _, _, fn, sn, sc in agg_items]
                    code_lines.extend([
                        f"if '{src_var}' in dir() and {src_var} is not None:",
                        f"    {op_alias} = pd.DataFrame({{",
                        ",\n".join(item_lines),
                        f"    }})",
                    ])
            else:
                code_lines.append(f"{op_alias} = {last_df_var or 'df'}")

        elif op_name == "Diff":
            dep_ids = params.get("depends_on", [])
            src_vars = []
            for dep_id in dep_ids:
                dep_op = op_map.get(dep_id)
                if dep_op:
                    src_vars.append(alias_vars.get(dep_op.get('output_alias'), f'df_{dep_id}'))
            if len(src_vars) < 2:
                all_vars = list(alias_vars.values())
                src_vars = all_vars[-2:] if len(all_vars) >= 2 else (src_vars + ['df', 'df'])
            left_var, right_var = (src_vars[0], src_vars[1]) if len(src_vars) >= 2 else ('df', 'df')
            # v3.1: 同 Merge，Diff 两边同源时回退到不同数据源
            if left_var == right_var:
                all_aliases = list(alias_vars.values())
                alt = None
                for v in all_aliases:
                    if v != left_var:
                        alt = v
                        break
                if alt:
                    right_var = alt
            keys = params.get("keys", params.get("on", [])) or []
            if isinstance(keys, str):
                keys = [keys]
            col_a = str(params.get("col_a", "auto"))
            col_b = str(params.get("col_b", "auto"))
            tol_abs = params.get("tolerance_abs")
            tol_pct = params.get("tolerance_pct")
            try:
                tol_repr = repr(float(tol_abs)) if tol_abs is not None else "None"
            except (ValueError, TypeError):
                tol_repr = "None"
            try:
                pct_repr = repr(float(tol_pct)) if tol_pct is not None else "None"
            except (ValueError, TypeError):
                pct_repr = "None"
            # 按业务键 outer 对齐（indicator 标记仅左/仅右/双方），拒绝按行号比对
            # 修复：keys 为空时也能正常回退；merge 前统一日期类型，避免 datetime vs object 报错
            keys_repr = repr(list(keys))
            code_lines.extend([
                f"if '{left_var}' in dir() and '{right_var}' in dir():",
                f"    _keys = {keys_repr}",
                f"    _keys = [c for c in _keys if c in {left_var}.columns and c in {right_var}.columns and not _is_meaningless_key(c, {left_var})]",
                f"    if not _keys:",
                f"        _keys = [c for c in list({left_var}.columns) if c in {right_var}.columns and not _is_meaningless_key(c, {left_var})]",
                f"        print('[Diff] 未指定有效键，回退公共有意义键: ' + str(_keys))",
                f"    if not _keys:",
                f"        print('[Diff] ❌ 两表无公共有意义键，拒绝按行号比对！左列: ' + str(list({left_var}.columns)) + ' | 右列: ' + str(list({right_var}.columns)))",
                f"        {op_alias} = pd.DataFrame()",
                f"    else:",
                f"        # 日期类型归一化，避免 merge 时 datetime 与 object 冲突",
                f"        _df_left = {left_var}.copy()",
                f"        _df_right = {right_var}.copy()",
                f"        for _k in _keys:",
                f"            try:",
                f"                _l_is_dt = pd.api.types.is_datetime64_any_dtype(_df_left[_k])",
                f"                _r_is_dt = pd.api.types.is_datetime64_any_dtype(_df_right[_k])",
                f"                if _l_is_dt or _r_is_dt:",
                f"                    _df_left[_k] = pd.to_datetime(_df_left[_k], errors='coerce')",
                f"                    _df_right[_k] = pd.to_datetime(_df_right[_k], errors='coerce')",
                f"            except Exception as _e:",
                f"                pass",
                f"        try:",
                f"            _m = pd.merge(_df_left, _df_right, on=_keys, how='outer', indicator=True, suffixes=('_A', '_B'))",
                f"        except Exception as _e:",
                f"            print('[Diff] merge 失败（可能列类型不一致）: ' + str(_e))",
                f"            {op_alias} = pd.DataFrame()",
                f"        else:",
                f"            _ca, _cb = {col_a!r}, {col_b!r}",
                f"            if _ca == 'auto':",
                f"                _na = [c for c in {left_var}.columns if pd.api.types.is_numeric_dtype({left_var}[c]) and not _is_meaningless_key(c, {left_var})]",
                f"                _ca = _na[0] if _na else None",
                f"            if _cb == 'auto':",
                f"                _nb = [c for c in {right_var}.columns if pd.api.types.is_numeric_dtype({right_var}[c]) and not _is_meaningless_key(c, {right_var})]",
                f"                _cb = _nb[0] if _nb else None",
                f"            _ca2 = _ca if _ca in _m.columns else (str(_ca) + '_A' if str(_ca) + '_A' in _m.columns else None)",
                f"            _cb2 = _cb if _cb in _m.columns else (str(_cb) + '_B' if str(_cb) + '_B' in _m.columns else None)",
                f"            if _ca2 and _cb2:",
                f"                _m['差异'] = pd.to_numeric(_m[_ca2], errors='coerce') - pd.to_numeric(_m[_cb2], errors='coerce')",
                f"                _tol_abs, _tol_pct = {tol_repr}, {pct_repr}",
                f"                if _tol_abs is not None:",
                f"                    _m['差异超限'] = _m['差异'].abs() > _tol_abs",
                f"                    print('[Diff] 容差: ±' + str(_tol_abs) + ' 元（绝对）')",
                f"                elif _tol_pct is not None:",
                f"                    _base = pd.to_numeric(_m[_cb2], errors='coerce').abs().clip(lower=0.01)",
                f"                    _m['差异超限'] = (_m['差异'].abs() / _base * 100) > _tol_pct",
                f"                    print('[Diff] 容差: ' + str(_tol_pct) + '%（仅汇总层面适用，逐笔核对请用 tolerance_abs=0.01）')",
                f"                else:",
                f"                    _m['差异超限'] = _m['差异'].abs() > 0.01",
                f"                    print('[Diff] 未指定容差，默认 ±0.01 元（精确到分）')",
                f"            else:",
                f"                print('[Diff] ⚠ 未找到可比数值列: ' + str((_ca, _cb)))",
                f"            _m['对齐状态'] = _m['_merge'].map({{'both': '双方都有', 'left_only': '仅左表有', 'right_only': '仅右表有'}})",
                f"            {op_alias} = _m.drop(columns=['_merge'])",
                f"            print('[Diff] 键=' + str(_keys) + ', 仅左=' + str(int(({op_alias}['对齐状态'] == '仅左表有').sum())) + ', 仅右=' + str(int(({op_alias}['对齐状态'] == '仅右表有').sum())) + ', 双方=' + str(int(({op_alias}['对齐状态'] == '双方都有').sum())))",
                f"else:",
                f"    {op_alias} = pd.DataFrame()",
            ])
            alias_vars[op_alias] = op_alias
            last_df_var = op_alias

        elif op_name == "Reconcile":
            # v3.4: DAG Reconcile calls run_bank_reconciliation directly (not _reconcile_lite)
            dep_ids = params.get("depends_on", [])
            src_vars = []
            for dep_id in dep_ids:
                dep_op = op_map.get(dep_id)
                if dep_op:
                    src_vars.append(alias_vars.get(dep_op.get('output_alias'), f'df_{dep_id}'))
            if len(src_vars) < 2:
                all_vars = list(alias_vars.values())
                src_vars = all_vars[-2:] if len(all_vars) >= 2 else (src_vars + ['df', 'df'])
            left_var, right_var = (src_vars[0], src_vars[1]) if len(src_vars) >= 2 else ('df', 'df')
            try:
                rc_tol = float(params.get("tolerance_abs", 0.01))
            except (ValueError, TypeError):
                rc_tol = 0.01
            try:
                rc_dw = int(params.get("date_window_days", 3))
            except (ValueError, TypeError):
                rc_dw = 3
            _base = str(BASE_DIR)
            code_lines.extend([
                f"import sys; sys.path.insert(0, r'{_base}')",
                "from core.bank_reconcile_engine import run_bank_reconciliation",
                f"if '{left_var}' in dir() and '{right_var}' in dir() and not {left_var}.empty and not {right_var}.empty:",
                f"    _rec_res = run_bank_reconciliation({left_var}, {right_var}, {{'amount_tolerance': {rc_tol}, 'date_window_days': {rc_dw}}})",
                f"    {op_alias} = _rec_res['book_std'].copy()",
                f"    print('[Reconcile] book_match=' + str(_rec_res['stats']['book_match_rate']) + '% bank_match=' + str(_rec_res['stats']['bank_match_rate']) + '%')",
                f"    print('[Reconcile] L1=' + str(_rec_res['stats']['matched_L1']) + ' L2=' + str(_rec_res['stats']['matched_L2']) + ' L3=' + str(_rec_res['stats']['matched_L3_groups']) + ' L4=' + str(_rec_res['stats']['review_L4']))",
                f"    print('[Reconcile] red_flags=' + str(_rec_res['stats']['red_flag_count']) + ' timing=' + str(_rec_res['stats']['timing_categories']))",
                f"else:",
                f"    print('[Reconcile] empty input')",
                f"    {op_alias} = pd.DataFrame()",
            ])
            alias_vars[op_alias] = op_alias
            last_df_var = op_alias

        elif op_name == "AuditAdjustment":
            # 审计调整建议：基于 Diff 差异列生成（科目需人工确认，平台不臆造）
            dep_ids = params.get("depends_on", [])
            src_var = last_df_var
            for dep_id in dep_ids:
                dep_op = op_map.get(dep_id)
                if dep_op:
                    src_var = alias_vars.get(dep_op.get('output_alias'), last_df_var)
                    break
            src_var = src_var or 'df'
            code_lines.extend([
                f"if '{src_var}' in dir() and {src_var} is not None and not {src_var}.empty and '差异' in {src_var}.columns:",
                f"    _adj = {src_var}[pd.to_numeric({src_var}['差异'], errors='coerce').abs() > 0.01].copy()",
                f"    _adj['调整建议'] = '差异 ' + pd.to_numeric(_adj['差异'], errors='coerce').round(2).astype(str) + ' 元；借/贷科目待审计师确认'",
                f"    {op_alias} = _adj",
                f"    print('[AuditAdjustment] 调整建议 ' + str(len(_adj)) + ' 条（科目需人工确认）')",
                f"else:",
                f"    print('[AuditAdjustment] 输入无差异列或为空，输出空调整表')",
                f"    {op_alias} = pd.DataFrame(columns=['调整建议'])",
            ])
            alias_vars[op_alias] = op_alias
            last_df_var = op_alias

        else:
            code_lines.append(f"# [跳过] 未实现算子: {op_name}")
            print(f"[警告] 算子 {op_name} 未实现，已跳过——请确认这是预期行为")

        code_lines.append("")

    # 格式化短摘要，防止越界或空值崩溃
    safe_explanation = str(objective)[:30].replace('"', "'")

    code_lines.extend([
        "# 生成实际输出文件（基于计算结果，非硬编码）",
        f"_final_df = {last_df_var} if '{last_df_var}' in dir() and {last_df_var} is not None and not {last_df_var}.empty else None",
        "if _final_df is not None:",
        "    # 导出 CSV",
        "    _final_df.to_csv(os.path.join('outputs', 'analysis_result.csv'), index=False, encoding='utf-8-sig')",
        "    print('[Output] CSV 已导出: analysis_result.csv, rows=' + str(len(_final_df)))",
        "    # 同时导出 Excel（保留原表格式）",
        "    try:",
        "        _final_df.to_excel(os.path.join('outputs', 'analysis_result.xlsx'), index=False)",
        "        print('[Output] Excel 已导出: analysis_result.xlsx')",
        "    except Exception as _e:",
        "        print('[Output] Excel 导出失败（可能缺少 openpyxl）: ' + str(_e))",
        "    # 生成 JSON 摘要",
        "    _summary = {",
        "        'total_rows': len(_final_df),",
        "        'columns': list(_final_df.columns),",
        "        'dtypes': {str(k): str(v) for k, v in _final_df.dtypes.items()},",
        "    }",
        "    # 数值列统计",
        "    _num_cols = _final_df.select_dtypes(include='number').columns.tolist()",
        "    if _num_cols:",
        "        _summary['numeric_summary'] = {c: {'sum': float(_final_df[c].sum()), 'mean': float(_final_df[c].mean()), 'max': float(_final_df[c].max()), 'min': float(_final_df[c].min())} for c in _num_cols}",
        "    with open(os.path.join('outputs', 'journal_entries.json'), 'w', encoding='utf-8') as f:",
        "        json.dump(_summary, f, ensure_ascii=False, indent=2)",
        "    print('[Output] JSON 已导出: journal_entries.json')",
        "else:",
        "    print('[Output] 警告：无有效数据可导出')",
        "    with open(os.path.join('outputs', 'journal_entries.json'), 'w', encoding='utf-8') as f:",
        "        json.dump({'error': '无有效数据', 'columns': [], 'total_rows': 0}, f, ensure_ascii=False, indent=2)",
    ])
    code_lines.extend([
        "",
        "# === 数据质量层：只清理全空行/列；缺失值一律标记报告，绝不编造 ===",
        "# （审计红线：前向填充会把上一行客户名填到缺失行、金额填 0，属于伪造审计证据）",
        "_missing_report = {}",
        "for _name, _v in list(locals().items()):",
        "    if isinstance(_v, pd.DataFrame) and not _v.empty:",
        "        _v.dropna(how='all', inplace=True)",
        "        _v.dropna(axis=1, how='all', inplace=True)",
        "        _miss = {c: int(_v[c].isna().sum()) for c in _v.columns if int(_v[c].isna().sum()) > 0}",
        "        if _miss:",
        "            _missing_report[_name] = _miss",
        "if _missing_report:",
        "    with open(os.path.join('outputs', 'data_missing_report.json'), 'w', encoding='utf-8') as _f:",
        "        json.dump(_missing_report, _f, ensure_ascii=False, indent=2)",
        "    print('[数据质量] ⚠ 检测到缺失值，已输出 data_missing_report.json（缺失不填充、不编造）: ' + str(_missing_report)[:300])",
        "else:",
        "    print('[数据质量] 无缺失值')",
    ])
    return "\n".join(code_lines)


def _detect_reconcile_scenario(input_dir: Path):
    """检测 序时账×银行流水 组合，返回 (book_path, bank_path) 或 None。

    基于列特征确定性识别（凭证号/科目 → 序时账；对方/账号/余额 → 流水），
    支持 Excel/CSV 与 docx/pdf/md 文档中的表格。
    """
    try:
        from core.bank_reconcile_engine import detect_book_type, JOURNAL, BANK_STATEMENT
        from core.document_loader import load_tables
        journal = bank = None
        for f in sorted(input_dir.glob("*")):
            if f.suffix.lower() not in (".xlsx", ".xls", ".csv", ".docx", ".pdf", ".md"):
                continue
            try:
                tables = load_tables(f)
                if not tables:
                    continue
                t = detect_book_type(tables[0], f.name)
                if t in (JOURNAL, "generic_ledger") and journal is None:
                    journal = f  # 序时账或通用台账均可作为账方
                elif t == BANK_STATEMENT and bank is None:
                    bank = f
            except Exception:
                continue
            if journal and bank:
                break
        if journal and bank:
            # 汇总级台账（按年/分类汇总）不走逐笔快车道 → 交 LLM 按 summary_compare 规划
            try:
                from config.scenario_packs import is_detail_level
                if not is_detail_level(load_tables(journal)[0]):
                    print(f"[对账快车道] {journal.name} 为汇总级数据，交 LLM 汇总勾稽路径")
                    return None
            except Exception:
                pass
            return (journal, bank)
        return None
    except Exception as e:
        print(f"[对账快车道] 场景检测失败（按普通链路处理）: {e}")
        return None


def _account_from_intent(intent: str) -> str:
    """从用户意图提取银行账号线索（如 '农行5927' / 纯数字账号）"""
    m = re.search(r"(?:账号|账户|农行|工行|建行|中行|招行|交行|徽商|邮储|农商)(\D{0,4}?\d{3,})", intent)
    if m:
        return m.group(1)
    m2 = re.search(r"\b(\d{4,})\b", intent)
    return m2.group(1) if m2 else ""


async def _execute_in_sandbox(run_id: str, code: str, run_dir: Path) -> None:
    """在沙箱中执行审计代码（异步非阻塞 + 并发信号量控制）。
    策略：优先 Docker 容器（sandbox_v3），失败/不可用时降级为本地 subprocess。
    """
    import subprocess
    from time import time as _now
    _t0 = _now()
    output_dir = run_dir / "outputs"
    output_dir.mkdir(parents=True, exist_ok=True)
    input_dir = run_dir / "inputs"

    # ═══ 专业对账快车道：序时账×银行流水 → bank_reconcile_engine（平台层执行） ═══
    _rec0 = _get_snapshot_mgr().get_run(run_id)
    _intent0 = (_rec0.user_intent or "") if _rec0 else ""
    _preset0 = (_rec0.preset_button or "") if _rec0 else ""
    # 快车道触发条件：意图含对账关键词 OR 用户点了银行对账预设按钮
    _is_reconcile_intent = any(k in _intent0 for k in ("对账", "核对", "核账", "相符", "对一下"))
    _is_reconcile_preset = False
    if _preset0:
        try:
            from config.presets import normalize_preset_key
            _is_reconcile_preset = (normalize_preset_key(_preset0) == "银行对账")
        except Exception:
            pass
    if _is_reconcile_intent or _is_reconcile_preset:
        _pair = _detect_reconcile_scenario(input_dir)
        if _pair:
            try:
                from core.bank_reconcile_engine import reconcile_files
                _cfg = {}
                _acc = _account_from_intent(_intent0)
                if _acc:
                    _cfg["account"] = _acc

                # ── v3.4 资源护栏：快车道只在安全数据规模下启用 ──
                _FASTLANE_MAX_ROWS = 50000
                _FASTLANE_TIMEOUT = 1800  # 30分钟
                _skip_fastlane = False
                try:
                    import pandas as _pd
                    for _fp, _label in [(_pair[0], "账"), (_pair[1], "银")]:
                        _sz_mb = _fp.stat().st_size / (1024 * 1024)
                        # 文件大小超 50MB → 疑似超大数据，转沙箱
                        if _sz_mb > 50:
                            print(f"[对账快车道] {_label}方文件 {_sz_mb:.0f}MB > 50MB，转沙箱路径")
                            _skip_fastlane = True
                            break
                        # 用 openpyxl 读工作表维度（不加载数据，秒级）
                        _df = _pd.read_excel(_fp, nrows=0)
                        _est_rows = 0
                        try:
                            import openpyxl
                            _wb = openpyxl.load_workbook(_fp, read_only=True)
                            _est_rows = _wb.active.max_row or 0
                            _wb.close()
                        except Exception:
                            _est_rows = int(_sz_mb * 500)  # 降级：粗略估算 ~500行/MB
                        print(f"[对账快车道] {_label}方: {_fp.name} ({_sz_mb:.1f}MB, 估算 ~{_est_rows} 行)")
                        if _est_rows > _FASTLANE_MAX_ROWS:
                            print(f"[对账快车道] 估算行数 {_est_rows} > {_FASTLANE_MAX_ROWS}，转沙箱路径")
                            _skip_fastlane = True
                            break
                except Exception as _sz_e:
                    print(f"[对账快车道] 文件大小探测失败: {_sz_e}，继续快车道")

                if _skip_fastlane:
                    raise RuntimeError("数据规模超快车道阈值，转沙箱执行")

                print(f"[对账快车道] 序时账={_pair[0].name} × 流水={_pair[1].name}, cfg={_cfg}")
                # 进度回调：写入 execution_logs，前端轮询可见
                def _reconcile_progress(pct, step):
                    _save_logs(run_id, [f"[对账进度] {pct}% {step}"], append=True)
                _cfg["progress_callback"] = _reconcile_progress

                # ── 匹配逻辑说明（从场景注册表读取） ──
                try:
                    from config.scenario_packs import get_report_meta
                    _meta = get_report_meta("bank_reconcile_detail")
                    _plan = f"""## {_meta['name']}

### 数据源
- 账方：{_pair[0].name}
- 银方：{_pair[1].name}

### 匹配策略（五层递进）
1. L1 金额精确(±0.01元)+同日 → 自动确认
2. L2 金额精确+日期窗口(±3天) → 自动确认（跨期/在途）
3. L3 n:m 拆分合并 → 已核对（标注"拆分/合并入账"）
4. L3_fee 手续费差额(≤10元+摘要含"手续费") → 自动成组
4+. L3_fee_month 手续费月度聚合(两侧月度总额精确相等) → 自动成组
5. L4 摘要+对方户名模糊匹配 → 待人工复核

### 未匹配分类
- 四分类：银收企未收 / 银付企未付 / 企收银未收 / 企付银未付
- 利息/手续费/冲正 → 专项标记，不删除
- 其余 → 待人工核查（禁止默认洗白为未达）

### 红旗检测
一收一付同额、整数大额、期末负余额、分次转入转出、大额现金、余额不连贯(P5)

### 交付物
逐笔对账明细底稿、银行存款余额调节表、未达账项清单、异常资金交易清单、摘要JSON
"""
                    (_plan_path := run_dir / "match_explanation.txt").write_text(_plan, encoding="utf-8")
                    _save_logs(run_id, ["[匹配逻辑] 已生成详细匹配方案（五层递进+四分类+红旗）"], append=True)
                except Exception:
                    pass
                _res = await asyncio.wait_for(
                    asyncio.to_thread(reconcile_files, _pair[0], _pair[1], _cfg, output_dir),
                    timeout=_FASTLANE_TIMEOUT
                )
                _st = _res["stats"]
                _get_snapshot_mgr().update_status(run_id, "COMPLETED")
                _update_outputs(run_id, output_dir)
                _save_logs(run_id, [
                    f"[对账快车道] 匹配率 账={_st['book_match_rate']}% 银={_st['bank_match_rate']}% "
                    f"(L1={_st['matched_L1']}, L2={_st['matched_L2']}, L3组={_st['matched_L3_groups']}, L4待复核={_st['review_L4']})",
                    f"[对账快车道] 未达四分类: {_st['timing_categories']}",
                    f"[对账快车道] 交付物: {_res.get('output_files')}",
                ])
                try:
                    _generate_report_if_needed(run_id, output_dir, [])
                except Exception as _rg_e:
                    print(f"[对账快车道] 报告生成失败（不影响对账结果）: {type(_rg_e).__name__}: {_rg_e}")
                    _save_logs(run_id, [f"[报告生成] 失败但匹配已完成: {_rg_e}"], append=True)
                return
            except Exception as _e:
                # 状态护栏：已 COMPLETED 的 run 禁止回退覆盖
                _cur = _get_snapshot_mgr().get_run(run_id)
                if _cur and _cur.status == "COMPLETED":
                    print(f"[对账快车道] 异常发生在完成之后，保持 COMPLETED: {_e}")
                    return
                print(f"[对账快车道] 执行失败，回退常规链路: {type(_e).__name__}: {_e}")

    # ═══ 优先：Docker 沙箱（sandbox_v3 完整容器生灭） ═══
    # 双保险：任何路径下，已完成的 run 不再进沙箱
    _rec_chk = _get_snapshot_mgr().get_run(run_id)
    if _rec_chk and _rec_chk.status == "COMPLETED":
        return
    async with _SANDBOX_SEMAPHORE:  # 并发控制：最多 3 个同时运行
        try:
            from engine.sandbox_v3 import EphemeralSandbox
            sandbox = EphemeralSandbox(timeout=120)
            # asyncio.to_thread() 避免阻塞事件循环
            result = await asyncio.to_thread(
                sandbox.execute_with_retry, run_id, code, run_dir, 3
            )
            status = "COMPLETED" if result.status == "COMPLETED" else "FAILED"
            logs = [result.stdout[-2000:]] if result.stdout else []
            if result.stderr:
                logs.append("STDERR: " + result.stderr[-2000:])
            _get_snapshot_mgr().update_status(run_id, status)
            _save_logs(run_id, logs)
            if status == "COMPLETED":
                _update_outputs(run_id, output_dir)
                _generate_report_if_needed(run_id, output_dir, logs)
                _run_constraint_check(run_id, output_dir, logs)
            trace_record(run_id, "sandbox_run",
                         "OK" if status == "COMPLETED" else "FAIL",
                         (_now() - _t0) * 1000, f"retry={result.retry_count}")
            if result.retry_count:
                trace_record(run_id, "self_correct", "OK", 0,
                             f"自纠错 {result.retry_count} 轮")
            print(f"[Sandbox] Docker 执行完成 Run {run_id}: {status}")
            return
        except Exception as e:
            trace_record(run_id, "sandbox_fallback", "WARN",
                         (_now() - _t0) * 1000,
                         f"{type(e).__name__}: 降级本地subprocess")
            print(f"[Sandbox] Docker 执行失败 ({type(e).__name__})，降级为本地 subprocess")

    # ═══ 降级：本地 subprocess（Windows 无 Docker 时的兜底） ═══
    await _execute_in_sandbox_legacy(run_id, code, run_dir)


def _save_logs(run_id: str, logs: list, append: bool = False) -> None:
    try:
        import sqlite3, json
        from core.run_snapshot import DB_PATH
        with sqlite3.connect(str(DB_PATH)) as conn:
            if append:
                existing = conn.execute(
                    "SELECT execution_logs FROM runs WHERE run_id = ?", (run_id,)
                ).fetchone()
                if existing and existing[0]:
                    old_logs = json.loads(existing[0]) if isinstance(existing[0], str) else existing[0]
                    logs = (old_logs or []) + logs
            conn.execute("UPDATE runs SET execution_logs = ? WHERE run_id = ?",
                         (json.dumps(logs, ensure_ascii=False), run_id))
            conn.commit()
    except Exception:
        pass


def _update_outputs(run_id: str, output_dir: Path) -> None:
    output_files = [f.name for f in output_dir.iterdir() if f.is_file()]
    if output_files:
        _get_snapshot_mgr().update_outputs(
            run_id=run_id, output_files=output_files,
            validation_results=[{"check": "execution", "passed": True}],
            all_passed=True)


def _generate_report_if_needed(run_id: str, output_dir: Path, logs: list) -> None:
    try:
        record = _get_snapshot_mgr().get_run(run_id)
        if not record:
            return
        input_dir = record.run_dir / "inputs"
        input_names = [f.name for f in input_dir.iterdir() if f.is_file()] if input_dir.exists() else []
        dag_ops = (record.dag_blueprint or {}).get("operators", [])
        from core.report_generator import generate_audit_report
        rp = generate_audit_report(
            run_id=run_id, user_intent=record.user_intent or "",
            dag_operators=dag_ops, output_dir=output_dir,
            input_files=input_names, execution_logs=logs)
        print(f"[报告] Word 审计报告已生成: {rp.name}")
    except Exception as e:
        print(f"[报告] 生成失败（非致命）: {e}")


async def _call_dify_single_table(catalog_text, user_intent):
    if not DIFY_SINGLE_TABLE_KEY: return ""
    try:
        import httpx
        async with httpx.AsyncClient(timeout=httpx.Timeout(300.0, connect=10.0)) as client:
            resp = await client.post(
                f"{DIFY_BASE_URL}/v1/workflows/run",
                headers={"Authorization": f"Bearer {DIFY_SINGLE_TABLE_KEY}"},
                json={"inputs": {"catalog_text": catalog_text, "user_intent": user_intent}, "response_mode": "blocking", "user": "audit_platform"},
            )
            resp.raise_for_status()
            return resp.json().get("data", {}).get("outputs", {}).get("result", "")
    except: return ""


def _parse_single_table_result(result, catalog):
    from core.dag_compiler import DAGBlueprint, Operator
    import json as _j
    ops = []
    files = getattr(catalog, "files", [])
    fname = files[0].get("filename", "input.xlsx") if files else "input.xlsx"
    ops.append(Operator(id="load_0", name="Load", input_from=[], source_file=fname, output_alias="df_raw"))
    try:
        if result.strip().startswith("{"):
            data = _j.loads(result)
            if "operators" in data:
                return DAGBlueprint(blueprint_id="single_001", generated_at=datetime.now().isoformat(),
                    operators=[Operator(**o) if isinstance(o, dict) else o for o in data["operators"]],
                    objective=data.get("objective", "单表筛选"))
    except: pass
    ops.append(Operator(id="filter_0", name="ColumnFilter", input_from=["load_0"], output_alias="df_filtered",
        params={"column": "金额", "operator": ">", "value": 0}))
    ops.append(Operator(id="export_0", name="Export", input_from=["filter_0"], params={"filename": "筛选结果.xlsx"}))
    return DAGBlueprint(blueprint_id="single_fb", generated_at=datetime.now().isoformat(), operators=ops,
        objective="单表筛选", raw_intent="从单表筛选数据并导出")
DIFY_REFINE_API_KEY = os.environ.get("DIFY_REFINE_API_KEY", "")
DIFY_SINGLE_TABLE_KEY = os.environ.get("DIFY_SINGLE_TABLE_KEY", "")
DIFY_REPORT_GEN_KEY = os.environ.get("DIFY_REPORT_GEN_KEY", "")
DIFY_REPORT_REVIEW_KEY = os.environ.get("DIFY_REPORT_REVIEW_KEY", "")
DIFY_KNOWLEDGE_QA_KEY = os.environ.get("DIFY_KNOWLEDGE_QA_KEY", "")

def _run_constraint_check(run_id, output_dir, logs):
    try:
        record = _get_snapshot_mgr().get_run(run_id)
        if not record or not record.user_intent: return
        from core.constraint_engine import parse_constraints, format_constraint_report
        constraints = parse_constraints(record.user_intent)
        if not constraints: return
        excel_files = list(output_dir.glob("*.xlsx"))
        import pandas as pd
        all_passed = True
        for xf in excel_files:
            try:
                df = pd.read_excel(xf)
                for c in constraints:
                    if c.field == "差异金额" and "差异金额" in df.columns:
                        max_diff = df["差异金额"].abs().max()
                        c.actual_value = max_diff
                        c.satisfied = max_diff < c.value
                        if not c.satisfied: all_passed = False
            except: continue
        report = format_constraint_report(constraints)
        logs.append(f"[约束校验] {'PASS' if all_passed else 'FAIL'}: {report}")
        if not all_passed:
            _get_snapshot_mgr().update_status(run_id, "FAILED")
            _run_constraint_retry(run_id, record.user_intent, logs)
    except Exception as e:
        print(f"[约束校验] 校验失败(非致命): {e}")


def _write_audit_trace(run_id, user_intent, dag_blueprint):
    import json as _j
    p = Path(os.environ.get("AUDIT_TRACE_DIR", str(Path(__file__).resolve().parent.parent / "data" / "audit_traces")))
    p.mkdir(parents=True, exist_ok=True)
    try:
        with open(p / "v5_audit_traces.jsonl", "a", encoding="utf-8") as f:
            f.write(_j.dumps({"run_id": run_id, "timestamp": datetime.now().isoformat(), "instruction": user_intent or "", "output": _j.dumps(dag_blueprint or {}, ensure_ascii=False)}, ensure_ascii=False) + "\n")
    except Exception as e:
        print(f"[审计采集] 写入失败(非致命): {e}")


def _append_approval_hash(run_id, user_intent, dag_blueprint, confirmed):
    import hashlib, sqlite3, json as _j
    try:
        from core.run_snapshot import DB_PATH
        chain_data = _j.dumps({"run_id": run_id, "timestamp": datetime.now().isoformat(), "intent": user_intent[:200], "confirmed": confirmed}, ensure_ascii=False, sort_keys=True)
        current_hash = hashlib.sha256(chain_data.encode()).hexdigest()
        with sqlite3.connect(str(DB_PATH)) as conn:
            prev = conn.execute("SELECT approval_hash FROM approval_chain ORDER BY id DESC LIMIT 1").fetchone()
            conn.execute("CREATE TABLE IF NOT EXISTS approval_chain (id INTEGER PRIMARY KEY AUTOINCREMENT, run_id TEXT, hash TEXT, prev_hash TEXT, data TEXT, created_at TEXT)")
            conn.execute("INSERT INTO approval_chain (run_id, hash, prev_hash, data, created_at) VALUES (?,?,?,?,?)", (run_id, current_hash, prev[0] if prev else "GENESIS", chain_data, datetime.now().isoformat()))
            conn.commit()
    except Exception as e:
        print(f"[哈希链] 记录失败(非致命): {e}")


def _run_constraint_retry(run_id, user_intent, logs):
    try:
        from core.constraint_engine import parse_constraints
        constraints = parse_constraints(user_intent)
        failed = [c for c in constraints if hasattr(c, "satisfied") and not c.satisfied]
        if not failed: return False
        hint = "约束未满足："
        for c in failed: hint += f" {c.field} 实际={c.actual_value} vs 要求<={c.value};"
        logs.append(f"[约束重试] {hint}")
        return True
    except: return False


def _ensure_essential_operators(blueprint_dict, user_intent):
    ops = list(blueprint_dict.get("operators", []))
    existing = {op.get("name", "") for op in ops}
    max_id = max((int(op.get("id", "").replace("op_", "0")) for op in ops if op.get("id", "").replace("op_", "").isdigit()), default=0)
    last = [ops[-1]["id"]] if ops else []
    changed = False
    # load_count 必须先计算，后续规则都要用
    load_count = sum(1 for op in ops if op.get("name") in ("Load", "load"))
    try:
        from config.scenario_packs import detect_scenario as _ds2, required_ops_for
        _scn_req = required_ops_for(_ds2(user_intent or ""))
    except Exception:
        _scn_req = []
    if "Reconcile" in _scn_req and "Reconcile" not in existing and load_count >= 2:
        max_id += 1; load_ids2 = [op["id"] for op in ops if op.get("name") in ("Load", "load")]
        ops.append({"id": f"op_{max_id}", "name": "Reconcile", "input_from": load_ids2[:2],
                    "params": {"tolerance_abs": 0.01, "date_window_days": 3},
                    "output_alias": f"df_reconciled_{max_id}"})
        existing.add("Reconcile"); changed = True
    need_merge = any(k in user_intent for k in ["核对", "对账", "匹配", "比对", "两表"])
    need_diff = any(k in user_intent for k in ["差异", "对比", "比较", "差额"])
    need_group = any(k in user_intent for k in ["汇总", "分组", "按月"])
    need_sort = any(k in user_intent for k in ["排序", "从大到小", "从小到大"])
    need_save = any(k in user_intent for k in ["导出", "保存", "下载"])
    if need_merge and "Merge" not in existing and load_count >= 2:
        max_id += 1; load_ids = [op["id"] for op in ops if op.get("name") in ("Load", "load")]
        ops.append({"id": f"op_{max_id}", "name": "Merge", "input_from": load_ids, "params": {"how": "outer", "on": "auto"}, "output_alias": f"df_merged_{max_id}"}); last = [f"op_{max_id}"]; existing.add("Merge"); changed = True
    if need_group and "GroupBy" not in existing:
        max_id += 1
        ops.append({"id": f"op_{max_id}", "name": "GroupBy", "input_from": last, "params": {"by": ["auto"], "aggregations": {"auto": "sum"}}, "output_alias": f"df_grouped_{max_id}"}); last = [f"op_{max_id}"]; existing.add("GroupBy"); changed = True
    # 若 DAG 里已有 Reconcile（专业对账），不再因“比较”等关键字追加 Diff，避免对账+Diff 重复/冲突
    if need_diff and "Diff" not in existing and "Reconcile" not in existing and ("Merge" in existing or load_count >= 2):
        max_id += 1
        ops.append({"id": f"op_{max_id}", "name": "Diff", "input_from": last, "params": {"col_a": "auto", "col_b": "auto", "tolerance_abs": 0.01}, "output_alias": f"df_diff_{max_id}"}); last = [f"op_{max_id}"]; existing.add("Diff"); changed = True
    if need_sort and "Sort" not in existing:
        max_id += 1; sc = "auto"
        ops.append({"id": f"op_{max_id}", "name": "Sort", "input_from": last, "params": {"by": [sc], "ascending": False}, "output_alias": f"df_sorted_{max_id}"}); last = [f"op_{max_id}"]; existing.add("Sort"); changed = True
    if need_save and "Export" not in existing:
        max_id += 1
        ops.append({"id": f"op_{max_id}", "name": "Export", "input_from": last, "params": {"filename": "审计结果.xlsx"}, "output_alias": f"df_output_{max_id}"}); changed = True
    if changed: blueprint_dict["operators"] = ops
    return blueprint_dict


async def _execute_in_sandbox_legacy(run_id: str, code: str, run_dir: Path) -> None:
    try:
        output_dir = run_dir / "outputs"
        output_dir.mkdir(parents=True, exist_ok=True)
        input_dir = run_dir / "inputs"

        # 🔥 匹配引擎优先：如果检测到双文件匹配场景，调用匹配引擎
        record = _get_snapshot_mgr().get_run(run_id)
        intent = (record.user_intent or "") if record else ""
        input_files = list(input_dir.iterdir()) if input_dir.exists() else []
        excel_files = [f for f in input_files if f.suffix.lower()
                       in (".xlsx", ".xls", ".csv", ".docx", ".doc", ".pdf", ".md", ".txt")]

        # ═══ 专业对账快车道（legacy 路径）：序时账×流水 → bank_reconcile_engine ═══
        if any(k in intent for k in ("对账", "核对", "核账", "相符", "对一下")):
            _pair = _detect_reconcile_scenario(input_dir)
            if _pair:
                try:
                    from core.bank_reconcile_engine import reconcile_files
                    _cfg = {}
                    _acc = _account_from_intent(intent)
                    if _acc:
                        _cfg["account"] = _acc
                    print(f"[对账快车道] 序时账={_pair[0].name} × 流水={_pair[1].name}, cfg={_cfg}")
                    _res = reconcile_files(_pair[0], _pair[1], _cfg, output_dir)
                    _st = _res["stats"]
                    output_files_final = [f.name for f in output_dir.iterdir() if f.is_file()]
                    _get_snapshot_mgr().update_outputs(
                        run_id=run_id, output_files=output_files_final,
                        validation_results=[{"check": "bank_reconcile_engine", "passed": True}],
                        all_passed=True)
                    _get_snapshot_mgr().update_status(run_id, "COMPLETED")
                    _save_logs(run_id, [
                        f"[对账快车道] 匹配率 账={_st['book_match_rate']}% 银={_st['bank_match_rate']}%",
                        f"[对账快车道] 未达四分类: {_st['timing_categories']}",
                        f"[对账快车道] 交付物: {_res.get('output_files')}"])
                    if record:
                        _generate_report_if_needed(run_id, output_dir, [])
                    return
                except Exception as _e:
                    print(f"[对账快车道] 执行失败，回退常规链路: {_e}")

        if len(excel_files) >= 2 and ("匹配" in intent or "核对" in intent or "对账" in intent or "比对" in intent):
            print(f"[Sandbox] 检测到匹配场景（{len(excel_files)}个文件），调用匹配引擎...")
            logs = []
            status = "COMPLETED"
            try:
                # 历史文件保护：验证 dag_blueprint 中的文件引用与当前输入目录一致
                # 若不一致（可能是历史缓存），则仅使用当前输入目录的实际文件
                if record and record.dag_blueprint:
                    dag_bp = record.dag_blueprint
                    if isinstance(dag_bp, dict):
                        ops = dag_bp.get("operators", [])
                        actual_input_names = {f.name for f in input_files if f.is_file()}
                        for op in ops:
                            sf = op.get("source_file", "")
                            if sf and sf not in actual_input_names:
                                print(f"[Sandbox] 注意: dag_blueprint 中引用文件 '{sf}' 不在当前输入目录，将仅使用当前上传文件进行分析")
                                break

                from core.matching_engine import run_matching_pipeline
                # ── 三级关键词供给链（统一入口：core.keyword_resolver）───
                # ① DAG 蓝图已有 pattern → 直接用（LLM 编译阶段确定的）
                dag_patterns = _extract_patterns_from_dag(record.dag_blueprint) if record else ""
                kw_source = "dag" if dag_patterns else ""
                kw_version = ""
                kw_preview = None

                if not dag_patterns and record and record.user_intent:
                    # ② 词典命中 / 提案分流
                    from core.keyword_resolver import (
                        resolve_patterns_full as _kw_resolve,
                        backtest_patterns as _kw_backtest,
                        propose_via_search as _kw_propose,
                        save_proposal as _kw_save_proposal,
                    )
                    kw_res = _kw_resolve(record.user_intent)
                    if kw_res["status"] == "hit":
                        # 词典命中 → 直通（免确认）
                        dag_patterns = kw_res["patterns"]
                        kw_source = f"dictionary_{kw_res['version']}"
                        kw_version = kw_res["version"]
                        print(f"[关键词] 词典命中: {kw_res['category']} → {dag_patterns[:60]}")
                    else:
                        # ③ 提案路径：搜索 → LLM起草 → 预览 → 挂起等用户确认
                        print(f"[关键词] 词典未命中，触发拓荒提案: {record.user_intent[:60]}")
                        try:
                            proposal = _kw_propose(record.user_intent)
                        except Exception as _pe:
                            print(f"[关键词] 提案失败: {_pe}，回退 LLM 提取")
                            from core.matching_engine import _extract_patterns_via_llm
                            proposal = {"patterns": _extract_patterns_via_llm(record.user_intent),
                                        "依据摘要": "LLM 直接提取（搜索不可用）",
                                        "sources": []}

                        # 加载流水文件做预览
                        preview = {}
                        try:
                            bank_files = [f for f in input_files if f.is_file()
                                          and f.suffix.lower() in (".xlsx", ".xls", ".csv")]
                            if bank_files and proposal.get("patterns"):
                                import pandas as _pd2
                                _sample = _pd2.read_excel(bank_files[0], nrows=5000) \
                                    if bank_files[0].suffix.lower() != ".csv" \
                                    else _pd2.read_csv(bank_files[0], nrows=5000,
                                                        encoding="utf-8-sig")
                                preview = _kw_backtest(
                                    proposal["patterns"],
                                    _sample,
                                    ["摘要", "对方客户名称", "附言", "用途"],
                                )
                                print(f"[关键词预览] 命中 {preview['hit_count']} 行 "
                                      f"({preview['hit_rate']:.1%})，"
                                      f"排除TOP: {list(preview.get('excluded_top', {}).keys())[:3]}")
                        except Exception as _pv:
                            print(f"[关键词预览] 生成失败（非致命）: {_pv}")

                        # 保存提案 + 挂起等用户确认
                        _kw_save_proposal(run_id, {**proposal, "preview": preview})
                        _get_snapshot_mgr().update_status(run_id, "PENDING_KEYWORD_CONFIRM")
                        logs.append(f"[关键词提案] 已生成候选词条，等待用户确认")
                        print(f"[关键词] Run {run_id} 挂起，状态=PENDING_KEYWORD_CONFIRM")
                        return  # ← 挂起，等用户确认
                if not dag_patterns:
                    dag_patterns = ""  # 无关键词则全量匹配
                    kw_source = "fallback"
                print(f"[匹配引擎] pattern: '{dag_patterns[:80] if dag_patterns else '(全量)'}' "
                      f"(来源: {kw_source})")

                # 从用户意图中提取容差阈值（默认1%，用户可自然语言指定）
                from core.constraint_engine import extract_tolerance
                tolerance_pct = extract_tolerance(record.user_intent or "", default_pct=1.0)
                print(f"[Sandbox] 容差阈值: {tolerance_pct}%")

                match_result = run_matching_pipeline(input_dir, output_dir, patterns=dag_patterns,
                                                     kw_source=kw_source, kw_version=kw_version)
                logs.append("[匹配引擎] 匹配流水线执行完成")

                # 如果成功，更新输出文件列表
                output_files_final = [f.name for f in output_dir.iterdir() if f.is_file()]
                if output_files_final:
                    _get_snapshot_mgr().update_outputs(
                        run_id=run_id, output_files=output_files_final,
                        validation_results=[{"check": "matching_engine", "passed": True}],
                        all_passed=True,
                    )

                # 生成报告
                if record:
                    dag_ops = (record.dag_blueprint or {}).get("operators", [])
                    input_names = [f.name for f in input_files if f.is_file()]
                    from core.report_generator import generate_audit_report
                    # 收集匹配逻辑信息（含用户确认的自然语言说明 + 匹配引擎实际执行结果）
                    dag_bp = record.dag_blueprint or {}
                    explanation = dag_bp.get("match_explanation", "") if isinstance(dag_bp, dict) else ""
                    engine_match_logic = match_result.get("match_logic", {}) if match_result else {}
                    engine_stats = match_result.get("match_stats", {}) if match_result else {}
                    match_info = {
                        "patterns": dag_patterns or engine_match_logic.get("筛选模式", "未指定"),
                        "columns": engine_match_logic.get("筛选列", []),
                        "amount_column": engine_match_logic.get("金额列", ""),
                        "institution_columns": engine_match_logic.get("机构识别列", []),
                        "method": match_result.get("strategy_name", "多列联合匹配") if match_result else "多列联合匹配",
                        "explanation": explanation,
                        "strategy_comparison": match_result.get("all_strategies", []) if match_result else [],
                        # 关键词来源 + 预览（审计追溯用）
                        "kw_source": kw_source,
                        "kw_preview": kw_preview or {},
                        # 反向校验：未匹配行摘要聚类
                        "reverse_validation": _reverse_validation_for_report(
                            match_result, input_dir) if match_result else {},
                    }
                    rp = generate_audit_report(
                        run_id=run_id, user_intent=record.user_intent or "",
                        dag_operators=dag_ops, output_dir=output_dir,
                        input_files=input_names, execution_logs=logs,
                        match_logic=match_info,
                    )
                    output_files_final.append(rp.name)
                    print(f"[报告] Word 审计报告已生成: {rp.name}")

                if output_files_final:
                    _get_snapshot_mgr().update_outputs(
                        run_id=run_id, output_files=output_files_final,
                        validation_results=[{"check": "matching_engine", "passed": True}],
                        all_passed=True,
                    )
                _get_snapshot_mgr().update_status(run_id, "COMPLETED")
                print(f"[Sandbox] Run {run_id} 匹配完成")
                return

            except Exception as e:
                print(f"[Sandbox] 匹配引擎失败: {e}，回退到 DAG 执行")
                logs.append(f"[匹配引擎] 失败: {e}，回退到 DAG")
                # 不回退，继续走 DAG 方式
                pass

        # DAG 方式执行（带 OpenClaw 自纠错闭环，白皮书 §5.2：规则优先、LLM兜底）
        from engine.code_corrector import correct_code
        max_attempts = 3
        logs = []
        proc = None
        current_code = code
        for attempt in range(1, max_attempts + 1):
            suffix = "" if attempt == 1 else f"_retry{attempt}"
            script_path = run_dir / f"_run_script{suffix}.py"
            script_path.write_text(current_code, encoding="utf-8")
            proc = subprocess.run(
                [sys.executable, str(script_path)],
                cwd=str(run_dir),
                capture_output=True, text=True, timeout=120
            )
            if proc.returncode == 0:
                if attempt > 1:
                    logs.append(f"[自纠错] 第 {attempt} 次执行成功（经 {attempt - 1} 轮代码修正）")
                break
            logs.append(f"[自纠错] 第 {attempt} 次执行失败: {(proc.stderr or '')[-500:]}")
            if attempt < max_attempts:
                intent_ctx = (record.user_intent or "") if record else ""
                input_files = list(input_dir.iterdir()) if input_dir.exists() else []
                col_info = []
                for f in input_files[:5]:
                    try:
                        import pandas as pd
                        df = pd.read_excel(f, nrows=0) if f.suffix in ('.xlsx','.xls') else pd.read_csv(f, nrows=0)
                        col_info.append(f"{f.name}: {list(df.columns)}")
                    except: pass
                if col_info:
                    intent_ctx += f"\n数据列名: {'; '.join(col_info)}"
                fixed = correct_code(current_code, proc.stderr or "", attempt=attempt, context=intent_ctx)
                if not fixed:
                    logs.append("[自纠错] 无可用修正方案，停止重试")
                    break
                current_code = fixed
        status = "COMPLETED" if proc and proc.returncode == 0 else "FAILED"
        if proc and proc.stdout:
            logs.append(proc.stdout[-2000:])
        if proc and proc.stderr and proc.returncode != 0:
            logs.append("STDERR: " + proc.stderr[-2000:])
        _get_snapshot_mgr().update_status(run_id, status)
        # 保存日志到数据库
        try:
            from core.run_snapshot import DB_PATH
            with sqlite3.connect(str(DB_PATH)) as conn:
                conn.execute(
                    "UPDATE runs SET execution_logs = ? WHERE run_id = ?",
                    (json.dumps(logs, ensure_ascii=False), run_id)
                )
                conn.commit()
        except Exception:
            pass
        # 如果成功，更新输出文件列表
        if status == "COMPLETED":
            output_files = [f.name for f in output_dir.iterdir() if f.is_file()]
            if output_files:
                validations = [{"check": "local_execution", "passed": True}]
                _get_snapshot_mgr().update_outputs(
                    run_id=run_id,
                    output_files=output_files,
                    validation_results=validations,
                    all_passed=True,
                )
            # 生成 Word 审计报告
            try:
                record = _get_snapshot_mgr().get_run(run_id)
                if record:
                    input_names = [f.name for f in record.input_dir.iterdir() if f.is_file()] if record.input_dir.exists() else []
                    dag_ops = (record.dag_blueprint or {}).get("operators", [])
                    from core.report_generator import generate_audit_report
                    rp = generate_audit_report(
                        run_id=run_id,
                        user_intent=record.user_intent or "",
                        dag_operators=dag_ops,
                        output_dir=output_dir,
                        input_files=input_names,
                        execution_logs=logs,
                    )
                    output_files.append(rp.name)
                    print(f"[报告] Word 审计报告已生成: {rp.name}")
            except Exception as e:
                print(f"[报告] 生成失败（非致命）: {e}")
            # 更新输出列表（包含报告）
            if output_files:
                _get_snapshot_mgr().update_outputs(
                    run_id=run_id,
                    output_files=output_files,
                    validation_results=[{"check": "local_execution", "passed": True}],
                    all_passed=True,
                )
        print(f"[Sandbox] Run {run_id} 完成: status={status}, outputs={list(output_dir.iterdir()) if output_dir.exists() else []}")
    except Exception as e:
        print(f"[Sandbox] Run {run_id} 执行异常: {e}")
        _get_snapshot_mgr().update_status(run_id, "FAILED")


@router.get("/runs/{run_id}/quality", summary="数据质量报告（B4）")
async def get_quality_report(run_id: str):
    """返回 Run 的数据质量报告卡"""
    record = _get_snapshot_mgr().get_run(run_id)
    if not record:
        raise HTTPException(status_code=404, detail="Run 不存在")
    qp = record.run_dir / "quality_report.json"
    if qp.exists():
        return {"run_id": run_id,
                "report": json.loads(qp.read_text(encoding="utf-8"))}
    return {"run_id": run_id, "report": [], "message": "质量报告未生成（文件可能已删除）"}


# ── B5 意图澄清回路 ─────────────────────────────────

_STANDARD_INTENTS = {
    "对账": ["核对", "对账", "匹配", "比对", "核实", "核对"],
    "函证": ["函证", "询证", "确认函"],
    "抽样": ["抽样", "抽凭", "抽选"],
    "报告": ["报告", "底稿", "生成报告", "出具"],
    "分析": ["分析", "波动", "趋势", "异常"],
}

_SLOT_REQUIRED = {
    "核对": ["数据键"],
    "对账": ["银行账户", "对账期间", "数据键"],
    "匹配": ["匹配维度（金额/日期/对手方）"],
    "函证": ["阈值", "模板类型"],
    "抽样": ["抽样方法", "样本量"],
    "报告": ["底稿来源"],
    "分析": ["指标", "期间"],
}


@router.post("/runs/{run_id}/clarify", summary="意图澄清检查")
async def clarify_intent(run_id: str):
    """检查审计师指令是否有缺失槽位，返回待追问项。

    编译前调用此端点，若返回 missing_slots 非空，应弹窗追问；
    全部澄清后再触发编译。
    """
    record = _get_snapshot_mgr().get_run(run_id)
    if not record:
        raise HTTPException(status_code=404, detail="Run 不存在")
    intent = (record.user_intent or "").strip()
    if not intent:
        return {"run_id": run_id, "intent": "", "complete": False,
                "missing_slots": ["请描述审计任务"]}

    # 关键词匹配
    matched_scene = None
    for scene, keywords in _STANDARD_INTENTS.items():
        if any(k in intent for k in keywords):
            matched_scene = scene
            break

    if not matched_scene:
        return {"run_id": run_id, "intent": intent, "complete": True,
                "scene": "通用任务", "missing_slots": []}

    # 检查必要槽位
    missing = []
    for kw, slots in _SLOT_REQUIRED.items():
        if kw in intent:
            # 容差检查：有数字+单位词（万/元/%/以内/不超过）则视为已填
            if "容差" in slots and not re.search(r"\d+\s*(万|元|%|以内|不超过|以内)", intent):
                missing.append("容差阈值（如'5万以内'或'10%'；逐笔银行核对默认 ±0.01 元无需指定）")
            if "数据键" in slots and not re.search(r"按\s*\S+|根据\s*\S+|用\s*\S+", intent):
                missing.append("对账依据（如'按金额+日期'或'按凭证号'，不指定则自动识别）")
            if "银行账户" in slots and not re.search(r"(账号|账户|\d{6,})", intent):
                missing.append("对账银行账户/账号（如'农行5927'；流水含多账户时强烈建议指定）")
            if "对账期间" in slots and not re.search(r"(\d{4}\s*年|\d+\s*个?月|季度|期间|年度|\d{4}[-/]\d+)", intent):
                missing.append("对账期间（如'2026年1-3月'）")
            if "抽样方法" in slots and not re.search(r"(MUS|货币单位|随机|分层|系统选样|PPS)", intent, re.I):
                missing.append("抽样方法（MUS货币单位/简单随机/分层，默认 MUS）")
            if "阈值" in slots and not re.search(r"\d+\s*(万|元|以上|超过|大于)", intent):
                missing.append("筛选阈值（如'超过50万'）")
            if "样本量" in slots and not re.search(r"\d+\s*笔|抽\s*\d+|取\s*\d+", intent):
                missing.append("样本量（如'抽20笔'）")
            break

    return {
        "run_id": run_id,
        "intent": intent,
        "scene": matched_scene,
        "complete": len(missing) == 0,
        "missing_slots": missing,
        "hint": f"建议补充：{'；'.join(missing)}" if missing else "指令完整"
    }


# ═══════════════════════════════════════════════════════════════
# 链路可观测性 API（计划 B2）
# ═══════════════════════════════════════════════════════════════

@router.get("/runs/{run_id}/trace", summary="单 Run 链路瀑布")
async def get_run_trace(run_id: str):
    from core.pipeline_trace import get_trace, waterfall
    return {"run_id": run_id, "events": get_trace(run_id),
            "waterfall": waterfall(run_id)}


@router.get("/pipeline/stats", summary="链路聚合统计（近N个Run）")
async def get_pipeline_stats(last_n: int = 50):
    from core.pipeline_trace import stats
    return stats(last_n)


@router.get("/presets", summary="预设按钮统一注册表")
async def list_presets():
    """前端预设按钮列表（单一事实来源 config/presets.py；前端动态渲染用）"""
    try:
        from config.presets import public_list
        return {"success": True, "presets": public_list()}
    except Exception as e:
        return {"success": False, "error": str(e), "presets": []}


@router.get("/internal-kb/summary", summary="事务所内部知识库概览")
async def internal_kb_summary():
    """内部文件（SOP/底稿模板/询证函范本）文件名索引概览（不读文件内容）"""
    try:
        from core.internal_kb_registry import build_index, completion_checklist, summary
        idx = build_index()
        return {"success": True, "summary_text": summary(),
                "categories": idx.get("categories", {}),
                "templates_count": len(idx.get("templates", [])),
                "confirmation_forms": [t["rel"] for t in idx.get("confirmation_forms", [])],
                "completion_required": completion_checklist()}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ═══════════════════════════════════════════════════════════════
# RAG 知识库管理 API
# ═══════════════════════════════════════════════════════════════

@router.get("/rag/status")
async def rag_status():
    """获取 RAG 知识库完整状态（索引统计 + 目录扫描 + 新鲜度检查）"""
    try:
        from core.rag_admin import get_rag_status
        return {"success": True, "data": get_rag_status()}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/rag/search")
async def rag_search(query: str = Form(...), top_k: int = Form(5)):
    """RAG 知识库检索（供前端知识问答使用）"""
    try:
        results = await asyncio.to_thread(_rag_retrieve_best, query, top_k)
        return {"success": True, "data": results}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/rag/rebuild")
async def rag_rebuild(force: bool = True):
    """强制重建 RAG 索引（当审计师下载新法规后调用）"""
    try:
        from core.rag_admin import rebuild_index
        result = rebuild_index(force=force)
        return {"success": True, "data": result}
    except Exception as e:
        return {"success": False, "error": str(e)}



def _rag_retrieve_best(query: str, top_k: int) -> list:
    """检索质量优先级：混合检索（向量+关键词 RRF 融合）→ 纯 TF-IDF"""
    from core.rag_engine import hybrid_retrieve, retrieve
    try:
        results = hybrid_retrieve(query, top_k=top_k)
        if results:
            return results
    except Exception as e:
        print(f"[RAG] 混合检索不可用，降级 TF-IDF: {e}")
    return retrieve(query, top_k=top_k)


@router.post("/rag/qa")
async def rag_qa(
    query: str = Form(...),
    history: str = Form("[]"),
    top_k: int = Form(5),
):
    """知识问答（RAG 检索 + LLM 生成式回答，支持多轮追问）

    白皮书 §4.1 知识检索模式的完整实现：检索增强"生成"，而非裸检索。
    降级链：vLLM 合成回答 → 纯检索片段（AI 离线时前端仍可用）。
    """
    # 1) 检索（线程池执行，避免阻塞事件循环）
    try:
        chunks = await asyncio.to_thread(_rag_retrieve_best, query, top_k)
    except Exception as e:
        print(f"[RAG-QA] 检索失败: {e}")
        chunks = []

    # 2) 解析多轮对话历史（仅保留最近 6 条合法消息）
    try:
        hist_raw = json.loads(history or "[]")
    except Exception:
        hist_raw = []
    hist = [
        {"role": m["role"], "content": str(m["content"])[:2000]}
        for m in (hist_raw if isinstance(hist_raw, list) else [])
        if isinstance(m, dict) and m.get("role") in ("user", "assistant") and m.get("content")
    ][-6:]

    # 3) 组装生成上下文（片段 + 来源标注）
    context_blocks = [
        f"【片段{i} | 来源: {c.get('source', '未知')}】\n{(c.get('text') or '')[:800]}"
        for i, c in enumerate(chunks, 1)
    ]
    context_text = "\n\n".join(context_blocks) if context_blocks else "（知识库未检索到直接相关内容）"

    system_prompt = (
        "你是会计师事务所的审计智能助手，内嵌于本地审计平台。回答规则：\n"
        "1. 先直接回答问题（分步骤/分要点、可执行），不要只罗列法条原文；\n"
        "2. 引用知识库片段时标注（依据：文件名）；知识库覆盖不到时用审计通用实务回答，"
        "并注明（通用实务经验，请以准则原文为准）；\n"
        "3. 若问题涉及平台可代办的任务（对账/函证/底稿/报告生成/格式规范化），"
        "补充说明操作路径：在工作台上传数据 → 用一句话下达指令（如'根据底稿生成报告正文'）→ 审批后自动执行；\n"
        "4. 用中文回答，400 字以内。"
    )
    messages = [{"role": "system", "content": system_prompt}]
    messages.extend(hist)
    messages.append({"role": "user", "content": f"【知识库片段】\n{context_text}\n\n【问题】\n{query}"})

    # 4) LLM 合成回答（失败自动降级为纯检索）
    answer, engine = "", "retrieval_only"
    try:
        client = _get_http_client()
        resp = await client.post(
            VLLM_TUNNEL_URL,
            headers={"Authorization": f"Bearer {VLLM_API_KEY}"},
            json={"model": VLLM_MODEL, "messages": messages,
                  "temperature": 0.3, "max_tokens": 1024},
        )
        resp.raise_for_status()
        answer = resp.json()["choices"][0]["message"]["content"].strip()
        answer = _re.sub(r"<think>[\s\S]*?</think>", "", answer).strip()
        engine = "vllm_rag" if chunks else "vllm_general"
    except Exception as e:
        print(f"[RAG-QA] vLLM 不可用，降级纯检索: {e}")

    return {"success": True, "data": {
        "answer": answer,
        "engine": engine,
        "sources": [
            {"source": c.get("source", ""), "category": c.get("category", ""),
             "score": c.get("score", 0), "text": (c.get("text") or "")[:300]}
            for c in chunks
        ],
    }}


@router.get("/rag/freshness")
async def rag_freshness():
    """检查索引新鲜度（是否有新文件需要重建）"""
    try:
        from core.rag_admin import check_index_freshness
        return {"success": True, "data": check_index_freshness()}
    except Exception as e:
        return {"success": False, "error": str(e)}








# ═══════════════════════════════════════════════════════════════
# Tool-Use Agent API（一个端点处理一切文档操作）
# ═══════════════════════════════════════════════════════════════

@router.post("/agent/run")
async def agent_run_endpoint(
    message: str = Form(...),
    files: List[UploadFile] = File([]),
    background_tasks: BackgroundTasks = BackgroundTasks(),
):
    """
    Tool-Use Agent：上传任意文档 + 自然语言指令 → LLM 自行选择工具处理。

    支持的文件格式：md / docx / xlsx / xls / txt / csv
    支持的指令：解析、提取、填充、改写、合并、对比、格式转换、批量格式化、生成报告、知识检索

    例：
      "帮我把第三章提取出来" + 上传报告.md
      "用这份模板统一这五份Excel的格式" + 上传模板.xlsx + 五份文件
      "对比这两份报告，找出数字不一样的地方" + 上传两份.docx
    """
    try:
        import time, shutil

        # 保存上传文件到沙箱
        upload_dir = Path("data/uploads")
        upload_dir.mkdir(parents=True, exist_ok=True)

        file_infos = []
        for f in (files or []):
            if not f.filename: continue
            safe_name = f"{int(time.time())}_{f.filename}"
            dest = upload_dir / safe_name
            with open(dest, "wb") as df:
                df.write(await f.read())
            file_infos.append({
                "name": f.filename,
                "path": str(dest.relative_to("data")),
                "format": Path(f.filename).suffix.lower(),
            })

        # 调用 Agent（异步非阻塞）
        from core.agent import agent_run
        result = await agent_run(message, file_infos)

        return {"success": True, "reply": result, "files_processed": len(file_infos)}
    except Exception as e:
        return {"success": False, "error": str(e)}



# ══════════════════════════════════════════════════
# 多Agent协作管线 API
# ══════════════════════════════════════════════════

AGENT_PIPELINE_EXTENSIONS = {".md", ".txt", ".xlsx", ".xls", ".csv", ".docx", ".doc"}

@router.post("/agent/pipeline")
async def agent_pipeline(
    message: str = Form(...),
    stages: str = Form(""),
    files: List[UploadFile] = File([]),
):
    """多Agent协作管线：问题抽取 → 逻辑推理 → 法规检索 → 报告撰写。

    复杂问题自动拆解：四个专业 Agent 串行协作，前序输出注入后序上下文。
    stages: 逗号分隔阶段（issue_extractor,logic_reasoner,regulation_searcher,report_writer），
            留空 = 全部 4 个阶段。
    """
    import time as _time
    t0 = _time.time()
    try:
        upload_dir = Path("data/uploads")
        upload_dir.mkdir(parents=True, exist_ok=True)

        file_infos = []
        for f in (files or []):
            if not f or not f.filename:
                continue
            safe_name = Path(f.filename).name
            ext = Path(safe_name).suffix.lower()
            if safe_name != f.filename:
                raise HTTPException(status_code=400, detail=f"文件名包含非法路径字符: {f.filename}")
            if ext not in AGENT_PIPELINE_EXTENSIONS:
                raise HTTPException(
                    status_code=400,
                    detail=f"不支持的文件类型 {ext}（允许: {', '.join(sorted(AGENT_PIPELINE_EXTENSIONS))}）")
            content = await f.read()
            if len(content) > MAX_UPLOAD_SIZE:
                raise HTTPException(status_code=400, detail=f"文件过大: {safe_name}")
            stored = f"{int(_time.time())}_{safe_name}"
            (upload_dir / stored).write_bytes(content)
            # toolbox 沙箱根为 data/，Agent 调用 read_file 时使用相对 data/ 的路径
            file_infos.append({"name": safe_name, "path": f"uploads/{stored}", "format": ext})

        valid = {"issue_extractor", "logic_reasoner", "regulation_searcher", "report_writer"}
        stage_list = [s.strip() for s in (stages or "").split(",") if s.strip() in valid] or None

        from core.multi_agent import run_audit_pipeline
        result = await run_audit_pipeline(message, file_infos, stage_list)
        return {
            "success": True,
            "data": result,
            "files_processed": len(file_infos),
            "elapsed_seconds": round(_time.time() - t0, 1),
        }
    except HTTPException:
        raise
    except Exception as e:
        return {"success": False, "error": str(e),
                "elapsed_seconds": round(_time.time() - t0, 1)}


# ═══════════════════════════════════════════════════════════════
# 模板引擎 API
# ═══════════════════════════════════════════════════════════════

@router.get("/templates/list")
async def template_list():
    """列出所有可用模板（只返回文件名和目录，不返回内容）"""
    try:
        from core.template_engine import list_templates
        templates = [{"name": t.name, "format": t.format,
                      "category": t.category, "size_mb": t.size_mb}
                     for t in list_templates()]
        return {"success": True, "data": templates}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/templates/match")
async def template_match(intent: str = Form(...)):
    """根据用户意图自动匹配合适的模板"""
    try:
        from core.template_engine import match_template
        result = match_template(intent)
        return {"success": True, "template": result}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/templates/fill")
async def template_fill(
    template_path: str = Form(""),
    data_json: str = Form("{}"),
    user_intent: str = Form(""),
):
    """填充模板并另存，保留全部格式"""
    try:
        data = json.loads(data_json) if data_json else {}
        from core.template_engine import fill_template
        out = fill_template(template_path=template_path, data=data, user_intent=user_intent)
        return {"success": True, "file": out, "message": f"已生成: {Path(out).name}"}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/templates/batch-format")
async def template_batch_format(
    template: UploadFile = File(...),
    targets: List[UploadFile] = File(...),
    format_type: str = Form("auto"),
):
    """
    批量格式规范化：以第一个文件为模板，统一所有其他文件的格式。

    - template: 格式模板文件（Excel或Word）
    - targets: 需要调整格式的一批文件
    - format_type: "excel"/"word"/"auto"（自动识别）

    只改格式，不改数据。
    """
    try:
        import tempfile, shutil, time

        out_dir = str(Path("data/outputs").resolve())
        Path(out_dir).mkdir(parents=True, exist_ok=True)

        # 保存模板
        tmpl_ext = Path(template.filename).suffix.lower()
        tmpl_path = str(Path(out_dir) / f"_tmpl_{int(time.time())}{tmpl_ext}")
        with open(tmpl_path, "wb") as f:
            f.write(await template.read())

        # 保存目标文件
        target_paths = []
        for t in targets:
            tp = str(Path(out_dir) / f"_tgt_{int(time.time())}_{t.filename}")
            with open(tp, "wb") as f:
                f.write(await t.read())
            target_paths.append(tp)

        # 执行批量规范化
        if format_type == "auto":
            format_type = "excel" if tmpl_ext in (".xlsx", ".xls") else "word"

        if format_type == "excel":
            from core.format_engine import batch_normalize_excel
            results = batch_normalize_excel(tmpl_path, target_paths, out_dir)
        else:
            from core.format_engine import batch_normalize_word
            results = batch_normalize_word(tmpl_path, target_paths, out_dir)

        # 清理临时文件
        try: os.unlink(tmpl_path)
        except: pass
        for tp in target_paths:
            try: os.unlink(tp)
            except: pass

        return {
            "success": True,
            "total": len(results),
            "ok": sum(1 for r in results if r["status"] == "ok"),
            "failed": sum(1 for r in results if r["status"] != "ok"),
            "results": results,
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/templates/generate")
async def template_generate(
    data_json: str = Form("{}"),
    user_instruction: str = Form(...),
    template_path: str = Form(""),
    style_guides_json: str = Form("[]"),
):
    """
    规范驱动的报告生成。

    - 有模板时以模板为骨架填充，保留全部格式
    - 有排版规范时LLM严格遵规则生成
    - 两者可同时使用：模板提供结构，规范约束格式
    """
    try:
        data = json.loads(data_json) if data_json else {}
        style_guides = json.loads(style_guides_json) if style_guides_json else []

        from core.template_engine import generate_formatted_report
        out = generate_formatted_report(
            data=data,
            user_instruction=user_instruction,
            template_path=template_path,
            style_guides=style_guides,
        )
        return {"success": True, "file": out, "message": f"已生成: {Path(out).name}"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ═══════════════════════════════════════════════════════════════
# Markdown 文档引擎 API
# ═══════════════════════════════════════════════════════════════

@router.post("/md/parse")
async def md_parse(file: UploadFile = File(...)):
    """上传 md 文件 → 解析为结构化 JSON（标题树/段落/表格/列表）"""
    try:
        text = (await file.read()).decode("utf-8")
        from core.md_engine import parse_md_text
        doc = parse_md_text(text, file.filename)
        return {"success": True, "data": {
            "title": doc.title,
            "toc": [{"level": lv, "title": t} for lv, t in doc.toc],
            "headings": [{"level": n.level, "text": n.text} for n in doc.headings()],
            "tables": len(doc.tables()),
            "nodes": [{"type": n.type, "text": n.text[:200], "level": n.level} for n in doc.nodes[:50]],
            "front_matter": doc.front_matter,
        }}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/md/fill")
async def md_fill(
    template: UploadFile = File(...),
    variables: str = Form("{}"),
):
    """上传模板 + JSON变量 → {{ }} 占位符替换填充"""
    try:
        tmpl = (await template.read()).decode("utf-8")
        vars_dict = json.loads(variables)
        from core.md_engine import fill_template_text
        result = fill_template_text(tmpl, vars_dict)
        return {"success": True, "data": result}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/md/intelligent-fill")
async def md_intelligent_fill(
    template: UploadFile = File(...),
    instruction: str = Form(...),
    data_json: str = Form("{}"),
    out_format: str = Form("md"),
):
    """
    智能填充：上传任意 md 文档 + 自然语言指令 + 数据 → LLM 自动填充并另存。

    例：
      上传"审计报告模板.md"
      instruction="请用下面的差异明细填写第三章的审计结论和第四章的差异明细表"
      data_json='{"差异明细": [...], "结论": "..."}'
      out_format="docx"
      → 返回填充后的 docx 文件路径
    """
    try:
        tmpl = (await template.read()).decode("utf-8")
        data = json.loads(data_json) if data_json else {}
        from core.md_engine import intelligent_fill_text

        # 生成输出路径
        import time
        out_name = f"filled_{int(time.time())}.{out_format}"
        out_path = str(Path("data/downloads") / out_name)
        Path(out_path).parent.mkdir(parents=True, exist_ok=True)

        result = intelligent_fill_text(tmpl, instruction, data, out=out_path, out_format=out_format)
        return {
            "success": True,
            "file": out_path,
            "format": out_format,
            "data": result if out_format == "md" else None,
        }
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/md/generate")
async def md_generate(
    description: str = Form(...),
    data_json: str = Form("{}"),
    template_name: str = Form(""),
):
    """自然语言描述 + 数据 → LLM 动态生成 md 报告"""
    try:
        data = json.loads(data_json)
        from core.md_engine import generate_report
        result = generate_report(description, data, template=template_name)
        return {"success": True, "data": result}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/md/rewrite")
async def md_rewrite(
    file: UploadFile = File(...),
    instruction: str = Form("优化措辞，使其更专业"),
):
    """上传 md 文件 → LLM 改写 → 返回改写后的 md"""
    try:
        text = (await file.read()).decode("utf-8")
        from core.md_engine import rewrite_text
        result = rewrite_text(text, instruction)
        return {"success": True, "data": result}
    except Exception as e:
        return {"success": False, "error": str(e)}


@router.post("/md/convert")
async def md_convert(
    file: UploadFile = File(...),
    target_format: str = Form("docx"),
):
    """上传 md 文件 → 转为 docx/html"""
    try:
        import tempfile, os
        text = (await file.read()).decode("utf-8")
        # 写入临时文件
        with tempfile.NamedTemporaryFile(suffix=".md", delete=False, mode="w", encoding="utf-8") as tmp:
            tmp.write(text)
            tmp_path = tmp.name

        if target_format == "html":
            from core.md_engine import md_to_html
            result = md_to_html(tmp_path)
            os.unlink(tmp_path)
            return {"success": True, "data": result, "format": "html"}
        else:
            from core.md_engine import md_to_docx
            out = tmp_path.replace(".md", ".docx")
            md_to_docx(tmp_path, out)
            os.unlink(tmp_path)
            return {"success": True, "file": out, "format": "docx",
                    "message": "docx 已生成到服务器本地，请通过下载接口获取"}
    except Exception as e:
        return {"success": False, "error": str(e)}


# ═══════════════════════════════════════════════════════════════
# 应用入口
# ═══════════════════════════════════════════════════════════════

def create_app() -> FastAPI:
    app = FastAPI(title="智能审计大脑 v3.0", version="3.0.0")
    app.include_router(router)

    # ── 启动时后台异步预建 RAG 索引（不阻塞服务启动）──
    @app.on_event("startup")
    async def startup_prebuild_rag():
        import asyncio
        loop = asyncio.get_event_loop()
        print("[Startup] 后台预建 RAG 知识库索引（服务已就绪，无需等待）...")

        def _build():
            from core.rag_engine import build_index
            try:
                count = build_index(False)
                print(f"[Startup] RAG 索引就绪: {count} 个文本块")
            except Exception as e:
                print(f"[Startup] RAG 索引预建失败（非致命）: {e}")

        # fire-and-forget：不阻塞启动
        loop.run_in_executor(None, _build)

    return app


if __name__ == "__main__":
    import uvicorn

    app = create_app()
    uvicorn.run(app, host="0.0.0.0", port=8000)