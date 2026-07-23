# Inject docx/pdf helpers into _dag_to_python generated code
path = r"d:\Liu\ai_platform_code\phase2_v3\api\routes.py"
with open(path, "r", encoding="utf-8") as f:
    c = f.read()

# Add docx/pdf helper definitions before the code header
old = '"# === 文件追踪：按序分配 inputs 目录中的文件 ==="'
helpers = '''"",
        "# === docx/pdf 辅助函数 ===",
        "def _read_docx_tables(fp):",
        "    try:",
        "        from docx import Document; doc = Document(fp)",
        "        tables = []",
        "        for i, table in enumerate(doc.tables):",
        "            rows = [[cell.text for cell in row.cells] for row in table.rows]",
        "            if rows:",
        "                df = pd.DataFrame(rows[1:], columns=rows[0])",
        "                tables.append((f'Table_{i+1}', df))",
        "        return tables if tables else None",
        "    except: return None",
        "def _read_docx_text(fp):",
        "    try:",
        "        from docx import Document",
        "        return '\\n'.join(p.text for p in Document(fp).paragraphs if p.text.strip())",
        "    except: return ''",
        "def _read_pdf(fp):",
        "    try:",
        "        import pdfplumber",
        "        with pdfplumber.open(fp) as pdf:",
        "            return '\\n'.join(p.extract_text() or '' for p in pdf.pages)",
        "    except: return ''",
        "",'''

if '_read_docx_tables' in c:
    print("Already injected")
else:
    c = c.replace(old, helpers + old)
    with open(path, "w", encoding="utf-8") as f:
        f.write(c)
    print("Injected")
